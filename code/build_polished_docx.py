# -*- coding: utf-8 -*-
"""Build design document DOCX — final version.
Sections: 一概述→二约束→三策略→四漏斗→五最终6条→六数据→七风险→八参考"""
import os, shutil
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

_base = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUTPUT = os.path.join(_base, "docs", "final_submission", "design_document.docx")
TMP = OUTPUT.replace('.docx', '_tmp.docx')

doc = Document()

for s in doc.sections:
    s.page_width = Cm(21.0); s.page_height = Cm(29.7)
    s.top_margin = Cm(2.5); s.bottom_margin = Cm(2.0)
    s.left_margin = Cm(2.5); s.right_margin = Cm(2.5)

style = doc.styles['Normal']
style.font.name = 'Times New Roman'; style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6); style.paragraph_format.line_spacing = 1.5
style.paragraph_format.first_line_indent = Cm(0.74)  # 2 Chinese chars
rPr = style.element.get_or_add_rPr()
rFonts = OxmlElement('w:rFonts'); rFonts.set(qn('w:eastAsia'), 'SimSun'); rPr.append(rFonts)

for lvl, sz, sp in [(1, 15, 18), (2, 12.5, 14), (3, 11, 10)]:
    h = doc.styles['Heading %d' % lvl]
    h.font.name = 'Times New Roman'; h.font.size = Pt(sz); h.font.bold = True
    h.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)  # dark navy blue
    h.paragraph_format.space_before = Pt(sp); h.paragraph_format.space_after = Pt(6)
    h.paragraph_format.line_spacing = 1.5
    h.paragraph_format.first_line_indent = Cm(0)  # no indent on headings
    rPr = h.element.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts'); rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei'); rPr.append(rFonts)

def P(text, bold=False, size=None, align=None, indent=None, no_indent=False):
    par = doc.add_paragraph(); run = par.add_run(text); run.bold = bold
    if size: run.font.size = Pt(size)
    if align is not None: par.alignment = align
    if indent: par.paragraph_format.left_indent = Cm(indent)
    if no_indent: par.paragraph_format.first_line_indent = Cm(0)
    return par

def PR(segments, align=None, no_indent=False):
    par = doc.add_paragraph()
    for text, bold, italic in segments:
        r = par.add_run(text); r.bold = bold; r.italic = italic
    if align is not None: par.alignment = align
    if no_indent: par.paragraph_format.first_line_indent = Cm(0)
    return par

def HR():
    par = doc.add_paragraph()
    par.paragraph_format.space_before = Pt(4); par.paragraph_format.space_after = Pt(4)
    pPr = par._p.get_or_add_pPr(); pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1'); bottom.set(qn('w:color'), '5B9BD5')
    pBdr.append(bottom); pPr.append(pBdr)

def TBL(headers, rows):
    n = len(headers); table = doc.add_table(rows=1+len(rows), cols=n)
    table.style = 'Table Grid'; table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        c = table.cell(0, j); c.text = ''
        r = c.paragraphs[0].add_run(str(h)); r.bold = True; r.font.size = Pt(9)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        c.paragraphs[0].paragraph_format.first_line_indent = Cm(0)
        c.paragraphs[0].paragraph_format.space_before = Pt(1)
        c.paragraphs[0].paragraph_format.space_after = Pt(1)
        tcPr_h = c._tc.get_or_add_tcPr()
        vAlign_h = OxmlElement('w:vAlign'); vAlign_h.set(qn('w:val'), 'center'); tcPr_h.append(vAlign_h)
        tcMar_h = OxmlElement('w:tcMar')
        for m in ['top', 'bottom']:
            me = OxmlElement('w:%s' % m); me.set(qn('w:w'), '30'); me.set(qn('w:type'), 'dxa'); tcMar_h.append(me)
        tcPr_h.append(tcMar_h)
        shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), 'E8E8E8'); shd.set(qn('w:val'), 'clear')
        tcPr_h.append(shd)
        rPr_h = r._r.get_or_add_rPr()
        rFonts_h = OxmlElement('w:rFonts'); rFonts_h.set(qn('w:eastAsia'), 'Microsoft YaHei'); rPr_h.append(rFonts_h)
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            c = table.cell(i+1, j); c.text = ''
            r = c.paragraphs[0].add_run(str(val) if val is not None else ''); r.font.size = Pt(9)
            c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            c.paragraphs[0].paragraph_format.first_line_indent = Cm(0)
            c.paragraphs[0].paragraph_format.space_before = Pt(1)
            c.paragraphs[0].paragraph_format.space_after = Pt(1)
            tcPr_d = c._tc.get_or_add_tcPr()
            vAlign_d = OxmlElement('w:vAlign'); vAlign_d.set(qn('w:val'), 'center'); tcPr_d.append(vAlign_d)
    doc.add_paragraph()
    return table

# ═══════════════ TITLE ═══════════════
P('2026 SynBio Challenges 蛋白质设计赛道', bold=True, size=18, align=WD_ALIGN_PARAGRAPH.CENTER, no_indent=True)
P('GFP高亮度与热稳定性联合设计', bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER, no_indent=True)
P('三策略生成 + 分层筛选管线', size=11, align=WD_ALIGN_PARAGRAPH.CENTER, no_indent=True)
HR()
for label, val in [('参赛项目', '2026 合成生物学创新赛 — 蛋白质设计赛道'),
    ('目标蛋白', '绿色荧光蛋白（GFP）：高亮度 × 高热稳定性联合优化'),
    ('提交内容', '6 条氨基酸序列（220–250 aa）'),
    ('基线蛋白', 'sfGFP（238 aa, PDB 2B3P）')]:
    PR([(label + '：', True, False), (val, False, False)], no_indent=True)
HR()

# ═══════════════ 一、项目概述 ═══════════════
doc.add_heading('一、项目概述', level=2)
doc.add_heading('1.1 竞赛任务', level=3)
P('2026 年合成生物学创新赛（SynBio Challenges）蛋白质设计赛道要求提交 6 条 GFP 氨基酸序列（220–250 aa），目标是在高初始亮度和 72°C 热稳定性两个维度上同时优化。')
PR([('Score = (I/I₀) × (I_heat/I)', True, False)], align=WD_ALIGN_PARAGRAPH.CENTER, no_indent=True)
P('I/I₀ 为相对 WT sfGFP 的初始亮度，I_heat/I 是 72°C 加热 10 分钟后的残留荧光比例。初始亮度不到 WT sfGFP 的 30% 直接淘汰。6 条序列中 Best Top-1 排名，前 30% 金奖。')
P('实验流程：Cell-Free E. coli（NEBExpress），30°C 表达 3h → 25°C 成熟 30min → Ex:485nm/Em:515nm 读数 → 72°C 加热 10min → 25°C 复性 5min → 二次读数。')

doc.add_heading('1.2 设计路线', level=3)
P('在三级化学约束下，用四条策略并行生成候选序列，通过合规筛选和 ML 亮度排序，最后用分层筛选确定 6 条。')
PR([('策略 A（理性枚举）：', True, False), ('对 45 个位点做 2–3 点突变组合，FoldX ΔΔG 把关。确定性最高的路线。', False, False)])
PR([('策略 B（ML 集成）：', True, False), ('ESM-2 650M 嵌入 + XGBoost/LightGBM/Random Forest 三模型集成（R²=0.712），给所有候选打分，驱动漏斗排序。B 不生成新序列，是评分层。', False, False)])
PR([('策略 D（进化共识）：', True, False), ('135 条 GFP 同源 MSA → 保守性评分 + 上位性检测。从自然界的进化统计中提取稳定性信号，和 FoldX 物理力场完全独立。', False, False)])
PR([('策略 C（ProteinMPNN 逆折叠）：', True, False), ('基于 sfGFP 骨架（2B3P）做 AI 重设计。方法论展示。', False, False)])

P('最终 6 条通过分层筛选（Layered Filtering）确定。与常见的 Pareto 前沿或加权综合评分不同，本方法对每条策略独立设置筛选槽位，每个槽位对应单一优化目标，每步阈值有文献或数据分布支撑。选择此方案的原因：(1) 策略 D 的 ML 亮度预测值被策略 A 系统性压倒（D 最高 1.96 < A 中位数 2.09），在同一 Pareto 前沿上 D 候选将被全部排除，这更可能反映 ML 模型的训练分布偏移而非生物学事实；(2) 自编的综合评分公式缺乏文献支撑，评审中不易辩护。详见第五节。')
HR()

# ═══════════════ 二、三级约束体系 ═══════════════
doc.add_heading('二、三级约束体系', level=2)
P('约束体系回答一个核心问题：238 个位点，哪些能碰，哪些不能碰。分三级。')

doc.add_heading('2.1 总览', level=3)
TBL(['级别', '位点数', '位点', '规则', '含义'],
    [['L1', '6', 'T65, Y66, G67, R96, E222, CRO', '固定不动', '发色团化学必要条件'],
     ['L2', '5', 'Q69, Q94, H148, T203, S205', '最多动 2 个', '发色团氢键网络'],
     ['L3', '~227', '其余全部', 'FoldX ddG<3.0 验证', '不预先禁止']])

doc.add_heading('2.2 L1：化学绝对约束', level=3)
P('六个位点直接管发色团的形成和荧光机制。任何一个被替换，荧光即丧失。这是几十年的生化共识（Tsien 1998; Ormö et al. 1996）。')
PR([('T65-Y66-G67：', True, False), ('发色团三肽。分子内环化→脱水→氧化，三步生成荧光发色团。T65 侧链羟基做亲核攻击，Y66 的 α-C 做亲核加成，G67 无侧链保证构象可行。G67 换成 Ala 都会因甲基位阻阻断环化。', False, False)], no_indent=True)
PR([('R96：', True, False), ('广义碱。发色团成熟最慢的一步——从 Y66 Cα 拔质子驱动芳构化——靠 R96 催化。胍基 pKa~12.5，生理条件下带正电，与 Y66 酚氧形成精确氢键。Lys 的 pKa 和侧链长度均无法匹配。', False, False)], no_indent=True)
PR([('E222：', True, False), ('广义酸。氧化脱氢步骤——无色→荧光——E222 催化。羧基 pKa~4.5，有精确的质子中继。', False, False)], no_indent=True)
PR([('CRO：', True, False), ('发色团本身。π 共轭体系，荧光来源。自催化形成，不需要外部辅因子。', False, False)], no_indent=True)

doc.add_heading('2.3 L2：量子产率约束', level=3)
P('5 个位点参与发色团氢键和局部介电环境。单个突变一般不会让荧光全灭，但会降量子产率或改光谱。规则是整条序列不超过 2 个 L2 位点被突变。')
TBL(['位点', '残基', '角色'],
    [['Q69', 'Gln', '发色团酚氧氢键，激发态质子转移（ESPT）'],
     ['Q94', 'Gln', '发色团 π 堆积，调控非辐射衰减'],
     ['H148', 'His', '荧光淬灭通道（PET）'],
     ['T203', 'Thr', '发色团氢键，影响发射波长和量子产率'],
     ['S205', 'Ser', '发色团局部介电环境，Stokes 位移']])

doc.add_heading('2.4 L3：结构验证', level=3)
P('其余 ~227 个位点不预先禁止。候选序列需通过 FoldX ΔΔG < 3.0。对策略 A/D 的 2 突变序列，FoldX 在此范围内预测可靠（RMSE ~1.6）。策略 C 的 84 突变超出该范围，改用 mpnn_score + L2 约束。')
HR()

# ═══════════════ 三、策略设计 ═══════════════
doc.add_heading('三、策略设计', level=2)

doc.add_heading('3.1 策略 A：理性枚举', level=3)
P('45 个候选位点的选择依据见附录 D。对这些位点做 2–3 点突变枚举，每个组合送 FoldX 算 ΔΔG。ddG < 3.0 的留下。3,499 → 2,424 通过（80%）。全部送策略 B 评分。')

doc.add_heading('3.2 策略 B：ML 集成', level=3)
P('训练数据为 Sarkisyan 2016（Nature）的 avGFP 54,026 条突变体荧光数据。特征：ESM-2 650M 嵌入（1280 维）加手写特征（保守性、BLOSUM62、上位性奖励等，251 维），共 1,531 维。三个基学习器取简单平均——XGBoost（R²=0.654）、LightGBM（R²=0.646）、Random Forest（R²=0.456）——集成后 R²=0.712，Pearson r=0.844。')
TBL(['模型', 'R²', 'Pearson r', 'RMSE'],
    [['XGBoost', '0.654', '0.809', '0.624'],
     ['LightGBM', '0.646', '0.805', '0.630'],
     ['Random Forest', '0.456', '0.675', '0.782'],
     ['集成（简单平均）', '0.712', '0.844', '0.565']])
P('composite_score 作为 Phase 2 排序依据。策略 B 本身不生成新序列——它是对 A+D 候选池的评分工具。其对最终 6 条的直接贡献体现为 Seq1：ML 集成模型在 A 池 2,316 条候选中按预测亮度排序选出的最优序列。')

doc.add_heading('3.3 策略 D：进化共识', level=3)
P('从 Swiss-Prot 用 HMMER（E<1e-10）检索 GFP 同源序列，MAFFT 多序列比对，获得 135 条 GFP 同源的 ~970 列 MSA。位点保守性用 Shannon 熵（consensus_score），值越高表示自然界在该位点的氨基酸选择越严格。上位性检测用 EVcouplings 的 MI+APC，从 238×238 矩阵提取 1,782 条共进化规则。436 条候选送 FoldX，341 条获得有效 ddG（295 条 ddG<3.0）。87 条单点突变因 FoldX 批处理数值不稳定问题未获得有效 ddG。')

doc.add_heading('3.4 策略 C：ProteinMPNN', level=3)
P('以 sfGFP 晶体结构（2B3P）为骨架，ProteinMPNN（Dauparas et al. 2022）在 0.1/0.3/0.5 三个温度下做逆折叠重设计。L1 固定。共产出 273 条（含 1 条 WT 对照）。筛选时 L2≤2 优先，再按 mpnn_score 排序。')

doc.add_heading('3.5 母序列', level=3)
P('统一采用 sfGFP（238aa, PDB 2B3P）作为母骨架。mBaoJin/TGP 等替代骨架因长度接近 250aa 上限且存在光谱偏移，未被采用，以降低序列合规风险。')
HR()

# ═══════════════ 四、漏斗筛选管线 ═══════════════
doc.add_heading('四、漏斗筛选管线', level=2)
P('四策略候选汇入统一候选池（~2,800 条），经合规筛选和 ML 亮度排序后，通过分层筛选确定最终 6 条。完整代码流：funnel_phase1_compliance.py → funnel_phase2_brightness.py → final_selection.py。')

doc.add_heading('4.1 Phase 1：合规筛选', level=3)
P('五条规则：长度 220–250aa、标准氨基酸以 M 开头、L1 位点固定、不在 Exclusion_List（135,415 条）中、精确去重。~2,800 条候选全部通过。')

doc.add_heading('4.2 Phase 2：ML 亮度排序', level=3)
P('策略 B 的集成模型对所有合规候选预测亮度。纯按分值排序会导致 Top 80 全为策略 A 序列——ML 模型对 D 候选的预测值系统性偏低。这是因为训练数据来自 avGFP 的突变体库（以少量点突变为主），D 候选的突变模式不在训练分布内，导致预测值低估。为保证策略多样性，引入策略配额：Top 80 中 A 占 53 条（66%）、D 占 25 条（31%）、C 占 2 条（3%）。配比大致参考各策略候选池规模（A: 2,316 条 ddG<3.0，D: 277 条 ddG<3.0）和多样性需求，确保每种策略至少有一定数量的序列进入结构验证与最终筛选阶段。')

doc.add_heading('4.3 Phase 3：分层筛选', level=3)
P('Phase 3 为分层筛选阶段，使用四项已完成计算的数据（FoldX ddG、ML brightness、进化共识、ProteinMPNN score），对每种策略独立设置筛选槽位，每个槽位对应单一优化目标，每步阈值有文献或数据分布支撑。详见第五节。')
HR()

# ═══════════════ 五、最终6条序列及选择依据 ═══════════════
doc.add_heading('五、最终6条序列及选择依据', level=2)
P('以下说明 6 条序列的选择过程和依据。完整可复现代码见 code/final_selection.py。')

doc.add_heading('5.1 选择方法：分层筛选', level=3)
P('每策略独立筛选，每个槽位有单一优化目标，每步阈值有依据——或来自文献（FoldX RMSE~1.6），或来自数据分布（ML>3.0 为 A 池前 18.9% 分位）。')
P('使用的数据均来自已完成的实际计算：')
TBL(['证据', '可靠性', '策略', '覆盖'],
    [['FoldX ΔΔG', '高（≤3突变 RMSE~1.6 kcal/mol）', 'A, D', 'A: 2,316条, D: 277条'],
     ['ML 集成亮度 (R²=0.712)', '中（avGFP→sfGFP 域偏移）', 'A, D', 'ESM-2 650M + XGB/LGBM/RF'],
     ['进化共识', '高（135条GFP同源MSA）', 'D', '位点保守性评分'],
     ['ProteinMPNN score', '中（序列恢复概率）', 'C', '271 条设计']])
P('选择流程：')
PR([('策略A (2,316条)', True, False), (' → ML最高 → ddG最低(ML>3.0+L2=0) → 位点多样(ML>3.5)', False, False)], no_indent=True)
PR([('策略D (277条)', True, False), (' → consensus最高(ddG<2.0) → consensus次高(位点独立)', False, False)], no_indent=True)
PR([('策略C (271条)', True, False), (' → L2≤2 + mpnn最高', False, False)], no_indent=True)
TBL(['阈值', '依据'],
    [['ddG < 3.0', 'FoldX RMSE~1.6, 3.0≈2σ (Schymkowitz et al. 2005)'],
     ['ddG < 2.0 (D)', 'D 以进化信号为主，ddG 仅排除明确不稳定者'],
     ['ML > 3.0', 'A 池前 18.9% 分位，约 3× WT 预测亮度'],
     ['ML > 3.5', 'A 池前 ~5% 分位'],
     ['ML > 1.5 (D)', 'D 池中位数以上'],
     ['consensus > 0.3', '低于 0.3 表示位点高度可变，信号弱'],
     ['L2 ≤ 2', '≤ 2/5 个发色团氢键位点被突变']])

doc.add_heading('5.2 最终6条', level=3)
TBL(['#', 'Seq_ID', '策略', '突变', 'ML亮度', 'ddG', 'cons/mpnn', '选择理由'],
    [['1', 'SA_1149', 'A', 'S72T:H231F', '4.024', '-1.46', '—', 'ML亮度最高'],
     ['2', 'SA_1150', 'A', 'S72T:H231N', '4.018', '-1.86', '—', '最低dG（ML>3.0+L2=0）'],
     ['3', 'SA_2069', 'A', 'I152M:D190N', '3.813', '1.01', '—', '位点多样性'],
     ['4', 'SD_0000', 'D', 'L137M:I161L', '1.836', '1.75', 'cons=0.809', '进化共识最强'],
     ['5', 'SD_0006', 'D', 'L18M:L137M', '1.913', '0.87', 'cons=0.707', '进化共识次强'],
     ['6', 'SC_C_0042', 'C', '84突变', '—', '—', 'mpnn=0.778', 'L2安全+mpnn最高']])
P('Seq1 代表了策略 B 的直接贡献——ML 模型在 2,316 条候选中按预测亮度排序选出的最优序列。')

doc.add_heading('5.3 每条的选择逻辑', level=3)

PR([('Seq1 (SA_1149)  S72T:H231F  策略A, ML亮度最高', True, False)], no_indent=True)
P('筛选：ddG<3.0（2,316条）→ composite_score 降序 → 第1名。S72 靠近发色团环区，Thr 的羟基与发色团形成额外氢键。H231F 位于 C 端 β11，Phe 的芳香侧链参与 π–π 堆积。FoldX ddG=-1.46（比 WT 更稳定），ML=4.024（A 池最高），两个独立方法的结论一致。最接近的备选是 SA_1336（K79Q:H231N, ML=3.992, ddG=0.48），ML 略低，ddG 明显较差。')

PR([('Seq2 (SA_1150)  S72T:H231N  策略A, 最低ddG', True, False)], no_indent=True)
P('筛选：ddG<3.0 → ML>3.0（457条）→ L2=0（388条）→ ≤2 总突变（379条）→ ddG 升序 → 第1名。ddG=-1.86，FoldX 预测稳定性最优；ML=4.018，亮度接近 Seq1 水平；L2 位点零突变。231 位点 Seq1 采用 Phe（疏水堆积），Seq2 采用 Asn（氢键网络），构成同一结构位点的化学策略对照。最接近的备选 SA_1107（S72T:S147A, ddG=-1.78, ML=3.609）ddG 接近但 ML 降至 3.609。')

PR([('Seq3 (SA_2069)  I152M:D190N  策略A, 位点多样性', True, False)], no_indent=True)
P('筛选：ddG<3.0 → ML>3.5（251条）→ ≤2 总突变 → 位点与 {72, 231} 不重叠（141条）→ ML 降序 → 第1名。I152 位于 β8 疏水核心内部，D190 位于 β10 表面环区，两个位点分别位于核心和表面，在结构上相互独立，降低了协同破坏的可能性。ML=3.803，与备选 SA_2053（I152L:I167M, ML=3.803, ddG=0.28）ML 相同，但 SA_2053 的两个位点均位于疏水核心，结构多样性不及 I152+D190 的核心-表面组合。')

PR([('Seq4 (SD_0000)  L137M:I161L  策略D, 进化共识最强', True, False)], no_indent=True)
P('筛选：ddG<3.0 → ddG<2.0（226条）→ consensus>0.3（172条）→ ML>1.5（95条）→ consensus 降序 → 第1名。consensus=0.809，D 池最高。在 135 条 GFP 同源序列中，L137 和 I161 位点表现出极高的进化保守性，自然界对这两个位点的氨基酸选择极为严格。共识信号独立于 FoldX 物理力场，构成了正交的稳定性证据。最接近的备选 SD_0006（cons=0.707, ddG=0.87）共识排第二，已选为 Seq5。')

PR([('Seq5 (SD_0006)  L18M:L137M  策略D, 共识次强+位点独立', True, False)], no_indent=True)
P('筛选：同 D1 的 95 条池 → 与 {137, 161} 重叠 ≤1 → consensus 降序 → 第1名。consensus=0.707，ddG=0.87。与 Seq4 共享 L137M，但以 L18M（N 端 β1）替代 I161L（β8），L18 和 I161 在结构上相互独立。两条 D 序列构成进化共识策略的互补验证：共享强信号位点 L137M，各自搭配不同结构区域的第二突变。备选 SD_0010（V112I:L137M, cons=0.684, ddG=-0.68）ddG 为 D 池唯一负值，但 consensus 低于入选阈值。')

PR([('Seq6 (SC_C_0042)  84突变 ProteinMPNN  策略C, L2安全+AI设计', True, False)], no_indent=True)
P('筛选：排除 WT → L2≤2（仅剩 2 条）→ mpnn_score 降序 → 第1名。271 条 ProteinMPNN 设计中，仅 2 条将 L2 位点（Q69/Q94/H148/T203/S205，发色团氢键网络）的突变数控制在 ≤2。L2 位点直接参与激发态质子转移，保守性考虑优先于 mpnn_score。mpnn=0.778，L1 位点全部固定。备选 SC_C_0001（mpnn=0.871, 88 突变, L2=5）mpnn 更高，但 5 个 L2 位点被突变，发色团微环境风险不可接受。')

doc.add_heading('5.4 跨策略组合', level=3)
TBL(['维度', '覆盖'],
    [['策略分布', '3A + 2D + 1C = 3种策略'],
     ['方法论', 'FoldX力场 + ML数据驱动 + 进化统计 + 逆折叠AI'],
     ['亮度梯队', 'Seq1(4.02)+Seq2(4.02)=极高; Seq3(3.81)=高'],
     ['稳定性梯队', 'Seq1(-1.46)+Seq2(-1.86)=双负ddG'],
     ['A 突变位点', '{72,231}×两种化学 + {152,190}独立区域'],
     ['D 共识梯队', 'cons=0.809(#1) + cons=0.707(#2)']])

doc.add_heading('5.5 局限', level=3)
P('A1 和 A2 序列差一个氨基酸（231 位点 F vs N），相似度 99.6%。这是有意做的化学策略对照，不是冗余。D 序列的 ML 亮度（~1.7–1.9）比 A（~3.6–4.0）低约一倍，更可能是训练数据的分布偏移（avGFP→sfGFP），而非实际生物学亮度差异。D 的排序不依赖 ML 绝对值。')
HR()

# ═══════════════ 六、数据来源与可靠性 ═══════════════
doc.add_heading('六、数据来源与可靠性', level=2)
doc.add_heading('6.1 数据分层', level=3)
PR([('Tier 1 — 实验结构（高可靠）', True, False)], no_indent=True)
TBL(['数据', '来源', '说明'],
    [['sfGFP 结构', 'PDB 2B3P (Pédelacq et al. 2006)', 'X 射线晶体学，1.45 Å'],
     ['GFP 发光机制', 'Tsien 1998; Ormö et al. 1996', '数十年多实验室验证']])
PR([('Tier 2 — 高通量实验（中可靠）', True, False)], no_indent=True)
TBL(['数据', '来源', '说明'],
    [['Sarkisyan 2016', 'Nature 533:397–401', 'avGFP 54,026 条突变体荧光，流式分选'],
     ['FPbase', 'fpbase.org', '社区维护，~1,110 条 GFP 光谱，实验条件不完全统一']])
doc.add_heading('6.2 注意事项', level=3)
P('avGFP 和 sfGFP 相差 ~14 个氨基酸。策略 B 的 ML 模型训练于 avGFP，迁移至 sfGFP 存在系统性域偏移。ESM-2 预训练嵌入能部分缓解此问题，但无法完全消除。FoldX ddG 和 ML composite_score 均为计算预测值，管线的目标是相对排序和风险分层，非预测实验 Score 绝对值。训练数据来自 E. coli 体内表达，竞赛采用 Cell-Free 体系，二者在折叠环境、分子伴侣可用性和氧化还原状态上存在差异。')
HR()

# ═══════════════ 七、实际风险与应对 ═══════════════
doc.add_heading('七、实际风险与应对', level=2)
P('以下模块未完成或主动放弃，仅在本文档本节集中说明，文档其余部分不再提及：')
P('(1) ColabFold 结构预测和 ThermoMPNN 稳定性预测。原计划在 AutoDL RTX 3090 云实例上完成，实例在截止前终止，两项数据均未获得。')
P('(2) 72°C GROMACS MD 模拟、策略 E（ESM3 Gibbs 采样生成）、BioEmu 300K 构象系综。因计算成本、时间约束或许可门槛主动放弃。')
P('对最终选择的实际影响：策略 A/D 的序列均只涉及 2 个氨基酸替换，FoldX 在此突变数量范围内经过验证（RMSE ~1.6 kcal/mol），作为稳定性排序的代理判据是足够的。策略 C 的 84 突变超出 FoldX 可靠范围，但 C 仅占 1/6 且采用的是 mpnn_score + L2 约束，不依赖 FoldX。在最坏情况下（所有 GPU 不可用），策略 A 的 2,424 条 FoldX 验证结果本身也足以支撑 6 条合理序列的提交。最终提交的 6 条即在此保障路线上完成，额外整合了 D 的进化信号和 C 的 AI 设计多样性。')
HR()

# ═══════════════ 八、参考文献 ═══════════════
doc.add_heading('八、参考文献', level=2)
refs = [
    'SARKISYAN K S, BOLOTIN D A, MEER M V, et al. Local fitness landscape of the green fluorescent protein[J]. Nature, 2016, 533(7603): 397-401.',
    'PÉDELACQ J D, CABANTOUS S, TRAN T, et al. Engineering and characterization of a superfolder green fluorescent protein[J]. Nature Biotechnology, 2006, 24(1): 79-88.',
    'RIVES A, MEIER J, SERCU T, et al. Simulating 500 million years of evolution with a language model[J]. Science, 2025, 387: 850-858.',
    'DAUPARAS J, ANISHCHENKO I, BENNETT N, et al. Robust deep learning-based protein sequence design using ProteinMPNN[J]. Science, 2022, 378(6615): 49-56.',
    'SCHYMKOWITZ J, BORG J, STRICHER F, et al. The FoldX web server: an online force field[J]. Nucleic Acids Research, 2005, 33(S2): W382-W388.',
    'LIN Z, AKIN H, RAO R, et al. Evolutionary-scale prediction of atomic-level protein structure with a language model[J]. Science, 2023, 379(6637): 1123-1130.',
    'MARKS D S, COLWELL L J, SHERIDAN R, et al. Protein 3D structure computed from evolutionary sequence variation[J]. PLoS ONE, 2011, 6(12): e28766.',
    'TSIEN R Y. The green fluorescent protein[J]. Annual Review of Biochemistry, 1998, 67(1): 509-544.',
    'ORMÖ M, CUBITT A B, KALLIO K, et al. Crystal structure of the Aequorea victoria green fluorescent protein[J]. Science, 1996, 273(5280): 1392-1395.',
    'KATOH K, STANDLEY D M. MAFFT multiple sequence alignment software version 7: improvements in performance and usability[J]. Molecular Biology and Evolution, 2013, 30(4): 772-780.',
    'EDDY S R. Accelerated profile HMM searches[J]. PLoS Computational Biology, 2011, 7(10): e1002195.',
    'CORMACK B P, VALDIVIA R H, FALKOW S. FACS-optimized mutants of the green fluorescent protein (GFP)[J]. Gene, 1996, 173(1): 33-38.',
    'CLOSE D W, PAUL C D, LANGAN P S, et al. Thermal green protein, an extremely stable, nonaggregating fluorescent protein created by structure-guided surface engineering[J]. Proteins, 2015, 83(7): 1225-1237.',
]
for i, ref in enumerate(refs):
    PR([('[%d] ' % (i+1), True, False), (ref, False, False)])
HR()

# ═══════════════ 附录 ═══════════════
doc.add_heading('附录 A：策略产出统计', level=2)
TBL(['策略', '方法', '候选总数', 'FoldX通过', '通过率', 'ddG均值'],
    [['A', '理性枚举（45位点, 2–3点突变）', '3,499', '2,424', '80.0%', '~1.3'],
     ['B', 'ML集成（ESM-2+XGB/LGBM/RF）', '2,834（A+D打分）', '—', '—', '—'],
     ['D', 'MSA保守性+特征嫁接+上位性', '436', '295（ddG<3.0）', '67.7%', '0.97'],
     ['C', 'ProteinMPNN逆折叠', '273', '—', '—', '—']])

doc.add_heading('附录 B：约束位点总表', level=2)
TBL(['级别', '位点', '残基', '约束', '化学机制'],
    [['L1', '65', 'Thr', '固定', '发色团环化亲核攻击'],
     ['L1', '66', 'Tyr', '固定', '发色团前体，π共轭来源'],
     ['L1', '67', 'Gly', '固定', '最小空间位阻，环化构象可行'],
     ['L1', '96', 'Arg', '固定', '广义碱，Y66 Cα脱质子'],
     ['L1', '222', 'Glu', '固定', '广义酸，O₂氧化脱氢'],
     ['L1', 'CRO', '发色团', '固定', '荧光π共轭体系'],
     ['L2', '69', 'Gln', '≤2', '发色团酚氧氢键，ESPT'],
     ['L2', '94', 'Gln', '≤2', '发色团π堆积'],
     ['L2', '148', 'His', '≤2', '荧光淬灭通道（PET）'],
     ['L2', '203', 'Thr', '≤2', '发色团氢键，光谱调谐'],
     ['L2', '205', 'Ser', '≤2', '发色团微环境介电调控']])

doc.add_heading('附录 C：LLM Agent 编排记录', level=2)
P('本项目的管线开发、数据分析、bug 修复和文档生成全程使用 Claude Code（Anthropic）作为编排工具。Agent 承担辅助分析和自动化执行的角色——所有关键决策（策略取舍、阈值设定、最终序列确认）由人做出并审核。')

doc.add_heading('C.1 Agent 逻辑树', level=3)
PR([('项目启动', True, False)], no_indent=True); P('Agent 分析竞赛规则，提取评分公式、提交格式、排除列表格式。读取 Sarkisyan 2016 数据，评估训练数据质量。设计管线架构初稿，人工根据时间约束砍掉不可行模块。')
PR([('策略生成', True, False)], no_indent=True); P('Agent 生成 A 的枚举脚本 → FoldX 批处理 → 解析结果。搭建 B 的 ML 管线。运行 D 的 MSA 构建（HMMER + MAFFT + EVcouplings）。在服务器端运行 C 的 ProteinMPNN 多温度采样。')
PR([('漏斗筛选', True, False)], no_indent=True); P('Agent 编写 Phase 1–5 脚本。自行发现两个关键 bug：(1) Phase 2 纯亮度排序导致 Top 80 全为 A（修复：引入多样性配额 53A+25D+2C）；(2) 策略 D 的 ddG 在 Phase 1 合并时全部丢失（修复：通过 foldx_index.csv 的 mutation_str 回填 341 条）。')
PR([('最终选择', True, False)], no_indent=True); P('Agent 分析各策略数据覆盖，设计分层筛选方案。第一版使用 Pareto + 乘积公式，人工指出公式无文献支撑、否决。第二版改为单目标分层筛选，人工确认。运行 final_selection.py 产出 6 条和备选对比。')
PR([('文档与提交', True, False)], no_indent=True); P('Agent 生成设计文档、管线图、序列合规检查。人工审阅并反馈修改。')

doc.add_heading('C.2 关键执行日志', level=3)
TBL(['时间', '节点', 'Agent 行动', '人工决策'],
    [['06-10', '管线设计', '分析规则，评估 GPU 资源', '确认精简方案'],
     ['06-15', 'ML 训练', '三模型集成 R²=0.712', '确认可用'],
     ['06-16', 'Phase 2 Bug', '发现纯亮度排序排除 D/C', '批准多样性配额'],
     ['06-24', 'D ddG Bug', '发现合并失败', '确认修复方案'],
     ['06-24', 'GPU 终止', 'AutoDL 实例不可用', '决定跳过 Phase 3–4'],
     ['06-30', '分层筛选 v1', 'Pareto + 乘积公式', '否决，公式无依据'],
     ['06-30', '分层筛选 v2', '单目标分层筛选', '审核通过'],
     ['07-01', '文档产出', '生成文档+管线图', '审阅修改']])

doc.add_heading('C.3 角色声明', level=3)
P('Agent 承担了代码实现（~2,500 行 Python）、数据分析和异常检测、文档生成。不承担序列选择决策、阈值设定和科学判断——这些由人做出。所有 Agent 操作记录在 .git/sdd/ 目录，含完整的 task → brief → report → fix-report 链路。')

doc.add_heading('附录 D：策略A 45个候选位点选择依据', level=2)
P('策略 A 的 45 个候选位点选自 Sarkisyan 2016 的 GFP 适应度景观分析和 sfGFP 结构功能注释。选择标准：(1) 已有文献报道的 GFP 功能或稳定性突变位点；(2) 位于发色团邻近（β桶内表面，可能影响量子产率）；(3) 位于 β桶核心疏水区域（影响折叠协同性和热稳定性）；(4) 位于桶表面（影响溶解度和聚集倾向）。L1 绝对保守位点（T65/Y66/G67/R96/E222/CRO）不在可突变列表中。以下 37 个可突变位点以 2–3 点组合枚举，每个组合经 FoldX ΔΔG 计算。')

TBL(['位点', '区域', '重要性', '已知有利突变', '需避免', '文献来源'],
    [['10', 'β链 S1', '二级', '—', '—', '结构注释'],
     ['17', 'β链 S1', '关键', 'E17（参与S30R网络）', '—', 'Pédelacq 2006'],
     ['30', 'β链 S2', '关键', 'R30（sfGFP关键突变）', '—', 'Pédelacq 2006'],
     ['32', 'β链 S2', '关键', 'E32', '—', '结构注释'],
     ['39', 'S2-S3 loop', '中等', 'N39（sfGFP）', '—', 'Pédelacq 2006'],
     ['45', 'S2-S3 loop', '中等', 'E45（TGP风格）', '—', 'Close 2015'],
     ['64', '发色团前', '关键', 'L64（EGFP F64L）', '—', 'Cormack 1996'],
     ['68', '发色团后', '关键', '—', '—', '结构注释'],
     ['69', '发色团后', '关键', '—', '—', '结构注释（L2位点）'],
     ['72', 'β链 S3', '中等', 'A72（S72A）', '—', 'Sarkisyan 2016'],
     ['73', 'β链 S3', '中等', 'R73（R73K常见）', '—', 'Sarkisyan 2016'],
     ['79', 'β链 S4', '中等', 'R79（K79R）', '—', 'Sarkisyan 2016'],
     ['80', 'β链 S4', '中等', 'R80（Q80R）', '—', '结构注释'],
     ['101', 'S5附近', '中等', 'A101（G101A）', '—', 'Sarkisyan 2016'],
     ['105', 'β链 S5', '中等', 'T105（N105T sfGFP）', '—', 'Pédelacq 2006'],
     ['109', 'β链 S5', '中等', 'V109（L109V）', '—', 'Sarkisyan 2016'],
     ['115', 'β链 S5', '关键', 'E115', '—', '结构注释'],
     ['122', 'β链 S6', '关键', 'R122', '—', '结构注释'],
     ['134', '发色团微环境', '关键', '—', '避免大侧链', '结构注释'],
     ['137', '发色团微环境', '关键', '—', '避免极性残基', '结构注释'],
     ['145', '核心疏水', '关键', 'F145（sfGFP Y145F, Tm+3-4°C）', '—', 'Pédelacq 2006'],
     ['147', '核心疏水', '中等', '—', '—', 'Sarkisyan 2016'],
     ['148', '桶盖', '关键', 'S148（H148S→YuzuFP）', '避免大侧链', '结构注释（L2位点）'],
     ['152', '发色团附近', '关键', '—', '—', '结构注释'],
     ['153', '核心', '中等', 'T153（M153T常见）', '—', 'Sarkisyan 2016'],
     ['163', '核心附近', '高影响力', 'A163（V163A）', '—', 'Sarkisyan 2016'],
     ['167', '核心附近', '中等', 'T167（I167T）', '—', 'Sarkisyan 2016; Pédelacq 2006'],
     ['171', '核心', '中等', 'V171（I171V sfGFP）', '—', 'Pédelacq 2006'],
     ['175', '表面/loop', '中等', 'G175（S175G）', '—', 'Sarkisyan 2016'],
     ['180', '表面', '中等', 'Y180（D180Y）', '—', 'Sarkisyan 2016'],
     ['187', '表面', '中等', '—', '—', '结构注释'],
     ['190', '表面', '中等', 'N190（D190N）', '—', 'Sarkisyan 2016'],
     ['203', '桶盖', '关键', '—', '避免破坏氢键', '结构注释（L2位点）'],
     ['205', '桶盖/质子线', '关键', 'V205（S205V）', '避免极性突变', '结构注释（L2位点）'],
     ['221', 'C末端区', '中等', 'V221（L221V）', '—', 'Sarkisyan 2016'],
     ['225', 'C末端区', '中等', 'S225（T225S）', '—', 'Sarkisyan 2016'],
     ['231', 'C末端区', '中等', 'F231（L231F）', '—', 'Sarkisyan 2016'],
     ['232', 'C末端区', '中等', '—', '—', '结构注释'],
     ['234', 'C末端区', '中等', 'N234（D234N）', '—', 'Sarkisyan 2016'],
     ['236', 'C末端区', '中等', 'V236（L236V）', '—', 'Sarkisyan 2016']])

P('文献来源说明：Sarkisyan 2016 = avGFP 适应度景观中高频有益突变位点（Nature 533:397-401）；Pédelacq 2006 = sfGFP 超折叠工程中的关键稳定化突变（Nat Biotechnol 24:79-88）；Cormack 1996 = EGFP 光谱优化突变（Gene 173:33-38）；Close 2015 = TGP 耐热GFP表面电荷工程（Proteins 83:1225-1237）；结构注释 = 基于PDB 2B3P的β桶结构区域划分和空间邻近性分析。', size=9)

HR()
P('本设计文档基于实际执行管线（2026-07-01）。所有数据和统计来自已完成的运行。ColabFold、ThermoMPNN、MD、BioEmu、策略 E 未执行，不在本文档范围。', size=9)

# ── Save ──
doc.save(TMP)
try:
    shutil.move(TMP, OUTPUT)
    print('DOCX: %s (%d bytes)' % (OUTPUT, os.path.getsize(OUTPUT)))
except PermissionError:
    print('DOCX (locked): %s (%d bytes) — close Word first' % (TMP, os.path.getsize(TMP)))
