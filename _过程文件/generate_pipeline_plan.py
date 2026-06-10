# -*- coding: utf-8 -*-
"""生成 2026合成生物学创新赛 GFP蛋白质设计 管线设计与实施方案 Word文档"""

import datetime
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# ============================================================
# 文档设置
# ============================================================
doc = Document()

# -- 页面设置 A4 --
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)

# -- 样式设置 --
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.paragraph_format.space_before = Pt(3)
style.paragraph_format.space_after = Pt(3)
style.paragraph_format.line_spacing = 1.5

# ============================================================
# 辅助函数
# ============================================================

def set_cn_font(run, font_name='宋体'):
    """设置中文字体"""
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

def set_en_font(run, font_name='Times New Roman'):
    """设置英文字体"""
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def add_h1(text):
    """一级标题: 微软雅黑 小二(18pt) 粗体"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(text)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)
    p.style = doc.styles['Heading 1']
    return p

def add_h2(text):
    """二级标题: 微软雅黑 小三(15pt) 粗体"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(15)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 80, 130)
    p.style = doc.styles['Heading 2']
    return p

def add_h3(text):
    """三级标题: 微软雅黑 四号(14pt) 粗体"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(46, 117, 182)
    p.style = doc.styles['Heading 3']
    return p

def add_p(text, bold=False, indent=False):
    """正文段落: 中文宋体 英文TNR 小四(12pt)"""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    if indent:
        p.paragraph_format.first_line_indent = Pt(24)
    run = p.add_run(text)
    set_en_font(run)
    run.font.size = Pt(12)
    run.font.bold = bold
    return p

def add_rich_p(runs_list, indent=False):
    """富文本段落: runs_list = [(text, bold, color), ...]"""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    if indent:
        p.paragraph_format.first_line_indent = Pt(24)
    for item in runs_list:
        if isinstance(item, str):
            text, is_bold, color = item, False, None
        elif len(item) == 2:
            text, is_bold = item
            color = None
        else:
            text, is_bold, color = item
        run = p.add_run(text)
        set_en_font(run)
        run.font.size = Pt(12)
        run.font.bold = is_bold
        if color:
            run.font.color.rgb = RGBColor(*color)
    return p

def add_bullet(text, level=0):
    """项目符号"""
    p = doc.add_paragraph(style='List Bullet')
    p.clear()
    run = p.add_run(text)
    set_en_font(run)
    run.font.size = Pt(12)
    if level > 0:
        p.paragraph_format.left_indent = Cm(1.5 * (level + 1))
    return p

def add_numbered(text, level=0):
    """编号列表"""
    p = doc.add_paragraph(style='List Number')
    p.clear()
    run = p.add_run(text)
    set_en_font(run)
    run.font.size = Pt(12)
    return p

def add_bullet_rich(runs_list, level=0):
    """富文本项目符号"""
    p = doc.add_paragraph(style='List Bullet')
    p.clear()
    if level > 0:
        p.paragraph_format.left_indent = Cm(1.5 * (level + 1))
    for item in runs_list:
        if isinstance(item, str):
            text, is_bold, color = item, False, None
        elif len(item) == 2:
            text, is_bold = item
            color = None
        else:
            text, is_bold, color = item
        run = p.add_run(text)
        set_en_font(run)
        run.font.size = Pt(12)
        run.font.bold = is_bold
        if color:
            run.font.color.rgb = RGBColor(*color)
    return p

def add_code(text):
    """代码块: Consolas 小五(9pt) 灰底"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.2
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F5" w:val="clear"/>')
    p.paragraph_format.element.get_or_add_pPr().append(shading_elm)
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(51, 51, 51)
    return p

def add_page_break():
    """分页"""
    doc.add_page_break()

def make_table(headers, rows, col_widths=None):
    """创建表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(header)
        set_en_font(run)
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 51, 102)
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="D6E4F0" w:val="clear"/>')
        cell._element.get_or_add_tcPr().append(shading)

    # 数据行
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = ''
            run = cell.paragraphs[0].add_run(str(val))
            set_en_font(run)
            run.font.size = Pt(11)

    # 列宽
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)

    doc.add_paragraph()  # 表后空行
    return table


# ============================================================
# 封面
# ============================================================
for _ in range(6):
    doc.add_paragraph()

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_p.add_run('2026合成生物学创新赛')
set_cn_font(run, '微软雅黑')
run.font.size = Pt(26)
run.font.color.rgb = RGBColor(0, 51, 102)

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub_p.add_run('蛋白质设计赛道')
set_cn_font(run, '微软雅黑')
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(0, 80, 130)

doc.add_paragraph()

main_p = doc.add_paragraph()
main_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = main_p.add_run('GFP 高亮度与热稳定性联合设计\n管线设计与实施方案')
set_cn_font(run, '微软雅黑')
run.font.size = Pt(30)
run.font.bold = True
run.font.color.rgb = RGBColor(0, 51, 102)

doc.add_paragraph()
doc.add_paragraph()

subtitle_p = doc.add_paragraph()
subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle_p.add_run('（含零基础预实验 · 30分钟快速上手）')
set_cn_font(run, '微软雅黑')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(80, 80, 80)

for _ in range(4):
    doc.add_paragraph()

info_p = doc.add_paragraph()
info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info_p.add_run(f'版本：v2.0  |  生成日期：{datetime.date.today().strftime("%Y年%m月%d日")}')
set_cn_font(run, '微软雅黑')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(80, 80, 80)

deadline_p = doc.add_paragraph()
deadline_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = deadline_p.add_run('竞赛截止日期：2026年7月1日  |  剩余时间：24天')
set_cn_font(run, '微软雅黑')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(204, 0, 0)

add_page_break()

# ============================================================
# 目录页
# ============================================================
add_h1('目录')

toc_items = [
    ('第〇部分：零基础预实验 —— 30分钟快速上手', ''),
    ('    0.1  预实验目标', ''),
    ('    0.2  环境准备（一次性，约15分钟）', ''),
    ('    0.3  预实验运行（约10分钟）', ''),
    ('    0.4  预实验常见问题排查', ''),
    ('    0.5  预实验交付物', ''),
    ('第一部分：项目背景与竞赛约束', ''),
    ('    1.1  竞赛基本信息', ''),
    ('    1.2  评分机制深度解析', ''),
    ('    1.3  关键生物学约束', ''),
    ('    1.4  已有工作基础', ''),
    ('第二部分：正式管线总体架构', ''),
    ('    2.1  设计哲学：三路线并行 + 集成筛选', ''),
    ('    2.2  管线总体流程（五阶段）', ''),
    ('    2.3  技术栈总览', ''),
    ('第三部分：分步骤详细实施方案', ''),
    ('    阶段一：数据准备与增强（第1-3天）', ''),
    ('    阶段二A：路线A —— PLM嵌入 + ML回归 + 约束采样（第2-6天）', ''),
    ('    阶段二B：路线B —— ProteinMPNN结构条件设计（第3-7天）', ''),
    ('    阶段二C：路线C —— 进化分析引导理性设计（第2-5天）', ''),
    ('    阶段三：多维度虚拟筛选（第6-10天）', ''),
    ('    阶段四：Top候选精选（第10-13天）', ''),
    ('    阶段五：最终验证与提交准备（第13-18天）', ''),
    ('第四部分：24天时间规划', ''),
    ('    4.1  甘特图式时间表', ''),
    ('    4.2  关键时间节点', ''),
    ('    4.3  资源需求预估', ''),
    ('第五部分：风险识别与应对策略', ''),
    ('附录', ''),
    ('    附录A：关键命令速查表', ''),
    ('    附录B：项目文件结构建议', ''),
    ('    附录C：核心参考文献速查', ''),
]

for item, _ in toc_items:
    p = doc.add_paragraph()
    run = p.add_run(item)
    set_cn_font(run, '微软雅黑' if not item.startswith('    ') else '宋体')
    run.font.size = Pt(14) if not item.startswith('    ') else Pt(12)
    if not item.startswith('    '):
        run.font.bold = True
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)

add_page_break()

# ============================================================
# 第〇部分：零基础预实验
# ============================================================
add_h1('第〇部分：零基础预实验 —— 30分钟快速上手')

add_h2('0.1 预实验目标')
add_p('在进入正式管线设计之前，通过一个极简流程完成以下目标：')
add_bullet('理解竞赛数据格式（训练数据、野生型序列、排除列表）')
add_bullet('运行一次完整的"嵌入 → 训练 → 预测 → 输出"流程')
add_bullet('熟悉 Python 环境和依赖库的安装')
add_bullet('获得第一组基线预测结果（为后续优化提供参照）')
add_bullet('验证 GPU/CPU 环境是否正常工作')

add_h2('0.2 环境准备（一次性，约15分钟）')

add_h3('0.2.1 安装 Python 依赖')
add_p('打开终端（Anaconda Prompt 或 PowerShell），依次执行以下命令：')
add_code('# 1. 创建并激活虚拟环境（推荐）')
add_code('conda create -n gfp_design python=3.10 -y')
add_code('conda activate gfp_design')
add_code('')
add_code('# 2. 安装 PyTorch（CPU 版本，零基础友好；有 NVIDIA GPU 则去掉 --index-url）')
add_code('pip install torch --index-url https://download.pytorch.org/whl/cpu')
add_code('')
add_code('# 3. 安装其余依赖')
add_code('pip install fair-esm openpyxl scikit-learn pandas numpy')
add_code('pip install transformers sentencepiece    # SaProt 备用')

add_h3('0.2.2 验证安装')
add_p('在 Python 中逐行运行以下代码，确认无报错：')
add_code('import torch')
add_code('import esm')
add_code('import pandas as pd')
add_code('import numpy as np')
add_code('from sklearn.ensemble import RandomForestRegressor')
add_code("print('All imports successful!')")
add_code(f"print(f'PyTorch version: {{torch.__version__}}')")
add_code(f"print(f'CUDA available: {{torch.cuda.is_available()}}')")

add_h3('0.2.3 确认数据文件')
add_p('确保以下文件存在于项目目录中（默认路径 ./data/）：')
make_table(
    ['文件名', '说明', '必需'],
    [
        ['GFP data.xlsx', 'GFP 突变体亮度训练数据（含 brightness sheet）', '是'],
        ['WildType AAseqs of 4 GFP proteins.txt', '4 种 GFP 的野生型氨基酸序列（FASTA）', '是'],
        ['Exclusion_List.csv', '禁止提交的序列黑名单', '是'],
    ],
    [4.5, 7.5, 2.0]
)

add_h2('0.3 预实验运行（约10分钟）')

add_h3('0.3.1 最小化运行命令')
add_p('使用最保守的参数运行管线，确保能在普通笔记本上完成：')
add_code('cd D:\\蛋白质设计-合成生物学创新赛-Claude')
add_code('')
add_code('python 代码/gfp_design.py \\')
add_code('    --data-dir ./data \\')
add_code('    --method esm \\')
add_code('    --esm-model esm2_t6_8M_UR50D \\')
add_code('    --max-mutations 3 \\')
add_code('    --n-candidates 100 \\')
add_code('    --top-n 6 \\')
add_code('    --max-train-samples 1000 \\')
add_code('    --batch-size 4 \\')
add_code('    --seed 42')

add_p('')
add_rich_p([('参数说明：', True)])
add_bullet_rich([('esm2_t6_8M_UR50D', True), '：最小的 ESM 模型（~8M 参数），CPU 可运行，3-5 分钟完成'])
add_bullet_rich([('--max-mutations 3', True), '：每序列最多 3 个突变（低复杂度，快速生成）'])
add_bullet_rich([('--n-candidates 100', True), '：仅生成 100 个候选（而非正式运行的 500-1000）'])
add_bullet_rich([('--max-train-samples 1000', True), '：最多使用 1000 条训练数据'])

add_h3('0.3.2 预期输出')
add_p('运行成功后，终端将显示：')
add_code('============================================')
add_code('TOP 6 CANDIDATES (Predicted Brightness)')
add_code('============================================')
add_code('Sequence ID    Mutations              PredictedBrightness')
add_code('Candidate_1    F64L:V163A:S65T        3.4521')
add_code('Candidate_2    S65T:M153T:F64L        3.2108')
add_code('...')
add_code('')
add_code('Saved top 6 candidates to ./output/top_candidates.csv')
add_p('同时在 ./output/ 目录下生成 top_candidates.csv 文件。')

add_h3('0.3.3 预实验检查清单')
add_bullet('✓ 所有 import 无报错')
add_bullet('✓ ESM 模型成功下载并加载（模型文件自动缓存至 ~/.cache/torch/hub/）')
add_bullet('✓ 训练数据正确加载（日志显示 "Filtered down to XX avGFP entries"）')
add_bullet('✓ 嵌入生成完成（日志显示 "Embedding generation completed in Xs"）')
add_bullet('✓ 随机森林训练完成（显示 R² 值）')
add_bullet('✓ 候选序列生成并排序')
add_bullet('✓ CSV 文件成功保存')

add_h2('0.4 预实验常见问题排查')
make_table(
    ['问题', '可能原因', '解决方法'],
    [
        ['ImportError: No module named esm', 'fair-esm 未正确安装', 'pip install fair-esm --force-reinstall'],
        ['CUDA out of memory', 'GPU 显存不足', '使用 --batch-size 2 或更小值'],
        ['FileNotFoundError: GFP data.xlsx', '数据文件路径不正确', '检查 --data-dir 参数指向的目录'],
        ['Empty embeddings tensor', '序列长度异常或格式错误', '检查 WT 序列文件编码是否为 UTF-8'],
        ['R² 值极低 (< 0.1)', '训练数据量不足或质量差', '增加 --max-train-samples 参数'],
    ],
    [3.5, 4.5, 5.0]
)

add_h2('0.5 预实验交付物')
add_p('完成预实验后，你应获得：')
add_numbered('一份 top_candidates.csv（6 条预测序列及其突变描述）')
add_numbered('对管线各步骤运行时间的初步感知')
add_numbered('确认计算环境（CPU/GPU）的能力边界')
add_numbered('为正式管线运行积累的基线参数')

add_page_break()

# ============================================================
# 第一部分：项目背景与竞赛约束
# ============================================================
add_h1('第一部分：项目背景与竞赛约束')

add_h2('1.1 竞赛基本信息')
make_table(
    ['项目', '内容'],
    [
        ['竞赛名称', '2026 合成生物学创新赛 —— 蛋白质设计赛道'],
        ['任务目标', '运用计算方法设计兼具高荧光亮度和优良热稳定性的 GFP 变体'],
        ['提交内容', '6 条氨基酸序列（每条 220-250 aa，以 M 开头）'],
        ['截止日期', '2026 年 7 月 1 日'],
        ['剩余时间', '24 天'],
        ['宿主系统', '大肠杆菌（E. coli）表达系统'],
        ['反向翻译', '竞赛统一使用 DNAChisel 算法执行（参赛队伍无需自行操作）'],
    ],
    [3.5, 10.5]
)

add_h2('1.2 评分机制深度解析')
add_p('竞赛采用两阶段荧光测定评估：')
add_bullet_rich([('F_initial', True), '：表达后初始荧光强度（评估亮度）'])
add_bullet_rich([('F_final', True), '：72°C 热处理后的残余荧光强度（评估热稳定性）'])
add_bullet_rich([('综合得分', True), ' = (F_initial / F_initial_WT) × (F_final / F_initial)'])
add_bullet_rich([('淘汰线', True), '：F_initial < 0.3 × F_initial_WT 的序列直接淘汰'])

add_p('')
add_rich_p([('评分逻辑解读：', True)])
add_p('第一项 (F_initial / F_initial_WT) 衡量亮度是否优于野生型；第二项 (F_final / F_initial) 衡量热处理后荧光保留率（即热稳定性）。两项相乘意味着——既亮又稳定的序列才能获得高分。单独亮但不耐热（第二项低），或耐热但本身暗（第一项低），均无法胜出。这是典型的多目标联合优化问题。', indent=True)

add_h2('1.3 关键生物学约束')
add_p('基于 sfGFP 晶体结构（PDB: 2B3P）的位点功能注释（详见 sfGFP_氨基酸功能全注释.docx），我们将 238 个氨基酸位点分为四个约束等级：')

add_h3('约束等级一：绝对禁止突变（37 个位点）')
add_p('这些位点突变将导致荧光完全消失或蛋白无法折叠，任何候选序列必须保留野生型残基：')
add_bullet('发色团核心与成熟催化：Y66, G67, R96, E222, H148, Q69, Q94, T203, S205, E95')
add_bullet('疏水核心锚定（仅限芳香族保守替换）：F27, F46, F83, F71, F100, F130, F165, F223, W57, F166, F114, F8, F84')
add_bullet('β-转角关键 Gly：G31, G33, G35, G51, G91, G116, G127, G134, G160, G174, G189, G228')
add_bullet('β-转角关键 Pro：P54, P56, P58, P75, P89, P192')
add_bullet('发色团空间冲突 Gly：G91（任何侧链直接碰撞发色团）')

add_h3('约束等级二：严重受限（约 50 个位点）')
add_p('仅允许保守替换（如疏水 → 疏水、小残基 → 小残基）。包括疏水核心 Leu/Val/Ile 位点、发色团邻近疏水残基（V68, V150, V93, L60 等）、桶底部疏水密封残基。')

add_h3('约束等级三：可突变但需谨慎（约 60 个位点）')
add_p('包括表面暴露残基、环区残基、C 端柔性区域等，突变容忍度较高，是设计多样性的主要来源。')

add_h3('约束等级四：正向设计热点（约 20 个位点）')
add_p('已知的增强突变位点——这些是天然存在的 "Nature 已替你验证过" 的设计素材：')
add_bullet('F64L：增强 37°C 折叠效率（sfGFP/EGFP 共有核心突变）')
add_bullet('S65T：提升发色团成熟效率与亮度')
add_bullet('S30R：引入表面正电荷盐桥，增强折叠')
add_bullet('Y39N：改变桶内极性网络')
add_bullet('F99S：减少表面疏水斑块，抑制聚集')
add_bullet('N105T：减少表面柔性，提升折叠效率')
add_bullet('Y145F：消除不必要羟基，保留芳香环')
add_bullet('M153T：Cycle3 GFP 核心折叠增强突变')
add_bullet('V163A：减少内部空间位阻')
add_bullet('I171V：表面突变增强折叠')
add_bullet('A206V：增强 C 端疏水堆积，提升热稳定性')
add_bullet('H148S：YuzuFP 验证，亮度提升 1.5 倍')

add_h2('1.4 已有工作基础')
add_p('在正式管线设计之前，已完成以下准备工作：')
make_table(
    ['工作项', '内容', '状态'],
    [
        ['文献调研', '18 篇核心文献系统调研（ESM3、GeoEvoBuilder、SPURS 等）', '✓ 完成'],
        ['sfGFP 位点注释', '238 个氨基酸逐位功能注释，明确禁止/允许突变位点', '✓ 完成'],
        ['基线管线', 'ESM 嵌入 + 随机森林回归（gfp_design.py）', '✓ 完成'],
        ['训练数据', 'avGFP 突变体亮度数据集（含多种 GFP 类型）', '✓ 完成'],
        ['竞赛规则分析', '评分公式、DNAChisel 反向翻译、实验流程', '✓ 完成'],
        ['参考项目对标', 'synbiochallenges2025、SC2025、MCCOP、GGS 等分析', '✓ 完成'],
    ],
    [3.0, 7.5, 2.0]
)

add_page_break()

# ============================================================
# 第二部分：正式管线总体架构
# ============================================================
add_h1('第二部分：正式管线总体架构')

add_h2('2.1 设计哲学：三路线并行 + 集成筛选')
add_p('基于文献调研和竞赛约束，采取"三条互补技术路线并行生成候选序列，统一通过多维度筛选确定最终 6 条提交序列"的策略。三条路线各具优势，互为补充：')

make_table(
    ['路线', '核心方法', '优势', '风险', '优先级'],
    [
        ['路线 A', 'PLM 嵌入 + ML 回归 + MCMC 约束采样', '成熟稳定、计算成本低、可直接复用现有代码', '依赖于训练数据质量、探索空间有限', '★★★★★'],
        ['路线 B', 'ProteinMPNN + 结构约束 + AF2 验证', '结构感知、零样本能力、可大规模采样', '需要 GPU、AF2 预测耗时', '★★★★'],
        ['路线 C', '进化分析 + 共识设计 + 理性突变组合', '生物意义明确、可解释性强、低风险', '创新性可能不足、依赖 MSA 质量', '★★★'],
    ],
    [1.8, 4.0, 3.5, 2.8, 1.8]
)

add_h2('2.2 管线总体流程（五阶段）')

add_rich_p([('阶段一：数据准备与增强（第 1-3 天）', True)])
add_numbered('整理并清洗现有训练数据（去重、异常值检测、数据增强）')
add_numbered('构建多序列比对（MSA），提取进化耦合信息')
add_numbered('准备结构模板（sfGFP PDB: 2B3P 及其他高分辨率结构）')
add_numbered('编译禁止突变位点约束文件（37 个绝对禁止 + ~50 个严重受限）')

add_rich_p([('阶段二：多路线并行生成（第 2-7 天）', True)])
add_numbered('路线 A：升级版 ESM/SaProt + XGBoost 集成 + 约束 MCMC 采样')
add_numbered('路线 B：ProteinMPNN 结构条件生成 + AlphaFold2 结构验证')
add_numbered('路线 C：共识序列分析 + 已知增强突变组合设计')

add_rich_p([('阶段三：多维度虚拟筛选（第 6-10 天）', True)])
add_numbered('亮度预测筛选（ML 集成模型 + 物理化学特征）')
add_numbered('热稳定性预测筛选（SPURS / TemBERTure / ProtSSN）')
add_numbered('结构完整性验证（AlphaFold2 / ESMFold pLDDT + pTM）')
add_numbered('可表达性评估（溶解度、聚集倾向）')

add_rich_p([('阶段四：Top 候选精选（第 10-13 天）', True)])
add_numbered('三路线候选池合并去重（约 150-250 条）')
add_numbered('多目标 Pareto 前沿分析（亮度 vs 热稳定性）')
add_numbered('人工专家审核（检查突变位点是否违反约束）')
add_numbered('精选 6 条互补序列（覆盖不同突变策略）')

add_rich_p([('阶段五：最终验证与提交准备（第 13-18 天）', True)])
add_numbered('DNAChisel 反向翻译模拟验证')
add_numbered('序列合规性自动检查（长度、起始氨基酸、字符集）')
add_numbered('排除列表交叉验证')
add_numbered('提交文件格式化与备份')

add_h2('2.3 技术栈总览')
make_table(
    ['类别', '工具/模型', '用途', '安装方式'],
    [
        ['蛋白质语言模型', 'ESM-2 (35M/150M/650M)', '序列嵌入提取', 'pip install fair-esm'],
        ['蛋白质语言模型', 'SaProt (35M_AF2)', '结构感知序列嵌入', 'huggingface transformers'],
        ['结构预测', 'AlphaFold2 / ColabFold', '预测结构验证折叠', 'LocalColabFold 或 API'],
        ['结构预测', 'ESMFold', '快速结构预测（备用）', 'pip install esm'],
        ['逆折叠设计', 'ProteinMPNN', '结构条件序列生成', 'git clone + conda env'],
        ['稳定性预测', 'SPURS', 'L×20 全突变 ΔΔG 扫描', 'git clone + pip install'],
        ['稳定性预测', 'TemBERTure', 'Tm 预测 + 嗜热分类', 'huggingface'],
        ['稳定性预测', 'ProtSSN', '序列-结构融合零样本预测', 'git clone'],
        ['ML 模型', 'scikit-learn / XGBoost', '亮度回归预测', 'pip install xgboost'],
        ['序列分析', 'Biopython', '序列操作、MSA 解析', 'pip install biopython'],
        ['反向翻译', 'DNAChisel', '氨基酸 → DNA（本地验证用）', 'pip install dnachisel'],
        ['理化计算', 'Rosetta / FoldX', 'ΔΔG、溶剂可及面积等', '各自独立安装'],
    ],
    [3.0, 4.0, 4.0, 3.0]
)

add_page_break()

# ============================================================
# 第三部分：分步骤详细实施方案
# ============================================================
add_h1('第三部分：分步骤详细实施方案')

# ---- 阶段一 ----
add_h2('阶段一：数据准备与增强（第 1-3 天）')

add_h3('步骤 1.1：训练数据清洗与质量评估')
add_p('输入文件：data/GFP_training_data.csv')

add_rich_p([('① 去重处理：', True)])
add_p('检查 aaMutations 列的重复条目，对于同一突变有多个 Brightness 值的，取均值并记录标准差（评估实验重复性）。', indent=True)
add_code("import pandas as pd")
add_code("df = pd.read_csv('data/GFP_training_data.csv')")
add_code("# 按突变去重，取平均亮度")
add_code("df_clean = df.groupby(['GFP type', 'aaMutations'])['Brightness'] \\")
add_code("    .agg(['mean', 'std', 'count']).reset_index()")

add_rich_p([('② 异常值检测：', True)])
add_bullet('对于 F64L 突变，确认 Brightness 是否在合理范围（avGFP 中 F64L 亮度约 1.8-2.4 倍 WT）')
add_bullet('Y66F、Y66H 等发色团突变亮度应显著偏离 WT（颜色改变而非亮度增强）')
add_bullet('剔除标注为 "WT" 但在不同 GFP type 下亮度差异 > 3σ 的异常条目')
add_bullet('使用 IQR 方法：超出 Q1 - 1.5×IQR 至 Q3 + 1.5×IQR 范围的标记为可疑')

add_rich_p([('③ 数据增强（可选但强烈推荐）：', True)])
add_bullet('从文献（Sarkisyan et al. Nature 2016）补充 GFP 适应度景观数据（~51,000 条 avGFP 突变数据）')
add_bullet('从 Nature Communications 2023 (htFuncLib) 补充设计的 GFP 功能变体数据')
add_bullet('合并不同来源数据时注意标注来源，后续可做加权训练')

add_h3('步骤 1.2：构建禁止突变约束矩阵')
add_p('输入文件：sfGFP_禁止突变位点.csv')
add_p('创建 238×20 的二进制约束矩阵 constraint[position][AA]，将每个位点-氨基酸组合分类为：')
add_bullet('0 = 允许（该位置可突变为该氨基酸）')
add_bullet('1 = 禁止（绝对禁止，37 个位点的非保守替换）')
add_bullet('2 = 仅限保守替换（疏水 ↔ 疏水、芳香族 ↔ 芳香族等）')

add_p('')
add_p('保守替换判定规则：')
add_bullet('疏水组：L, I, V, M 之间互相替换')
add_bullet('芳香族组：F, Y, W 之间互相替换')
add_bullet('小残基组：G, A, S 之间互相替换')
add_bullet('正电荷组：K, R, H 之间互相替换')
add_bullet('负电荷组：D, E 之间互相替换')
add_bullet('极性组：N, Q, S, T 之间互相替换')

add_code('''def is_mutation_allowed(wt_seq, position_1based, new_aa):\n    """检查某位点突变是否在允许范围内"""\n    pos_idx = position_1based - 1\n    aa_idx = 'ACDEFGHIKLMNPQRSTVWY'.index(new_aa)\n    level = CONSTRAINT_MATRIX[pos_idx][aa_idx]\n    if level == 1:      # 绝对禁止\n        return False\n    if level == 2:      # 仅保守替换\n        return is_conservative(wt_seq[pos_idx], new_aa)\n    return True''')

add_h3('步骤 1.3：多序列比对（MSA）构建')
add_p('目标：从进化角度识别保守位点和共进化位点对。')
add_numbered('收集 GFP 家族序列：从 UniProt 检索 GFP-like 蛋白，使用 jackhmmer 迭代搜索（3 轮，E-value < 1e-10），目标收集 500-2000 条同源序列')
add_numbered('MSA 质量控制：移除 gap 比例 > 50% 的序列；移除序列一致性 < 20% 或 > 98% 的序列；最终保留 200-800 条高质量同源序列')
add_numbered('从 MSA 提取：位点保守性评分（Shannon 熵）、共进化残基对（CCMpred/GREMLIN 耦合矩阵）、与已知禁止突变位点交叉验证')

add_page_break()

# ---- 阶段二A：路线A ----
add_h2('阶段二A：路线 A —— PLM 嵌入 + ML 回归 + 约束采样（第 2-6 天）')

add_h3('步骤 2A.1：嵌入提取升级')
add_p('在预实验的基础上进行以下升级：')

add_rich_p([('模型升级：', True)])
add_bullet('从 esm2_t6_8M 升级到 esm2_t30_150M_UR50D（更好表达能力，GPU 推荐）或 esm2_t33_650M_UR50D（最佳，需 16GB+ 显存）')
add_bullet('同步使用 SaProt（结构感知嵌入），与 ESM 嵌入拼接或分别建模后集成')
add_bullet('在发色团口袋残基（65-67, 96, 222, 148, 203, 205 等）附近使用局部注意力加权嵌入')

add_rich_p([('嵌入缓存策略：', True)])
add_bullet('对整个训练集提取嵌入后保存为 .npy 文件，避免重复计算')
add_code('X_esm = get_esm_embeddings(sequences, ...)\nnp.save(\'cache/esm_embeddings_150M.npy\', X_esm.numpy())')

add_h3('步骤 2A.2：ML 模型升级')
add_p('从单一随机森林升级到集成模型：')

add_rich_p([('模型组合策略：', True)])
add_bullet_rich([('XGBoost + 随机森林 + KNN Ridge', True), '：三模型集成，取均值预测'])
add_bullet_rich([('特征工程增强', True), '：除 PLM 嵌入外，增加以下特征：'])
add_bullet('突变数量（num_mutations）', level=1)
add_bullet('突变位点的平均保守性评分（来自 MSA）', level=1)
add_bullet('突变引入的疏水性变化（Eisenberg scale ΔG）', level=1)
add_bullet('突变引入的体积变化（侧链体积差）', level=1)
add_bullet('突变位点是否位于已知增强热点', level=1)
add_bullet('突变位点距离发色团的最短距离（来自 PDB 结构）', level=1)

add_rich_p([('训练策略：', True)])
add_bullet('5 折交叉验证（确保模型稳定性）')
add_bullet('留一 GFP 类型交叉验证（如用 avGFP+egGFP 训练，预测 sfGFP 变体）')
add_bullet('记录每折 R²、Pearson r、RMSE')
add_bullet('目标：R² > 0.75，Pearson r > 0.85')

add_h3('步骤 2A.3：约束 MCMC 序列采样')
add_p('替代当前随机采样策略，使用 MCMC 进行有偏向的序列空间探索：')

add_p('MCMC 采样框架：')
add_bullet('初始状态：sfGFP 野生型序列')
add_bullet('提议分布：从允许突变位点池中随机选择 1-K 个位置进行突变')
add_bullet('接受准则（Metropolis-Hastings）：P_accept = min(1, exp(-(pred_score_new - pred_score_old)/T))')
add_bullet('T（温度参数）控制探索 vs 利用的平衡——高 T 广泛探索，低 T 精细搜索')
add_bullet('退火策略：T 从 2.0 逐步降至 0.1')

add_p('多链并行 + 去重策略：')
add_bullet('运行 10-20 条独立 MCMC 链（不同随机种子），每条链采样 500-1000 步')
add_bullet('合并所有链的样本并去重，生成候选池（预期 5000-10000 条唯一序列）')

add_h3('步骤 2A.4：路线 A 最终候选筛选')
add_numbered('使用集成 ML 模型预测所有 MCMC 生成序列的亮度')
add_numbered('排除预测亮度 < 1.0×WT 的序列')
add_numbered('按预测亮度排序，取 Top 100 进入总候选池')
add_numbered('确保 Top 100 的突变多样性（来自相同突变组合的序列仅保留最高分）')

add_page_break()

# ---- 阶段二B：路线B ----
add_h2('阶段二B：路线 B —— ProteinMPNN 结构条件设计（第 3-7 天）')

add_h3('步骤 2B.1：环境配置')
add_code('# 克隆 ProteinMPNN 仓库')
add_code('git clone https://github.com/dauparas/ProteinMPNN.git')
add_code('cd ProteinMPNN')
add_code('pip install torch numpy')
add_code('# 模型权重自动下载至 vanilla_model_weights/ 和 soluble_model_weights/')

add_h3('步骤 2B.2：结构模板准备')
add_p('首选模板：sfGFP 晶体结构 PDB 2B3P（分辨率 1.45Å，链 A 含完整发色团）。')
add_p('PDB 预处理：提取单链（链 A）→ 移除水分子和异源分子（保留发色团 CRO）→ 重新编号残基（从 1 开始）。')

add_rich_p([('固定残基设置（Tied positions）—— 以下残基在 ProteinMPNN 采样中固定不变：', True)])
add_bullet('发色团核心：Y66, G67（发色团三肽第 2、3 位）')
add_bullet('发色团成熟催化：R96, E222')
add_bullet('发色团氢键网络：Q69, Q94, E95, H148, T203, S205')
add_bullet('疏水核心锚定：F27, F46, F83, F71, F100, F130, F165, F223, W57 等')
add_bullet('关键 Gly/Pro 转角残基')
add_p('共计固定约 50-60 个残基，允许 ProteinMPNN 设计其余 ~180 个残基。')

add_h3('步骤 2B.3：ProteinMPNN 采样运行')
add_code('python protein_mpnn_run.py \\')
add_code('    --pdb_path ./proteinmpnn_input.pdb \\')
add_code('    --out_folder ./proteinmpnn_output/ \\')
add_code('    --num_seq_per_target 500 \\')
add_code('    --sampling_temp "0.1 0.2 0.3 0.5" \\')
add_code('    --seed 42 \\')
add_code('    --batch_size 8 \\')
add_code('    --tied_positions "65 66 67 96 222 69 94 95 148 203 205 ..." \\')
add_code('    --use_soluble_model')
add_p('参数说明：--num_seq_per_target 500（每温度采样 500 条）、--sampling_temp（多温度采样，低温接近 WT，高温更多样化）、--tied_positions（固定残基列表）、--use_soluble_model（使用可溶性优化版模型）。')

add_h3('步骤 2B.4：AlphaFold2 结构验证')
add_p('对 ProteinMPNN 生成的所有序列进行快速结构预测验证：')
add_code('# 安装 ColabFold 本地版（参考 github.com/sokrypton/ColabFold）')
add_code('colabfold_batch proteinmpnn_sequences.fasta \\')
add_code('    proteinmpnn_af2_results/ \\')
add_code('    --num-recycle 3 \\')
add_code('    --use-gpu-relax')

add_p('筛选标准：')
add_bullet_rich([('pLDDT > 85', True), '（全局折叠置信度高）'])
add_bullet_rich([('pTM > 0.75', True), '（模板建模评分高）'])
add_bullet_rich([('发色团区域 pLDDT > 90', True), '（活性中心结构可靠）'])
add_bullet('RMSD to WT < 2.0Å（整体结构不偏离野生型）')
add_bullet('排除预测结构中发色团口袋有明显空间冲突的序列')
add_p('随后使用 CD-HIT 聚类（sequence identity cutoff 90%），从每个聚类中选取最高 pLDDT 代表序列，目标获得 50-100 条通过结构验证的候选。')

add_page_break()

# ---- 阶段二C：路线C ----
add_h2('阶段二C：路线 C —— 进化分析引导的理性设计（第 2-5 天）')

add_h3('步骤 2C.1：共识序列分析')
add_p('基于 MSA（步骤 1.3 的结果）计算每个位点的共识氨基酸：')
add_code('from Bio.Align import AlignInfo\nsummary = AlignInfo.SummaryInfo(msa_alignment)\nconsensus = summary.dumb_consensus(threshold=0.3)\npssm = summary.pos_specific_score_matrix()')
add_p('将 sfGFP 序列与共识序列逐位比对，标记 sfGFP 残基频率 < 30% 的位点（可能是非最优残基），对这类位点优先突变为共识氨基酸。')

add_h3('步骤 2C.2：已知增强突变组合设计')
add_p('综合文献报道的 GFP 增强突变，设计组合方案。以 sfGFP 为骨架（已含多个增强突变），额外叠加 2-4 个增强突变，确保突变总数 ≤ 6。使用 Rosetta ddG 或 FoldX 快速筛选稳定组合。预期生成 50-100 条理性组合序列。')

add_p('')
add_rich_p([('亮度增强突变池：', True)])
add_bullet('F64L（sfGFP/EGFP 共有，增强 37°C 折叠）')
add_bullet('S65T（发色团成熟效率提升，核心亮度突变）')
add_bullet('H148S（YuzuFP 验证，亮度提升 1.5 倍）')

add_rich_p([('折叠增强突变池（Cycle3 GFP 来源）：', True)])
add_bullet('F99S + M153T + V163A（三突变组合显著增强折叠）')

add_rich_p([('热稳定性增强突变池（sfGFP 来源）：', True)])
add_bullet('S30R + Y39N + N105T + Y145F + I171V + A206V')

add_h3('步骤 2C.3：共进化引导的突变设计')
add_p('使用 CCMpred 分析 MSA 中的共进化残基对。若某位点发生突变，其共进化配对位点可能需要补偿突变。识别 sfGFP 中违背共进化模式的位点对，设计补偿突变以恢复共进化模式。只考虑耦合强度 top 100 的残基对，并与已知禁止突变位点交叉验证。')

add_page_break()

# ---- 阶段三 ----
add_h2('阶段三：多维度虚拟筛选（第 6-10 天）')

add_h3('步骤 3.1：候选池合并与去重')
add_p('合并三条路线的候选序列（路线 A ~100 条、路线 B ~50-100 条、路线 C ~50-100 条），合并候选池预期规模 150-250 条唯一序列。')
add_p('去重与预筛选流程：精确序列匹配去重 → 序列一致性 > 98% 的聚类（CD-HIT，每类保留代表性序列）→ 与排除列表交叉比对 → 与禁止突变约束矩阵交叉验证 → 去除含 > 6 个突变的序列。')

add_h3('步骤 3.2：亮度多模型预测')
add_numbered('使用路线 A 训练的集成 ML 模型预测亮度')
add_numbered('使用 METL（如可获取）进行零样本功能预测')
add_numbered('使用 GLEAM 风格全局-局部注意力模型（可选）')
add_numbered('三模型预测取中位数 ± 标准差作为亮度估计及不确定性')
add_numbered('排除预测亮度 < 1.5×WT 的序列（保守阈值）')

add_h3('步骤 3.3：热稳定性多模型预测')
add_p('对通过亮度筛选的候选序列进行热稳定性评估：')

add_rich_p([('SPURS（推荐优先）：', True)])
add_bullet('单次前向传播即可预测 L×20 所有单突变 ΔΔG')
add_bullet('对候选序列的每个突变位点查询 SPURS 预测的 ΔΔG，累积 ΣΔΔG < -2.0 kcal/mol 视为优良')

add_rich_p([('TemBERTure：', True)])
add_bullet('Tm 回归预测 + 嗜热/非嗜热分类')
add_bullet('高 Tm（> 65°C）或嗜热分类得分 > 0.7 视为优良')

add_rich_p([('ProtSSN（备选）：', True)])
add_bullet('序列-结构融合零样本预测，无需 MSA，ProteinGym 基准上表现优异')

add_rich_p([('稳定性综合评分：', True)])
add_bullet('将 SPURS ΔΔG、TemBERTure Tm、ProtSSN 评分标准化至 [0,1]')
add_bullet('取三模型加权平均（SPURS 0.4, TemBERTure 0.3, ProtSSN 0.3）')
add_bullet('综合评分 < 0.5 的序列淘汰')

add_h3('步骤 3.4：结构完整性与可表达性')
add_numbered('AlphaFold2 快速预测 Top 50 候选的结构（pLDDT > 85, pTM > 0.75）')
add_numbered('可溶性预测（ProteinMPNN soluble model 评分，或 CamSol 等工具）')
add_numbered('聚集倾向预测（TANGO/AGGRESCAN，排除含强聚集倾向区域的序列）')

add_page_break()

# ---- 阶段四 ----
add_h2('阶段四：Top 候选精选（第 10-13 天）')

add_h3('步骤 4.1：多目标 Pareto 前沿分析')
add_p('对通过所有筛选的候选序列（预期 30-50 条），进行多目标优化分析：')
add_numbered('构建二维目标空间：X 轴 = 亮度预测值，Y 轴 = 热稳定性预测值')
add_numbered('识别 Pareto 前沿——在两个目标上均不被其他序列支配的序列')
add_numbered('Pareto 前沿序列优先入选')
add_numbered('使用 matplotlib 生成可视化图辅助决策')

add_h3('步骤 4.2：突变策略多样性保障')
add_p('为确保 6 条提交序列覆盖不同的设计策略，按以下维度分组并各选最优代表：')
make_table(
    ['策略类型', '描述', '亮度贡献', '稳定性贡献'],
    [
        ['折叠增强型', '以折叠效率突变为主（F64L, Cycle3 等）', '★★★', '★★'],
        ['发色团优化型', '优化发色团微环境（H148S 等）', '★★★★', '★'],
        ['表面优化型', '表面电荷/疏水优化抑制聚集', '★★', '★★★★'],
        ['核心稳定型', '疏水核心/β-桶几何优化', '★', '★★★★★'],
        ['多策略混合型', '综合多种策略', '★★★', '★★★'],
        ['探索创新型', 'ProteinMPNN 生成的新颖序列', '??', '??'],
    ],
    [3.0, 5.5, 2.5, 2.5]
)

add_h3('步骤 4.3：人工专家审核')
add_p('对 6 条候选序列进行逐突变审查：')
add_bullet('每个突变是否在允许位点池内？是否违反约束矩阵（等级 1）？')
add_bullet('保守替换（等级 2）是否确实为保守替换？')
add_bullet('突变数量是否合理（2-6 个）？')
add_bullet('是否有已知的负面突变组合（如拮抗上位效应）？')
add_bullet('参考 sfGFP_氨基酸功能全注释.docx 确认每个突变位点功能')

add_h3('步骤 4.4：替补方案')
add_p('除 6 条提交序列外，额外准备 4-6 条替补序列：')
add_bullet('2 条高亮度策略（亮度优先但稳定性可能边际）')
add_bullet('2 条高稳定性策略（稳定性优先但亮度可能边际）')
add_bullet('2 条保守策略（从训练集中验证过的已知突变组合，可靠性最高）')

add_page_break()

# ---- 阶段五 ----
add_h2('阶段五：最终验证与提交准备（第 13-18 天）')

add_h3('步骤 5.1：DNAChisel 反向翻译验证')
add_p('虽然竞赛统一执行反向翻译，但本地验证可提前发现潜在问题：')
add_code('pip install dnachisel')
add_code('')
add_code('from dnachisel import reverse_translate, CodonOptimize')
add_code('')
add_code('dna = reverse_translate(aa_sequence, organism="E. coli")')
add_code('')
add_code('problem = CodonOptimize(')
add_code('    sequence=dna, species="e_coli", method="use_best_codon"')
add_code(')')
add_code('optimized_dna = problem.resolve()')

add_p('检查要点：DNA 长度（aa 长度 × 3）、限制性酶切位点（避免 EcoRI/BamHI）、GC 含量（40-60%）、polyA/polyT 转录终止风险、起始/终止密码子。')

add_h3('步骤 5.2：序列合规性自动化检查')
add_code('''def validate_submission(sequences, exclusion_set, constraint_matrix):
    """自动化检查提交序列的合规性"""
    errors = []
    for i, seq in enumerate(sequences):
        if not (220 <= len(seq) <= 250):
            errors.append(f"Seq{i}: length {len(seq)} out of [220,250]")
        if not seq.startswith('M'):
            errors.append(f"Seq{i}: does not start with M")
        valid_aa = set('ACDEFGHIKLMNPQRSTVWY')
        if not set(seq).issubset(valid_aa):
            errors.append(f"Seq{i}: invalid characters")
        if seq in exclusion_set:
            errors.append(f"Seq{i}: in exclusion list")
        for pos, aa in enumerate(seq, 1):
            aa_idx = 'ACDEFGHIKLMNPQRSTVWY'.index(aa)
            if CONSTRAINT_MATRIX[pos-1][aa_idx] == 1:
                errors.append(f"Seq{i}: forbidden mut at pos {pos}")
    return errors''')

add_h3('步骤 5.3：提交文件准备')
add_numbered('参考 submission_template.csv 格式')
add_numbered('填写：Sequence ID（Candidate_1 至 Candidate_6）、Mutations（突变描述）、Sequence（完整序列）')
add_numbered('保存为 CSV（UTF-8 编码，无 BOM）')
add_numbered('另存 Excel 版本（含元数据：预测亮度、预测 Tm、pLDDT、筛选路线来源）作为内部记录')
add_numbered('备份所有中间文件（候选池、筛选日志、模型权重）')

add_page_break()

# ============================================================
# 第四部分：24天时间规划
# ============================================================
add_h1('第四部分：24 天时间规划')

add_h2('4.1 甘特图式时间表')
make_table(
    ['阶段', '任务', '天数', '日期范围', '里程碑'],
    [
        ['预实验', '环境搭建 + 管线试运行', '0.5', '6/7', '✓ 基线结果产出'],
        ['阶段一', '数据清洗 + 约束矩阵 + MSA', '2', '6/7 - 6/9', '✓ 清洁数据集锁定'],
        ['阶段二A', '路线 A：ML 升级 + MCMC', '4', '6/8 - 6/12', '100 条路线 A 候选'],
        ['阶段二B', '路线 B：ProteinMPNN + AF2', '4', '6/9 - 6/13', '100 条路线 B 候选'],
        ['阶段二C', '路线 C：共识 + 理性组合', '3', '6/8 - 6/11', '100 条路线 C 候选'],
        ['阶段三', '多维度虚拟筛选', '3', '6/13 - 6/16', '30-50 条精选候选'],
        ['阶段四', 'Top 候选精选 + 专家审核', '3', '6/16 - 6/19', '6 条最终候选确定'],
        ['阶段五', '最终验证 + 提交准备', '3', '6/19 - 6/22', '提交文件就绪'],
        ['缓冲期', '修改优化 + 文档完善', '5-9', '6/22 - 7/1', '提交确认'],
    ],
    [2.0, 4.5, 1.2, 2.5, 3.5]
)

add_h2('4.2 关键时间节点')
add_bullet_rich([('6/9（第 3 天）', True), '：数据准备完成，三条路线同时启动生成'])
add_bullet_rich([('6/13（第 7 天）', True), '：三条路线候选生成完成，进入筛选阶段'])
add_bullet_rich([('6/16（第 10 天）', True), '：第一轮筛选完成，候选池缩小至 30-50 条'])
add_bullet_rich([('6/19（第 13 天）', True), '：6 条最终候选序列确定'])
add_bullet_rich([('6/22（第 16 天）', True), '：提交文件准备完成，进入缓冲期审查'])
add_bullet_rich([('6/29（第 23 天）', True), '：建议提前 2 天提交（避免截止日期前系统拥堵）'])
add_bullet_rich([('7/1（第 25 天）', True, (204, 0, 0)), '：最终截止日期'])

add_h2('4.3 资源需求预估')
make_table(
    ['资源', '最低配置', '推荐配置', '用途'],
    [
        ['GPU', 'NVIDIA RTX 3060 (12GB)', 'NVIDIA RTX 4090 (24GB)', 'ESM 嵌入、AF2 预测'],
        ['CPU 模式', '可全流程 CPU（较慢）', 'GPU 加速 10-50 倍', '后备方案'],
        ['存储', '50 GB 可用', '100 GB 可用', '模型权重 + 嵌入缓存 + AF2'],
        ['内存', '16 GB', '32 GB+', 'MSA 构建、大数据加载'],
        ['网络', '稳定互联网', '高速 + 镜像代理', '模型下载、数据库检索'],
        ['云 GPU（可选）', 'AutoDL / 恒源云', 'A100 40GB 按需', 'AF2 大批量预测'],
    ],
    [2.5, 4.0, 4.0, 3.5]
)

add_page_break()

# ============================================================
# 第五部分：风险识别与应对策略
# ============================================================
add_h1('第五部分：风险识别与应对策略')

make_table(
    ['风险', '概率', '影响', '应对策略'],
    [
        ['GPU 不可用/显存不足', '中', '高',
         '①降级到 esm2_t12_35M 模型（CPU 可运行）；②使用 Colab 免费 GPU；③租用 AutoDL 云 GPU（约 2-5 元/小时）'],
        ['AF2 预测耗时长', '高', '中',
         '①优先使用 ESMFold（快 5-10 倍）；②分批次预测，先 Top 50；③ColabFold --num-recycle 1 加速'],
        ['ML 模型 R² < 0.6', '中', '高',
         '①补充 Sarkisyan 2016 数据集增强训练；②特征工程（理化特征 + 嵌入）；③使用预训练 METL 模型'],
        ['ProteinMPNN 折叠率低', '中', '中',
         '①提高采样温度范围；②减少固定残基数量；③使用多个结构模板'],
        ['三条路线候选重叠严重', '低', '低',
         '①增加路线 C 突变组合多样性；②路线 B 提高采样温度；③路线 A 扩大 MCMC 探索范围'],
        ['发色团邻近突变破坏荧光', '中', '高',
         '①严格执行禁止突变约束；②对邻近发色团位点（距离 < 8Å）进行额外 MD 验证'],
        ['稳定性预测与实际不一致', '中', '中',
         '①多模型交叉验证；②参考 PROSS 设计结果；③优先选择 SPURS+TemBERTure 双模型一致预测'],
        ['时间不足', '中', '高',
         '①路线 C 可最早产出（并行推进）；②如紧迫，优先完成路线 A+C，路线 B 作为补充；③缓冲期预留充足'],
    ],
    [3.5, 1.0, 1.0, 8.0]
)

add_page_break()

# ============================================================
# 附录
# ============================================================
add_h1('附录')

add_h2('附录 A：关键命令速查表')
make_table(
    ['操作', '命令'],
    [
        ['预实验运行', 'python 代码/gfp_design.py --data-dir ./data --method esm --esm-model esm2_t6_8M_UR50D --max-mutations 3 --n-candidates 100 --top-n 6'],
        ['正式 ESM 运行', 'python 代码/gfp_design.py --data-dir ./data --method esm --esm-model esm2_t30_150M_UR50D --max-mutations 6 --n-candidates 1000 --top-n 100'],
        ['SaProt 运行', 'python 代码/gfp_design.py --data-dir ./data --method saprot --saprot-model westlake-repl/SaProt_35M_AF2 --max-mutations 6 --n-candidates 1000'],
        ['ProteinMPNN 采样', 'python protein_mpnn_run.py --pdb_path input.pdb --num_seq_per_target 500 --sampling_temp "0.1 0.2 0.3 0.5"'],
        ['ColabFold 批量预测', 'colabfold_batch sequences.fasta results/ --num-recycle 3 --use-gpu-relax'],
        ['DNAChisel 反向翻译', 'from dnachisel import reverse_translate; dna = reverse_translate(aa_seq, organism="E. coli")'],
        ['CD-HIT 去重', 'cd-hit -i candidates.fasta -o candidates_nr.fasta -c 0.95 -n 5'],
    ],
    [4.0, 10.0]
)

add_h2('附录 B：推荐项目文件结构')
add_code('蛋白质设计-合成生物学创新赛-Claude/')
add_code('├── data/                    # 原始数据')
add_code('│   ├── GFP data.xlsx')
add_code('│   ├── GFP_training_data.csv')
add_code('│   └── Exclusion_List.csv')
add_code('├── 代码/                    # Python 脚本')
add_code('│   ├── gfp_design.py        # 主管线（ESM + RF）')
add_code('│   ├── gfp_design_v2.py     # 升级版（XGBoost + MCMC）')
add_code('│   └── utils.py             # 辅助函数库（约束检查等）')
add_code('├── _过程文件/               # 中间过程脚本和临时文件')
add_code('├── cache/                   # 缓存')
add_code('│   ├── esm_embeddings/')
add_code('│   └── af2_results/')
add_code('├── 运行结果/                # 输出结果')
add_code('│   ├── route_a_candidates.csv')
add_code('│   ├── route_b_candidates.csv')
add_code('│   ├── merged_candidates.csv')
add_code('│   ├── final_6_submission.csv')
add_code('│   └── submission_metadata.xlsx')
add_code('├── logs/                    # 运行日志')
add_code('├── sfGFP_氨基酸功能全注释.docx')
add_code('├── sfGFP_禁止突变位点.csv')
add_code('└── 参考文献/')

add_h2('附录 C：核心参考文献速查')
add_p('详见"GFP蛋白质设计文献与工具调研报告.docx"，以下为最核心的 5 篇：')
add_bullet_rich([('[1] ESM3/esmGFP', True), ' — Science 387:850-858 (2025). DOI: 10.1126/science.ads0018'])
add_bullet_rich([('[2] GeoEvoBuilder', True), ' — PNAS 122:e2504117122 (2025). DOI: 10.1073/pnas.2504117122'])
add_bullet_rich([('[3] ProteinMPNN for FP', True), ' — Protein Science 33:e70002 (2024). DOI: 10.1002/pro.70002'])
add_bullet_rich([('[4] SPURS', True), ' — Nature Communications (2025). DOI: 10.1038/s41467-025-56475-9'])
add_bullet_rich([('[5] YuzuFP', True), ' — Communications Chemistry 8:174 (2025). DOI: 10.1038/s42004-025-01573-4'])
add_bullet_rich([('[6] Sarkisyan et al.', True), ' — Nature 533:397-401 (2016). GFP 适应度景观基准数据集'])

# -- 尾注 --
doc.add_paragraph()
doc.add_paragraph()
end_p = doc.add_paragraph()
end_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = end_p.add_run('—— 文档结束 ——')
set_cn_font(run, '微软雅黑')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(150, 150, 150)

# ============================================================
# 保存
# ============================================================
output_path = r'D:\蛋白质设计-合成生物学创新赛-Claude\文档\GFP蛋白质设计_管线设计与实施方案_v2.0.docx'
os.makedirs(os.path.dirname(output_path), exist_ok=True)
doc.save(output_path)
print(f'Document saved to: {output_path}')
print(f'File size: {os.path.getsize(output_path) / 1024:.1f} KB')
