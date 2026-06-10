# -*- coding: utf-8 -*-
"""读取5.25管线规划文档，提取全文结构"""
from docx import Document
import os

path = r'D:\蛋白质设计-合成生物学创新赛-Claude\文档\From Claude\下一步设计思路与管线规划_2026ProteinDesign 2026.5.25.docx'
doc = Document(path)

with open(r'D:\蛋白质设计-合成生物学创新赛-Claude\_过程文件\plan_525.txt', 'w', encoding='utf-8') as f:
    f.write(f"段落总数: {len(doc.paragraphs)}\n")
    f.write(f"表格总数: {len(doc.tables)}\n\n")
    f.write("=" * 70 + "\n")
    f.write("全文内容 (非空段落)\n")
    f.write("=" * 70 + "\n\n")

    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if t:
            style = p.style.name if p.style else 'Normal'
            f.write(f"[{style}] {t[:200]}\n\n")

    # 表格内容
    if doc.tables:
        f.write("=" * 70 + "\n表格内容\n" + "=" * 70 + "\n\n")
        for ti, table in enumerate(doc.tables):
            f.write(f"\n--- Table {ti+1} ---\n")
            for ri, row in enumerate(table.rows):
                cells = [cell.text.strip()[:80] for cell in row.cells]
                f.write(f"  Row {ri}: {' | '.join(cells)}\n")

print(f"Done. Written {len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables.")
