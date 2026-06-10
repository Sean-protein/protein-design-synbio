# -*- coding: utf-8 -*-
"""读取项目中的关键设计文档并输出摘要到文本文件"""

import os, sys

BASE = r'D:\蛋白质设计-合成生物学创新赛-Claude'

def read_docx(path):
    """读取docx文档的全部文本"""
    from docx import Document
    doc = Document(path)
    return '\n'.join([p.text for p in doc.paragraphs if p.text.strip()])

def read_pdf(path, max_pages=10):
    """读取PDF文本"""
    try:
        import pdfplumber
        with pdfplumber.open(path) as pdf:
            texts = []
            for i, page in enumerate(pdf.pages):
                if i >= max_pages:
                    break
                t = page.extract_text()
                if t:
                    texts.append(t)
            return '\n'.join(texts)
    except Exception as e:
        return f"PDF读取失败: {e}"

output_lines = []
output_lines.append("=" * 80)
output_lines.append("项目设计文档摘要汇总")
output_lines.append("=" * 80)

# 1. 读取已有的下一步设计思路文档
plan_path = os.path.join(BASE, '下一步设计思路与管线规划_2026ProteinDesign 2026.5.25.docx')
if os.path.exists(plan_path):
    output_lines.append("\n\n### [1] 下一步设计思路与管线规划 (2026.5.25)")
    output_lines.append("-" * 60)
    try:
        text = read_docx(plan_path)
        output_lines.append(text[:8000])
        output_lines.append(f"\n... (总长度: {len(text)} 字符)")
    except Exception as e:
        output_lines.append(f"读取失败: {e}")

# 2. 读取sfGFP注释文档
gfp_anno_path = os.path.join(BASE, 'sfGFP_氨基酸功能全注释.docx')
if os.path.exists(gfp_anno_path):
    output_lines.append("\n\n### [2] sfGFP氨基酸功能全注释")
    output_lines.append("-" * 60)
    try:
        text = read_docx(gfp_anno_path)
        output_lines.append(text[:5000])
        output_lines.append(f"\n... (总长度: {len(text)} 字符)")
    except Exception as e:
        output_lines.append(f"读取失败: {e}")

# 3. 读取竞赛PDF
pdf_path = os.path.join(BASE, '2026Protein Design', '2026Protein Design in Synbio challenges.pdf')
if os.path.exists(pdf_path):
    output_lines.append("\n\n### [3] 竞赛官方说明 PDF")
    output_lines.append("-" * 60)
    try:
        text = read_pdf(pdf_path, max_pages=5)
        output_lines.append(text[:5000])
    except Exception as e:
        output_lines.append(f"读取失败: {e}")

# 4. 读取禁止突变位点CSV
csv_path = os.path.join(BASE, 'sfGFP_禁止突变位点.csv')
if os.path.exists(csv_path):
    output_lines.append("\n\n### [4] 禁止突变位点CSV (前50行)")
    output_lines.append("-" * 60)
    try:
        with open(csv_path, 'r', encoding='utf-8') as f:
            lines = f.readlines()[:50]
            output_lines.append(''.join(lines))
    except Exception as e:
        output_lines.append(f"读取失败: {e}")

# 5. 读取训练数据CSV
train_csv = os.path.join(BASE, 'data', 'GFP_training_data.csv')
if os.path.exists(train_csv):
    output_lines.append("\n\n### [5] 训练数据CSV (前30行)")
    output_lines.append("-" * 60)
    try:
        with open(train_csv, 'r', encoding='utf-8') as f:
            lines = f.readlines()[:30]
            output_lines.append(''.join(lines))
    except Exception as e:
        output_lines.append(f"读取失败: {e}")

# 6. 读取生成的管线方案
gen_path = os.path.join(BASE, '文档', 'GFP蛋白质设计_管线设计与实施方案_v2.0.docx')
if os.path.exists(gen_path):
    output_lines.append("\n\n### [6] 生成的管线方案文档结构")
    output_lines.append("-" * 60)
    try:
        from docx import Document
        doc = Document(gen_path)
        for p in doc.paragraphs:
            if p.text.strip() and (p.style.name.startswith('Heading') or len(p.text.strip()) > 30):
                output_lines.append(f"[{p.style.name}] {p.text.strip()[:120]}")
        output_lines.append(f"\n总段落: {len(doc.paragraphs)}, 总表格: {len(doc.tables)}")
    except Exception as e:
        output_lines.append(f"读取失败: {e}")

# 输出
out_path = os.path.join(BASE, '_过程文件', 'design_summary.txt')
with open(out_path, 'w', encoding='utf-8') as f:
    f.write('\n'.join(output_lines))

print(f'Done. Written to {out_path}')
print(f'Total length: {len(output_lines)} lines')
