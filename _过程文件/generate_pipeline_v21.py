# -*- coding: utf-8 -*-
"""生成 v2.1 管线文档——整合5.25全部内容 + DBLT + Pareto + Agent + esmGFP修正 + 预备方案"""

import datetime, os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

doc = Document()

# ── 页面设置 ──
for s in doc.sections:
    s.top_margin = Cm(2.5); s.bottom_margin = Cm(2.5)
    s.left_margin = Cm(2.5); s.right_margin = Cm(2.5)

# ── 样式 (不设Heading避免TOC复杂度) ──
def cn(run, font='微软雅黑'):
    if run._element.rPr is None:
        from lxml import etree
        nsmap = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        rPr = etree.SubElement(run._element, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rPr')
    else:
        rPr = run._element.rPr
    rPr_rFonts = rPr.find(qn('w:rFonts'))
    if rPr_rFonts is None:
        from lxml import etree
        rPr_rFonts = etree.SubElement(rPr, qn('w:rFonts'))
    rPr_rFonts.set(qn('w:eastAsia'), font)

def en(run, font='Times New Roman'):
    run.font.name = font
    if run._element.rPr is None:
        from lxml import etree
        nsmap = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        rPr = etree.SubElement(run._element, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rPr')
    else:
        rPr = run._element.rPr
    rPr_rFonts = rPr.find(qn('w:rFonts'))
    if rPr_rFonts is None:
        from lxml import etree
        rPr_rFonts = etree.SubElement(rPr, qn('w:rFonts'))
    rPr_rFonts.set(qn('w:eastAsia'), '宋体')

def h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24); p.paragraph_format.space_after = Pt(12)
    r = p.add_run(text); cn(r,'微软雅黑'); r.font.size=Pt(18); r.bold=True; r.font.color.rgb=RGBColor(0,51,102)

def h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18); p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text); cn(r,'微软雅黑'); r.font.size=Pt(15); r.bold=True; r.font.color.rgb=RGBColor(0,80,130)

def h3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text); cn(r,'微软雅黑'); r.font.size=Pt(14); r.bold=True; r.font.color.rgb=RGBColor(46,117,182)

def bp(text, bold=False):
    """正文段落"""
    p = doc.add_paragraph(); p.paragraph_format.line_spacing = 1.5
    r = p.add_run(text); en(r); r.font.size=Pt(12); r.bold=bold
    return p

def rp(items, indent=False):
    """富文本段落 items=[(text, bold, color), ...]"""
    p = doc.add_paragraph(); p.paragraph_format.line_spacing = 1.5
    if indent: p.paragraph_format.first_line_indent = Pt(24)
    for x in items:
        if isinstance(x,str): t,b,c = x,False,None
        elif len(x)==2: t,b,c = x[0],x[1],None
        else: t,b,c = x
        r = p.add_run(t); en(r); r.font.size=Pt(12); r.bold=b
        if c: r.font.color.rgb=RGBColor(*c)
    return p

def bl(text, level=0):
    """项目符号"""
    p = doc.add_paragraph(style='List Bullet'); p.clear()
    r = p.add_run(text); en(r); r.font.size=Pt(12)
    if level>0: p.paragraph_format.left_indent=Cm(1.5*(level+1))
    return p

def br(items, level=0):
    """富文本项目符号"""
    p = doc.add_paragraph(style='List Bullet'); p.clear()
    if level>0: p.paragraph_format.left_indent=Cm(1.5*(level+1))
    for x in items:
        if isinstance(x,str): t,b,c = x,False,None
        elif len(x)==2: t,b,c = x[0],x[1],None
        else: t,b,c = x
        r = p.add_run(t); en(r); r.font.size=Pt(12); r.bold=b
        if c: r.font.color.rgb=RGBColor(*c)
    return p

def nl(text):
    """编号列表"""
    p = doc.add_paragraph(style='List Number'); p.clear()
    r = p.add_run(text); en(r); r.font.size=Pt(12)
    return p

def code(text):
    """代码块"""
    p = doc.add_paragraph(); p.paragraph_format.left_indent=Cm(1.0)
    p.paragraph_format.space_before=Pt(2); p.paragraph_format.space_after=Pt(2)
    p.paragraph_format.line_spacing=1.2
    s = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F5" w:val="clear"/>')
    p.paragraph_format.element.get_or_add_pPr().append(s)
    r = p.add_run(text); r.font.name='Consolas'; r._element.rPr.rFonts.set(qn('w:eastAsia'),'宋体')
    r.font.size=Pt(9); r.font.color.rgb=RGBColor(51,51,51)

def pb(): doc.add_page_break()

def table(hds, rows, ws=None):
    """创建表格"""
    t = doc.add_table(rows=1+len(rows), cols=len(hds)); t.style='Table Grid'
    for i,h in enumerate(hds):
        c = t.rows[0].cells[i]; c.text=''
        r = c.paragraphs[0].add_run(h); en(r); r.font.size=Pt(10); r.bold=True; r.font.color.rgb=RGBColor(0,51,102)
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="D6E4F0" w:val="clear"/>')
        c._element.get_or_add_tcPr().append(shd)
    for ri,row in enumerate(rows):
        for ci,val in enumerate(row):
            c = t.rows[ri+1].cells[ci]; c.text=''
            r = c.paragraphs[0].add_run(str(val)); en(r); r.font.size=Pt(10)
    if ws:
        for i,w in enumerate(ws):
            for row in t.rows: row.cells[i].width=Cm(w)
    doc.add_paragraph()

# ════════════════════════════════════════════════════════════
# 封面
# ════════════════════════════════════════════════════════════
for _ in range(6): doc.add_paragraph()

p = doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('2026合成生物学创新赛 · 蛋白质设计赛道'); cn(r,'微软雅黑'); r.font.size=Pt(26); r.font.color.rgb=RGBColor(0,51,102)

p = doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('GFP 高亮度与热稳定性联合设计\n管线设计与实施方案 v2.1'); cn(r,'微软雅黑')
r.font.size=Pt(30); r.bold=True; r.font.color.rgb=RGBColor(0,51,102)

doc.add_paragraph()
p = doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Design-Build-Test-Learn 迭代闭环 · Pareto 双目标优化 · LLM Agent 编排'); cn(r,'微软雅黑')
r.font.size=Pt(14); r.font.color.rgb=RGBColor(80,80,80)

for _ in range(3): doc.add_paragraph()
p = doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(f'{datetime.date.today().strftime("%Y年%m月%d日")}  |  截止: 2026年7月1日  |  有效工作日: 18天  |  考试期: 6/25-7/1'); cn(r,'微软雅黑')
r.font.size=Pt(14); r.font.color.rgb=RGBColor(204,0,0)
pb()

# ════════════════════════════════════════════════════════════
# 目录
# ════════════════════════════════════════════════════════════
h1('目录')
toc = [
    '一、设计哲学与核心修正（v2.1 相比此前版本的关键变更）',
    '二、数据可靠性声明（哪些能用、哪些不能用、哪些需验证）',
    '三、约束体系：从禁止名单到化学固定+协同设计+验证',
    '四、四条设计策略路线（整合5.25方案精华）',
    '五、DBLT 迭代闭环：Design→Build→Test→Learn',
    '六、双目标优化：Pareto 前沿 + 贝叶斯引导搜索',
    '七、LLM Agent 编排架构（竞赛加分项）',
    '八、五级漏斗筛选详细设计',
    '九、18天时间规划 + 四级预备方案',
    '十、风险矩阵与应对策略',
    '附录A：关键命令速查表',
    '附录B：45个候选设计位点全表',
    '附录C：参考文献与工具速查',
]
for item in toc: bp(item)
pb()

# ════════════════════════════════════════════════════════════
# 一、设计哲学与核心修正
# ════════════════════════════════════════════════════════════
h1('一、设计哲学与核心修正')
bp('本方案 v2.1 基于 2026.5.25 版规划文档整合，并根据 6月6-7日两轮深度讨论做了以下关键修正：')

table(
    ['修正点', 'v2.0/5.25 方案', 'v2.1 方案', '依据'],
    [
        ['约束体系', '四等级"禁止突变"名单(37绝对+50严重受限)', '化学固定~6位点(G67/Y66/R96/E222等) + 其余协同设计+AF2验证', 'esmGFP: 96/229突变→荧光保留 → 单点禁忌≠协同不可行'],
        ['管线架构', '线性漏斗: 生成→筛选→精选→提交(一次性)', 'DBLT闭环: Design→Build→Test→Learn, 2-3轮迭代', '竞赛评分逻辑: 需实验反馈 → 用MD/BioEmu作为计算代理'],
        ['优化目标', '乘法加权评分(亮度30%+稳定性30%+...)', 'Pareto前沿 + 乘积最优2条+亮度极端2条+稳定极端1条+探索1条', '乘积脆弱性: 一个维度预测错→乘积全错; Pareto分散风险'],
        ['MD角色', '最终验证Top 10-15, 一次性', 'BioEmu预筛→GROMACS MD金标准验证→结果回传贝叶斯优化', 'BioEmu(Science 2025)速度提升4-5数量级; MD结果应驱动下一轮'],
        ['时间线', '21天全天投入', '18天有效(6/7-24) + 7天考试微调(6/25-7/1) + 四级预备方案', '考试约束: 6/25-7/1无法全力工作'],
        ['数据态度', '未区分AI生成vs实验验证', '明确标注每个文档/数据的可靠性等级', 'sfGFP注释为AI推断(Un-named作者); Doubao文档禁止使用'],
        ['Agent', '未涉及', '三层Agent架构(编排+决策+执行) + 决策逻辑树 + 执行日志', '竞赛明确鼓励LLM Agent, 客观得分相近时优先展示'],
    ],
    [2.0, 4.0, 4.5, 3.5]
)

bp('核心设计原则（不变，来自5.25方案）：')
bl('广度优先(Exploration): 多路线并行覆盖序列空间')
bl('深度验证(Exploitation): 结构预测+MD模拟逐级筛选')
bl('联合决策(Integration): 综合多管线Pareto前沿选6条互补序列')
bl('风险分散(Diversification): 覆盖不同"亮度×热稳定性"区域, 降低单一策略失败风险')
bl('6中取1即可登顶 → 鼓励多样化探索')

pb()

# ════════════════════════════════════════════════════════════
# 二、数据可靠性声明
# ════════════════════════════════════════════════════════════
h1('二、数据可靠性声明')

bp('管道使用以下数据源, 按可靠性分级:')

table(
    ['等级', '数据/文档', '来源', '使用方式'],
    [
        ['★★★ 金标准', 'Sarkisyan 2016 Nature (~54K GFP变体+亮度)', '同行评审期刊, 实验测定', '训练核心数据集'],
        ['★★★ 金标准', 'PDB 2B3P (sfGFP 1.45Å)', 'X射线晶体学', 'ProteinMPNN结构模板; MD初始结构'],
        ['★★★ 金标准', 'Pédelacq 2006 (sfGFP创始论文)', '同行评审期刊', 'sfGFP核心折叠增强突变依据'],
        ['★★★ 金标准', 'esmGFP (Science 2025)', '同行评审期刊', '协同设计范式依据; 约束体系修正依据'],
        ['★★★ 金标准', 'BioEmu (Science 2025)', '同行评审期刊', 'MD加速方案; 构象系综预筛'],
        ['★★☆ 验证后可用', 'FPbase GFP数据 (~1,110条)', '社区数据库', '补充光学性质数据; 需逐条核实'],
        ['★★☆ 验证后可用', 'ProTherm热力学数据', '社区数据库', '热稳定性训练数据; 条件与72°C不同, 需校准'],
        ['★☆☆ 方向参考', 'sfGFP_氨基酸功能全注释.docx', 'AI生成 (Un-named, 2026-06-06)', '约束方向性参考; 具体数值不可直接引用; 需原始文献核实'],
        ['✗ 禁止使用', 'From Doubao 全部文档', 'AI生成, 未校验', '任何参数/命令/模型选择均不可直接使用; 仅作概念参考'],
    ],
    [1.5, 4.5, 3.5, 4.5]
)

bp('', bold=True)
rp([('特别声明: ', True), ('sfGFP_氨基酸功能全注释.docx 由 Claude 于2026年6月6日生成(作者"Un-named"), 引用了Pédelacq 2006和Sarkisyan 2016两篇论文, 但逐位点的具体约束等级和数值是AI基于PDB结构推断的产物, 未经实验验证。其方向性正确（发色团三肽不可动、催化残基关键）, 但"R96E残留1-5%荧光"这类具体数值不可直接引用。')])
pb()

# ════════════════════════════════════════════════════════════
# 三、约束体系
# ════════════════════════════════════════════════════════════
h1('三、约束体系：从禁止名单到化学固定+协同设计+验证')

h2('3.1 esmGFP 的根本性启示')
bp('ESM3生成的esmGFP与最近天然GFP(tagRFP)仅58%序列一致, 229个残基中96个不同, 但仍正常折叠并发出与EGFP相当的荧光。这证明了:')
bl('"单独突变危险" ≠ "协同突变不可行": 疏水核心/β-转角/表面电荷可以同时重写')
bl('约束应来自化学第一性原理, 而非单一结构的"不可突变"推论')
bl('ProteinMPNN/ESM3这类协同设计工具是正确方向——给定骨架结构, 全局重设计')

h2('3.2 修正后的约束体系')

table(
    ['约束等级', '位点数量', '判定依据', '执行方式'],
    [
        ['化学绝对约束', '~6', 'GFP发色团自催化环化-脱水-氧化的有机化学反应机理', 'ProteinMPNN中tied positions, MCMC中直接拒绝非法突变'],
        ['结构协同约束', '其余~232', '单独突变可能危险, 但可通过协同突变补偿', '不预先禁止; 用AF2 pLDDT + MD验证替代黑名单; 约束从"禁止"变为"验证"'],
    ],
    [2.5, 3.0, 5.0, 3.5]
)

h3('化学绝对约束位点(6个)——在ProteinMPNN中固定')
bl('G67: 绝对不可替换。发色团环化步骤中G67酰胺氮亲核进攻T65羰基碳——这是SN2类亲核加成反应, 任何侧链(包括H原子)都会空间阻断环化。esmGFP必定保留了Gly67。')
bl('Y66: 仅限芳香族替换(Y/F/H/W)。酚羟基参与发色团π共轭体系, 无芳香环→无可见光吸收→无荧光。可改变颜色(Y66H→BFP蓝光, Y66W→CFP青光), 但不可替换为非芳香族。')
bl('T65: 可保守小残基替换(G/A/C/S)。avGFP天然为S65, sfGFP为T65(亮度更高)。不能突变为大侧链(阻断环化), 也不能突变为带电残基(改变发色团静电环境)。')
bl('R96: 大概率需保留。催化发色团环化脱水步骤, 与咪唑啉酮环形成氢键稳定过渡态。理论上若周围残基协同重排可能由其他残基替代催化, 但风险极高。建议固定。')
bl('E222: 大概率需保留。氧化步骤的广义碱催化剂, 催化Y66 Cα-Cβ脱氢。E222D部分保留功能, 其他替换基本失活。建议固定。')
bl('发色团氢键网络(Q69/Q94/H148/T203/S205): 这些位点的氢键对量子产率至关重要。单点突变大幅降亮度, 但理论上协同重排氢键网络可能找到替代方案。建议在路线B中尝试, 路线A中保守保留。')

h3('结构协同约束(其余位点)——验证替代禁止')
bp('对剩余的~232个位点, 不预先禁止任何突变, 而是通过以下验证管线确保设计质量:')
bl('AF2/ESMFold pLDDT > 80, pTM > 0.75 → 折叠可行性保证')
bl('发色团区域pLDDT > 85 → 活性中心结构可靠')
bl('WT sfGFP MD基线对比 → RMSD/Rg/SASA相对变化量判断')
bl('上位效应过滤(EVcouplings/GREMLIN) → 排除违反共进化的组合')

pb()

# ════════════════════════════════════════════════════════════
# 四、四条设计策略路线
# ════════════════════════════════════════════════════════════
h1('四、四条设计策略路线（整合5.25方案精华）')

bp('以下四条策略继承自5.25方案, 并在DBLT闭环中并行运行。每条策略生成~2,000-5,000条候选序列, 合并去冗余后进入漏斗筛选。')

h2('策略A: 理性工程优化（保守稳健型, ★★★★★ 最高优先级）')
rp([('核心思路: ', True), ('以sfGFP为设计起点, 保留6个核心折叠增强突变(S30R/F64L/S65T/F99S/M153T/V163A), 从45个候选设计位点中组合亮度增强突变(2-3个)+热稳定性增强突变(2-3个)。最稳健的策略, 自带"保底"属性。')])
rp([('技术栈: ', True), ('Biopython序列操作 + 组合枚举 + 上位效应过滤(EVcouplings) + FoldX ΔΔG快速筛选')])
rp([('预期产出: ', True), ('~2,000-3,000条组合序列, 预期贡献1-2条保险方案')])

h2('策略B: PLM深度学习驱动（AI驱动型, ★★★★）')
rp([('核心思路: ', True), ('ESM-2 650M嵌入 + SaProt结构感知嵌入 → XGBoost+LightGBM+RF三模型集成亮度预测。贝叶斯优化(Gaussian Process + Expected Improvement)引导MCMC在嵌入空间中搜索。可降级(无需ESM-3)。')])
rp([('技术栈: ', True), ('ESM-2 650M/SaProt + XGBoost/LightGBM/RF + GPyTorch + 约束MCMC采样')])
rp([('预期产出: ', True), ('~3,000-5,000条AI生成序列, 预期贡献1-2条创新方案')])
rp([('降级预案: ', True), ('如GPU不足 ESM-2 650M → ESM-2 150M; 如R²<0.6 → 增加Sarkisyan 2016数据增强; 最坏情况→策略A+C撑场')])

h2('策略C: ProteinMPNN结构协同设计（结构导向型, ★★★★）')
rp([('核心思路: ', True), ('以sfGFP晶体结构(PDB 2B3P)为模板, 固定~6个化学绝对位点, ProteinMPNN多温度采样(0.1/0.2/0.3/0.5)协同设计其余~232个位点。AF2/ESMFold验证折叠, CD-HIT去冗余。')])
rp([('技术栈: ', True), ('ProteinMPNN + ESMFold/ColabFold + CD-HIT + BioEmu构象系综预筛')])
rp([('预期产出: ', True), ('~2,000-4,000条结构感知序列, 预期贡献1-2条结构验证方案')])

h2('策略D: 进化共识与高性能特征迁移（进化导向型, ★★★）')
rp([('核心思路: ', True), ('构建GFP深度MSA(jackhmmer 3轮迭代) → 共进化分析(EVcouplings) → 识别高性能GFP共有特征模式(sfGFP/mBaoJin/TGP/StayGold)。将mBaoJin的H148D、T203I等"特征模块"嫁接到sfGFP骨架。')])
rp([('技术栈: ', True), ('jackhmmer + MAFFT + EVcouplings/GREMLIN + 共进化引导的组合突变生成')])
rp([('预期产出: ', True), ('~1,000-2,000条进化约束序列, 提供上位效应过滤规则（所有策略共用）')])

h2('四条策略的交叉验证节点')
bl('结构可行性: 策略B产出 → 策略C的AF2/ESMFold验证')
bl('突变合理性: 策略A突变选择 → 策略D保守性评分过滤')
bl('稳定性评估: 所有策略 → 统一的Phase 4热稳定性评估管线')
bl('最终决策: 四条策略Pareto前沿合并 → 6条序列选择')

pb()

# ════════════════════════════════════════════════════════════
# 五、DBLT迭代闭环
# ════════════════════════════════════════════════════════════
h1('五、DBLT 迭代闭环：Design → Build → Test → Learn')

bp('这是v2.1最核心的架构升级。不再是一次性漏斗筛选, 而是2-3轮迭代循环, 每轮用上一轮的MD/BioEmu验证结果更新贝叶斯优化代理模型, 引导下一轮序列生成方向。')

h2('5.1 闭环架构')
code('                 ┌─────────── DESIGN ───────────┐')
code('                 │ 四策略并行生成候选              │')
code('                 │ GP采集函数引导MCMC方向          │')
code('                 │ 从多进化起点出发(sfGFP/TGP/     │')
code('                 │   mBaoJin/StayGold)            │')
code('                 └────────────┬──────────────────┘')
code('                              ↓')
code('                 ┌─────────── BUILD ─────────────┐')
code('                 │ 亮度集成预测(3模型 → μ±σ)      │')
code('                 │ 稳定性分层预测                  │')
code('                 │   ThermoMPNN → FoldX → BioEmu  │')
code('                 │ 上位效应过滤(EVcouplings)       │')
code('                 └────────────┬──────────────────┘')
code('                              ↓')
code('                 ┌─────────── TEST ──────────────┐')
code('                 │ AF2/ESMFold pLDDT/pTM         │')
code('                 │ BioEmu 构象系综预筛(50-100条)  │')
code('                 │ GROMACS MD金标准(Top 10-15条)  │')
code('                 │ 72°C热冲击+复性模拟            │')
code('                 │ RMSD/Rg/SASA/氢键维持率         │')
code('                 └────────────┬──────────────────┘')
code('                              ↓')
code('                 ┌─────────── LEARN ─────────────┐')
code('                 │ MD结果回传GP代理模型            │')
code('                 │ 更新Expected Improvement        │')
code('                 │ 收敛判断(EI<阈值? 3轮完成?)     │')
code('                 │ → 如未收敛: 回到DESIGN         │')
code('                 │ → 如收敛: 进入精选阶段          │')
code('                 └───────────────────────────────┘')

h2('5.2 每轮迭代的规模')
table(
    ['轮次', 'Design生成数', 'Build筛选后', 'Test验证数', 'MD验证数', '目的'],
    [
        ['Round 1', '~10,000 (四策略合并)', '~1,000', 'BioEmu ~100, AF2 ~500', 'Top 10', '广泛探索 + 建立GP初始模型'],
        ['Round 2', '~5,000 (GP引导MCMC)', '~500', 'BioEmu ~50, AF2 ~200', 'Top 10', '聚焦高EI区域 + 更新GP'],
        ['Round 3', '~2,000 (精细搜索)', '~200', 'BioEmu ~30, AF2 ~100', 'Top 10', '收敛确认 + 精选'],
    ],
    [1.5, 3.0, 2.5, 3.0, 2.0, 2.0]
)

h2('5.3 贝叶斯优化核心组件')
rp([('高斯过程代理模型(GP): ', True), ('使用Matern 5/2核。输入=序列嵌入(ESM+SaProt拼接的PCA降维50维), 输出=乘积得分(Brightness×Stability)。GP不仅预测得分均值μ, 还给出不确定性σ。数据密集区域σ低("已知"), 数据稀疏区域σ高("未知")。')])
rp([('Expected Improvement采集函数: ', True), ('EI(x) = E[max(0, f(x) - f_best)]。在μ高且σ也高的区域EI最大——"这里可能有金矿, 而且我还不太确定"。每轮选EI最大的区域进行MCMC采样。')])
rp([('MD结果校准: ', True), ('MD的RMSD/氢键保留率作为"金标准"稳定标签, 校准快速预测器(ThermoMPNN/FoldX)的系统偏差。每轮MD验证后用新数据重新训练GP。')])

h2('5.4 收敛判断')
bl('EI_max连续2轮 < 0.05 → 收敛, 进入精选阶段')
bl('已完成3轮迭代 → 强制进入精选 (时间约束)')
bl('MD结果与GP预测Spearman r > 0.7 → 模型可靠, 可提前收敛')
bl('MD结果与预测矛盾 → 扩大搜索范围, 降低GP先验权重')

pb()

# ════════════════════════════════════════════════════════════
# 六、双目标优化
# ════════════════════════════════════════════════════════════
h1('六、双目标优化：Pareto前沿 + 贝叶斯引导搜索')

h2('6.1 为什么不用乘积最大化作为唯一目标')
bp('乘积目标函数 Score = Brightness × Stability 的致命缺陷是: 预测误差在乘法的两端都会被放大。如果亮度被高估20%, 乘积直接虚高20%。而如果另一个维度的预测也有误差, 乘积的误差是复合的。')
bp('正确的做法是构建Pareto前沿, 然后从前沿的四个区域各选序列:')

code('          亮度↑')
code('           │    ★ 亮度极端区 (2条)')
code('           │          ★ Pareto前沿')
code('           │               ★ 乘积最优区 (2条)')
code('           │                      ★ 探索区(1条, EI最大)')
code('           │                              ★ 稳定极端区 (1条)')
code('           └────────────────────────────────────→ 热稳定性')

h2('6.2 6条序列分配策略')
table(
    ['分组', '数量', '选择标准', '预期角色'],
    [
        ['乘积最优组', '2条', 'Pareto前沿上 Score=Brightness×Stability 的top 2', '主攻综合评分榜(金银奖)'],
        ['亮度极端组', '2条', 'Brightness > 80%分位 AND Stability > 中位数', '竞逐最佳亮度单项奖'],
        ['稳定性极端组', '1条', 'Stability > 80%分位 AND Brightness > 中位数', '竞逐最佳热稳定单项奖'],
        ['高不确定性组', '1条', 'GP Expected Improvement最大(预测不确定性最高)', '覆盖未知的高潜力区域'],
    ],
    [2.5, 1.2, 5.5, 4.0]
)

h2('6.3 序列多样性检查')
bl('两两序列相似度 < 90% (如违反, 从同一Pareto区域选择下一个候补)')
bl('6条序列至少覆盖3种设计策略 (避免"把鸡蛋放一个篮子")')
bl('风险分层: 2条低风险(策略A) + 2条中风险(策略C/D) + 2条中高风险(策略B)')

pb()

# ════════════════════════════════════════════════════════════
# 七、Agent编排
# ════════════════════════════════════════════════════════════
h1('七、LLM Agent 编排架构（竞赛加分项）')

bp('竞赛规则明确: "鼓励并欢迎使用各种大模型与代码Agent"、"若使用智能体设计, 需在文档中展示Agent的逻辑树和关键执行日志"。以下架构在18天内可落地, 不依赖外部AutoProteinEngine。')

h2('7.1 三层Agent架构')

table(
    ['层', '角色', '实现方式', '输出'],
    [
        ['编排层', '管线总控: 决定什么时候跑什么, 4策略→DBLT调度', 'Claude Code + 结构化管理', '执行计划 + Pipeline状态记录'],
        ['决策层(5节点)', '关键决策点的智能判断: 模型选型/参数自适应/结构筛选/Pareto选择/收敛判断', 'Python脚本 + LLM API调用 + 决策模板', '每节点的决策依据 + 备选方案 + if-then逻辑树'],
        ['执行层', '实际计算: ESM嵌入/ProteinMPNN/AF2/MD模拟', '现有Python管线脚本', 'CSV/PDB/MD轨迹'],
    ],
    [2.0, 4.5, 4.5, 3.0]
)

h2('7.2 五个决策节点逻辑树')

h3('节点1: 模型选型Agent')
code('输入: GPU状态 (torch.cuda检测)')
code('├─ 显存≥24GB → ESM-2 650M + SaProt并行 (推荐)')
code('├─ 显存≥12GB → ESM-2 150M + SaProt')
code('├─ 显存≥8GB  → ESM-2 150M')
code('└─ 仅CPU     → ESM-2 35M (降级)')
code('决策记录: [时间戳] 检测到RTX 3090 24GB → 选择ESM-2 650M+SaProt')
code('备选: 下载失败 → 降级到本地缓存ESM-2 150M')

h3('节点2: 训练质量自适应Agent')
code('输入: 5-fold CV R², Pearson r')
code('├─ R²>0.75, r>0.85 → 直接进入生成阶段')
code('├─ R² 0.6-0.75     → 增加理化特征+MSA保守性(特征工程)')
code('├─ R² 0.4-0.6      → 数据增强(补充Sarkisyan 2016全量)')
code('└─ R²<0.4          → 触发预备Plan C: 切换到纯理性设计')
code('决策记录: R²=0.68 → 中等, 启动特征工程+数据增强')

h3('节点3: 结构验证筛选Agent')
code('输入: AF2/ESMFold pLDDT, pTM, RMSD')
code('├─ pLDDT>85 AND pTM>0.75 AND RMSD<2.0 → 通过')
code('├─ pLDDT 70-85 → AF2二次验证')
code('├─ pLDDT<70 OR RMSD>3.0 → 淘汰')
code('└─ 发色团区域pLDDT<80 → 淘汰(即使全局通过)')
code('决策记录: R2_047: pLDDT=76.3→AF2验证→pLDDT=72.1→淘汰')
code('原因: 发色团区域pLDDT=68.3, H148-发色团距离>4Å')

h3('节点4: Pareto前沿选择Agent')
code('输入: 所有候选(亮度预测μ±σ, 稳定性预测μ±σ)')
code('1. 构建二维Pareto前沿')
code('2. 从前沿4个区域按策略选6条')
code('3. 序列多样性检查(CD-HIT 90%)')
code('4. 输出每条序列的入选理由+Pareto图上位置+替代方案')

h3('节点5: 收敛判断Agent')
code('输入: 最近2轮GP变化, EI_max趋势')
code('├─ EI_max<0.05 (连续2轮) → 收敛→精选')
code('├─ 已3轮迭代 → 强制精选(时间约束)')
code('├─ MD与预测Spearman r>0.7 → 模型可靠→可提前收敛')
code('└─ MD与预测矛盾 → 重新校准GP, 降低先验权重, 再跑一轮')

h2('7.3 提交文档中的Agent展示')
bp('最终提交的PDF将包含:')
bl('Agent逻辑树图 (Mermaid流程图, 展示5个节点的决策逻辑)')
bl('关键决策日志表 (每轮DBLT迭代的代表性决策记录, 含时间戳/输入/决策/依据/人工干预)')
bl('人工干预记录 (Agent建议被人工覆盖的实例, 说明覆盖原因)')
bl('这不需要额外的AutoProteinEngine——管线运行时自动记录的结构化日志整理即可')

pb()

# ════════════════════════════════════════════════════════════
# 八、五级漏斗筛选详细设计
# ════════════════════════════════════════════════════════════
h1('八、五级漏斗筛选详细设计')

table(
    ['阶段', '输入→输出', '核心方法', '筛选标准', '每序列耗时'],
    [
        ['Phase 1: 生成', '4策略 → ~10,000条', 'ESM嵌入+ProteinMPNN+理性组合+进化共识', '220-250aa, M开头, 20标准AA, 排除黑名单, CD-HIT 95%', '<1秒'],
        ['Phase 2: 亮度预测', '~10,000 → ~1,000', 'ESM-2 650M/SaProt + XGBoost+LightGBM+RF集成', '预测亮度 > sfGFP_WT 80%, 模型间σ/μ<0.3(低不确定性优先)', '~0.1秒'],
        ['Phase 3: 结构验证', '~1,000 → ~200', 'ESMFold → 低置信度AF2二次验证', 'pLDDT>80, pTM>0.75, RMSD<2.0Å, 发色团区pLDDT>85', 'ESMFold 60秒, AF2 10分钟'],
        ['Phase 4: 稳定性评估', '~200 → ~50', 'ThermoMPNN→FoldX→BioEmu→MD(三级递进)', 'ΣΔΔG<-2 kcal/mol, BioEmu native_contacts>0.65, MD RMSD<3Å', '秒级→分钟级→小时级'],
        ['Phase 5: Pareto精选', '~50 → 6(+6备选)', 'Pareto前沿+Agent节点4决策+专家审核', '4区分配+多样性+风险分层+理化合理性检查', '数据分析'],
    ],
    [2.0, 2.5, 4.0, 4.0, 1.8]
)

h2('Phase 4稳定性评估三级递进')
bl('第一级(ThermoMPNN+FoldX): ~200条, 秒级/条, 预测ΔTm+ΔΔG → 筛选至~50条')
bl('第二级(BioEmu构象系综): ~50条, 分钟级/条, native_contacts>0.65, 发色团RMSF<2Å → 筛选至~15条')
bl('第三级(GROMACS MD): ~10-15条, 小时级/条 × GPU并行, 72°C 10ns+25°C 5ns复性, RMSD/Rg/SASA/氢键维持率 → 金标准验证')
bp('MD协议(每个序列, CHARMM36m力场):')
code('NVT 25°C 100ps → NPT 25°C 100ps → NPT 72°C 10ns → NPT 25°C 5ns → 分析')
bp('MD是相对比较(wt sfGFP基线)而非绝对Tm预测。10ns看局部波动/氢键/膨胀, 看不完全去折叠。')

pb()

# ════════════════════════════════════════════════════════════
# 九、时间规划 + 预备方案
# ════════════════════════════════════════════════════════════
h1('九、18天时间规划 + 四级预备方案')

h2('9.1 时间线')

table(
    ['日期', '天数', '阶段', '关键任务', 'DBLT轮次'],
    [
        ['6/7', '0.5天', '预实验', '环境验证 + ESM基线运行 + BioEmu安装', '-'],
        ['6/7-6/9', '2天', '阶段一: 准备', '数据清洗+MSA+约束定稿+Agent节点框架', '-'],
        ['6/9-6/13', '4天', 'DBLT Round 1', '四策略生成→Build预测→BioEmu预筛→Top10 MD→Learn更新GP', 'R1'],
        ['6/13-6/17', '4天', 'DBLT Round 2', 'GP引导MCMC生成→Build预测→MD验证→收敛判断', 'R2'],
        ['6/17-6/19', '2天', 'DBLT Round 3', '精细搜索(如R2未收敛)→收敛确认→精选', 'R3(可选)'],
        ['6/19-6/22', '3天', '精选+审核', 'Pareto前沿→Agent节点4选6条→人工审核→替补准备', '-'],
        ['6/22-6/24', '2天', '验证+文档', 'DNAChisel验证→合规检查→Agent日志汇总→PDF→GitHub', '-'],
        ['6/24', '1天', '提交缓冲', '最后检查, ★建议此日前完成提交', '-'],
        ['6/25-7/1', '7天', '考试期', '仅紧急修改, 不新增计算任务, 伺机提交', '-'],
    ],
    [1.8, 1.0, 2.0, 5.0, 1.5]
)

h2('9.2 四级预备方案（考试期保险）')

rp([('核心理念: ', True), ('每个Plan独立准备并行推进, 确保6/24前至少有一份可提交的结果。Plan D最晚6/15启动, 作为死线兜底。')])

table(
    ['方案', '序列来源', '可靠性', '完成时间', '启用条件'],
    [
        ['Plan A (首选)', 'DBLT Round 2/3 Pareto前沿精选6条', '最高(多轮MD验证)', '6/22', '正常完成DBLT全流程'],
        ['Plan B', 'DBLT Round 1 候选池Pareto精选6条', '中高(一轮验证+BioEmu)', '6/17', '6/17 R2进度不足, 用R1结果精选'],
        ['Plan C', '策略A理性组合+策略D进化共识, 合并精选6条', '中(文献支撑, 无ML依赖性)', '6/15', '6/15 ML管线受阻, 切换到纯理性设计'],
        ['Plan D (兜底)', 'sfGFP骨架+2-4个最保守增强突变(F64L+S65T+M153T+V163A等)', '最高(实验验证过, 最低风险)', '6/15', '极端情况, 确保死线前必有一份'],
    ],
    [1.5, 3.5, 3.0, 2.0, 3.5]
)

h3('6条序列来源预案 (Plan A执行时)')
table(
    ['序列编号', '来源策略', 'Pareto区域', '风险等级', '设计特征'],
    [
        ['Seq 1', '策略A(理性工程)', '乘积最优', '低', 'sfGFP核心+亮度增强2-3位点+热稳增强2-3位点'],
        ['Seq 2', '策略B(AI驱动)', '乘积最优', '中', 'GP+BO引导的ESM嵌入空间最优序列'],
        ['Seq 3', '策略A(理性工程)', '亮度极端', '低', '牺牲部分热稳, 最大化亮度(竞逐最佳亮度奖)'],
        ['Seq 4', '策略C(ProteinMPNN)', '亮度极端', '中', '结构协同设计的亮度导向序列'],
        ['Seq 5', '策略C(ProteinMPNN)', '稳定性极端', '中', '结构协同设计的稳定性导向序列(竞逐最佳热稳定奖)'],
        ['Seq 6', '策略B(AI驱动)', '探索区(EI最大)', '中高', 'GP不确定性最高的区域——可能惊喜或翻车'],
    ],
    [1.5, 3.0, 2.5, 1.5, 5.0]
)

h2('9.3 6/24前必须完成的提交材料')
bl('提交CSV: 6条序列(Team_Name+Seq_ID+Sequence), UTF-8无BOM, 与Exclusion_List逐条比对')
bl('设计思路PDF: 管线架构+4策略详述+DBLT迭代过程+Agent逻辑树+关键决策日志+6条序列入选理由')
bl('GitHub仓库: README.md(环境配置+依赖+运行方式)+推理代码+模型权重链接')
bl('内部归档: metadata.xlsx(含预测亮度/稳定性/pLDDT/来源策略/选择依据)')

pb()

# ════════════════════════════════════════════════════════════
# 十、风险矩阵
# ════════════════════════════════════════════════════════════
h1('十、风险矩阵与应对策略')

table(
    ['风险', '概率', '影响', '应对'],
    [
        ['ML模型R²<0.6', '中30%', '高', '补充Sarkisyan 2016全量; 特征工程(理化+MSA); 降级到Plan C'],
        ['GPU不可用/显存不足', '中30%', '中', '降级ESM-2 150M/35M; Colab免费GPU; AutoDL云GPU ¥2-5/h'],
        ['AF2/ESMFold pLDDT普遍<70', '低15%', '高', '放宽RMSD标准; 检查是否破坏折叠的突变; 回到Phase 1重生成'],
        ['ThermoMPNN与FoldX预测矛盾', '中40%', '中', '以FoldX物理能量函数为主; 增加BioEmu/MD验证量'],
        ['MD中WT sfGFP本身不稳定', '低20%', '中', '延长平衡时间; 检查力场参数; 使用多个PDB模板'],
        ['上位效应: 多点组合远低于预期', '高50%', '高', '限制<6突变; 策略D共进化过滤; 保留策略A保守方案'],
        ['约束过度保守: 错失高潜力序列', '中30%', '中', '化学固定仅~6位点; ProteinMPNN自由探索其余位点'],
        ['数据泄露: 无意匹配排除列表', '低10%', '高', 'Phase 1和提交前分别严格比对Exclusion_List+FPbase'],
        ['时间不足: MD未完成或未收敛', '中25%', '中', '优先BioEmu; MD仅做Top 10; Plan C/D兜底; 6/24前提交'],
        ['考试期无法修改', '确定', '高', '6/24前完成提交; Plan D 6/15启动并行准备'],
    ],
    [3.0, 1.0, 1.0, 7.0]
)

pb()

# ════════════════════════════════════════════════════════════
# 附录
# ════════════════════════════════════════════════════════════
h1('附录A: 关键命令速查表')

table(
    ['操作', '命令'],
    [
        ['ESM-2嵌入生成', 'python gfp_design.py --method esm --esm-model esm2_t30_150M_UR50D'],
        ['SaProt嵌入生成', 'python gfp_design.py --method saprot --saprot-model westlake-repl/SaProt_35M_AF2'],
        ['ProteinMPNN采样', 'python protein_mpnn_run.py --pdb_path sfGFP.pdb --num_seq_per_target 500 --sampling_temp "0.1 0.2 0.3" --tied_positions "66 67 96 222"'],
        ['ESMFold预测', 'python -c "from transformers import EsmForProteinFolding; ..."'],
        ['ColabFold预测', 'colabfold_batch seqs.fasta results/ --num-recycle 3'],
        ['BioEmu构象采样', '(参考 github.com/microsoft/bioemu README)'],
        ['GROMACS MD', 'gmx mdrun -v -deffnm npt_72C -nb gpu'],
        ['DNAChisel反向翻译', 'python -c "from dnachisel import reverse_translate; ..."'],
        ['CD-HIT去重', 'cd-hit -i candidates.fasta -o nr.fasta -c 0.95 -n 5'],
        ['jackhmmer MSA', 'jackhmmer -N 3 --incE 1e-10 -A gfp_msa.sto sfGFP.fasta uniref90.fasta'],
        ['EVcouplings分析', 'python -c "from evcouplings import ..."'],
    ],
    [3.0, 11.0]
)

h1('附录B: 45个候选设计位点（来自5.25方案，保留证据等级标注）')

bp('亮度增强位点（8个核心位点）:')
table(
    ['位点', '推荐突变', '预期效果', '机制', '证据等级'],
    [
        ['S65', 'T(保留sfGFP)', '成熟效率+消光系数', '优化环化后氧化步骤', '★★★★★ EGFP/sfGFP核心'],
        ['F64', 'L(保留sfGFP)', '增强37°C折叠效率', '减少α-螺旋空间位阻', '★★★★★ EGFP/sfGFP核心'],
        ['H148', 'D', '量子产率→~0.8', '稳定发色团激发态, 减少非辐射衰变', '★★★★ mBaoJin验证'],
        ['T203', 'I/V/F', '调节激发/发射峰+亮度', '优化π-π堆积', '★★★★ 多种GFP验证'],
        ['S205', 'V/T/A', '稳定去质子化B态', '增强对Y66酚羟基质子接受', '★★★★ 多种变体验证'],
        ['E222', 'D/Q', '调节发色团静电微环境', '影响ESPT速率', '★★★★ avGFP关键催化'],
        ['I167', 'T', '光稳定性+折叠', '减少发色团光氧化', '★★★'],
        ['N149', 'K', '量子产率', '优化β-桶内部氢键网络', '★★★'],
    ],
    [1.2, 2.5, 3.5, 4.0, 2.5]
)

bp('')
bp('热稳定性增强位点（12个核心位点）:')
table(
    ['位点', '推荐突变', '预期效果', '机制', '证据等级'],
    [
        ['S30', 'R(保留sfGFP)', '表面盐桥', '与E17/E32形成静电网络', '★★★★★ sfGFP核心'],
        ['F99', 'S(保留sfGFP)', '改善表面亲水性', '减少疏水暴露, 抑制聚集', '★★★★★ sfGFP核心'],
        ['M153', 'T(保留sfGFP)', '折叠效率', '优化β-桶核心疏水堆积', '★★★★★ sfGFP核心'],
        ['V163', 'A(保留sfGFP)', '折叠效率', '核心堆积优化, 减少侧链冲突', '★★★★★ sfGFP核心'],
        ['G112', 'P', 'Loop刚性化', '限制β-折叠间环区构象自由度', '★★★★ TGP验证'],
        ['Q69', 'R/K', '分子内盐桥', '与E222形成盐桥, 稳定发色团区域', '★★★'],
        ['R122', 'E/D', '表面静电优化', '与K/R形成互补盐桥', '★★★'],
        ['L221', 'I/V', '核心疏水填充', '增强β-桶底部疏水堆积', '★★★'],
        ['H231', 'W/Y/F', '表面芳香堆积', '优化C端区域结构稳定性', '★★★'],
        ['E17', 'R/K', 'N端盐桥网络', '与D21/E32形成静电作用', '★★'],
        ['D76', 'N/S', '减少表面电荷排斥', '优化β-桶外侧静电分布', '★★'],
        ['V169+T203', 'C+C(双突变)', '链内二硫键', 'S-S共价交联β-折叠链, 极高热稳增益', '★★ 高风险高回报'],
    ],
    [2.2, 3.0, 3.0, 3.8, 2.0]
)

h1('附录C: 核心参考文献与工具速查')
bp('参考文献:')
bl('[1] Sarkisyan et al. (2016) Nature 533:397-401. GFP适应度景观. DOI:10.1038/nature17995')
bl('[2] Pédelacq et al. (2006) Nat Biotechnol 24:79-88. Superfolder GFP. DOI:10.1038/nbt1172')
bl('[3] Rives et al. (2025) Science 387:850-858. ESM3/esmGFP. DOI:10.1126/science.ads0018')
bl('[4] Lewis, Hempel et al. (2025) Science. BioEmu. DOI:10.1126/science.adq7214')
bl('[5] Zhang et al. (2024) Nat Methods 21:645-656. mBaoJin. DOI:10.1038/s41592-024-02203-y')
bl('[6] Close et al. (2015) Proteins 83:275-289. TGP. DOI:10.1002/prot.24699')
bl('[7] Montalvillo Ortega et al. (2025) ACS Central Science. LGL+MD协同管线.')

bp('')
bp('工具速查:')
bl('ESM-2: github.com/facebookresearch/esm | ProteinMPNN: github.com/dauparas/ProteinMPNN')
bl('ColabFold: github.com/sokrypton/ColabFold | BioEmu: github.com/microsoft/bioemu')
bl('FoldX: foldxsuite.crg.eu | GROMACS: gromacs.org | DNAChisel: github.com/Edinburgh-Genome-Foundry/DnaChisel')
bl('EVcouplings: github.com/debbiemarkslab/EVcouplings | CD-HIT: github.com/weizhongli/cdhit')
bl('云GPU: AutoDL autodl.com (A100约¥3-5/h) | FPbase: fpbase.org')

doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('—— 文档结束 ——'); cn(r, '微软雅黑'); r.font.size = Pt(12); r.font.color.rgb = RGBColor(150,150,150)

# ════════════════════════════════════════════════════════════
# 保存
# ════════════════════════════════════════════════════════════
out = r'D:\蛋白质设计-合成生物学创新赛-Claude\文档\GFP蛋白质设计_管线设计与实施方案_v2.1.docx'
os.makedirs(os.path.dirname(out), exist_ok=True)
doc.save(out)
print(f'Saved: {out}')
print(f'Size: {os.path.getsize(out)/1024:.1f} KB')
