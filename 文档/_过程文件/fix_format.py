"""
Fix v2.1 document formatting to match v2.0.

Key formatting from v2.0 (REFERENCE):
- Page: A4 (21cm x 29.7cm), margins 2.5cm all sides
- Normal style: Times New Roman (English) / 宋体 (Chinese), 12pt (小四)
- Heading 1: 微软雅黑, 18pt (小二), bold, color #003366
- Heading 2: 微软雅黑, 15pt (小三), bold, color #005082
- Heading 3: 微软雅黑, 14pt (四号), bold, color #4F81BD
- Body text: 宋体 (SimSun) for Chinese, Times New Roman for English, 12pt
- Code: Consolas, 9pt (小五)
- TOC main entries: 微软雅黑, 14pt, bold
- TOC sub-entries: 宋体, 12pt
- Title: 微软雅黑 with colors #003366, #005082, #505050, #CC0000
"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import copy
import re

v20_path = "D:/蛋白质设计-合成生物学创新赛-Claude/文档/GFP蛋白质设计_管线设计与实施方案_v2.0.docx"
v21_path = "D:/蛋白质设计-合成生物学创新赛-Claude/文档/GFP蛋白质设计_管线设计与实施方案_v2.1.docx"
output_path = "D:/蛋白质设计-合成生物学创新赛-Claude/文档/GFP蛋白质设计_管线设计与实施方案_v2.1.docx"

# Load v2.0 as reference
ref_doc = Document(v20_path)
# Load v2.1 as target
doc = Document(v21_path)

# ============================================================
# 1. Fix page size to A4 (21cm x 29.7cm)
# ============================================================
for section in doc.sections:
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ============================================================
# 2. Fix Normal style font
# ============================================================
normal_style = doc.styles['Normal']
normal_font = normal_style.font
normal_font.name = 'Times New Roman'
normal_font.size = Pt(12)  # 小四
# Set East Asian font to 宋体
rPr = normal_style.element.get_or_add_rPr()
rFonts = rPr.find(qn('w:rFonts'))
if rFonts is None:
    rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} />')
    rPr.insert(0, rFonts)
rFonts.set(qn('w:eastAsia'), '宋体')
rFonts.set(qn('w:ascii'), 'Times New Roman')
rFonts.set(qn('w:hAnsi'), 'Times New Roman')

# ============================================================
# 3. Fix Heading 1 style
# ============================================================
h1_style = doc.styles['Heading 1']
h1_font = h1_style.font
h1_font.name = '微软雅黑'
h1_font.size = Pt(18)  # 小二
h1_font.bold = True
h1_font.color.rgb = RGBColor(0x00, 0x33, 0x66)
# Set East Asian font
h1_rPr = h1_style.element.get_or_add_rPr()
h1_rFonts = h1_rPr.find(qn('w:rFonts'))
if h1_rFonts is None:
    h1_rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} />')
    h1_rPr.insert(0, h1_rFonts)
h1_rFonts.set(qn('w:eastAsia'), '微软雅黑')
h1_rFonts.set(qn('w:ascii'), '微软雅黑')
h1_rFonts.set(qn('w:hAnsi'), '微软雅黑')
# Set paragraph spacing
h1_pPr = h1_style.element.get_or_add_pPr()
h1_spacing = h1_pPr.find(qn('w:spacing'))
if h1_spacing is None:
    h1_spacing = parse_xml(f'<w:spacing {nsdecls("w")} />')
    h1_pPr.append(h1_spacing)
h1_spacing.set(qn('w:after'), '240')

# ============================================================
# 4. Fix Heading 2 style
# ============================================================
try:
    h2_style = doc.styles['Heading 2']
except KeyError:
    # Create Heading 2 if it doesn't exist
    from docx.enum.style import WD_STYLE_TYPE
    h2_style = doc.styles.add_style('Heading 2', WD_STYLE_TYPE.PARAGRAPH)
    h2_style.base_style = doc.styles['Normal']

h2_font = h2_style.font
h2_font.name = '微软雅黑'
h2_font.size = Pt(15)  # 小三
h2_font.bold = True
h2_font.color.rgb = RGBColor(0x00, 0x50, 0x82)
h2_rPr = h2_style.element.get_or_add_rPr()
h2_rFonts = h2_rPr.find(qn('w:rFonts'))
if h2_rFonts is None:
    h2_rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} />')
    h2_rPr.insert(0, h2_rFonts)
h2_rFonts.set(qn('w:eastAsia'), '微软雅黑')
h2_rFonts.set(qn('w:ascii'), '微软雅黑')
h2_rFonts.set(qn('w:hAnsi'), '微软雅黑')
h2_pPr = h2_style.element.get_or_add_pPr()
h2_spacing = h2_pPr.find(qn('w:spacing'))
if h2_spacing is None:
    h2_spacing = parse_xml(f'<w:spacing {nsdecls("w")} />')
    h2_pPr.append(h2_spacing)
h2_spacing.set(qn('w:before'), '240')
h2_spacing.set(qn('w:after'), '120')

# ============================================================
# 5. Fix Heading 3 style
# ============================================================
h3_style = doc.styles['Heading 3']
h3_font = h3_style.font
h3_font.name = '微软雅黑'
h3_font.size = Pt(14)  # 四号
h3_font.bold = True
h3_font.color.rgb = RGBColor(0x4F, 0x81, 0xBD)
h3_rPr = h3_style.element.get_or_add_rPr()
h3_rFonts = h3_rPr.find(qn('w:rFonts'))
if h3_rFonts is None:
    h3_rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} />')
    h3_rPr.insert(0, h3_rFonts)
h3_rFonts.set(qn('w:eastAsia'), '微软雅黑')
h3_rFonts.set(qn('w:ascii'), '微软雅黑')
h3_rFonts.set(qn('w:hAnsi'), '微软雅黑')
h3_pPr = h3_style.element.get_or_add_pPr()
h3_spacing = h3_pPr.find(qn('w:spacing'))
if h3_spacing is None:
    h3_spacing = parse_xml(f'<w:spacing {nsdecls("w")} />')
    h3_pPr.append(h3_spacing)
h3_spacing.set(qn('w:before'), '200')
h3_spacing.set(qn('w:after'), '100')

# ============================================================
# 6. Fix List Bullet style
# ============================================================
try:
    lb_style = doc.styles['List Bullet']
    lb_font = lb_style.font
    lb_font.name = 'Times New Roman'
    lb_font.size = Pt(12)
    lb_rPr = lb_style.element.get_or_add_rPr()
    lb_rFonts = lb_rPr.find(qn('w:rFonts'))
    if lb_rFonts is None:
        lb_rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} />')
        lb_rPr.insert(0, lb_rFonts)
    lb_rFonts.set(qn('w:eastAsia'), '宋体')
    lb_rFonts.set(qn('w:ascii'), 'Times New Roman')
    lb_rFonts.set(qn('w:hAnsi'), 'Times New Roman')
except KeyError:
    pass

# ============================================================
# 7. Fix paragraph-level formatting in document body
#    - Replace MS Gothic with 微软雅黑 (for headings/titles) or 宋体 (for body text)
#    - Fix font sizes to match v2.0 conventions
# ============================================================

def is_chinese_char(ch):
    """Check if character is Chinese."""
    return '一' <= ch <= '鿿' or '㐀' <= ch <= '䶿'

def fix_run_font(run, default_east_asian='宋体', default_ascii='Times New Roman', size=None):
    """Fix a run's font to match v2.0 conventions."""
    font = run.font

    # If font is MS Gothic, replace it
    current_font = font.name
    if current_font == 'MS Gothic':
        # Check if this text is part of a heading/title (bold/large) or body text
        if font.bold or (font.size and font.size >= Pt(14)):
            font.name = '微软雅黑'
        else:
            font.name = default_ascii

    # Fix East Asian font in XML
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is not None:
        east = rFonts.get(qn('w:eastAsia'))
        if east == 'MS Gothic':
            # Check context
            if font.bold or (font.size and font.size >= Pt(14)):
                rFonts.set(qn('w:eastAsia'), '微软雅黑')
                rFonts.set(qn('w:ascii'), '微软雅黑')
                rFonts.set(qn('w:hAnsi'), '微软雅黑')
            else:
                rFonts.set(qn('w:eastAsia'), default_east_asian)
                rFonts.set(qn('w:ascii'), default_ascii)
                rFonts.set(qn('w:hAnsi'), default_ascii)

    if size is not None:
        font.size = size

# Process each paragraph
import re

for para in doc.paragraphs:
    style_name = para.style.name if para.style else 'Normal'
    text = para.text.strip()

    # Determine paragraph type
    is_heading1 = style_name == 'Heading 1'
    is_heading2 = style_name == 'Heading 2'
    is_heading3 = style_name == 'Heading 3'
    is_toc = style_name == 'Heading 3'  # TOC items use Heading 3 in v2.1
    is_list_bullet = style_name == 'List Bullet'

    for run in para.runs:
        # Replace MS Gothic everywhere
        fix_run_font(run)

        # Fix heading font colors and sizes
        if is_heading1:
            run.font.color.rgb = RGBColor(0x00, 0x33, 0x66)
            run.font.size = Pt(18)
            run.font.name = '微软雅黑'
            # Fix XML level
            rPr = run._element.get_or_add_rPr()
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is None:
                rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} />')
                rPr.insert(0, rFonts)
            rFonts.set(qn('w:eastAsia'), '微软雅黑')
            rFonts.set(qn('w:ascii'), '微软雅黑')
            rFonts.set(qn('w:hAnsi'), '微软雅黑')

        elif is_heading2:
            run.font.color.rgb = RGBColor(0x00, 0x50, 0x82)
            run.font.size = Pt(15)
            run.font.name = '微软雅黑'
            rPr = run._element.get_or_add_rPr()
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is None:
                rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} />')
                rPr.insert(0, rFonts)
            rFonts.set(qn('w:eastAsia'), '微软雅黑')
            rFonts.set(qn('w:ascii'), '微软雅黑')
            rFonts.set(qn('w:hAnsi'), '微软雅黑')

        elif is_heading3:
            run.font.color.rgb = RGBColor(0x4F, 0x81, 0xBD)
            run.font.size = Pt(14)
            run.font.name = '微软雅黑'
            rPr = run._element.get_or_add_rPr()
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is None:
                rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} />')
                rPr.insert(0, rFonts)
            rFonts.set(qn('w:eastAsia'), '微软雅黑')
            rFonts.set(qn('w:ascii'), '微软雅黑')
            rFonts.set(qn('w:hAnsi'), '微软雅黑')

    # Fix paragraph spacing
    pf = para.paragraph_format
    if is_toc and not is_heading1 and not is_heading2:
        # TOC spacing in v2.0: before=80, after=80 (in twips)
        from docx.shared import Twips
        pf.space_before = Twips(80)
        pf.space_after = Twips(80)

# ============================================================
# 8. Fix the title page (first few paragraphs)
# ============================================================
# v2.0 title page format:
# - P6 (main title): 微软雅黑, 26pt, color 003366, center
# - P7 (subtitle): 微软雅黑, 22pt, color 005082, center
# - P9 (doc title): 微软雅黑, 30pt, bold, color 003366, center
# - P12 (subtitle desc): 微软雅黑, 16pt, color 505050, center
# - P17 (version): 微软雅黑, 14pt, color 505050, center
# - P18 (deadline): 微软雅黑, 14pt, color CC0000, center

# Count title paragraphs
title_paras = []
for i, para in enumerate(doc.paragraphs):
    if para.style.name == 'Heading 1' and i < 5:
        title_paras.append(i)

print(f"Title paragraphs: {title_paras}")

# Fix the first heading 1 paragraphs (title)
for i, para in enumerate(doc.paragraphs):
    if para.style.name == 'Heading 1':
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in para.runs:
            run.font.color.rgb = RGBColor(0x00, 0x33, 0x66)
            run.font.size = Pt(26) if i < 20 else Pt(18)  # title gets 26pt, other H1 get 18pt
            run.font.name = '微软雅黑'
            run.font.bold = True
            # Fix font in XML
            rPr = run._element.get_or_add_rPr()
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is None:
                rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} />')
                rPr.insert(0, rFonts)
            rFonts.set(qn('w:eastAsia'), '微软雅黑')
            rFonts.set(qn('w:ascii'), '微软雅黑')
            rFonts.set(qn('w:hAnsi'), '微软雅黑')

# ============================================================
# 9. Save the fixed document
# ============================================================
doc.save(output_path)
print(f"Saved fixed document to: {output_path}")
print("Done!")
