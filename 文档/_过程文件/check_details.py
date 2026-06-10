"""Check remaining format differences."""
from docx import Document
from docx.shared import Pt

ref = Document("D:/蛋白质设计-合成生物学创新赛-Claude/文档/GFP蛋白质设计_管线设计与实施方案_v2.0.docx")
doc = Document("D:/蛋白质设计-合成生物学创新赛-Claude/文档/GFP蛋白质设计_管线设计与实施方案_v2.1.docx")

# Check line spacing patterns
print("=== Line Spacing Patterns ===")
for label, d in [("v2.0 REF", ref), ("v2.1 FIXED", doc)]:
    ls_counts = {}
    for para in d.paragraphs:
        ls = para.paragraph_format.line_spacing
        ls_counts[ls] = ls_counts.get(ls, 0) + 1
    print(f"\n{label}:")
    for ls, count in sorted(ls_counts.items(), key=lambda x: -x[1])[:5]:
        print(f"  line_spacing={ls}: {count} paragraphs")

# Check font table
print("\n=== Font Table ===")
for label, d in [("v2.0 REF", ref), ("v2.1 FIXED", doc)]:
    fonts = set()
    for para in d.paragraphs:
        for run in para.runs:
            if run.font.name:
                fonts.add(run.font.name)
    print(f"\n{label}: Fonts used: {sorted(fonts)}")

# Compare body paragraph formats
print("\n=== Body Paragraph Format Comparison ===")
# Find first body paragraph with substantial text
for i, (rp, dp) in enumerate(zip(ref.paragraphs, doc.paragraphs)):
    rt = rp.text.strip()
    dt = dp.text.strip()
    if len(rt) > 80 and len(dt) > 80:
        rpf = rp.paragraph_format
        dpf = dp.paragraph_format
        print(f"v2.0 P{i}: line_spacing={rpf.line_spacing}, space_before={rpf.space_before}, space_after={rpf.space_after}")
        print(f"v2.1 P{i}: line_spacing={dpf.line_spacing}, space_before={dpf.space_before}, space_after={dpf.space_after}")
        print(f"  v2.0 text: {rt[:60]}...")
        print(f"  v2.1 text: {dt[:60]}...")

        # Check if v2.0 runs have explicit font
        for j, (rr, dr) in enumerate(zip(rp.runs[:3], dp.runs[:3])):
            print(f"  v2.0 R{j}: font={rr.font.name}, size={rr.font.size}")
            print(f"  v2.1 R{j}: font={dr.font.name}, size={dr.font.size}")
        break

# Check numbering
print("\n=== Numbering ===")
for label, d in [("v2.0 REF", ref), ("v2.1 FIXED", doc)]:
    num_count = 0
    for para in d.paragraphs:
        if para.paragraph_format.first_line_indent:
            num_count += 1
    print(f"{label}: paragraphs with first_line_indent: {num_count}")
