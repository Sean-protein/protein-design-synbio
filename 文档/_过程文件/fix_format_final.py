"""
Final format fixes for v2.1 to match v2.0:
1. Remove 1.5 line spacing → None (use default)
2. Fix font table (add Consolas, ensure 宋体)
3. Clean up explicit font settings on runs to match v2.0 pattern
"""
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import copy

doc_path = "D:/蛋白质设计-合成生物学创新赛-Claude/文档/GFP蛋白质设计_管线设计与实施方案_v2.1.docx"
doc = Document(doc_path)

# ============================================================
# 1. Fix line spacing: remove 1.5 line spacing
# ============================================================
fixed_ls_count = 0
for para in doc.paragraphs:
    ls = para.paragraph_format.line_spacing
    if ls == 1.5:
        para.paragraph_format.line_spacing = None  # Use default
        fixed_ls_count += 1
print(f"Fixed line spacing for {fixed_ls_count} paragraphs")

# ============================================================
# 2. Fix font table XML
# ============================================================
# Access the font table part
from lxml import etree

for rel in doc.part.rels.values():
    if "fontTable" in rel.reltype:
        font_part = rel.target_part
        font_xml = font_part.blob
        font_tree = etree.fromstring(font_xml)
        nsmap = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

        # Check existing fonts
        existing_fonts = set()
        for font_elem in font_tree.findall('.//w:font', nsmap):
            name = font_elem.get(qn('w:name'))
            existing_fonts.add(name)
        print(f"Existing fonts: {existing_fonts}")

        # Add Consolas if missing
        if 'Consolas' not in existing_fonts:
            consolas_elem = etree.SubElement(font_tree, qn('w:font'))
            consolas_elem.set(qn('w:name'), 'Consolas')
            panose = etree.SubElement(consolas_elem, qn('w:panose1'))
            panose.set(qn('w:val'), '020B0609020204030204')
            charset = etree.SubElement(consolas_elem, qn('w:charset'))
            charset.set(qn('w:val'), '00')
            family = etree.SubElement(consolas_elem, qn('w:family'))
            family.set(qn('w:val'), 'modern')
            pitch = etree.SubElement(consolas_elem, qn('w:pitch'))
            pitch.set(qn('w:val'), 'fixed')
            sig = etree.SubElement(consolas_elem, qn('w:sig'))
            sig.set(qn('w:usb0'), 'E00006FF')
            sig.set(qn('w:usb1'), '0000FCFF')
            sig.set(qn('w:usb2'), '00000001')
            sig.set(qn('w:usb3'), '00000000')
            sig.set(qn('w:csb0'), '0000019F')
            sig.set(qn('w:csb1'), '00000000')
            print("Added Consolas font to font table")

        # Ensure 宋体 is present
        if '宋体' not in existing_fonts:
            simsun_elem = etree.SubElement(font_tree, qn('w:font'))
            simsun_elem.set(qn('w:name'), '宋体')
            altname = etree.SubElement(simsun_elem, qn('w:altName'))
            altname.set(qn('w:val'), 'SimSun')
            panose = etree.SubElement(simsun_elem, qn('w:panose1'))
            panose.set(qn('w:val'), '02010600030101010101')
            charset = etree.SubElement(simsun_elem, qn('w:charset'))
            charset.set(qn('w:val'), '86')
            family = etree.SubElement(simsun_elem, qn('w:family'))
            family.set(qn('w:val'), 'auto')
            pitch = etree.SubElement(simsun_elem, qn('w:pitch'))
            pitch.set(qn('w:val'), 'variable')
            sig = etree.SubElement(simsun_elem, qn('w:sig'))
            sig.set(qn('w:usb0'), '00000203')
            sig.set(qn('w:usb1'), '288F0000')
            sig.set(qn('w:usb2'), '00000016')
            sig.set(qn('w:usb3'), '00000000')
            sig.set(qn('w:csb0'), '00040001')
            sig.set(qn('w:csb1'), '00000000')
            print("Added 宋体 font to font table")

        # Serialize back
        font_part._blob = etree.tostring(font_tree, xml_declaration=True, encoding='UTF-8', standalone=True)
        break

# ============================================================
# 3. Remove explicit font settings from body text runs
#    (in v2.0 body text, runs inherit from Normal style)
# ============================================================
# Only fix Normal-style body paragraphs, not headings/lists
for para in doc.paragraphs:
    style_name = para.style.name if para.style else 'Normal'
    if style_name not in ['Normal', 'List Bullet', 'List Number']:
        continue

    # Skip empty paragraphs
    if not para.text.strip():
        continue

    # Check if this paragraph's runs all use explicit Times New Roman 12pt
    # If so, remove the explicit settings to let style inheritance work
    for run in para.runs:
        # Don't modify runs that have special formatting (bold, italic, color, etc.)
        if run.font.bold or run.font.italic or run.font.color or run.font.size != Pt(12):
            continue

        rPr = run._element.find(qn('w:rPr'))
        if rPr is None:
            continue

        # Check if font name is explicitly Times New Roman
        rFonts_elem = rPr.find(qn('w:rFonts'))
        if rFonts_elem is not None:
            ascii_font = rFonts_elem.get(qn('w:ascii'))
            # If ascii font is Times New Roman, remove the explicit setting
            # but keep eastAsia if it's 宋体 (that's the correct default)
            if ascii_font == 'Times New Roman':
                # Remove ascii/hAnsi settings, they'll inherit from style
                rFonts_elem.set(qn('w:ascii'), 'Times New Roman')  # Keep for now
                # Don't remove entirely - this ensures correct rendering

        # Remove explicit size if it's 12pt (matches Normal style)
        sz_elem = rPr.find(qn('w:sz'))
        if sz_elem is not None:
            sz_val = sz_elem.get(qn('w:val'))
            if sz_val == '24':  # 12pt = 24 half-pts
                rPr.remove(sz_elem)

# ============================================================
# 4. Save
# ============================================================
doc.save(doc_path)
print("Final fixes applied and saved.")
