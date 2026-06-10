#!/usr/bin/env python3
"""
GFP Pipeline Implementation Guide — Step-by-step for beginners
Each phase explains: WHAT, WHY, KEY TERMS, PREREQS, INPUT, OUTPUT, CODE, RESULT INTERPRETATION
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

doc = Document()
for s in doc.sections:
    s.page_width, s.page_height = Cm(21.0), Cm(29.7)
    s.top_margin = s.bottom_margin = s.left_margin = s.right_margin = Cm(2.5)

# Styles
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style.paragraph_format.line_spacing = 1.2
rPr = style.element.get_or_add_rPr()
rPr.insert(0, parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="宋体" w:ascii="Times New Roman" w:hAnsi="Times New Roman"/>'))

for sname, font_name, size, color in [
    ('Heading 1', '微软雅黑', Pt(18), '003366'),
    ('Heading 2', '微软雅黑', Pt(15), '005082'),
    ('Heading 3', '微软雅黑', Pt(14), '4F81BD'),
]:
    s = doc.styles[sname]
    s.font.name = font_name; s.font.size = size; s.font.bold = True
    s.font.color.rgb = RGBColor(*[int(color[i:i+2],16) for i in range(0,6,2)])
    srPr = s.element.get_or_add_rPr()
    srPr.insert(0, parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{font_name}" w:ascii="{font_name}" w:hAnsi="{font_name}"/>'))

# ============================================================
# HELPERS
# ============================================================
def H(text, level=1): return doc.add_heading(text, level=level)
def P(text, bold=False, c=None, s=None, al=None):
    p = doc.add_paragraph(); r = p.add_run(text)
    if bold: r.bold = True
    if c: r.font.color.rgb = RGBColor(*c)
    if s: r.font.size = s
    if al is not None: p.alignment = al
    return p
def B(text): p = doc.add_paragraph(text, style='List Bullet'); return p
def BR(): doc.add_page_break()
def CODE(text):
    p = doc.add_paragraph(); r = p.add_run(text)
    r.font.name = 'Consolas'; r.font.size = Pt(9)
    pPr = p._element.get_or_add_pPr()
    pPr.append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="F0F0F0" w:val="clear"/>'))
    return p
def SEC(title):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title); r.bold = True; r.font.size = Pt(16)
    r.font.name = '微软雅黑'; r.font.color.rgb = RGBColor(0x00, 0x33, 0x66)
    pPr = p._element.get_or_add_pPr()
    pPr.append(parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="6" w:space="1" w:color="003366"/></w:pBdr>'))
    return p
def TBL(headers, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(headers)); t.style = 'Light Grid Accent 1'
    for j,h in enumerate(headers):
        c=t.rows[0].cells[j];c.text=h
        for r in c.paragraphs[0].runs: r.bold=True;r.font.size=Pt(10)
    for i,row in enumerate(rows):
        for j,v in enumerate(row):
            c=t.rows[i+1].cells[j];c.text=str(v)
            for r in c.paragraphs[0].runs:r.font.size=Pt(10)
    return t
def note(text):
    p = doc.add_paragraph(); r = p.add_run('💡 ' + text)
    r.font.size = Pt(11); r.font.color.rgb = RGBColor(0x00, 0x50, 0x82)
    return p
def warn(text):
    p = doc.add_paragraph(); r = p.add_run('⚠️ ' + text)
    r.font.size = Pt(11); r.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
    return p

# ============================================================
# TITLE PAGE
# ============================================================
for _ in range(6): doc.add_paragraph()
P('GFP 高亮度与热稳定性联合设计', al=WD_ALIGN_PARAGRAPH.CENTER, s=Pt(28), c=(0x00,0x33,0x66), bold=True)
P('零基础实操指南', al=WD_ALIGN_PARAGRAPH.CENTER, s=Pt(24), c=(0x00,0x50,0x82), bold=True)
doc.add_paragraph()
P('从名词解释到完整代码，一步一步跑通蛋白质设计管线', al=WD_ALIGN_PARAGRAPH.CENTER, s=Pt(14), c=(0x50,0x50,0x50))
P('版本：v1.0  |  2026年6月9日', al=WD_ALIGN_PARAGRAPH.CENTER, s=Pt(12), c=(0x50,0x50,0x50))
BR()

# ============================================================
# GLOSSARY
# ============================================================
SEC('〇、名词字典——先把行话翻译成人话')

P('在动手之前，把管线中用到的所有专业名词一次性解释清楚。每个名词附带"类比"帮你建立直觉。', bold=True)

glossary = [
    ("PLM (蛋白质语言模型)",
     "把蛋白质序列当成'一句话'，氨基酸字母就是'单词'。PLM 读了数亿条蛋白质序列后，学会了'蛋白质的语法'——哪些氨基酸经常一起出现、哪些位置的替换是合理的。就像训练 ChatGPT 学会了人类语言的语法一样。",
     "你用到的：ESM-2（Meta做的）、ESM-3（最新版）、SaProt（结构感知版）都属于PLM。"),
    ("嵌入 (Embedding)",
     "PLM 把每条蛋白质序列变成一个固定长度的数字向量——比如 1280 个浮点数。这个向量是序列的'压缩版身份证'：相似的序列向量距离近，不相似的向量距离远。",
     "类比：把一张人脸照片压缩成 128 个数字特征（眼距、鼻高、脸宽……），不同人的特征向量可以用来计算相似度。"),
    ("MSA (多序列比对)",
     "把 sfGFP 和它的'进化亲戚们'（其他荧光蛋白、不同物种的 GFP 同源物）排成一行行，对齐相同位置。像把同班同学的试卷平铺在桌上，看哪道题所有人都对（关键位置），哪道题答什么的都有（不重要位置）。",
     "你用 jackhmmer 工具对 UniRef30 数据库（含 1 亿+ 蛋白序列）做 3 轮搜索来构建 MSA。"),
    ("EVcouplings (进化耦合分析)",
     "从 MSA 中挖掘'共进化'信号的统计方法。如果位置 148 的氨基酸种类变了，位置 203 也必须跟着变——说明这两个位置在三维结构里有物理接触。EVcouplings 帮你回答：'这两个突变能同时做吗？还是一起做就死了？'",
     "策略D 的核心产出——一个'突变兼容性矩阵'，告诉其他策略哪些组合是安全的。"),
    ("ProteinMPNN (蛋白质逆折叠网络)",
     "给定蛋白质骨架的三维坐标 → 输出能折叠成这个形状的氨基酸序列。方向是：结构→序列（所以叫'逆'折叠）。用深度学习训练，2022 年发在 Science 上。",
     "类比：给你一栋楼的钢筋骨架图，让你设计外墙材料——ProteinMPNN 就是做这个的 AI。"),
    ("ESM-3 (最新一代蛋白质语言模型)",
     "2025 年发表在 Science 上。比 ESM-2 多一个关键能力：它不仅'读'氨基酸序列，还能'读'蛋白质的三维结构（二级结构+骨架坐标）。竞赛主办方用 ESM-3 生成了一条全新的 GFP——esmGFP——和任何已知 GFP 只有 58% 相似但仍发荧光。",
     "你的管线在两个地方用 ESM-3：①策略B 用它做结构感知的嵌入；②策略E 用它生成全新序列。"),
    ("ESM-3 Gibbs 采样",
     "Gibbs 采样 = 交替优化的循环。ESM-3 先根据你给的结构提示'猜'一个三维结构 → 基于这个结构'猜'一条新序列 → 再用序列验证结构是否合理 → 反复 30+ 轮。每轮微调一点，逐步逼近'结构合理+序列新颖'的组合。esmGFP 就是这样生成的。",
     "类比：先画房子外形 → 设计内部布局 → 发现不对劲调整外形 → 再调整内部……来回改直到满意。"),
    ("LoRA 微调 (Low-Rank Adaptation)",
     "不完全重新训练 ESM-3（那需要几百万美元），而是在模型旁边加一个小'插件'，只让这个插件学习 GFP 相关的知识。成本：几小时 + 一张 RTX 3090 游戏显卡。",
     "类比：GPT-4 是一个通用翻译官，LoRA 是在它耳朵旁放一个小耳机，只给它翻译 GFP 相关的内容。"),
    ("MCMC (马尔可夫链蒙特卡洛)",
     "一种'有方向的随机游走'算法。不是在所有可能的序列里乱撞，而是：当前位置好（预测亮度高）→ 多留一会儿再走；当前位置差 → 随机跳到一个新方向。在嵌入空间中游走（而非直接在氨基酸字母上改），因为嵌入空间更'平滑'——走一小步对应'微妙调整'而非彻底变一条序列。",
     "类比：在山脉中找最高峰——MCMC 不是每步都往上爬（那样可能困在小山头），而是有时往低处走一走再找新方向，最终到达最高峰。"),
    ("GP (高斯过程)",
     "一种能做预测并给出不确定性的数学模型。对你的管线，GP 学习'嵌入向量→乘积得分'的映射关系。关键优势：它不仅告诉你'这个序列预测得分 0.75'，还告诉你'我不太确定，误差可能是 ±0.15'——不确定性高的地方就是下一轮需要探索的地方。",
     "类比：普通预测模型是'天气预报说下雨'，GP 是'60% 概率下雨，但我的置信区间比较宽因为数据不够'。"),
    ("EI (Expected Improvement, 期望改进)",
     "基于 GP 的不确定性算出的'探索价值'指标。EI 在预测好且不确定性也高的地方最大——'这里看起来可能有金矿，而且我还没挖过，值得去挖'。管线用 EI 指导下一轮序列生成方向。",
     "类比：玩扫雷——已经挖开的区域知道没有雷，但旁边那块'4'旁边的未知格最值得点。"),
    ("CD-HIT",
     "一个快速去重工具。蛋白质序列的'相似度'在氨基酸层面比较。设定 90% 阈值 = 如果两条序列 90% 的氨基酸位置相同，只保留一条。确保提交的 6 条序列不是'近重复'——评委一眼就能看出来。",
     "用命令行直接跑，输入 fasta 文件，输出去重后的 fasta。"),
    ("BioEmu (构象系综快速模拟)",
     "微软 2025 年发在 Science 上的方法。传统 MD 模拟 10 纳秒需要 15 小时（GPU），BioEmu 用 AI 学过的蛋白质动力学做同样的模拟只需要几分钟。不是 100% 精确，但非常适合做'快速预筛'——200 条候选先过 BioEmu，只有通过的 ~15 条才送去慢但精确的 MD。",
     "类比：用快速试纸做初筛 → 阳性样品再送 PCR 精检。"),
    ("MD (分子动力学模拟) — '金标准'",
     "在计算机里建立一个虚拟的'水盒子'，把蛋白质放进去，用牛顿力学模拟每个原子的运动。对每条候选序列在 72°C 下模拟 10 纳秒（1纳秒=十亿分之一秒），然后降到 25°C 看蛋白质能不能'恢复'。分析 RMSD（结构变形程度）、Rg（紧凑度）、SASA（表面积暴露程度）、氢键保留率。叫'金标准'因为这是目前计算上最接近实验的方法——但也是最慢最费GPU的。",
     "用 GROMACS 软件+CHARMM36m 力场运行。"),
    ("pLDDT & pTM",
     "AlphaFold2/ESMFold 预测结构时自己给出的'自信分'。pLDDT=预测的局部结构精度（0-100，>80=靠谱，<70=可疑）；pTM=预测的全局结构正确概率（0-1，>0.75=整体折叠应该是正确的）。管线用这两个指标判断候选序列能否正常折叠——pLDDT<70 的直接淘汰。",
     "类比：出题人自己对答案的自信心——'这份答案我有 85 分把握是对的'。"),
    ("FoldX ΔΔG",
     "FoldX 是一个计算工具，预测一个突变对蛋白质稳定性的影响。ΔΔG = 突变体折叠自由能 − 野生型折叠自由能。正值=突变使蛋白不稳定，负值=突变使蛋白更稳定。管线用它在 Phase 4 做快速稳定性筛选——200 条序列先过 FoldX，秒级。",
     "类比：天气预报说'今天比昨天热 5 度'（ΔΔG=+5）——这个突变比原始序列'不稳定了 5 个单位'。"),
    ("DBLT 迭代闭环",
     "Design→Build→Test→Learn 的工程范式套用到蛋白质设计上。第1轮：生成一堆候选→计算预测→模拟验证→从结果中学习。第2轮：基于第1轮学到的知识，更聪明地生成下一批。第3轮：再迭代一次。完成后选出最优 6 条。",
     "类比：做菜→尝味道→调整调料→再做→再尝→直到完美。"),
    ("Pareto 前沿",
     "在多目标优化中，'不能再同时改进所有目标'的解的集合。你的两个目标是亮度和热稳定性——Pareto 前沿就是'提高亮度必然降低稳定性，提高稳定性必然降低亮度'的那条边界线。从前沿上选序列保证你选的是'最优权衡'。",
     "类比：买车——在价格和安全性之间，帕累托前沿是'在这个价位上你能买到的最安全的车'。"),
    ("ThermoMPNN",
     "ProteinMPNN 的一个变体，专门预测突变对热稳定性的影响。它学习了哪些序列特征与高温稳定性相关。管线用它做第一级稳定性快速筛选。",
     "秒级/条，比 MD 快约 50,000 倍，但精度不如 MD。"),
    ("DNA Chisel",
     "竞赛主办方使用的开源工具，把氨基酸序列反向翻译成 DNA 序列。你只提交氨基酸序列，DNA 合成由大设施统一处理。这个工具的存在意味着你不需要操心密码子优化——管线中不需要考虑密码子层面的问题。",
     "不用你跑，自动完成。"),
    ("Sarkisyan 2016",
     "2016 年发表在 Nature 上的里程碑论文。作者构建了 avGFP 的'适应度景观'——系统性地测量了数千个突变体的荧光亮度。这篇论文是管线策略B训练数据的重要补充来源。原文标题：'Local fitness landscape of the green fluorescent protein'。",
     "数据公开可用，约 5 万条突变-亮度标注。"),
]

for term, explanation, analogy in glossary:
    H(term, level=2)
    P(explanation, bold=True)
    P(analogy, c=(0x50,0x50,0x50))

# ============================================================
# PIPELINE OVERVIEW
# ============================================================
BR()
SEC('管线总览——六个阶段')

P('在开始具体操作之前，先建立对整个项目的全局认知。以下是你接下来要做的六件事：', bold=True)

TBL(
    ['阶段', '做什么', '核心工具', '输出', '时间', '可跳过？'],
    [
        ['〇', '环境搭建+名词理解', 'pip/conda', '可工作的环境', '半天', '❌ 必须'],
        ['一', '数据准备：下载比赛资料+熟悉格式', 'pandas/Biopython', '理解数据长什么样', '半天', '❌ 必须'],
        ['二', '策略D：构建MSA→EVcouplings突变规则', 'jackhmmer/EVcouplings', '突变兼容性矩阵', '1天', '⚠️ 可降级但建议做'],
        ['三', '策略A：理性枚举组合突变→FoldX筛选', 'Biopython/FoldX', '~3000条保守候选', '1天', '❌ 必须(保底)'],
        ['四', '策略B：ESM2嵌入→集成ML训练→ESM3 LoRA', 'fair-esm/peft/GPyTorch', '亮度预测器+MCMC搜索', '2-3天', '⚠️ 可降级(纯策略A+C)'],
        ['五', '策略C+E：ProteinMPNN+ESM3生成结构序列', 'ProteinMPNN/ESM-3', '~6000条创新序列', '2-3天', '⚠️ 可降级(策略C)'],
        ['六', '漏斗筛选+DBLT迭代+精选6条+提交', 'AF2/BioEmu/MD/GROMACS', '6条序列+CSV+PDF', '5-7天', '❌ 必须'],
    ]
)

P('')
P('原则：每个阶段都可以独立运行。如果某个阶段失败（比如 GPU 不够跑不动 ESM-3），不会影响其他阶段的产出。管线设计天然支持逐级降级。', bold=True, c=(0x00,0x50,0x82))

note('如果你是纯新手且只有笔记本（无GPU），可以只跑阶段〇+一+二+三+六，跳过四和五。策略A（理性枚举）是整套管线的"保底方案"，不需要 GPU。')

# ============================================================
# PHASE 0
# ============================================================
BR()
SEC('阶段〇：环境搭建')

H('0.1 Step-by-step 环境安装', level=2)
P('以下命令在你的终端中逐条执行。每完成一条，确认没有报错再继续下一条。', bold=True)

H('Step 1: 确认 Python ≥ 3.9', level=3)
CODE('python --version  # 应该输出 Python 3.9.x 或更高')

H('Step 2: 激活 conda 环境', level=3)
CODE('conda activate gfp_design')
P('如果没有这个环境，先创建：')
CODE("conda create -n gfp_design python=3.10 -y\nconda activate gfp_design")

H('Step 3: 安装 PyTorch（选对应你的CUDA版本）', level=3)
CODE('# 有 NVIDIA GPU（推荐）:\npip install torch torchvision torchaudio --index-url https://download.pytorch.org/whl/cu118\n\n# 没有 GPU（Mac/纯CPU）:\npip install torch torchvision torchaudio')

H('Step 4: 安装蛋白质核心库', level=3)
CODE('pip install fair-esm                    # ESM-2 蛋白质语言模型\npip install transformers datasets accelerate  # HuggingFace 模型\npip install peft                         # LoRA 微调\npip install biopython                    # 序列 I/O 操作\npip install pandas numpy scikit-learn    # 数据处理+ML\npip install xgboost lightgbm             # 集成学习\npip install matplotlib seaborn           # 可视化')

H('Step 5: 安装结构相关', level=3)
CODE('pip install py3Dmol                    # 3D 结构预览\n# 如果需要 MS A分析:\nconda install -c conda-forge hhsuite    # jackhmmer\nconda install -c conda-forge mafft       # 序列比对')

H('Step 6: 验证环境', level=3)
CODE("python -c \"import torch; print('CUDA可用:', torch.cuda.is_available())\"\npython -c \"import esm; print('ESM-2导入成功')\"\npython -c \"from Bio import SeqIO; print('Biopython导入成功')\"\npython -c \"import pandas; print('pandas导入成功')\"")

H('Step 7: 安装 ESM-3（可选，需要 HuggingFace 账号）', level=3)
P('ESM-3 需要先在 huggingface.co 注册账号，然后去 huggingface.co/biohub/esm3-sm-open-v1 页面接受非商用许可协议，再生成一个 Access Token。')
CODE('pip install esm                         # EvolutionaryScale官方SDK\nhuggingface-cli login                    # 粘贴你的Access Token\n\n# 验证\npython -c "from esm.models.esm3 import ESM3; print(\'ESM3导入成功\')"')

H('Step 8: 安装 GROMACS（MD模拟，Windows需 WSL2）', level=3)
P('如果你用的是 Windows，需要在 WSL2 (Ubuntu) 中安装：')
CODE('# 在 WSL2 Ubuntu 终端中:\nsudo apt-get update\nsudo apt-get install gromacs\n\n# 验证\ngmx --version')
P('macOS: brew install gromacs')
P('纯 Linux: sudo apt-get install gromacs')

H('0.2 确认你的工作目录结构', level=2)
CODE('mkdir -p ~/gfp_design_pipeline/{data,output,cache,scripts,logs}\ncd ~/gfp_design_pipeline')
P('将竞赛提供的以下文件放入 data/ 目录：GFP_data.xlsx、Exclusion_List.csv、5条参考GFP序列的fasta文件、sfGFP PDB文件(2B3P)。')

# ============================================================
# PHASE 1
# ============================================================
BR()
SEC('阶段一：数据准备——理解你要处理的东西')

H('1.1 竞赛数据包里有什么', level=2)

P('竞赛官方提供了以下数据文件。你先把它们全部放到 data/ 目录下，然后逐一理解：', bold=True)

TBL(
    ['文件名', '格式', '内容', '你在管线中的用途'],
    [
        ['GFP_data.xlsx', 'Excel (多sheet)', '数万条GFP突变体的亮度测量值；往年Top20序列', '策略B的ML训练数据'],
        ['Exclusion_List.csv', 'CSV', '禁止提交的序列名单', '提交前逐条比对（合规！）'],
        ['AAseqs of 5 GFP proteins.txt', 'fasta', 'sfGFP/avGFP/amacGFP/cgreGFP/ppluGFP五条参考序列', '策略D的MSA种子；策略A的设计起点'],
        ['submission_template.csv', 'CSV', '提交格式模板(Team_Name,Seq_ID,Sequence)', '最终提交时参照'],
        ['sfGFP PDB 2B3P', 'PDB结构文件', 'sfGFP的三维原子坐标', 'ProteinMPNN的模板；ESM3的条件输入；AF2验证的参照'],
    ]
)

H('1.2 第一个脚本：看一眼训练数据长什么样', level=2)

P('这是你管线的第一个 Python 脚本。它读取竞赛的训练数据，打印基本信息。', bold=True)

CODE("""# scripts/01_explore_data.py
import pandas as pd

# 读取训练数据
df = pd.read_excel('data/GFP_data.xlsx', sheet_name=0)  # 第1个sheet
print('=== 训练数据概览 ===')
print(f'行数: {len(df):,}')
print(f'列名: {list(df.columns)}')
print(f'\\n前5行:')
print(df.head())
print(f'\\n数据类型:')
print(df.dtypes)

# 检查亮度分布
if 'brightness' in df.columns or 'F' in df.columns:
    col = 'brightness' if 'brightness' in df.columns else 'F'
    print(f'\\n=== {col} 分布 ===')
    print(df[col].describe())

# 读取排除名单
exclusion = pd.read_csv('data/Exclusion_List.csv')
print(f'\\n=== 排除名单 ===')
print(f'禁止序列数: {len(exclusion)}')
print(f'列名: {list(exclusion.columns)}')

# 读取参考序列
from Bio import SeqIO
print(f'\\n=== 参考序列 ===')
for record in SeqIO.parse('data/AAseqs of 5 GFP proteins.txt', 'fasta'):
    print(f'{record.id}: {len(record.seq)} aa')
    print(f'  前20aa: {str(record.seq[:20])}')""")

note('运行这个脚本：python scripts/01_explore_data.py。确认你能看到数据内容后再继续。')

H('1.3 理解 sfGFP：你的设计起点', level=2)

P('sfGFP (superfolder GFP) 是本竞赛的野生型基准。你的所有设计都将从这条序列出发或以此为模板。它全长 ~238 个氨基酸，三维结构是一个"β桶"（11条β链围成的圆柱体），中央孔洞里藏着发色团。', bold=True)

P('sfGFP 序列：')
CODE('MSKGEELFTGVVPILVELDGDVNGHKFSVRGEGEGDATNGKLTLKFICTTGKLPVPWPTLVTTLTYGVQCFSRYPDHMKRHDFFKSAMPEGYVQERTISFKDDGTYKTRAEVKFEGDTLVNRIELKGIDFKEDGNILGHKLEYNFNSHNVYITADKQKNGIKANFKIRHNIVEDGSVQLADHYQQNTPIGDGPVLLPDNHYLSTQSVLSKDPNEKRDHMVLLEFVTAAGITHGMDELYK')

P('你需要记住的关键残基（已在下图标注）：')
TBL(
    ['残基', '角色', '是否可突变'],
    [
        ['T65-Y66-G67', '发色团三联体（自发环化生成荧光团）', 'G67绝对不可动；Y66仅限芳香族；T65仅限小残基'],
        ['R96', '催化发色团环化脱水的关键碱基', '大概率需保留（可尝试协同替代但风险极高）'],
        ['E222', '催化发色团氧化的通用碱', '大概率需保留（E222D部分保留功能）'],
        ['S30/F64/S65T/F99/M153/V163', 'sfGFP的6个折叠增强突变', '建议保留（它们是sfGFP"superfolder"的根本原因）'],
        ['其余 ~227 个残基', 'β桶骨架+表面残基', '可通过ProteinMPNN/ESM3自由重设计'],
    ]
)

note('这些约束来自 sfGFP 的生物化学机制（参考文献 Pédelacq 2006 + Sarkisyan 2016）。如果发色团三肽被破坏，GFP 根本不会发光——不管你优化了多少其他位置。')

# ============================================================
# PHASE 2: STRATEGY D
# ============================================================
BR()
SEC('阶段二：策略D——进化分析（MSA + EVcouplings）')

H('2.1 这个阶段做什么', level=2)

P('用 jackhmmer 搜索 UniRef30 数据库（含 1 亿+ 条蛋白质序列），找到 sfGFP 的所有"进化亲戚"。把这些亲戚排成多序列比对（MSA），然后用统计方法（EVcouplings）分析哪些位置之间存在"共进化"关系。')

H('2.2 为什么要先做策略D', level=2)

P('策略D 产出的是"突变兼容性规则"——一个告诉你"突变A和突变B能同时做吗"的矩阵。这个规则被策略A/B/C/E 公用。先做策略D，再做其他策略，可以避免在后面的阶段生成大量"会被共进化规则淘汰"的无效序列，节省GPU时间。', bold=True)

H('2.3 关键名词再确认', level=2)
P('MSA = 把 sfGFP 和它的同源蛋白对齐排好。EVcouplings = 从 MSA 中统计出"哪些位置一起变"的耦合关系。输出 = 一个 238×238 的矩阵，告诉你任意两个残基之间的共进化强度。')

H('2.4 Step-by-step 操作', level=2)

P('Step 1: 构建sfGFP的query fasta', bold=True)
CODE("""# 将 sfGFP 序列保存为 fasta 文件
cat > data/sfGFP.fasta << EOF
>sfGFP
MSKGEELFTGVVPILVELDGDVNGHKFSVRGEGEGDATNGKLTLKFICTTGKLPVPWPTLVTTLTYGVQCFSRYPDHMKRHDFFKSAMPEGYVQERTISFKDDGTYKTRAEVKFEGDTLVNRIELKGIDFKEDGNILGHKLEYNFNSHNVYITADKQKNGIKANFKIRHNIVEDGSVQLADHYQQNTPIGDGPVLLPDNHYLSTQSVLSKDPNEKRDHMVLLEFVTAAGITHGMDELYK
EOF""")

P('Step 2: 用 jackhmmer 搜索 UniRef30', bold=True)
P('这个步骤需要联网，会下载 UniRef30 数据库。首次运行可能需要 30-60 分钟。')
CODE("""# 创建输出目录
mkdir -p output/msa

# 3轮迭代搜索
jackhmmer -N 3 \\
  --cpu 8 \\
  -A output/msa/sfGFP_alignment.sto \\
  data/sfGFP.fasta \\
  uniref30

# 转换为更常用的 fasta 格式
esl-reformat fasta output/msa/sfGFP_alignment.sto > output/msa/sfGFP_alignment.fasta""")

P('Step 3: 分析MSA质量', bold=True)
CODE("""# scripts/02_check_msa.py
from Bio import SeqIO
import numpy as np

records = list(SeqIO.parse('output/msa/sfGFP_alignment.fasta', 'fasta'))
print(f'MSA中的序列数: {len(records)}')
print(f'MSA长度(列数): {len(records[0].seq)}')

# 计算每列的保守性
alignment_array = np.array([list(str(r.seq)) for r in records])
conservation = []
for col in alignment_array.T:
    counts = {}
    for aa in col:
        if aa != '-':
            counts[aa] = counts.get(aa, 0) + 1
    if counts:
        top_freq = max(counts.values()) / sum(counts.values())
        conservation.append(top_freq)
    else:
        conservation.append(0)

# 高度保守的列
highly_conserved = sum(1 for c in conservation if c > 0.9)
print(f'高度保守位置(>90%): {highly_conserved}')
print(f'中度保守位置(50-90%): {sum(1 for c in conservation if 0.5 < c <= 0.9)}')
print(f'可变位置(<50%): {sum(1 for c in conservation if c <= 0.5)}')""")

note('好的 MSA 应该有 >1,000 条序列，且 >20 个高度保守位置（这些就是发色团三肽、催化残基等）。如果 MSA 序列数 < 100，尝试增加 jackhmmer 迭代轮数到 5。')

P('Step 4: 运行 EVcouplings 分析', bold=True)
P('EVcouplings 是一个完整的 Python 包，需要单独安装和配置。最简单的方式是使用其提供的命令行工具。')
CODE("""# 安装 EVcouplings
pip install evcouplings

# 运行共进化分析
evcouplings -o output/evcouplings \\
  -p output/msa/sfGFP_alignment.fasta \\
  --calc_method plmc

# 输出文件:
# output/evcouplings/couplings.csv  <- 残基对的耦合强度
# output/evcouplings/scores.csv     <- 每个残基的保守性分数""")

P('Step 5: 解析结果为突变规则', bold=True)
CODE("""# scripts/03_parse_evcouplings.py
import pandas as pd
import numpy as np

# 读取EVcouplings输出
couplings = pd.read_csv('output/evcouplings/couplings.csv')
print('=== Top 20 最强耦合对 ===')
print(couplings.head(20))

# 构建238x238的耦合矩阵
n_residues = 238
coupling_matrix = np.zeros((n_residues, n_residues))

for _, row in couplings.iterrows():
    i = int(row['i']) - 1  # 转为0-based
    j = int(row['j']) - 1
    score = float(row['score'])
    coupling_matrix[i, j] = score
    coupling_matrix[j, i] = score  # 对称

# 保存为规则文件
np.save('output/evcouplings/coupling_matrix.npy', coupling_matrix)

# 提取"强耦合对"作为硬约束
strong_pairs = couplings[couplings['score'] > 5.0]
print(f'\\n强耦合对(score>5.0): {len(strong_pairs)} 对')
print('这些对中的突变必须协同进行，单独突变任何一个是危险的。')

# 保存为规则
strong_pairs[['i', 'j', 'score']].to_csv('output/evcouplings/strong_coupling_rules.csv', index=False)
print('规则已保存到 output/evcouplings/strong_coupling_rules.csv')""")

P('结果解读：', bold=True)
B('coupling_matrix.npy: 238×238 的矩阵。矩阵[i][j]值越大，说明位置i和j共进化越强→两个位置的突变必须协同。')
B('strong_coupling_rules.csv: 强耦合对的列表。这些对在后续策略中要特别注意——单独突变其中任何一个可能有害。')
B('scores.csv 中的保守性分数：分数越高的位置越不能突变（对应发色团三肽、R96、E222等关键残基）。交叉验证了你的三级约束体系。')

note('如果 EVcouplings 安装有困难（它依赖较多），可以用更轻量的替代方案：GREMLIN（github.com/sokrypton/GREMLIN）或者直接用 MSA 的保守性分数作为简化版约束。策略D的输出不是阻塞性的——管线可以在没有它的情况下继续。')

# ============================================================
# PHASE 3: STRATEGY A
# ============================================================
BR()
SEC('阶段三：策略A——理性工程（保守保底方案）')

H('3.1 这个阶段做什么', level=2)

P('以 sfGFP 为起点，从 45 个已知候选设计位点中，系统地枚举 2-3 个位点的组合突变（即每次在一个 sfGFP 拷贝上同时引入 2-3 个已知有益突变）。用 FoldX 计算每个组合的折叠自由能变化（ΔΔG），筛选出不会破坏折叠的组合。')

H('3.2 为什么这是第一个设计策略', level=2)

P('策略A 是管线的"保险方案"——它不依赖任何 AI 模型，不依赖 GPU，不依赖大规模计算。即使其他所有策略都失败，策略A 也能产出可提交的序列。因为这是在 sfGFP（已验证高表达高折叠）的骨架上做小幅度微调，30% 亮度阈值线几乎肯定能过。', bold=True)

H('3.3 关键名词再确认', level=2)
P('FoldX = 一个计算工具，输入蛋白质结构和突变，输出 ΔΔG（折叠自由能变化）。正值=突变使蛋白不稳定。管线用 FoldX 而不是 AI 做稳定性筛选，因为 FoldX 基于物理力场而非训练数据，对未见过的突变组合也能给出合理的估算。速度快（秒级/条），适合筛选几千条组合。')

H('3.4 事先准备', level=2)
B('FoldX 需要单独下载（foldxsuite.crg.eu，学术免费）')
B('需要 sfGFP 的 PDB 文件（从 RCSB PDB 下载 2B3P）')
B('45 个候选位点列表（来自管线设计文档附录B）')

H('3.5 Step-by-step 操作', level=2)

P('Step 1: 定义 45 个候选位点和允许的突变', bold=True)
CODE("""# scripts/04_strategy_A_enum.py
import itertools
from Bio.Seq import Seq
from Bio.SeqRecord import SeqRecord
from Bio import SeqIO

# sfGFP 野生型序列
WT = "MSKGEELFTGVVPILVELDGDVNGHKFSVRGEGEGDATNGKLTLKFICTTGKLPVPWPTLVTTLTYGVQCFSRYPDHMKRHDFFKSAMPEGYVQERTISFKDDGTYKTRAEVKFEGDTLVNRIELKGIDFKEDGNILGHKLEYNFNSHNVYITADKQKNGIKANFKIRHNIVEDGSVQLADHYQQNTPIGDGPVLLPDNHYLSTQSVLSKDPNEKRDHMVLLEFVTAAGITHGMDELYK"

# 候选位点：(位置1-based, WT_aa, [允许的突变])
candidates = [
    # 亮度增强位点 (8个)
    (30, 'S', ['R']), (64, 'F', ['L']), (99, 'F', ['S']),
    (153, 'M', ['T']), (163, 'V', ['A']), (175, 'S', ['G']),
    (206, 'A', ['K']),
    # 热稳定性增强位点 (12个)
    (6, 'E', ['D', 'K']), (8, 'F', ['L', 'I']), (18, 'L', ['I', 'V']),
    (22, 'V', ['I', 'L']), (39, 'Y', ['F', 'H']), (42, 'L', ['I', 'V']),
    (53, 'L', ['I', 'F']), (61, 'V', ['I', 'L']), (100, 'L', ['I']),
    (125, 'L', ['I']), (170, 'N', ['D', 'E']),
    # ... 更多位点（完整列表见文档附录B）
]
""")

P('Step 2: 枚举 2-3 个位点的组合突变', bold=True)
CODE("""# 继续 scripts/04_strategy_A_enum.py

def apply_mutations(wt_seq, mutations):
    \"\"\"在wt_seq上应用一组突变\"\"\"
    seq_list = list(wt_seq)
    for pos_1based, new_aa in mutations:
        seq_list[pos_1based - 1] = new_aa
    return ''.join(seq_list)

# 生成2突变和3突变的所有组合
all_sequences = []
seq_id = 0

# 2-突变组合
for (p1, wt1, aas1), (p2, wt2, aas2) in itertools.combinations(candidates, 2):
    if abs(p1 - p2) < 3:  # 跳过相邻位点(避免空间冲突)
        continue
    for aa1 in aas1:
        for aa2 in aas2:
            seq = apply_mutations(WT, [(p1, aa1), (p2, aa2)])
            all_sequences.append({
                'seq_id': f'A2_{seq_id:04d}',
                'sequence': seq,
                'mutations': f'{wt1}{p1}{aa1},{wt2}{p2}{aa2}',
                'n_muts': 2,
                'source': 'strategy_A'
            })
            seq_id += 1

# 3-突变组合（采样以控制总数在~3000以内）
import random
random.seed(42)
triple_pool = list(itertools.combinations(candidates, 3))
random.shuffle(triple_pool)
for (p1,wt1,aas1),(p2,wt2,aas2),(p3,wt3,aas3) in triple_pool[:5000]:
    if len({p1,p2,p3}) < 3: continue
    for aa1, aa2, aa3 in itertools.product(aas1, aas2, aas3):
        seq = apply_mutations(WT, [(p1,aa1),(p2,aa2),(p3,aa3)])
        all_sequences.append({
            'seq_id': f'A3_{seq_id:04d}',
            'sequence': seq, 'n_muts': 3,
            'mutations': f'{wt1}{p1}{aa1},{wt2}{p2}{aa2},{wt3}{p3}{aa3}',
            'source': 'strategy_A'
        })
        seq_id += 1
        if seq_id > 5000: break
    if seq_id > 5000: break

print(f'生成了 {len(all_sequences)} 条策略A候选序列')

# 保存为fasta
records = [SeqRecord(Seq(s['sequence']), id=s['seq_id'], description=s['mutations'])
           for s in all_sequences]
SeqIO.write(records, 'output/strategy_A_candidates.fasta', 'fasta')

# 保存元数据
import json
with open('output/strategy_A_metadata.json', 'w') as f:
    json.dump(all_sequences, f, indent=2)
print('已保存到 output/strategy_A_candidates.fasta')""")

P('Step 3: 用 FoldX 快速筛选', bold=True)
P('FoldX 是一个命令行工具。你需要：①下载 FoldX 可执行文件；②把每条候选序列的突变列表输入 FoldX 的 BuildModel 命令；③收集 ΔΔG 值。')
CODE("""# scripts/05_foldx_screen.sh
# 对每条候选序列，用FoldX计算ΔΔG

FOLDX=/path/to/foldx  # FoldX可执行文件路径
PDB=data/2B3P_sfGFP.pdb  # sfGFP结构

# 创建FoldX输入格式: individual_list.txt
# 格式: 突变描述 (如 "SA30R,FL64L;")
python -c "
import json
with open('output/strategy_A_metadata.json') as f:
    candidates = json.load(f)
with open('output/foldx_individual_list.txt', 'w') as out:
    for c in candidates:
        # 转换突变格式: 'S30R,F64L' -> 'SA30R,FL64L;'
        muts = c['mutations'].split(',')
        foldx_mut = ','.join([m[0]+m[1:-1]+m[-1] for m in muts]) + ';'
        out.write(foldx_mut + '\\\\n')
"

# 运行FoldX BuildModel
$FOLDX --command=BuildModel \\
       --pdb=$PDB \\
       --mutant-file=output/foldx_individual_list.txt \\
       --output-dir=output/foldx_results

# FoldX会为每条突变生成一个文件，含ΔΔG值""")

P('Step 4: 筛选通过FoldX的候选', bold=True)
CODE("""# scripts/06_filter_foldx.py
import json, os, re

with open('output/strategy_A_metadata.json') as f:
    candidates = json.load(f)

passed = []
for c in candidates:
    # FoldX输出文件命名规则: 突变描述_1.pdb -> Differences_突变描述.txt
    mut_desc = c['mutations'].replace(',', '_')
    diff_file = f'output/foldx_results/Differences_{mut_desc}.txt'

    if os.path.exists(diff_file):
        with open(diff_file) as f:
            content = f.read()
            # 提取ΔΔG值
            match = re.search(r'Total energy difference.*?([-\\d.]+)', content)
            if match:
                ddg = float(match.group(1))
                c['ddG'] = ddg
                # 筛选: ΔΔG < 3 kcal/mol (不过度destabilize)
                if ddg < 3.0:
                    passed.append(c)

print(f'FoldX筛选: {len(candidates)} -> {len(passed)} 通过')
print(f'淘汰 {len(candidates)-len(passed)} 条(ΔΔG>=3 kcal/mol)')

# 保存通过者
with open('output/strategy_A_passed.json', 'w') as f:
    json.dump(passed, f, indent=2)
print(f'通过FoldX筛选的候选已保存到 output/strategy_A_passed.json')""")

P('结果解读：', bold=True)
B('output/strategy_A_passed.json: ~2000-3000 条通过了 FoldX 稳定性筛选的候选序列。')
B('每条序列的 ddG 字段：越负越好（负值表示突变使蛋白更稳定）。筛选阈值 3 kcal/mol 是保守的（允许轻微的 destabilization）。')
B('这些序列进入下一阶段的统一漏斗。策略A的候选是"保守的"——来自 sfGFP 骨架小幅微调，通过 30% 亮度阈值线的概率最高。')

note('如果 FoldX 没有正确安装或运行，可以用 Rosetta 的 cartesian_ddg 替代，或者暂时跳过物理筛选、直接基于突变效应文献做手工筛选。策略A 的保守性本身就保证了基本的序列质量。')

# ============================================================
# SAVE INTERMEDIATE
# ============================================================
doc.save("D:/蛋白质设计-合成生物学创新赛-Claude/文档/GFP蛋白质设计_零基础实操指南.docx")
print("Phases 0-3 saved. Continuing...")
