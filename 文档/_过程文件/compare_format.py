"""Compare formatting between v2.0 and v2.1 documents."""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
import json

def analyze_paragraph(para):
    """Extract formatting info from a paragraph."""
    info = {
        'text': para.text[:80] if para.text else '(empty)',
        'style': para.style.name if para.style else None,
        'alignment': str(para.alignment) if para.alignment else None,
        'runs': []
    }

    # Paragraph format
    pf = para.paragraph_format
    info['space_before'] = pf.space_before
    info['space_after'] = pf.space_after
    info['line_spacing'] = pf.line_spacing
    info['first_line_indent'] = pf.first_line_indent

    # Run format
    for run in para.runs:
        rf = run.font
        run_info = {
            'text': run.text[:40] if run.text else '(empty)',
            'font_name': rf.name,
            'size': str(rf.size) if rf.size else None,
            'bold': rf.bold,
            'color': str(rf.color.rgb) if rf.color and rf.color.rgb else None,
        }
        info['runs'].append(run_info)

    return info

def analyze_document(filepath, label):
    """Analyze document formatting."""
    doc = Document(filepath)
    print(f"\n{'='*80}")
    print(f"Document: {label}")
    print(f"{'='*80}")

    # Page setup (first section)
    for i, section in enumerate(doc.sections):
        print(f"\nSection {i}:")
        print(f"  Page: {section.page_width} x {section.page_height}")
        print(f"  Margins: top={section.top_margin}, bottom={section.bottom_margin}, "
              f"left={section.left_margin}, right={section.right_margin}")

    # Print styles used
    print(f"\nStyles in use:")
    styles_seen = set()
    for para in doc.paragraphs[:200]:  # Check first 200 paragraphs
        if para.style and para.style.name not in styles_seen:
            styles_seen.add(para.style.name)
            s = para.style
            try:
                font = s.font
                print(f"  Style '{s.name}': font={font.name}, size={font.size}, bold={font.bold}, color={font.color.rgb if font.color and font.color.rgb else None}")
            except:
                print(f"  Style '{s.name}': (could not read font)")

    # Analyze paragraphs
    print(f"\nFirst 60 paragraphs:")
    for i, para in enumerate(doc.paragraphs[:60]):
        info = analyze_paragraph(para)
        if info['text'].strip():  # Skip empty paragraphs
            print(f"\n  P{i}: [{info['style']}] align={info['alignment']} "
                  f"space_before={info['space_before']} space_after={info['space_after']} "
                  f"line_spacing={info['line_spacing']}")
            print(f"    Text: {info['text']}")
            for j, run_info in enumerate(info['runs']):
                if run_info['text'].strip():
                    print(f"    R{j}: font={run_info['font_name']} size={run_info['size']} "
                          f"bold={run_info['bold']} color={run_info['color']} text='{run_info['text']}'")

    # Count total paragraphs
    print(f"\nTotal paragraphs: {len(doc.paragraphs)}")

    return doc

# Analyze v2.0 (reference format)
doc_v20 = analyze_document(
    "D:/蛋白质设计-合成生物学创新赛-Claude/文档/GFP蛋白质设计_管线设计与实施方案_v2.0.docx",
    "v2.0 (REFERENCE)"
)

# Analyze v2.1 (to be adjusted)
doc_v21 = analyze_document(
    "D:/蛋白质设计-合成生物学创新赛-Claude/文档/GFP蛋白质设计_管线设计与实施方案_v2.1.docx",
    "v2.1 (TO FIX)"
)
