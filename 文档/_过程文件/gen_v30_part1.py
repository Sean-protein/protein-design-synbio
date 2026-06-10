#!/usr/bin/env python3
"""
Generate GFP Pipeline Design v3.0 — Complete Implementation Plan
Incorporates judge review feedback + ESM3 LoRA integration
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import datetime

# ============================================================
# Document Setup
# ============================================================
doc = Document()

# --- Page Setup ---
for section in doc.sections:
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# --- Style Definitions ---
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style.paragraph_format.line_spacing = 1.2
rPr = style.element.get_or_add_rPr()
rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="宋体" w:ascii="Times New Roman" w:hAnsi="Times New Roman"/>')
rPr.insert(0, rFonts)

for (sname, font_name, size, color_hex, space_before, space_after) in [
    ('Heading 1', '微软雅黑', Pt(18), '003366', Pt(12), Pt(6)),
    ('Heading 2', '微软雅黑', Pt(15), '005082', Pt(10), Pt(4)),
    ('Heading 3', '微软雅黑', Pt(14), '4F81BD', Pt(8), Pt(4)),
]:
    s = doc.styles[sname]
    s.font.name = font_name
    s.font.size = size
    s.font.bold = True
    s.font.color.rgb = RGBColor(*[int(color_hex[i:i+2],16) for i in range(0,6,2)])
    s.paragraph_format.space_before = space_before
    s.paragraph_format.space_after = space_after
    srPr = s.element.get_or_add_rPr()
    srFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{font_name}" w:ascii="{font_name}" w:hAnsi="{font_name}"/>')
    srPr.insert(0, srFonts)

# ============================================================
# Helper Functions
# ============================================================
def add_para(text, style='Normal', bold=False, alignment=None, font_name=None, font_size=None, color=None, space_before=None, space_after=None):
    """Add a paragraph with optional formatting."""
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    if bold:
        run.bold = True
    if font_name:
        run.font.name = font_name
    if font_size:
        run.font.size = font_size
    if color:
        run.font.color.rgb = RGBColor(*color)
    if alignment is not None:
        p.alignment = alignment
    if space_before is not None:
        p.paragraph_format.space_before = space_before
    if space_after is not None:
        p.paragraph_format.space_after = space_after
    return p

def add_heading(text, level=1):
    """Add a heading."""
    return doc.add_heading(text, level=level)

def add_bullet(text, level=0):
    """Add a bullet point."""
    p = doc.add_paragraph(text, style='List Bullet')
    if level > 0:
        p.paragraph_format.left_indent = Cm(1.5 * level)
    return p

def add_table(headers, rows, col_widths=None):
    """Add a formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    # Headers
    for j, header in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = header
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
    # Data
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i+1].cells[j]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    return table

def add_code(text):
    """Add a code block."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    # Gray background via shading
    pPr = p._element.get_or_add_pPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F0F0F0" w:val="clear"/>')
    pPr.append(shd)
    return p

def page_break():
    doc.add_page_break()

def add_section_title(text):
    """Add a centered section divider."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = '微软雅黑'
    run.font.color.rgb = RGBColor(0x00, 0x33, 0x66)
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(12)
    # Bottom border
    pPr = p._element.get_or_add_pPr()
    pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="6" w:space="1" w:color="003366"/></w:pBdr>')
    pPr.append(pBdr)
    return p

# ============================================================
# TITLE PAGE
# ============================================================
for _ in range(6):
    doc.add_paragraph()

add_para('2026合成生物学创新赛', alignment=WD_ALIGN_PARAGRAPH.CENTER,
         font_name='微软雅黑', font_size=Pt(26), color=(0x00, 0x33, 0x66))
add_para('蛋白质设计赛道', alignment=WD_ALIGN_PARAGRAPH.CENTER,
         font_name='微软雅黑', font_size=Pt(22), color=(0x00, 0x50, 0x82))

doc.add_paragraph()

add_para('GFP 高亮度与热稳定性联合设计', alignment=WD_ALIGN_PARAGRAPH.CENTER,
         font_name='微软雅黑', font_size=Pt(30), bold=True, color=(0x00, 0x33, 0x66))
add_para('管线设计与实施方案', alignment=WD_ALIGN_PARAGRAPH.CENTER,
         font_name='微软雅黑', font_size=Pt(30), bold=True, color=(0x00, 0x33, 0x66))

doc.add_paragraph()
doc.add_paragraph()

add_para('（含 ESM3 集成 · LoRA 微调 · 评委反馈改进）', alignment=WD_ALIGN_PARAGRAPH.CENTER,
         font_name='微软雅黑', font_size=Pt(16), color=(0x50, 0x50, 0x50))

doc.add_paragraph()
doc.add_paragraph()

add_para('版本：v3.0  |  生成日期：2026年06月09日', alignment=WD_ALIGN_PARAGRAPH.CENTER,
         font_name='微软雅黑', font_size=Pt(14), color=(0x50, 0x50, 0x50))
add_para('竞赛截止日期：2026年7月1日  |  有效工作日：18天', alignment=WD_ALIGN_PARAGRAPH.CENTER,
         font_name='微软雅黑', font_size=Pt(14), color=(0xCC, 0x00, 0x00))

page_break()

# ============================================================
# TABLE OF CONTENTS placeholder
# ============================================================
add_heading('目录', level=1)

toc_items = [
    ('一、设计哲学与核心原则', True),
    ('二、数据可靠性声明', True),
    ('三、三级约束体系（v3.0 改进）', True),
    ('  3.1 化学绝对约束', False),
    ('  3.2 量子产率优化约束', False),
    ('  3.3 结构协同验证约束', False),
    ('四、五条设计策略（含 ESM3 策略E）', True),
    ('  4.1 策略A：理性工程优化', False),
    ('  4.2 策略B：PLM深度学习驱动（ESM2+SaProt+ESM3 LoRA集成）', False),
    ('  4.3 策略C：ProteinMPNN结构协同设计', False),
    ('  4.4 策略D：进化共识与特征迁移', False),
    ('  4.5 策略E：ESM3生成式管线（esmGFP式Gibbs采样）', False),
    ('五、DBLT迭代闭环与贝叶斯优化', True),
    ('六、双目标优化：Pareto前沿策略', True),
    ('七、LLM Agent编排架构', True),
    ('八、五级漏斗筛选详细设计', True),
    ('九、ESM3集成与LoRA微调方案', True),
    ('十、18天时间规划与GPU预算', True),
    ('十一、答辩预案：评委可能提出的关键问题', True),
    ('十二、风险矩阵与应对策略', True),
    ('十三、6/24提交材料清单', True),
    ('附录A：关键命令速查表', True),
    ('附录B：45个候选设计位点', True),
    ('附录C：参考文献与工具速查', True),
]

for item_text, is_main in toc_items:
    p = doc.add_paragraph()
    run = p.add_run(item_text)
    if is_main:
        run.bold = True
        run.font.name = '微软雅黑'
        run.font.size = Pt(14)
    else:
        run.font.name = '宋体'
        run.font.size = Pt(12)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)

page_break()

print("Title page + TOC done. Building main content...")

# ============================================================
# SECTION 1: DESIGN PHILOSOPHY
# ============================================================
add_section_title('一、设计哲学与核心原则')

add_para('本方案 v3.0 在 v2.1 基础上，整合了评委视角审阅反馈和ESM3集成方案，核心升级包括：'
         '① 约束体系从两级升级为三级（化学绝对→量子产率→结构验证）；'
         '② 新增策略E：ESM3生成式管线；③ 策略B集成ESM3 LoRA微调嵌入；'
         '④ 补充GPU/时间预算表；⑤ 细化贝叶斯优化实现路径；⑥ 增加答辩预案。')

add_heading('1.1 核心设计原则（继承自 v2.1，不变）', level=2)

add_bullet('广度优先 (Exploration)：五条策略并行覆盖序列空间的不同区域——理性组合、AI嵌入搜索、逆折叠设计、进化引导、ESM3生成。')
add_bullet('深度验证 (Exploitation)：结构预测（AF2/ESMFold）+ 多级稳定性评估（ThermoMPNN→FoldX→BioEmu→MD）逐级筛选。')
add_bullet('联合决策 (Integration)：五条策略的候选在Pareto前沿上合并，按四个区域分配6条提交序列。')
add_bullet('风险分散 (Diversification)：覆盖不同"亮度×热稳定性"策略组合，6条序列来自至少3种不同策略。')
add_bullet('迭代学习 (Iteration)：DBLT闭环中的贝叶斯优化从MD"金标准"数据中学习，引导后续生成方向。')
add_bullet('6中取1即可登顶 → 鼓励多样化探索而非单一策略all-in。')

add_heading('1.2 v3.0 相比 v2.1 的关键变更', level=2)

add_table(
    ['变更项', 'v2.1', 'v3.0', '变更原因'],
    [
        ['约束体系', '两级（绝对+结构）', '三级（化学绝对+量子产率+结构验证）', '评委反馈：氢键网络约束等级模糊'],
        ['设计策略数', '4条（A/B/C/D）', '5条（A/B/C/D/E）', '新增ESM3生成式管线策略E'],
        ['策略B嵌入', 'ESM-2 + SaProt拼接', 'ESM-2 + SaProt + ESM3 LoRA集成', 'ESM3结构感知嵌入提升预测精度'],
        ['贝叶斯优化', '概念级描述', '含GP核选择/MCMC参数/训练集构建', '评委反馈：欧阳颀教授会深度审查数学细节'],
        ['GPU预算', '未提供', '按GPU型号的时间预估矩阵', '评委反馈：需评估18天内可行性'],
        ['Agent日志', '仅描述格式', '含模板示例+模拟决策记录', '评委反馈：需证明可运行性'],
        ['答辩准备', '无', '6个核心评委模拟问答', '针对评委团构成（浙大4人+蛋白质工程专家）'],
        ['MD协议', '承认10ns局限性', '补充时间尺度论证+文献支撑', '评委反馈：于浩然教授可能追问'],
    ]
)

# ============================================================
# SECTION 2: DATA RELIABILITY
# ============================================================
page_break()
add_section_title('二、数据可靠性声明')

add_para('管线使用以下数据源，按可靠性分级管理。这一声明本身是学术诚信的体现——评委明确表示"数据可靠性自我声明在所有评委中都会产生好感"（评委审阅报告，第8节）。')

add_heading('2.1 数据可靠性分级', level=2)

add_table(
    ['等级', '数据源', '可靠性', '用途', '注意事项'],
    [
        ['A', 'sfGFP晶体结构 (PDB 2B3P)', '实验验证', '结构模板、ProteinMPNN输入', 'Gold standard'],
        ['A', 'GFP_data.xlsx (官方训练数据)', '实验验证', 'ML模型训练', '注意avGFP→sfGFP分布偏移'],
        ['A', 'Exclusion_List.csv (官方)', '实验验证', '提交序列比对', '逐条完全匹配检查'],
        ['B', 'Sarkisyan et al. 2016 (Nature)', '同行评审', '数据增强、适应度景观参考', 'avGFP骨架，与sfGFP有序列差异'],
        ['B', 'Pédelacq et al. 2006 (Nat Biotechnol)', '同行评审', 'sfGFP突变效应参考', '位点特异性，非全序列'],
        ['B', 'ESM3/esmGFP (Rives et al. 2025)', '同行评审', '设计哲学启发、策略E模板', '使用7B API模型，开源1.4B效果待验证'],
        ['C', 'FoldX ΔΔG预测', '计算预测', '快速稳定性筛选', '系统偏差需MD校准'],
        ['C', 'ThermoMPNN预测', '计算预测', '热稳定性初筛', '训练数据不含72°C条件'],
        ['D', 'sfGFP_氨基酸功能全注释.docx', 'AI推断', '位点约束方向性参考', '具体数值不可直接引用；方向性正确'],
    ]
)

add_heading('2.2 特别声明', level=2)

add_para('sfGFP_氨基酸功能全注释.docx 由 Claude 于2026年6月6日生成（作者"Un-named"），引用了 Pédelacq 2006 和 Sarkisyan 2016 两篇同行评审论文，但逐位点的具体约束等级和数值是AI基于PDB结构推断的产物，未经实验验证。其方向性正确（发色团三肽不可动、催化残基关键），但"R96E残留1-5%荧光"这类具体数值不可直接引用为设计约束。该文档用作启发式参考，所有约束均需通过本管线第3节的三级体系重新审查。', bold=False)

# ============================================================
# SECTION 3: THREE-TIER CONSTRAINT SYSTEM
# ============================================================
page_break()
add_section_title('三、三级约束体系（v3.0 核心改进）')

add_para('v2.1的约束体系将发色团氢键网络(Q69/Q94/H148/T203/S205)混在"化学绝对约束"中，造成内部不一致（"绝对约束"但又"可在路线B中尝试"）。v3.0修正为明确的三级体系，每级有不同的约束强度和验证要求。')

add_heading('3.1 第一级：化学绝对约束（6个位点，ProteinMPNN/ESM3中固定）', level=2)

add_para('这些位点的约束来自化学反应机理（SN2亲核加成、广义酸碱催化、π共轭体系），不可妥协。', bold=True)

add_table(
    ['位点', '约束', '化学依据', '允许的变异', '跨物种保守性'],
    [
        ['G67', '绝对不可替换', '酰胺氮SN2亲核进攻T65羰基碳。任何侧链(含H原子)空间阻断环化', '无。Gly的φ/ψ角灵活性也是必需的', '100%保守于所有已知GFP'],
        ['Y66', '仅限芳香族(Y/F/H/W)', '酚羟基参与发色团π共轭。无芳香环→无可见光吸收→无荧光', 'Y66F(光谱蓝移)、Y66H(BFP)、Y66W(CFP)', '>99%保守为芳香族'],
        ['T65', '仅限小残基(G/A/C/S)', '侧链体积限制环化空间。大侧链/带电残基阻断环化', 'S65(avGFP天然)、A65、C65', '>95%保守为小残基'],
        ['R96', '大概率需保留', '催化发色团环化脱水。与咪唑啉酮环形成氢键稳定过渡态', '理论上协同重排可替代。风险极高，建议固定', '>90%保守为Arg/Lys'],
        ['E222', '大概率需保留', '广义碱催化Y66 Cα-Cβ脱氢。E222D部分保留功能', 'E222D(部分活性)。其他替换基本失活', '>85%保守为Glu/Asp'],
        ['发色团π堆积', '保持芳香环境', 'Y66周围需芳香/疏水环境维持量子产率', '允许疏水替换(F/L/I/V)，禁止带电残基', '结构保守，序列可变'],
    ]
)

add_heading('3.2 第二级：量子产率优化约束（5个位点，可探索但需补偿）', level=2)

add_para('这是v3.0新增的约束等级，专门针对v2.1中约束模糊的氢键网络位点。这些位点突变不会导致零荧光，但单点突变会显著降低量子产率。理论上协同重排氢键网络可能找到替代方案。', bold=True)

add_table(
    ['位点', '氢键角色', '单点突变影响', '探索策略', '验证要求'],
    [
        ['Q69', '与Y66酚羟基形成氢键', '量子产率降低50-80%', '路线B/C/E中尝试替代', '发色团区域pLDDT>85 + 氢键分析'],
        ['Q94', '稳定发色团咪唑啉酮环', '量子产率降低30-60%', '路线B/C/E中尝试替代', 'MD中发色团RMSF<2Å'],
        ['H148', '与发色团π堆积+氢键', '量子产率降低40-70%', '路线B/E中尝试H148D(mBaoJin特征)', 'AF2验证H148D-发色团距离'],
        ['T203', '与Y66酚氧形成氢键', '量子产率降低20-50%', '路线B/D中尝试T203I(StayGold特征)', '允许保守替换'],
        ['S205', '稳定β桶C端构象', '量子产率降低20-40%', '允许小残基替换', 'AF2 pLDDT>80'],
    ]
)

add_para('关键原则：第二级位点突变必须伴随补偿性突变（协同设计），且通过MD验证氢键网络完整性。单独突变任一第二级位点应触发Agent预警。', bold=True)

add_heading('3.3 第三级：结构协同验证约束（其余约227个位点）', level=2)

add_para('对剩余的约227个位点，不预先禁止任何突变。约束通过以下验证管线事后实施：')

add_bullet('AF2/ESMFold折叠验证：全局pLDDT>80, pTM>0.75 → 折叠可行性保证')
add_bullet('发色团区域验证：局部pLDDT>85 → 活性中心结构可靠')
add_bullet('MD动力学验证：WT sfGFP基线对比 → RMSD/Rg/SASA/氢键维持率相对变化量判断')
add_bullet('上位效应过滤：EVcouplings/GREMLIN → 排除违反共进化的组合')
add_bullet('序列多样性：CD-HIT 90%去冗余 → 确保6条序列非近重复')

# ============================================================
# SECTION 4: FIVE DESIGN STRATEGIES
# ============================================================
page_break()
add_section_title('四、五条设计策略')

add_para('以下五条策略并行运行，每条策略生成约1,000-5,000条候选序列。合并去冗余后进入统一的五级漏斗筛选。策略E（ESM3生成式管线）为v3.0新增，直接回应评委"为何不利用ESM3"的预期。')

# Strategy A
add_heading('4.1 策略A：理性工程优化（保守稳健型）', level=2)

add_table(
    ['属性', '内容'],
    [
        ['优先级', '★★★★★（最高，保险方案）'],
        ['核心思路', '以sfGFP为设计起点，保留6个核心折叠增强突变(S30R/F64L/S65T/F99S/M153T/V163A)，从45个候选设计位点中组合亮度增强突变(2-3个)+热稳定性增强突变(2-3个)。最稳健的策略，自带"保底"属性。'],
        ['技术栈', 'Biopython序列操作 + 组合枚举 + 上位效应过滤(EVcouplings) + FoldX ΔΔG快速筛选'],
        ['位点来源', '附录B：45个候选设计位点（来自5.25方案，保留证据等级标注）'],
        ['预期产出', '~2,000-3,000条组合序列，预期贡献1-2条保险方案（确保过30%亮度阈值）'],
        ['设计约束', '固定6个化学绝对位点(G67/Y66/T65/R96/E222/发色团π堆积)；第二级位点保守处理（Q69→Q/E, H148→H, T203→T/S, S205→S/T, Q94→Q/E）'],
    ]
)

# Strategy B
add_heading('4.2 策略B：PLM深度学习驱动（AI驱动型，v3.0升级）', level=2)

add_table(
    ['属性', '内容'],
    [
        ['优先级', '★★★★'],
        ['核心思路', 'ESM-2 650M嵌入 + SaProt结构感知嵌入 + ESM3 LoRA微调嵌入 → 三模型加权集成→ XGBoost+LightGBM+RF三模型亮度/稳定性联合预测。贝叶斯优化(GP+EI)引导MCMC在嵌入空间中搜索。'],
        ['v3.0升级', '新增ESM3 LoRA微调嵌入（见第9节），预期R²提升0.05-0.10；三嵌入加权融合权重由5-fold CV确定'],
        ['技术栈', 'ESM-2 650M/SaProt/ESM3-sm-open(LoRA) + XGBoost/LightGBM/RF + GPyTorch + 约束MCMC采样'],
        ['预期产出', '~3,000-5,000条AI生成序列，预期贡献1-2条创新方案'],
        ['降级预案', 'GPU≥24GB→ESM-2 650M+SaProt+ESM3 LoRA；GPU≥12GB→ESM-2 150M+SaProt(无ESM3)；GPU≥8GB→ESM-2 150M；仅CPU→ESM-2 35M'],
        ['质量阈值', '5-fold CV R²>0.65→正常生成；R² 0.5-0.65→增加特征工程；R²<0.5→触发Plan C（纯理性设计）'],
    ]
)

# Strategy C
add_heading('4.3 策略C：ProteinMPNN结构协同设计（结构导向型）', level=2)

add_table(
    ['属性', '内容'],
    [
        ['优先级', '★★★★'],
        ['核心思路', '以sfGFP晶体结构(PDB 2B3P)为模板，固定6个化学绝对位点(第一级约束)，多温度采样(0.1/0.2/0.3/0.5)协同设计其余约232个位点。AF2/ESMFold自验证折叠，CD-HIT去冗余。'],
        ['技术栈', 'ProteinMPNN + ESMFold/ColabFold + CD-HIT + BioEmu构象系综预筛'],
        ['预期产出', '~2,000-4,000条结构感知序列，预期贡献1-2条结构验证方案'],
        ['与策略E的关系', 'ProteinMPNN是确定性逆折叠（速度快、探索保守），ESM3策略E是概率性生成（探索广、新颖性高）。两者在漏斗中合并比较。'],
    ]
)

# Strategy D
add_heading('4.4 策略D：进化共识与特征迁移（进化导向型）', level=2)

add_table(
    ['属性', '内容'],
    [
        ['优先级', '★★★'],
        ['核心思路', '构建GFP深度MSA(jackhmmer对UniRef30做3轮迭代)→共进化分析(EVcouplings)→识别高性能GFP共有特征模式。将mBaoJin的H148D、T203I等"特征模块"通过三种方式嫁接到sfGFP骨架。'],
        ['嫁接方案', '方案1(保守)：基于MSA共进化位点，识别共存突变模式，在sfGFP背景中引入协同突变对。方案2(结构)：使用US-align结构比对mBaoJin/TGP/StayGold→sfGFP，识别可移植的局部loop+表面电荷模块。方案3(统计)：以EVcouplings上位性矩阵为指导，为每个设计位点生成条件突变谱。'],
        ['技术栈', 'jackhmmer + MAFFT + EVcouplings/GREMLIN + 共进化引导的组合突变生成'],
        ['预期产出', '~1,000-2,000条进化约束序列；上位效应过滤规则供所有策略共用'],
    ]
)

# Strategy E (NEW)
add_heading('4.5 策略E：ESM3生成式管线（v3.0新增，esmGFP式Gibbs采样）', level=2)

add_para('这是v3.0最重大的新增策略——直接利用ESM3的多模态生成能力，复现esmGFP论文的Gibbs采样管线。竞赛参考文献[3]明确引用ESM3/esmGFP (Rives et al. 2025, Science)，评委将期待看到对这一前沿方法的实际应用。', bold=True)

add_table(
    ['属性', '内容'],
    [
        ['优先级', '★★★★（创新加分，但高计算成本）'],
        ['模型选择', 'ESM3-sm-open (1.4B参数，HuggingFace biohub/esm3-sm-open-v1)。注：esmGFP论文使用7B API模型，1.4B开源模型生成质量可能略低，但LoRA微调可缩小差距。'],
        ['Step 1: 结构提示', '以sfGFP (PDB 2B3P)为模板，编码关键区域的结构token：(a)发色团区域 a.a.62-67(T65/Y66/G67)的backbone坐标；(b)R96和E222的侧链坐标；(c)β桶核心疏水残基的backbone。其余所有位置mask。'],
        ['Step 2: 结构生成', 'track="structure", temperature=1.0→0.3退火, num_steps=序列长度。筛选：发色团模板RMSD<1.5Å AND 整体backbone RMSD>1.5Å（确保新颖）。'],
        ['Step 3: 序列生成', '以生成的结构为条件，track="sequence", temperature=1.0, num_steps=20。ESM3自验证折叠(track="structure", num_steps=1, temperature=0)。'],
        ['Step 4: Gibbs迭代', '结构→序列→结构→序列……交替30+轮。每轮：(a)Negative Local Sequence Guidance推动远离局部序列logits（增加多样性）；(b)PSSM Bias用71条天然GFP的position-specific scoring matrix偏置（来自策略D的MSA）；(c)Temperature annealing从1.0退火到0.1。'],
        ['筛选标准', '发色团RMSD<1.5Å + pLDDT>85 + 序列与Exclusion_List完全比对 + CD-HIT 90%。'],
        ['预期产出', '~500-1,500条ESM3生成序列。由于计算成本高（每条需多轮Gibbs采样），建议生成较少但更高质量的候选。'],
        ['计算需求', '单条序列完整Gibbs管线约需5-10 GPU分钟(RTX 3090)。500条×8min≈67 GPU小时。建议使用AutoDL A100减少至~20h。'],
        ['与策略B的LoRA互补', '策略B的LoRA微调提升ESM3嵌入质量（用于预测），策略E直接使用ESM3的生成能力（用于设计）。两者独立运行、结果互补。'],
    ]
)

add_heading('4.6 五条策略的交叉验证矩阵', level=2)

add_table(
    ['验证维度', '策略A', '策略B', '策略C', '策略D', '策略E'],
    [
        ['结构可行性', 'FoldX预筛', '策略C的AF2验证', 'AF2/ESMFold自验证', 'EVcouplings过滤', 'ESM3自验证(温度=0)'],
        ['亮度预测', '基于已知突变效应', '三模型集成(含ESM3 LoRA)', '不直接预测', '不直接预测', '不直接预测'],
        ['稳定性评估', 'FoldX ΔΔG', 'ThermoMPNN→FoldX→BioEmu→MD', 'BioEmu+MD', '基于进化保守性推断', 'BioEmu+MD'],
        ['新颖性评估', '低（sfGFP骨架）', '中（嵌入空间采样）', '高（全局重设计）', '中（特征嫁接）', '最高（Gibbs探索新区域）'],
        ['风险等级', '低（保底）', '中高', '中', '中低', '高（计算+质量不确定性）'],
    ]
)

# ============================================================
# SAVE (Intermediate)
# ============================================================
doc.save("D:/蛋白质设计-合成生物学创新赛-Claude/文档/GFP蛋白质设计_管线设计与实施方案_v3.0.docx")
print("Part 1 saved (Sections 1-4). Continuing...")
