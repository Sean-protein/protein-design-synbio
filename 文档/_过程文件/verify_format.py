"""Verify v2.1 formatting after fix."""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor

doc = Document("D:/蛋白质设计-合成生物学创新赛-Claude/文档/GFP蛋白质设计_管线设计与实施方案_v2.1.docx")

print("=== Page Setup ===")
for i, section in enumerate(doc.sections):
    w_cm = section.page_width / 360000
    h_cm = section.page_height / 360000
    t_cm = section.top_margin / 360000
    b_cm = section.bottom_margin / 360000
    l_cm = section.left_margin / 360000
    r_cm = section.right_margin / 360000
    print(f"Section {i}: {w_cm:.1f}cm x {h_cm:.1f}cm, margins: T={t_cm:.1f} B={b_cm:.1f} L={l_cm:.1f} R={r_cm:.1f}")

print("\n=== Styles ===")
for sname in ['Normal', 'Heading 1', 'Heading 2', 'Heading 3', 'List Bullet']:
    try:
        style = doc.styles[sname]
        font = style.font
        print(f"Style '{sname}': font={font.name}, size={font.size}, bold={font.bold}, color={font.color.rgb if font.color and font.color.rgb else None}")
    except KeyError:
        print(f"Style '{sname}': NOT FOUND")

print("\n=== First 30 paragraphs ===")
for i, para in enumerate(doc.paragraphs[:30]):
    text = para.text[:80] if para.text else '(empty)'
    style = para.style.name if para.style else 'None'

    # Check for MS Gothic
    has_ms_gothic = False
    fonts_used = set()
    for run in para.runs:
        if run.font.name:
            fonts_used.add(run.font.name)
        rPr = run._element.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rPr')
        if rPr is not None:
            rFonts = rPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rFonts')
            if rFonts is not None:
                east = rFonts.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia')
                if east:
                    fonts_used.add(f"XML:{east}")
                if east == 'MS Gothic':
                    has_ms_gothic = True

    ms_warning = " ⚠️ MS GOTHIC!" if has_ms_gothic else ""
    if text.strip():
        print(f"P{i}: [{style}] fonts={fonts_used}{ms_warning}")
        print(f"    '{text}'")

print(f"\n=== Total paragraphs: {len(doc.paragraphs)} ===")

# Check for MS Gothic in entire document
ms_gothic_count = 0
for para in doc.paragraphs:
    for run in para.runs:
        rPr = run._element.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rPr')
        if rPr is not None:
            rFonts = rPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rFonts')
            if rFonts is not None:
                east = rFonts.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia')
                if east == 'MS Gothic':
                    ms_gothic_count += 1
        if run.font.name == 'MS Gothic':
            ms_gothic_count += 1

print(f"MS Gothic references remaining: {ms_gothic_count}")
