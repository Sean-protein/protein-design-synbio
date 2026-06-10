#!/usr/bin/env python3
"""
Generate GFP Pipeline Design v3.0 — Part 2
Sections 5-13 + Appendices
"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import datetime

doc = Document("D:/蛋白质设计-合成生物学创新赛-Claude/文档/GFP蛋白质设计_管线设计与实施方案_v3.0.docx")

# ============================================================
# Helper Functions (same as part 1)
# ============================================================
def add_para(text, style='Normal', bold=False, alignment=None, font_name=None, font_size=None, color=None, space_before=None, space_after=None):
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    if bold: run.bold = True
    if font_name: run.font.name = font_name
    if font_size: run.font.size = font_size
    if color: run.font.color.rgb = RGBColor(*color)
    if alignment is not None: p.alignment = alignment
    if space_before is not None: p.paragraph_format.space_before = space_before
    if space_after is not None: p.paragraph_format.space_after = space_after
    return p

def add_heading(text, level=1):
    return doc.add_heading(text, level=level)

def add_bullet(text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    if level > 0:
        p.paragraph_format.left_indent = Cm(1.5 * level)
    return p

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    for j, header in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = header
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i+1].cells[j]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    return table

def add_code(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    pPr = p._element.get_or_add_pPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F0F0F0" w:val="clear"/>')
    pPr.append(shd)
    return p

def page_break():
    doc.add_page_break()

def add_section_title(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = '微软雅黑'
    run.font.color.rgb = RGBColor(0x00, 0x33, 0x66)
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(12)
    pPr = p._element.get_or_add_pPr()
    pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="6" w:space="1" w:color="003366"/></w:pBdr>')
    pPr.append(pBdr)
    return p

# ============================================================
# SECTION 5: DBLT ITERATIVE CLOSED LOOP
# ============================================================
page_break()
add_section_title('五、DBLT迭代闭环与贝叶斯优化（v3.0细化）')

add_para('这是v3.0最核心的架构。v2.1的GP描述偏概念化——评委审阅明确指出"欧阳颀教授会深度审查数学细节"。v3.0补充了GP核函数选择、MCMC采样参数、训练集构建和收敛判据的理论依据。')

add_heading('5.1 GP代理模型的数学规范', level=2)

add_table(
    ['参数', '选择', '依据'],
    [
        ['核函数', 'Matern 5/2', '蛋白质适应度景观具有局部平滑性+中程崎岖性。Matern 5/2允许二阶可微（比RBF核更灵活），适合建模既有平滑趋势又有局部波动的响应面。参考文献：Sarkisyan 2016展示GFP适应度景观的崎岖性。'],
        ['输入维度', '50维（ESM+SaProt+ESM3 LoRA嵌入拼接后PCA降维）', '嵌入原始维度~2560(1280+1280+1536拼接)→PCA降维到50维保留~90%方差。降维减少GP的O(n³)计算负担。'],
        ['先验均值', '零均值函数', '标准做法。在缺少先验知识时，零均值函数的GP预测完全由数据和核函数决定。无偏先验避免引入错误假设。'],
        ['噪声模型', '异方差高斯噪声（Heteroscedastic）', '不同策略生成的序列预测不确定性差异大（策略A低，策略E高）。homoscedastic GP会低估高风险区域的σ。'],
        ['超参数优化', '最大化对数边际似然 (L-BFGS)', '每轮迭代后用新MD数据重新优化核函数长度尺度和信号方差。在scikit-learn/GPyTorch中可直接调用。'],
    ]
)

add_heading('5.2 MCMC采样设计', level=2)

add_para('MCMC用于在GP的Expected Improvement(EI)高值区域采样新的候选序列。参数设计如下：')

add_bullet('采样器：Metropolis-Hastings with Gaussian proposal（实现简单，在50维空间中可接受）')
add_bullet('Proposal标准差：自适应调节——目标接受率0.23-0.30（50维空间的最佳接受率理论值，Gelman 1997）')
add_bullet('链数：5条独立链，每条1,000步（burn-in 200步），共~4,000个有效样本')
add_bullet('起始点：从当前Pareto前沿的5个不同区域各选一个起始点（确保覆盖）')
add_bullet('收敛诊断：Gelman-Rubin R-hat < 1.1（5条链之间的一致性）')
add_bullet('备选方案：如MCMC在50维空间混合过慢（R-hat>1.2 after 2,000 steps），切换到Replica-exchange MCMC（提高温度阶梯克服局部势垒）')

add_heading('5.3 DBLT迭代闭环流程', level=2)

add_para('以下为每轮DBLT迭代的具体操作（共2-3轮）：', bold=True)

add_table(
    ['阶段', '操作', '输入', '输出', '时间(单GPU)'],
    [
        ['DESIGN', '五策略并行生成候选 + GP-EI引导MCMC方向', 'GP代理模型 + 结构模板', '~8,000-14,000条候选', '2-3天'],
        ['BUILD', '嵌入→亮度集成预测→三级稳定性评估→上位效应过滤', '候选序列 + 预训练模型', '~200条通过序列', '1-2天'],
        ['TEST', 'AF2/ESMFold→BioEmu→GROMACS MD(Top 10-15)', '~200条→~50条→~15条', 'MD金标准标签', '2-4天'],
        ['LEARN', 'MD结果回传GP→更新EI→收敛判断', 'MD标签 + GP模型', '下一轮采样方向OR收敛', '数小时'],
    ]
)

add_heading('5.4 收敛判断（v3.0细化的四条件）', level=2)

add_bullet('EI_max连续2轮 < 0.05 → 收敛，进入精选阶段（说明GP认为"已知区域的好序列已被充分挖掘"）')
add_bullet('已完成3轮迭代 → 强制进入精选（时间约束，18天内最多3轮）')
add_bullet('MD结果与GP预测 Spearman r > 0.7 → 模型可靠，可提前收敛（说明快速预测器与金标准高度一致）')
add_bullet('MD结果与GP预测矛盾（r < 0.3）→ 重新校准GP，降低先验权重，扩大搜索范围（说明模型失效，需要更多探索）')

add_heading('5.5 初始训练集构建策略', level=2)

add_para('第1轮迭代前，GP需要初始训练数据来"预热"。构建方案：')

add_bullet('从5条策略各取约10条代表性序列（共~50条），确保覆盖：① 策略A的保守突变；② 策略B的AI高分序列；③ 策略C的ProteinMPNN低/中/高温度样本；④ 策略D的进化嫁接序列；⑤ 策略E的ESM3生成序列（如可用）')
add_bullet('用FoldX ΔΔG + ThermoMPNN做快速稳定性标注（无需MD，可数分钟内完成）')
add_bullet('用策略B的集成预测器做亮度标注')
add_bullet('这50条构成GP的先验训练集。第1轮MD验证后替换为金标准标签。')

# ============================================================
# SECTION 6: DUAL-OBJECTIVE OPTIMIZATION
# ============================================================
page_break()
add_section_title('六、双目标优化：Pareto前沿策略')

add_heading('6.1 为什么不用乘积最大化', level=2)

add_para('竞赛最终排名取决于乘积得分 Score = Brightness_relative × Stability_retention。但将乘积作为唯一优化目标有方法论缺陷：')
add_bullet('预测误差放大：亮度被高估20%且稳定性被高估20%→乘积虚高44%。如果两个维度的预测误差正相关（都来自模型偏差），乘积的复合误差无法通过Pareto前沿分散。')
add_bullet('Pareto前沿包含了乘积最优解（前沿上乘积最大的点自然就是乘积最优），但额外保留了"亮度极端"和"稳定极端"的备选，提供风险对冲。')

add_heading('6.2 6条序列分配策略', level=2)

add_table(
    ['序列#', 'Pareto区域', '来源策略', '风险等级', '设计目标'],
    [
        ['Seq 1', '亮度极端区（高亮度×中稳定）', '策略A', '低', '保险方案1：确保过30%阈值，争取高亮度'],
        ['Seq 2', '稳定极端区（中亮度×高稳定）', '策略A/C', '低', '保险方案2：确保过30%阈值，争取高稳定'],
        ['Seq 3', '乘积最优区（高亮度×高稳定）', '策略B/C', '中', '主力竞争序列：AI+结构联合优化'],
        ['Seq 4', '乘积最优区（高亮度×高稳定）', '策略E', '高', '冲金序列：ESM3生成的新颖序列'],
        ['Seq 5', '探索区（EI最大）', '策略D/E', '中高', '高风险高回报：GP认为探索不足的区域'],
        ['Seq 6', '平衡区（备选）', '最优可用', '中低', '冗余保障：替换任何不合格的前5条'],
    ]
)

add_heading('6.3 序列多样性检查', level=2)

add_bullet('CD-HIT 90%阈值：6条序列两两序列相似度均 < 90%。如违反，从同一Pareto区域选择下一个候补。')
add_bullet('策略多样性：6条序列至少覆盖3种设计策略（避免"把鸡蛋放一个篮子"）。')
add_bullet('Exclusion_List逐条比对：每条序列与排除名单做全局比对（非仅局部），任何完全一致即替换。另参考FPbase排除已知全长变体。')
add_bullet('长度合规：每条序列长度严格在220-250 aa之间，且以M开头（竞赛强制要求）。')

# ============================================================
# SECTION 7: LLM AGENT ORCHESTRATION
# ============================================================
page_break()
add_section_title('七、LLM Agent编排架构（竞赛加分项）')

add_para('竞赛规则明确："鼓励并欢迎使用各种大模型与代码Agent，若使用智能体设计，需在文档中展示Agent的逻辑树和关键执行日志"。以下架构在18天内可落地。')

add_heading('7.1 三层Agent架构', level=2)

add_table(
    ['层级', 'Agent名称', '职责', '触发条件', '输出'],
    [
        ['L1: 资源层', '模型选型Agent', '检测GPU→选择模型配置→环境验证', '管线启动时', '模型配置方案+环境验证报告'],
        ['L2: 执行层', '训练质量Agent', '评估预测模型表现→决定是否继续/降级/切换', '每次ML训练后', 'R²/r报告+Go/No-Go决策'],
        ['L2: 执行层', '结构验证Agent', 'AF2/ESMFold验证→分类(通过/需复核/淘汰)', '每批候选序列生成后', '分类标签+淘汰原因'],
        ['L3: 决策层', 'Pareto选择Agent', '构建Pareto前沿→分区选择→多样性检查→输出6条', '漏斗筛选完成后', '6条序列+入选理由+替代方案'],
        ['L3: 决策层', '收敛判断Agent', '分析GP趋势→决定继续迭代/收敛/强制终止', '每轮DBLT Learn阶段后', '收敛决策+下一轮参数调整建议'],
    ]
)

add_heading('7.2 五个决策节点逻辑树', level=2)

add_para('节点1：模型选型Agent', bold=True)
add_code("""输入: GPU状态 (torch.cuda检测)
├─ 显存≥24GB → ESM-2 650M + SaProt + ESM3 LoRA (推荐)
├─ 显存≥12GB → ESM-2 150M + SaProt
├─ 显存≥8GB  → ESM-2 150M
└─ 仅CPU     → ESM-2 35M (降级)
决策记录: [2026-06-10 09:15] 检测到RTX 3090 24GB → 选择ESM-2 650M+SaProt+ESM3 LoRA
备选: ESM3下载失败 → 降级到本地缓存ESM-2 150M""")

add_para('节点2：训练质量自适应Agent', bold=True)
add_code("""输入: 5-fold CV R², Pearson r
├─ R²>0.65, r>0.80 → 直接进入生成阶段
├─ R² 0.50-0.65     → 增加理化特征+MSA保守性(特征工程)
├─ R² 0.40-0.50      → 数据增强(补充Sarkisyan 2016全量) + ESM3 LoRA微调
└─ R²<0.40          → 触发预备Plan C: 切换到纯理性设计
决策记录: [2026-06-12 14:30] R²=0.63 → 中等, 启动特征工程+ESM3 LoRA微调
人工干预: [2026-06-13 10:00] LoRA微调后R²=0.71 → 进入生成阶段""")

add_para('节点3：结构验证筛选Agent', bold=True)
add_code("""输入: AF2/ESMFold pLDDT, pTM, 发色团RMSD
├─ pLDDT>85 AND pTM>0.75 AND 发色团RMSD<1.5 → 通过
├─ pLDDT 70-85 → AF2二次验证(不同随机种子)
├─ pLDDT<70 OR 发色团RMSD>3.0 → 淘汰
└─ 发色团区域pLDDT<80 → 淘汰(即使全局通过)
决策记录: R2_047: pLDDT=76.3→AF2验证→pLDDT=72.1→淘汰
原因: 发色团区域pLDDT=68.3, H148-发色团距离>4Å""")

add_para('节点4：Pareto前沿选择Agent', bold=True)
add_code("""输入: 所有候选(亮度预测μ±σ, 稳定性预测μ±σ)
1. 构建二维Pareto前沿 (非支配排序)
2. 从前沿4个区域按第6.2节策略选6条
3. 序列多样性检查(CD-HIT 90%)
4. 输出每条序列的入选理由+Pareto图上位置+替代方案
决策记录: [2026-06-20 16:00] Seq3选自乘积最优区, 预测Brightness=1.35±0.12,
 Stability=0.78±0.08, 来源=策略B+ESM3 LoRA, Pareto rank=1""")

add_para('节点5：收敛判断Agent', bold=True)
add_code("""输入: 最近2轮GP变化, EI_max趋势, MD vs GP Spearman r
├─ EI_max<0.05 (连续2轮) → 收敛→精选
├─ 已3轮迭代 → 强制精选(时间约束)
├─ MD与预测Spearman r>0.7 → 模型可靠→可提前收敛
└─ MD与预测矛盾(r<0.3) → 重新校准GP, 降低先验权重, 再跑一轮
决策记录: [2026-06-18 20:00] 第2轮: EI_max=0.032, Spearman r=0.73→
模型未收敛但预测可靠→建议再跑1轮, 重点探索EI高值区域""")

add_heading('7.3 Agent日志模板示例', level=2)

add_para('以下为最终提交PDF中将展示的日志格式。每条日志约3-5行，在管线运行过程中自动记录。', bold=True)

add_table(
    ['时间戳', 'Agent节点', '输入摘要', '决策', '依据', '人工干预'],
    [
        ['2026-06-10 09:15', '节点1', 'GPU: RTX 3090 24GB', 'ESM-2 650M+SaProt+ESM3 LoRA', '显存≥24GB阈值', '无'],
        ['2026-06-12 14:30', '节点2', '5-fold CV R²=0.63, r=0.78', '特征工程+ESM3 LoRA微调', 'R²在0.50-0.65区间', '无'],
        ['2026-06-13 10:00', '节点2', 'LoRA微调后 R²=0.71', '进入生成阶段', 'R²突破0.65阈值', '确认LoRA训练loss收敛'],
        ['2026-06-15 11:20', '节点3', '候选R2_047: pLDDT=76.3', '淘汰', '二次验证pLDDT=72.1, 发色团pLDDT=68.3', '确认淘汰理由充分'],
        ['2026-06-20 16:00', '节点4', '238条Pareto前沿候选', '选定6条', 'Pareto分区+多样性>3策略', 'Seq6从策略C改为策略E生成序列(ESM3 Gibbs第28轮产出, 预测乘积更高)'],
        ['2026-06-18 20:00', '节点5', '第2轮: EI_max=0.032, r=0.73', '继续迭代1轮', '未完全收敛但预言可靠', '确认GPU时间余量可支撑第3轮'],
    ]
)

# ============================================================
# SECTION 8: FIVE-STAGE FUNNEL
# ============================================================
page_break()
add_section_title('八、五级漏斗筛选详细设计')

add_heading('8.1 漏斗总览', level=2)

add_table(
    ['阶段', '候选数(输入→输出)', '方法', '筛选标准', '时间/序列', '累计淘汰率'],
    [
        ['Phase 1: 粗筛', '~14,000 → ~3,000', '长度筛选(220-250aa) + 序列合规 + Exclusion_List比对', 'M开头, 20种标准aa, 无终止密码子', '<1秒', '~79%'],
        ['Phase 2: 亮度预测', '~3,000 → ~600', '策略B三模型集成亮度预测(ESM2+SaProt+ESM3 LoRA加权)', '预测亮度 > WT sfGFP的50%（安全阈值，30%为淘汰线）', '~1秒', '~80%'],
        ['Phase 3: 结构验证', '~600 → ~200', 'AF2/ESMFold快速折叠预测', 'pLDDT>80, pTM>0.75, 发色团pLDDT>85', '~5分钟', '~67%'],
        ['Phase 4: 稳定性三级递进', '~200 → ~15', 'ThermoMPNN→FoldX→BioEmu→MD', '见8.2节', '秒→分→时', '~92%'],
        ['Phase 5: 精选决策', '~15 → 6', 'Agent节点4+人工复核', 'Pareto分区+多样性+6.2节策略', '数小时', '~60%'],
    ]
)

add_heading('8.2 Phase 4 稳定性评估三级递进', level=2)

add_para('Phase 4是整个漏斗中计算最密集的阶段，采用三级递进策略平衡精度与效率：')

add_para('第一级：ThermoMPNN + FoldX（~200条，秒级/条）', bold=True)
add_bullet('ThermoMPNN：预测ΔTm（熔解温度变化），筛选 ΔTm > -5°C（较WT不显著降低）')
add_bullet('FoldX：计算ΔΔG（折叠自由能变化），筛选 ΔΔG < 3 kcal/mol（不过度 destabilize）')
add_bullet('通过标准：ThermoMPNN ΔTm > -5°C AND FoldX ΔΔG < 3 kcal/mol → 筛选至~50条')

add_para('第二级：BioEmu构象系综预筛（~50条，分钟级/条）', bold=True)
add_bullet('BioEmu (Lewis, Hempel et al. 2025, Science)：微软开发的构象系综采样方法，比MD快约100倍')
add_bullet('筛选标准：(a) native_contacts > 0.65（核心折叠保持）；(b) 发色团区域RMSF < 2Å（活性位点刚性）；(c) 整体RMSD < 3Å（vs WT sfGFP）')
add_bullet('通过BioEmu筛选 ~50→15条进入MD金标准验证')

add_para('第三级：GROMACS MD金标准（~15条，小时级/条 × GPU并行）', bold=True)
add_para('这是整个管线的"金标准"验证步骤。MD结果不仅用于筛选，还用于校准上游快速预测器(GP代理模型)的系统偏差。', bold=False)

add_heading('8.3 MD协议详细规范', level=2)

add_table(
    ['参数', '设定', '依据/说明'],
    [
        ['力场', 'CHARMM36m (2021修订版)', '针对IDP和折叠蛋白均优化；在300-400K温度范围内经过系统验证(Huang et al. 2017, Nat Methods)'],
        ['水模型', 'TIP3P显式溶剂', '与CHARMM36m配套的标准水模型'],
        ['盒子', '十二面体盒子，蛋白-盒壁距离≥1.2nm', '标准做法，避免周期性镜像相互作用'],
        ['离子', '0.15M NaCl + 中和电荷', '模拟生理条件离子强度'],
        ['平衡 protocol', 'NVT 25°C 100ps → NPT 25°C 100ps', '逐步释放约束；position restraints on heavy atoms'],
        ['生产 protocol', 'NPT 72°C 10ns → NPT 25°C 5ns (复性)', '生产运行；时间步长2fs；H-bonds用LINCS约束'],
        ['恒温/恒压', 'V-rescale (τ=0.1ps) / Parrinello-Rahman (τ=2ps)', 'GROMACS推荐的生产设置'],
        ['分析指标', 'RMSD (Cα)、Rg (回转半径)、SASA、二级结构含量(DSSP)、发色团氢键维持率', '多指标综合判断；RMSD/Rg/SASA相对WT的变化量为首要筛选依据'],
        ['GPU并行', 'RTX 3090: ~15h/序列 (10ns 72°C + 5ns 25°C)；A100: ~4h/序列', '使用GROMACS 2024 GPU加速(COMM_GPU+GMX_GPU_PME)'],
        ['时间尺度讨论', '10ns MD不能模拟完全去折叠(需要μs-ms)，目标是捕捉"早期去折叠事件"(局部二级结构丧失、疏水核心膨胀、关键氢键断裂)。这些事件通常在ns-μs时间尺度发生，10ns可捕捉到趋势性变化。MD结果用于相对排序(与WT sfGFP同条件对比)而非绝对Tm预测。引文: Childers & Daggett (2018) 验证ns级MD对突变体稳定性排序的有效性。'],
    ]
)

add_heading('8.4 MD筛选阈值', level=2)

add_table(
    ['指标', '通过阈值', '警告阈值', '淘汰阈值'],
    [
        ['Cα RMSD (72°C vs WT)', '< 1.5x WT RMSD', '1.5-2.5x', '> 2.5x WT RMSD'],
        ['Rg变化 (72°C vs 25°C)', '< 1.2x WT ΔRg', '1.2-1.5x', '> 1.5x WT ΔRg'],
        ['SASA变化', '< 1.15x WT ΔSASA', '1.15-1.3x', '> 1.3x WT ΔSASA'],
        ['发色团氢键维持率', '> 80%', '60-80%', '< 60%'],
        ['复性后RMSD (25°C)', '< 1.2x WT', '1.2-1.5x', '> 1.5x WT'],
    ]
)

# ============================================================
# SECTION 9: ESM3 INTEGRATION & LORA FINETUNING
# ============================================================
page_break()
add_section_title('九、ESM3集成与LoRA微调方案（v3.0新增）')

add_para('本节回应评委审阅报告的核心建议："论证esmGFP启发性时引用ESM-3，实际设计时却不用，是否存在利用不足？"v3.0通过策略E（ESM3生成式管线）和策略B（ESM3 LoRA嵌入）在两个层面集成了ESM3。')

add_heading('9.1 ESM3在管线中的两个角色', level=2)

add_table(
    ['角色', '集成位置', '模型', '目的', '时间估算(RTX 3090)'],
    [
        ['嵌入器（预测）', '策略B：替代ESM-2+SaProt拼接', 'ESM3-sm-open (1.4B) + LoRA', '结构感知序列嵌入，提升亮度/稳定性预测精度', '3-6h (5000条序列嵌入)'],
        ['生成器（设计）', '策略E：esmGFP式Gibbs采样', 'ESM3-sm-open (1.4B)', '条件于sfGFP结构模板生成新颖序列', '60-70h (500条 × 8min Gibbs)'],
    ]
)

add_heading('9.2 LoRA微调协议', level=2)

add_para('LoRA (Low-Rank Adaptation) 是2025年ESM模型微调的事实标准。EiRA论文(2025)直接比较了ESM3的各种PEFT方法，结论是vanilla LoRA在结构任务上的泛化能力优于AdaLoRA/LoKr/hydraLoRA等更复杂变体。')

add_table(
    ['参数', '设定', '说明'],
    [
        ['基础模型', 'ESM3-sm-open (1.4B, biohub/esm3-sm-open-v1)', 'HuggingFace开源，需接受非商用许可'],
        ['LoRA Rank', 'r=16', '在参数效率和表达能力之间的平衡。rank越大容量越高但训练越慢。'],
        ['LoRA Alpha', 'α=32', 'scaling factor = α/r = 2。适中的缩放因子。'],
        ['目标模块', 'q_proj, v_proj, o_proj (注意力层)', '仅微调注意力层，冻结FFN和嵌入层。'],
        ['Dropout', '0.1', '正则化防止过拟合。'],
        ['可训练参数', '~800万 (总参数1.4B的0.57%)', 'RTX 3090 24GB可轻松训练。'],
        ['训练数据', 'GFP_data.xlsx(完整) + Sarkisyan 2016(完整) + 往年Top20序列', '共约5-6万条标注序列，其中高亮度/稳定序列作为正样本增强。'],
        ['训练目标', 'MSE(亮度回归) + MSE(稳定性回归) + Margin Ranking Loss(亮度排序)', '多任务学习：回归+排序联合优化。'],
        ['优化器', 'AdamW, lr=1e-4, cosine schedule, batch_size=16', '标准微调设置。'],
        ['训练时间', '2-4小时 (RTX 3090, ~10 epochs)', '快速迭代，可在管线运行中重训练。'],
    ]
)

add_heading('9.3 ESM3条件生成细节（策略E）', level=2)

add_para('以下是策略E的完整技术实现路径，基于esmGFP论文的官方教程：', bold=True)

add_code("""# Step 1: 构建条件提示 (以sfGFP PDB 2B3P为模板)
from esm.sdk.api import ESMProtein, GenerationConfig

prompt = ESMProtein.from_pdb("2B3P_sfGFP.pdb")

# 固定发色团区域的结构token
FIXED_POSITIONS = [62, 63, 64, 65, 66, 67, 93, 219]  # T65/Y66/G67区域+R96+E222
for pos in range(len(prompt.sequence)):
    if pos not in FIXED_POSITIONS:
        prompt.structure[pos] = None  # mask所有非关键位置
        prompt.sequence[pos] = None   # mask所有非关键序列

# Step 2: 生成结构
structure_gen = model.generate(
    prompt,
    GenerationConfig(
        track="structure",
        num_steps=230,        # sfGFP全长约230aa
        temperature=1.0,      # 初始高温→多样性
        top_p=0.95,
    )
)

# Step 3: 条件于结构生成序列
sequence_gen = model.generate(
    structure_gen,
    GenerationConfig(
        track="sequence",
        num_steps=20,         # 序列生成步数
        temperature=1.0,
    )
)

# Step 4: Gibbs采样迭代 (重复30+次)
for iteration in range(30):
    # 降温退火
    temp = 1.0 - (iteration / 30) * 0.9  # 1.0 → 0.1

    # 结构→序列
    sequence_gen = model.generate(
        structure_gen,
        GenerationConfig(track="sequence", temperature=temp)
    )

    # 序列→结构 (自验证)
    structure_gen = model.generate(
        sequence_gen,
        GenerationConfig(track="structure", temperature=0.0)  # 确定性折叠
    )

    # 发色团RMSD检查
    if chromophore_rmsd(structure_gen, template) > 1.5:
        continue  # 偏离模板, 跳过此样本""")

add_heading('9.4 ESM3 license 合规', level=2)

add_bullet('ESM3-sm-open-v1使用EvolutionaryScale非商业许可，竞赛（学术用途）符合条件。')
add_bullet('提交GitHub仓库时需在README.md中标注"Built with ESM"并包含完整许可证声明。')
add_bullet('引用Rives et al. (2025) Science论文作为方法学引用。')

# ============================================================
# SECTION 10: 18-DAY TIMELINE + GPU BUDGET
# ============================================================
page_break()
add_section_title('十、18天时间规划与GPU预算')

add_para('评委审阅报告指出："18天时间线与计算需求之间存在张力……需提供按GPU算力的时间预估矩阵。"以下为详细的时间规划和GPU预算。')

add_heading('10.1 GPU小时预算矩阵', level=2)

add_table(
    ['任务', 'RTX 3090 24GB', 'A100 80GB', 'A100×2', '备注'],
    [
        ['ESM-2 650M嵌入 (~10K seq)', '8h', '3h', '2h', '策略B'],
        ['ESM3 LoRA微调', '4h', '2h', '1.5h', '策略B升级'],
        ['ESM3嵌入 (~5K seq)', '6h', '2h', '1.5h', '策略B升级'],
        ['ProteinMPNN (~4K seq)', '4h', '1.5h', '1h', '策略C'],
        ['AF2/ESMFold (~600 seq)', '12h', '4h', '2.5h', 'Phase 3'],
        ['BioEmu (~50 seq)', '3h', '1h', '0.5h', 'Phase 4第二级'],
        ['MD 72°C+25°C (15 seq)', '225h', '60h', '30h', 'Phase 4第三级（关键瓶颈）'],
        ['ESM3 Gibbs生成 (500 seq)', '67h', '20h', '12h', '策略E（v3.0新增）'],
        ['其他（数据处理/日志/Agent）', '20h', '10h', '8h', ''],
        ['总计（串行等价）', '~349h ≈ 14.5天', '~103.5h ≈ 4.3天', '~59h ≈ 2.5天', ''],
        ['总计（并行优化后）', '~12-14天', '~3-4天', '~2-3天', '关键：MD 15条可并行(如GPU充足)'],
    ]
)

add_heading('10.2 逐日时间线', level=2)

add_table(
    ['日期段', '天数', '核心任务', '里程碑', '预备方案触发条件'],
    [
        ['6/10-6/11', 'Day 1-2', '环境搭建：ESM-2/SaProt/ProteinMPNN/ESM3下载+验证；策略A组合枚举启动', '所有模型可用；策略A首批序列生成', '模型下载失败→使用本地缓存版本'],
        ['6/12-6/13', 'Day 3-4', 'ESM3 LoRA微调训练；策略C ProteinMPNN生成；策略D MSA构建', 'LoRA微调完成(R²>0.65?)；策略C首批序列', 'LoRA R²<0.5→降级到无ESM3方案'],
        ['6/14-6/15', 'Day 5-6', '策略B嵌入完成；AF2/ESMFold批量验证；BioEmu预筛', '五策略候选池建成(~14K)→漏斗Phase 1-3', 'BioEmu不可用→跳过第二级直接MD'],
        ['6/16-6/18', 'Day 7-9', 'MD第1轮(重点：策略A/C/D产出的Top 5+策略B/E各3条)；DBLT第1轮Learn', '第1轮MD完成；GP初始训练', 'MD时间不够→减少至Top 8-10条'],
        ['6/19-6/20', 'Day 10-11', 'DBLT第2轮Design(基于第1轮GP引导)；ESM3 Gibbs采样(策略E)', '第2轮候选生成；Go/No-Go决策点', 'EI_max<0.05→收敛跳过第3轮'],
        ['6/21-6/22', 'Day 12-13', 'DBLT第2轮MD+Built+Test；策略E序列进入漏斗', '第2轮MD完成', '时间紧张→MD仅做10条'],
        ['6/23-6/24', 'Day 14-15', 'Pareto精选+Agent决策+6条序列最终选定；内部交叉验证', '6条序列确定+metadata.xlsx归档', '如6条中<3条通过所有验证→启动Plan B'],
        ['6/25-7/1', 'Day 16-18', '考试期：提交材料准备（CSV/PDF/GitHub）+ 内部归档 + 预演', '提交完成', '如主力方案失败→Plan C/D兜底'],
    ]
)

add_heading('10.3 Go/No-Go决策点', level=2)

add_para('第10天（6/19）是关键的Go/No-Go决策点：', bold=True)
add_bullet('✅ GO (继续第2-3轮)：策略B R²>0.6 AND MD第1轮完成≥8条 AND 至少4条候选通过Phase 4 Stability → 继续DBLT迭代')
add_bullet('⚠️ PARTIAL (减少轮次)：R² 0.5-0.6 OR MD完成<8条 → 跳过第3轮，直接精选现有候选')
add_bullet('❌ ABORT (切换预备方案)：R²<0.5 OR MD完成<3条 OR 任何MD出现发色团解体 → 切换Plan B(策略A+C)')

# ============================================================
# SECTION 11: DEFENSE PREPARATION
# ============================================================
page_break()
add_section_title('十一、答辩预案：评委可能提出的关键问题')

add_para('以下基于对评审团成员专业背景的深入分析（详见评委审阅报告第0节），预判6位核心评委可能提出的最尖锐问题，并准备回答策略。')

add_heading('11.1 欧阳颀教授（浙大，生物物理）', level=2)

add_para('预期问题："你用Matern 5/2核做GP代理模型，这个选择有什么理论依据？你考虑过蛋白质适应度景观的崎岖性对GP预测的影响吗？在50维的PCA嵌入空间中，MCMC的混合时间你估计过吗？"', bold=True)

add_para('回答策略：(1) Matern 5/2是Snoek et al. (2012)贝叶斯优化基准中推荐的默认核，适用于既有平滑趋势又有局部波动的黑箱函数。(2) GFP适应度景观的崎岖性(Sarkisyan 2016)正是选择Matern 5/2而非RBF的原因——RBF假设无限可微（过强），Matern 5/2允许二阶可微（更灵活）。(3) MCMC在50维空间混合的挑战是真实的——我们准备了Replica-exchange MCMC作为备选，且50维是PCA降维后的维度（保留~90%方差），实际有效维度可能更低。(4) 如果MCMC效率不佳，可以用CMA-ES替代——它是一个在50维空间更高效的进化策略算法。')

add_heading('11.2 吴边教授（中科院微生物所，蛋白质工程）', level=2)

add_para('预期问题："你固定了6个化学绝对位点。但在ProteinMPNN或ESM3全局重设计周围~227个残基后，新的序列环境是否会改变固定位点的催化效率？比如R96的催化活性依赖于周围的精确几何环境，重排后pKa可能发生偏移。"', bold=True)

add_para('回答策略：(1) 这是一个我们意识到并设计了针对性验证的真实风险。(2) 化学绝对约束保护的是反应路径的"存在性"（G67空间阻断→环化不可能），而非反应的"最优性"。(3) 管线通过两个机制检测这个问题：①发色团区域pLDDT>85确保局部结构精度；②MD中的发色团氢键维持率直接测量重排后微环境的变化。(4) 策略A（理性微调，不全局重设计）作为"保守对照"——如果策略C/E的全局重设计序列出现活性问题，可以通过与策略A的对比来诊断是否为"重排后微环境变化"所致。')

add_heading('11.3 鲍泽华教授（浙大，AI+生物智造）', level=2)

add_para("预期问题：\"你的五个Agent决策节点看起来完整。但18天内这些Agent的代码真的能写出来并稳定运行吗？Agent决策失误——比如错误淘汰了'预测不佳但实际很好的序列'——的风险你怎么控制？\"", bold=True)

add_para('回答策略：(1) Agent不追求完全自主——每个决策节点都保留人工干预选项（这是v3.0对v2.1的关键升级）。(2) Agent的角色是"自动执行明确可程序化的决策+标记边界情况供人工判断"。(3) 举例：结构验证Agent在pLDDT=83时（通过阈值85附近）标记"人工复核"而非自动淘汰。(4) 所有淘汰决策都保留在日志中，人工可在任何时刻回溯并覆盖。(5) 这是一个务实的"human-in-the-loop"设计哲学——Agent提升效率，但不替代判断。')

add_heading('11.4 戴磊教授（中科院深圳，进化生物学）', level=2)

add_para('预期问题："你策略D用jackhmmer做3轮MSA，但GFP家族在UniRef中的序列多样性很高。你是以sfGFP全长做query还是仅用保守区域？MSA的深度和覆盖度能不能支撑EVcouplings的统计效力？"', bold=True)

add_para('回答策略：(1) 以sfGFP全长(238aa)为query，对UniRef30做3轮jackhmmer迭代。(2) GFP家族在UniRef30中非常丰富——3轮迭代后预期获得>10,000条同源序列，远超EVcouplings的最低有效序列数(~100条)。(3) MSA深度用Effective Sequence Number(Neff)量化——预期Neff>100。(4) 但承认MSA偏向性（β-barrel荧光蛋白在自然界中以珊瑚/水母特定进化支为主）可能影响共进化分析的普适性——这正是策略D的输出作为"过滤规则"（而非"生成规则"）的原因。')

add_heading('11.5 于浩然教授（浙大，计算化学）', level=2)

add_para('预期问题："你的MD协议用CHARMM36m力场在72°C下模拟蛋白质。CHARMM36m在这个温度下蛋白质行为的预测能力经过验证吗？10ns的模拟中发色团区域的局部波动能代表10分钟实验中发生的不可逆变性吗？"', bold=True)

add_para('回答策略：(1) CHARMM36m在300-400K温度范围内经过系统验证（Huang et al. 2017, Nat Methods；该论文已引用>5,000次）。(2) 10ns MD不能模拟完全去折叠（需要μs-ms），但我们的目标不是预测绝对Tm——目标是"早期去折叠事件"的捕捉：局部二级结构丧失（ns尺度）、疏水核心膨胀（ns-μs尺度）、关键氢键断裂（ns尺度）。(3) 这些早期事件是72°C下不可逆变性的"先兆"——在10ns内即可观察到。(4) MD结果仅用于"相对排序"（与WT sfGFP同条件对比），而非绝对值预测。(5) 如需进一步验证，可引用Childers & Daggett (2018)验证ns级MD对突变体稳定性排序有效性的工作。')

add_heading('11.6 司同教授（中科院深圳，合成生物学）', level=2)

add_para("预期问题：\"Cell-Free体系与体内表达有本质区别。你管线的所有预测模型(ESM-2/SaProt/ThermoMPNN)的训练数据基本来自体内实验。这个'in vivo→in vitro'的迁移你怎么处理？会不会出现'在模型里看起来很好但在Cell-Free里完全不行'的情况？\"", bold=True)

add_para('回答策略：(1) 承认这是系统性挑战——所有队伍面临相同问题（因为公开训练数据几乎全来自体内实验）。(2) 竞赛统一使用Cell-Free NEExpress体系意味着所有提交序列在同一条件下比较，相对优势能保持。(3) 管线中的策略A从sfGFP出发微调（sfGFP本身在Cell-Free中良好表达——否则不会被选为WT基准），序列背景与Cell-Free兼容性最高。(4) 策略C(ProteinMPNN)和策略E(ESM3)基于结构设计而非体内适应度，对表达体系相对不敏感。(5) 不使用"mRNA二级结构"、"密码子偏好"等体内特异性特征——因为主办方使用DNA Chisel统一反向翻译，消除了密码子层面的体系差异。')

# ============================================================
# SECTION 12: RISK MATRIX
# ============================================================
page_break()
add_section_title('十二、风险矩阵与应对策略')

add_table(
    ['风险ID', '风险描述', '概率', '影响', '等级', '预防措施', '应对措施'],
    [
        ['R1', '策略B AI模型R²<0.5，无法有效预测', '中(30%)', '高', '🔴', '多模型集成(3+2个)；ESM3 LoRA微调提升0.05-0.10 R²', '切换Plan B：策略A+C撑场，AI仅做粗略排序'],
        ['R2', 'MD时间不足，15条无法在18天内完成', '高(50%)', '中', '🟡', 'GPU预算表(10.1节)；AutoDL A100租用', '减少MD至8-10条；跳过策略D/E的MD，优先策略A/C'],
        ['R3', 'ESM3-sm-open生成质量远低于7B API模型', '中高(40%)', '中', '🟡', 'LoRA微调缩小差距；策略E与策略C并行', '策略E降级为"探索性补充"；主力仍为策略A/C/B'],
        ['R4', 'GP贝叶斯优化在高维空间中失效', '低(20%)', '中', '🟢', 'PCA降维至50维；Replica-exchange MCMC备选', '切换到CMA-ES替代MCMC；退化为随机采样+EI排序'],
        ['R5', '提交序列中出现Exclusion_List匹配', '低(10%)', '致命', '🟡', '逐条全局比对+FPbase交叉验证', '替换为同一Pareto区域的下一个候补(有充足备选)'],
        ['R6', 'Cell-Free体系中多条序列<30% WT亮度', '中(20%)', '高', '🔴', '策略A保守序列作为保底(in vivo已验证的sfGFP骨架)', 'Plan D：手工精选6条中取最保守的提交'],
        ['R7', 'Agent决策失误导致优秀序列被误淘汰', '中(20%)', '中', '🟡', 'human-in-the-loop设计；所有淘汰决策可回溯', '保留Top 20候选至最后决策日，人工复核所有淘汰原因'],
        ['R8', 'ESM3许可合规问题', '低(5%)', '低', '🟢', '非商用许可；GitHub标注Built with ESM', '降级到ESM-2（非ESM3功能不受影响）'],
        ['R9', '答辩时被问到超出准备范围的技术问题', '中(30%)', '低', '🟢', '第11节答辩预案6个核心问题覆盖', '坦诚"这是开放挑战"+引用备选方案+承认局限性'],
    ]
)

# ============================================================
# SECTION 13: SUBMISSION CHECKLIST
# ============================================================
page_break()
add_section_title('十三、6/24前必须完成的提交材料')

add_heading('13.1 提交清单', level=2)

add_table(
    ['#', '材料', '格式/要求', '负责人', '截止', '验证方法'],
    [
        ['1', '设计序列CSV', 'Team_Name, Seq_ID(1-6), Sequence; UTF-8 无BOM; 220-250aa, M开头', '管线最终输出', '6/24', '逐条比对Exclusion_List + 长度检查 + 字符合规'],
        ['2', '设计思路PDF', '管线架构+5策略详述+DBLT迭代+Agent逻辑树+关键日志+6条理由', 'Agent日志→人工润色', '6/24', '内部同行评审(2人交叉检查)'],
        ['3', 'GitHub仓库', 'README.md(环境+依赖+运行方式)+推理代码+模型权重链接', '代码整理→Push', '6/24', '从另一台机器git clone→按README运行'],
        ['4', 'metadata.xlsx', '6条序列+预测亮度±σ+预测稳定性±σ+pLDDT+来源策略+选择依据', '管线运行自动生成', '6/24', '交叉验证预测值vs MD实测值'],
        ['5', '排除名单比对报告', '每条序列vs Exclusion_List的全局比对结果(截图/日志)', '脚本自动→人工确认', '6/24', '逐条review比对输出'],
    ]
)

add_heading('13.2 最后24小时检查清单', level=2)

add_bullet('CSV文件：用pandas读取→验证列名→验证序列长度→验证字符集→验证行数=6')
add_bullet('PDF文件：确认所有图表正常显示→确认无乱码→确认页数→确认文件大小<限制')
add_bullet('GitHub：确认repo为public→确认README可正常渲染→确认requirements.txt完整→确认推理代码可运行')
add_bullet('metadata.xlsx：6条序列的预测值vs MD实测值交叉验证→确认一致性')
add_bullet('最终review：打印CSV序列→逐条与Exclusion_List手工比对→FPbase搜索确认非已知变体')

# ============================================================
# APPENDIX A
# ============================================================
page_break()
add_section_title('附录A：关键命令速查表')

add_heading('A.1 环境安装', level=2)
add_code("""# 核心依赖
pip install torch torchvision torchaudio --index-url https://download.pytorch.org/whl/cu118
pip install fair-esm transformers datasets peft accelerate
pip install git+https://github.com/facebookresearch/esm.git  # ESM-2/ESM3
pip install git+https://github.com/dauparas/ProteinMPNN.git
pip install git+https://github.com/sokrypton/ColabFold.git
pip install biopython pandas numpy scikit-learn xgboost lightgbm
pip install gpytorch botorch  # Bayesian Optimization
pip install gromacswrapper mdanalysis  # MD analysis
pip install cdhit-py  # CD-HIT""")

add_heading('A.2 ESM3 LoRA微调', level=2)
add_code("""# 加载ESM3 + 应用LoRA
from esm.models.esm3 import ESM3
from peft import LoraConfig, get_peft_model

model = ESM3.from_pretrained("esm3-sm-open-v1").to("cuda")
lora_config = LoraConfig(r=16, lora_alpha=32,
    target_modules=["q_proj", "v_proj", "o_proj"], lora_dropout=0.1)
model = get_peft_model(model, lora_config)

# 添加预测头
import torch.nn as nn
class GFPHead(nn.Module):
    def __init__(self, dim=1536):
        super().__init__()
        self.brightness = nn.Sequential(nn.Linear(dim,256), nn.ReLU(), nn.Dropout(0.2), nn.Linear(256,1))
        self.stability = nn.Sequential(nn.Linear(dim,256), nn.ReLU(), nn.Dropout(0.2), nn.Linear(256,1))
    def forward(self, x): return self.brightness(x), self.stability(x)""")

add_heading('A.3 MD运行（GROMACS）', level=2)
add_code("""# 准备→运行→分析一条突变体
gmx pdb2gmx -f mutant.pdb -o processed.gro -water tip3p -ff charmm36-mar2019
gmx editconf -f processed.gro -o box.gro -c -d 1.2 -bt dodecahedron
gmx solvate -cp box.gro -cs spc216.gro -o solvated.gro -p topol.top
gmx grompp -f ions.mdp -c solvated.gro -p topol.top -o ions.tpr
gmx genion -s ions.tpr -o neutral.gro -p topol.top -pname NA -nname CL -neutral
# 平衡+生产
gmx grompp -f nvt_25C.mdp -c neutral.gro -p topol.top -o nvt.tpr
gmx mdrun -deffnm nvt -v -nb gpu
gmx grompp -f npt_72C_10ns.mdp -c nvt.gro -p topol.top -o npt_72C.tpr
gmx mdrun -deffnm npt_72C -v -nb gpu -pme gpu
gmx grompp -f npt_25C_5ns_renat.mdp -c npt_72C.gro -p topol.top -o renat.tpr
gmx mdrun -deffnm renat -v -nb gpu -pme gpu
# 分析
echo "C-alpha C-alpha" | gmx rms -s nvt.tpr -f npt_72C.xtc -o rmsd_72C.xvg
gmx gyrate -s nvt.tpr -f npt_72C.xtc -o rg_72C.xvg
gmx sasa -s nvt.tpr -f npt_72C.xtc -o sasa_72C.xvg""")

# ============================================================
# APPENDIX B: 45 CANDIDATE POSITIONS
# ============================================================
page_break()
add_section_title('附录B：45个候选设计位点（来自5.25方案）')

add_para('以下45个候选设计位点继承自5.25版规划方案，保留原始证据等级标注。在v3.0中，这些位点被重新分类到三级约束体系中。', bold=True)

add_heading('B.1 亮度增强位点（8个核心位点）', level=2)

add_table(
    ['位点', 'WT', '推荐突变', '预期效应', '证据等级', 'v3.0约束等级'],
    [
        ['S30', 'S', 'R', '增强β1-β2 loop刚性→提升折叠效率', 'A (sfGFP)', '第三级'],
        ['F64', 'F', 'L', '稳定发色团前体构象→加速成熟', 'A (sfGFP)', '第三级'],
        ['S65', 'S', 'T', '增强发色团吸光效率→提升亮度', 'A (sfGFP)', '第一级（小残基约束）'],
        ['F99', 'F', 'S', '减少疏水聚集→提升可溶性', 'A (sfGFP)', '第三级'],
        ['M153', 'M', 'T', '稳定β桶构象→减少非特异性聚集', 'A (sfGFP)', '第三级'],
        ['V163', 'V', 'A', '减少侧链碰撞→增强折叠协同性', 'A (sfGFP)', '第三级'],
        ['S175', 'S', 'G', '增加发色团环境极性→红移发射', 'B (文献)', '第三级'],
        ['A206', 'A', 'K', '表面电荷优化→增强溶解度和稳定性', 'B (文献)', '第三级'],
    ]
)

add_heading('B.2 热稳定性增强位点（12个核心位点）', level=2)

add_table(
    ['位点', 'WT', '推荐突变', '预期效应', '证据等级', 'v3.0约束等级'],
    [
        ['E6', 'E', 'D/K', 'N端帽H-bond优化→提升热稳定性', 'B', '第三级'],
        ['F8', 'F', 'L/I', '疏水核心填充→减少热涨落', 'B', '第三级'],
        ['L18', 'L', 'I/V', '疏水核心优化', 'B', '第三级'],
        ['V22', 'V', 'I/L', 'β桶疏水稳定化', 'B', '第三级'],
        ['Y39', 'Y', 'F/H', '芳香环π-π堆叠增强', 'B', '第三级'],
        ['L42', 'L', 'I/V', '疏水核心优化', 'B', '第三级'],
        ['L53', 'L', 'I/F', '疏水核心优化', 'B', '第三级'],
        ['V61', 'V', 'I/L', '疏水核心优化', 'B', '第三级'],
        ['L100', 'L', 'I', 'β桶疏水稳定化', 'B', '第三级'],
        ['L125', 'L', 'I', 'β桶疏水稳定化', 'B', '第三级'],
        ['T153', 'T', 'S', '与M153T协同→稳定β桶', 'B', '第三级'],
        ['N170', 'N', 'D/E', '表面盐桥增强→热稳定性', 'B', '第三级'],
    ]
)

# ============================================================
# APPENDIX C: REFERENCES
# ============================================================
page_break()
add_section_title('附录C：核心参考文献与工具速查')

add_heading('C.1 核心参考文献', level=2)

refs = [
    '[1] Sarkisyan KS, et al. (2016) "Local fitness landscape of the green fluorescent protein." Nature 533:397-401. DOI:10.1038/nature17995 — GFP适应度景观。管线设计哲学的核心参考文献。',
    '[2] Pédelacq J-D, et al. (2006) "Engineering and characterization of a superfolder green fluorescent protein." Nat Biotechnol 24:79-88. DOI:10.1038/nbt1172 — sfGFP。策略A的骨架和折叠增强突变的来源。',
    '[3] Rives A, Hayes T, et al. (2025) "Simulating 500 million years of evolution with a language model." Science 387:850-858. DOI:10.1126/science.ads0018 — ESM3/esmGFP。策略E的方法学模板和管线设计哲学的核心启示。',
    '[4] Lewis S, Hempel T, et al. (2025) "Scalable emulation of protein conformational ensembles." Science. DOI:10.1126/science.adq7214 — BioEmu。Phase 4稳定性评估第二级的核心工具。',
    '[5] Zhang H, et al. (2024) "Bright and stable monomeric green fluorescent protein derived from StayGold." Nat Methods 21:657-665. DOI:10.1038/s41592-024-02203-y — mBaoJin。策略D特征嫁接的关键参考序列。',
    '[6] Close DW, et al. (2015) "Thermal green protein, an extremely stable, nonaggregating fluorescent protein." Proteins 83:1225-1237. DOI:10.1002/prot.24699 — TGP。策略D热稳定性特征参考。',
    '[7] Dauparas J, et al. (2022) "Robust deep learning-based protein sequence design using ProteinMPNN." Science 378:49-56. — ProteinMPNN。策略C的核心工具。',
    '[8] Snoek J, Larochelle H, Adams RP (2012) "Practical Bayesian Optimization of Machine Learning Algorithms." NeurIPS. — 贝叶斯优化方法论。GP核函数选择和EI采集函数的参考文献。',
    '[9] Huang J, et al. (2017) "CHARMM36m: an improved force field for folded and intrinsically disordered proteins." Nat Methods 14:71-73. — MD力场选择的参考文献。',
    '[10] Childers MC, Daggett V (2018) "Validating Molecular Dynamics Simulations against Experimental Observables." J Phys Chem B. — MD时间尺度论证的参考文献。',
]
for ref in refs:
    add_bullet(ref)

add_heading('C.2 关键工具与链接', level=2)

tools = [
    'ESM-2: github.com/facebookresearch/esm | ESM3 (HuggingFace): huggingface.co/biohub/esm3-sm-open-v1',
    'ProteinMPNN: github.com/dauparas/ProteinMPNN | ColabFold: github.com/sokrypton/ColabFold',
    'BioEmu: github.com/microsoft/bioemu | FoldX: foldxsuite.crg.eu',
    'GROMACS: gromacs.org | CHARMM36m: mackerell.umaryland.edu/charmm_ff.shtml',
    'EVcouplings: github.com/debbiemarkslab/EVcouplings | CD-HIT: github.com/weizhongli/cdhit',
    'jackhmmer (HMMER3): hmmer.org | MAFFT: mafft.cbrc.jp/alignment/software/',
    'GPyTorch: gpytorch.ai | BoTorch: botorch.org | PeFT (LoRA): github.com/huggingface/peft',
    'AutoDL (Cloud GPU): autodl.com (A100约¥3-5/h) | FPbase: fpbase.org',
    'DNA Chisel: github.com/Edinburgh-Genome-Foundry/DnaChisel | Biopython: biopython.org',
]
for tool in tools:
    add_bullet(tool)

add_heading('C.3 官方竞赛资源', level=2)

add_bullet('2026 SynBio Challenges 官方说明 (PDF)：2026Protein Design in Synbio challenges.pdf')
add_bullet('GFP训练数据：data/GFP_data.xlsx (含beforetopseqs sheet：历年Top20序列)')
add_bullet('排除名单：data/Exclusion_List.csv')
add_bullet('5条参考GFP序列：AAseqs of 5 GFP proteins.txt')
add_bullet('提交模板：submission_template.csv')
add_bullet('基础教程：bohrium.com/notebooks/36787332597 (2025版，需自行适配2026格式)')

# ============================================================
# FINAL PAGE
# ============================================================
page_break()
doc.add_paragraph()
doc.add_paragraph()

add_para('—— 文档结束 ——', alignment=WD_ALIGN_PARAGRAPH.CENTER,
         font_name='微软雅黑', font_size=Pt(14), color=(0x50, 0x50, 0x50))

doc.add_paragraph()
add_para('本方案 v3.0 基于以下来源整合而成：', font_size=Pt(10), color=(0x80, 0x80, 0x80))
add_para('· v2.0 管线设计（2026.06.07 版本）', font_size=Pt(10), color=(0x80, 0x80, 0x80))
add_para('· v2.1 管线设计（2026.06.07 版本，4策略+DBLT+Agent）', font_size=Pt(10), color=(0x80, 0x80, 0x80))
add_para('· 评委视角审阅报告（2026.06.09，含评委团画像分析）', font_size=Pt(10), color=(0x80, 0x80, 0x80))
add_para('· ESM3集成与LoRA微调方案（2026.06.09，含EiRA论文参考）', font_size=Pt(10), color=(0x80, 0x80, 0x80))
add_para('· 2026 SynBio Challenges 蛋白质设计赛道官方规则（14页PDF）', font_size=Pt(10), color=(0x80, 0x80, 0x80))

# ============================================================
# SAVE
# ============================================================
doc.save("D:/蛋白质设计-合成生物学创新赛-Claude/文档/GFP蛋白质设计_管线设计与实施方案_v3.0.docx")
print("v3.0 document saved successfully!")
