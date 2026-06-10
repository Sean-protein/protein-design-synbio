"""Final verification of v2.1 format matching v2.0."""
from docx import Document
from docx.shared import Pt, Cm

ref = Document("D:/蛋白质设计-合成生物学创新赛-Claude/文档/GFP蛋白质设计_管线设计与实施方案_v2.0.docx")
doc = Document("D:/蛋白质设计-合成生物学创新赛-Claude/文档/GFP蛋白质设计_管线设计与实施方案_v2.1.docx")

print("=" * 60)
print("FINAL FORMAT VERIFICATION")
print("=" * 60)

# 1. Page size
print("\n1. PAGE SETUP")
for label, d in [("v2.0 REF", ref), ("v2.1 OUT", doc)]:
    s = d.sections[0]
    w = s.page_width / 360000
    h = s.page_height / 360000
    t = s.top_margin / 360000
    print(f"  {label}: {w:.1f}cm x {h:.1f}cm, margins={t:.1f}cm")

# 2. Styles
print("\n2. STYLES")
for sname in ['Normal', 'Heading 1', 'Heading 2', 'Heading 3']:
    rs = ref.styles[sname]
    ds = doc.styles[sname]
    rf = rs.font
    df = ds.font
    r_color = str(rf.color.rgb) if rf.color and rf.color.rgb else 'none'
    d_color = str(df.color.rgb) if df.color and df.color.rgb else 'none'
    match = "✅" if (rf.size == df.size and rf.bold == df.bold and r_color == d_color) else "❌"
    print(f"  {sname}: {match}")
    print(f"    v2.0: font={rf.name}, size={rf.size}, bold={rf.bold}, color={r_color}")
    print(f"    v2.1: font={df.name}, size={df.size}, bold={df.bold}, color={d_color}")

# 3. Fonts used
print("\n3. FONTS IN USE")
for label, d in [("v2.0 REF", ref), ("v2.1 OUT", doc)]:
    fonts = set()
    ms_gothic = 0
    for para in d.paragraphs:
        for run in para.runs:
            if run.font.name:
                fonts.add(run.font.name)
            # Check XML for MS Gothic
            from docx.oxml.ns import qn
            rPr = run._element.find(qn('w:rPr'))
            if rPr is not None:
                rFonts = rPr.find(qn('w:rFonts'))
                if rFonts is not None:
                    east = rFonts.get(qn('w:eastAsia'))
                    if east == 'MS Gothic':
                        ms_gothic += 1
    print(f"  {label}: {sorted(fonts)}")
    print(f"    MS Gothic refs: {ms_gothic}")

# 4. Line spacing distribution
print("\n4. LINE SPACING")
for label, d in [("v2.0 REF", ref), ("v2.1 OUT", doc)]:
    counts = {}
    for para in d.paragraphs:
        ls = para.paragraph_format.line_spacing
        key = str(ls) if ls else 'default'
        counts[key] = counts.get(key, 0) + 1
    print(f"  {label}: {counts}")

# 5. Random spot check of body paragraphs
print("\n5. BODY TEXT SPOT CHECK")
for i, (rp, dp) in enumerate(zip(ref.paragraphs[100:120], doc.paragraphs[80:100])):
    rt = rp.text.strip()
    dt = dp.text.strip()
    if rt and dt and len(rt) > 20:
        # Compare run font settings
        if rp.runs and dp.runs:
            rr = rp.runs[0]
            dr = dp.runs[0]
            print(f"  v2.0: font={rr.font.name}, size={rr.font.size}, bold={rr.font.bold}")
            print(f"  v2.1: font={dr.font.name}, size={dr.font.size}, bold={dr.font.bold}")
            break

print("\n" + "=" * 60)
print("VERIFICATION COMPLETE")
print("=" * 60)
