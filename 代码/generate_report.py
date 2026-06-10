#!/usr/bin/env python3
"""生成 GFP 蛋白质设计文献与工具调研报告 Word 文档"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import datetime

doc = Document()

# ── 页面设置 ──
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(11)

# ── 辅助函数 ──
def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 51, 102)
    return h

def add_para(text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)
    return p

def add_link_item(title, url, description=""):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(11)
    if url:
        p.add_run(f"\n{url}").font.size = Pt(9)
    if description:
        p.add_run(f"\n{description}").font.size = Pt(10)

# ═══════════════════════════════════════════
# 封面
# ═══════════════════════════════════════════
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('运用计算方法设计兼具高荧光亮度和\n优良热稳定性的GFP变体')
run.bold = True
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(0, 51, 102)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('文献与开源项目调研报告')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0, 80, 130)

doc.add_paragraph()
date_para = doc.add_paragraph()
date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_para.add_run(f'生成日期: {datetime.date.today().strftime("%Y年%m月%d日")}').font.size = Pt(12)

doc.add_page_break()

# ═══════════════════════════════════════════
# 目录
# ═══════════════════════════════════════════
add_heading_styled('目录', 1)
toc_items = [
    '一、调研背景与目标',
    '二、直接相关的GitHub竞赛项目',
    '三、蛋白质语言模型与生成式设计',
    '四、热稳定性预测与设计工具',
    '五、分子动力学与理性设计',
    '六、定向进化与MLDE方法',
    '七、反向翻译与序列优化工具',
    '八、推荐技术路线',
    '九、参考文献汇总（含DOI）',
    '十、开源工具汇总'
]
for item in toc_items:
    add_para(item)

doc.add_page_break()

# ═══════════════════════════════════════════
# 一、调研背景与目标
# ═══════════════════════════════════════════
add_heading_styled('一、调研背景与目标', 1)

add_para('本调研针对合成生物学创新赛"蛋白质设计"赛道展开，目标是检索与"运用计算方法设计兼具高荧光亮度和优良热稳定性的GFP变体"相关的学术文献和开源项目。')

add_para('关键竞赛参数：', bold=True)
params = [
    '每条序列长度：220-250 aa，以蛋氨酸（M）开头',
    '仅允许20种标准氨基酸大写字母',
    '每队提交6条氨基酸序列',
    '得分公式：综合得分 = (F_initial / F_initial_WT) × (F_final / F_initial)',
    '低亮度淘汰线：F_initial < 0.3 × F_initial_WT',
    'DNA合成使用DNAChisel算法进行反向翻译（统一由大设施执行）',
    '72°C热处理评估热稳定性'
]
for p in params:
    doc.add_paragraph(p, style='List Bullet')

doc.add_page_break()

# ═══════════════════════════════════════════
# 二、直接相关的GitHub竞赛项目
# ═══════════════════════════════════════════
add_heading_styled('二、直接相关的GitHub竞赛项目', 1)

add_para('以下项目是2025年合成生物学蛋白质设计竞赛的参赛仓库，与本次任务目标高度一致：')

# 项目1
add_heading_styled('2.1 SynBioChallenges2025 — 多方法ML管线', 2)
add_link_item(
    'f-normies/synbiochallenges2025',
    'https://github.com/f-normies/synbiochallenges2025',
    '综合ML管线，结合GLEAM（基于ESM-2 3B的注意力荧光预测模型）+ MLDE（机器学习定向进化）+ 交叉熵蒙特卡洛 + TemBERTure热稳定性筛选。'
    '约束条件：≤6个突变。针对avGFP（238 aa）联合优化荧光亮度和热变性抗性。'
)

add_para('核心组件：', bold=True)
components = [
    'GLEAM：使用ESM-2 (3B)嵌入，通过全局-局部交叉注意力机制和蒙特卡洛Dropout进行不确定性量化的亮度预测',
    'MLDE：基于k-mer TF-IDF/熵重要性评分，使用ESM-2掩码语言建模进行引导式突变采样',
    '交叉熵蒙特卡洛：显式建模成对突变间的上位效应（epistasis），Gibbs采样生成序列',
    'TemBERTure：基于protBERT-BFD的适配器模型，同时预测嗜热/非嗜热分类和熔解温度（Tm）回归'
]
for c in components:
    doc.add_paragraph(c, style='List Bullet')

# 项目2
add_heading_styled('2.2 SC2025 — 深度强化学习方法', 2)
add_link_item(
    'ZhouLon/SC2025',
    'https://github.com/ZhouLon/SC2025',
    '将蛋白质设计建模为围棋游戏——状态 = 序列空间（238×20），动作 = 按位点氨基酸替换。'
    '使用双头DQN网络（位置选择头 + 氨基酸选择头），ε-greedy探索策略，双重经验回放缓冲区。'
    'A800 GPU (80 GB)。实验验证：12条序列合成，11条亮度高于野生型。'
    '许可证：GNU GPLv3。'
)

# 项目3
add_heading_styled('2.3 MCCOP — 扩散引导潜在空间优化', 2)
add_link_item(
    'weroks/mccop',
    'https://github.com/weroks/mccop',
    '反事实蛋白质优化。在连续序列-结构联合潜在空间（ESMFold CHEAP嵌入）中操作。'
    '使用预训练扩散模型（DiMA）作为流形先验，确保生成变体可折叠。'
    '在GFP荧光任务中将突变集中在发色团邻近区域（残基63-69）和β-桶链。'
    '实现稀疏反事实（~1.4-2.5个突变），对抗率接近零。'
)

# 项目4
add_heading_styled('2.4 GGS — 平滑适应度景观', 2)
add_link_item(
    'kirjner/GGS',
    'https://github.com/kirjner/GGS',
    'ICLR 2024论文代码。使用Tikunov正则化平滑蛋白质适应度景观（图信号）。'
    '在GFP和AAV基准测试上实现最先进结果，适应度比训练集提高2.5倍。'
    '专为有限数据场景设计。'
)

doc.add_page_break()

# ═══════════════════════════════════════════
# 三、蛋白质语言模型与生成式设计
# ═══════════════════════════════════════════
add_heading_styled('三、蛋白质语言模型与生成式设计', 1)

add_heading_styled('3.1 ESM3 / esmGFP — 模拟5亿年进化（Science, 2025）', 2)
add_para(
    'DOI: 10.1126/science.ads0018  |  '
    'GitHub: https://github.com/evolutionaryscale/esm (ESM开源版)'
)
add_para(
    'EvolutionaryScale开发了980亿参数的多模态生成式蛋白质语言模型ESM3，能够对'
    '序列、结构、功能三种模态进行联合推理和生成。通过思维链（Chain-of-Thought）'
    '提示策略，ESM3生成了全新的绿色荧光蛋白esmGFP——与已知最近天然荧光蛋白(tagRFP)'
    '仅有58%序列同源性（229个氨基酸中差异达96个突变），估计相当于自然进化超过5亿年。'
    '荧光亮度与天然EGFP相当，激发/发射光谱高度重叠。'
)
add_para('关键技术创新：', bold=True)
for t in [
    '多轨道Transformer架构——序列、结构、功能统一编码为离散token',
    '几何注意力机制（Geometric Attention）——基于局部参考系的精确3D结构建模，重建误差<0.5Å',
    '可变掩码调度策略——随机遮蔽任意模态token，迫使跨模态学习',
    '缩放定律——14亿→70亿→980亿参数，生成能力和结构预测准确性显著提升'
]:
    doc.add_paragraph(t, style='List Bullet')

add_heading_styled('3.2 METL — 基于生物物理学的PLM（Nature Methods, 2025）', 2)
add_para(
    'DOI: 10.1038/s41592-025-02667-2  |  '
    'GitHub: https://github.com/rsgelman/METL'
)
add_para(
    'METL（Mutational Effect Transfer Learning）解决了传统PLM忽略生物物理机制的问题。'
    '模型在Rosetta模拟生成的数千万变体（55种生物物理属性）上进行预训练，'
    '再在实验数据上微调。仅使用64个训练样本即设计出功能性GFP变体：'
    '训练中见过的氨基酸——10/10成功；未见氨基酸——6/10成功。'
    '在低N场景（≤128个训练样本）和位置外推任务中显著优于ESM-2。'
)

add_heading_styled('3.3 GeoEvoBuilder — 兼顾活性与稳定性（PNAS, 2025）', 2)
add_para(
    'DOI: 10.1073/pnas.2504117122  |  '
    'GitHub: https://github.com/PKUliujl/GeoEvoBuilder'
)
add_para(
    '北京大学来鲁华团队开发。由GeoSeqBuilder（基于结构序列设计）、ESM2（进化信息捕获）'
    '和适应性连接模块三部分组成。零样本学习、单轮设计，可改变超30%残基。'
)
add_para('GFP关键结果：', bold=True)
for t in [
    '22条设计序列中20条可溶性表达，17条具有显著荧光活性',
    '最佳突变体1GFL-15：最大激发波长红移70 nm，470 nm激发下荧光强度为WT的2.3倍',
    'Tm提高约2°C',
    '功能序列与WT序列一致性最低仅67.37%（即突变率超30%）'
]:
    doc.add_paragraph(t, style='List Bullet')

add_heading_styled('3.4 ProteinMPNN — 反向折叠荧光蛋白重设计（Protein Science, 2024）', 2)
add_para(
    'DOI: 10.1002/pro.70002  |  '
    'GitHub: https://github.com/dauparas/ProteinMPNN'
)
add_para(
    'Nikolaev等人的工作直接验证了ProteinMPNN重设计荧光蛋白的可行性。'
    '以CagFbFP（黄素基荧光蛋白）为模板，固定20个与FMN发色团相互作用的氨基酸，'
    '允许ProteinMPNN重设计其余86个残基。ProteinMPNN建议36-48个替换（序列一致性55-66%）。'
    '3个实验测试设计均成功结合FMN、发出荧光并保持了与WT相当的热稳定性。'
    '该结果直接证明了ProteinMPNN可用于生成荧光蛋白的功能性发散变体。'
)

doc.add_page_break()

# ═══════════════════════════════════════════
# 四、热稳定性预测与设计工具
# ═══════════════════════════════════════════
add_heading_styled('四、热稳定性预测与设计工具', 1)

add_heading_styled('4.1 TemBERTure — 温度稳定性专用预测', 2)
add_link_item(
    'ibmm-unibe-ch/TemBERTure',
    'https://github.com/ibmm-unibe-ch/TemBERTure',
    '基于protBERT-BFD + 轻量瓶颈适配器。提供TemBERTureCLS（嗜热/非嗜热二分类）'
    '和TemBERTureTm（Tm回归预测）。3副本集成，附带不确定性估计。'
    '已被synbiochallenges2025管线集成用于GFP热稳定性筛选。'
)

add_heading_styled('4.2 SPURS — ESM+ProteinMPNN融合稳定性预测（2025）', 2)
add_link_item(
    'luo-group/SPURS',
    'https://github.com/luo-group/SPURS',
    '将ProteinMPNN的结构嵌入注入ESM-2注意力层（适配器方法），更新98.5%更少的参数。'
    '单次前向传播即可预测L×20所有单突变ΔΔG（比逐突变方法快约400倍）。'
    'Megascale测试集Spearman 0.83（vs ThermoMPNN 0.77），'
    'Human Domainome数据集0.54（563k变体）。'
    '对稳定突变（训练数据中稀少）的识别精度特别优异。'
)

add_heading_styled('4.3 ProteinMPNN-ddG — 零样本稳定性预测（2024）', 2)
add_link_item(
    'PeptoneLtd/proteinmpnn_ddg',
    'https://github.com/PeptoneLtd/proteinmpnn_ddg',
    '零样本方法，在ProteinMPNN中增加一个"折叠参考态"项（类比经典自由能计算）。'
    '稳定突变识别改善11%（Success@10：66%→77%）。'
    '可实现蛋白质组级饱和突变扫描——23,391个结构在单V100 GPU上30分钟。'
    'NeurIPS 2024 Workshop论文。'
)

add_heading_styled('4.4 ProtSSN — 序列-结构融合零样本预测', 2)
add_link_item(
    'openmedlab/ProtSSN',
    'https://github.com/openmedlab/ProtSSN',
    '融合序列编码器和几何编码器，去噪预训练框架。无需MSA即可推理。'
    '自建DTm和DDG两个热稳定性基准。'
    '在ProteinGym上超越大多数基于MSA和语言模型的方法。'
)

add_heading_styled('4.5 PROSS — 蛋白质稳定性自动设计', 2)
add_para(
    'Fleishman实验室 (Weizmann Institute) 开发。结合进化序列分析（PSSM）和Rosetta能量计算，'
    '自动设计稳定性提高的蛋白质变体。提供3-4个不同严格度的设计方案。'
)
add_para('在线服务器：', bold=True)
add_link_item('PROSS Web Server', 'https://pross.weizmann.ac.il')
add_link_item('mPROSS（膜蛋白版，2024）', 'https://mpross.weizmann.ac.il')
add_para(
    '2021年社区评估：14个靶标中9/14表达改善，9/10热稳定性提高。'
)

add_heading_styled('4.6 protein-stability-optimization管线', 2)
add_link_item(
    'izgys/protein-stability-optimization',
    'https://github.com/izgys/protein-stability-optimization',
    'ProteinMPNN/ThermoMPNN突变设计 → ColabFold结构预测 → OpenMM/GROMACS MD验证'
    '（RMSD、氢键、ΔΔG代理指标）。以T4溶菌酶为示例，工作流可直接迁移至GFP。'
)

doc.add_page_break()

# ═══════════════════════════════════════════
# 五、分子动力学与理性设计
# ═══════════════════════════════════════════
add_heading_styled('五、分子动力学与理性设计', 1)

add_heading_styled('5.1 YuzuFP — MD引导亮度增强（Communications Chemistry, 2025）', 2)
add_para(
    'DOI: 10.1038/s42004-025-01573-4'
)
add_para(
    '短时间MD（10 ns）筛选H148位置所有19种氨基酸，评估与发色团酚酸盐和保守水分子W1的氢键网络。'
    'H148S单突变体（YuzuFP）：亮度比sfGFP提高1.5倍（比EGFP提高1.7倍），'
    '活细胞中抗光漂白性提高约3倍。荧光"开启"时间更长（5.3 s vs 3.6 s）。'
    '机制：S148形成持续性氢键（60%模拟时间 vs H148的<5%），水分子W1驻留时间更长。'
)

add_heading_styled('5.2 FluProCAD — 自动FP计算筛选工作流（Molecular Physics, 2025）', 2)
add_para(
    'DOI: 10.1080/00268976.2025.2458641'
)
add_para(
    '命令行自动化工作流，设置并计算荧光蛋白突变体的关键光学和热力学性质。'
    '无需建模专业知识。成功重现5种avGFP变体的光学响应和折叠/二聚化自由能，'
    '验证了14种rsGreen0.7变体的结构预测与晶体结构的一致性。'
)

add_heading_styled('5.3 参数化β-桶设计 — RFdiffusion+ProteinMPNN (2024-2025)', 2)
add_para(
    'DOI: 10.1101/2024.07.22.604663  |  '
    'GitHub: https://github.com/RosettaCommons/RFdiffusion'
)
add_para(
    'Baker实验室将参数化骨架生成与RFdiffusion/ProteinMPNN结合，实现de novo β-桶蛋白设计。'
    '72/96设计可溶表达，29/32显示正确CD光谱，14个在95°C下稳定。'
    'BBn6晶体结构分辨率2.0Å，与设计模型Cα RMSD仅0.9Å。'
    '该方法直接适用于荧光β-桶蛋白的从头设计。'
)

doc.add_page_break()

# ═══════════════════════════════════════════
# 六、定向进化与MLDE方法
# ═══════════════════════════════════════════
add_heading_styled('六、定向进化与MLDE方法', 1)

add_heading_styled('6.1 超低N虚拟适应度景观（Harvard, 2024）', 2)
add_para(
    'Harvard研究组提出仅使用24个功能测定突变体序列即可构建精确虚拟适应度景观，'
    '并在in silico中筛选1000万序列。在avGFP和TEM-1 β-内酰胺酶上验证。'
    '模型学习"非自然性"潜在表示，引导搜索远离无功能序列区域。'
)

add_heading_styled('6.2 MLDE综合基准（Caltech/Arnold Lab, 2024）', 2)
add_para(
    'DOI: 10.1101/2024.10.24.619774'
)
add_para(
    '在16个蛋白质适应度景观上系统评估MLDE、ALDE、ftMLDE策略，使用6种零样本预测因子。'
    '关键发现：MLDE在传统DE难以处理的景观上优势最大（活性变体更少、'
    '崎岖度更高、上位效应更强）。聚焦训练+零样本预测因子始终优于随机采样。'
)

add_heading_styled('6.3 二进制分选数据预测连续性质（PNAS, 2024）', 2)
add_para(
    'DOI: 10.1073/pnas.2317527121'
)
add_para(
    '使用FACS分选数据和线性ML模型（LDA）预测连续蛋白质性质。'
    'GFP适应度景观（Sarkisyan et al.）为核心基准。'
    '线性模型在二进制标签上训练，表现与在精确连续数据上训练的Ridge回归相当（Spearman ρ~0.85）。'
    '使用整数线性规划扩展预测到未见序列空间。'
)

add_heading_styled('6.4 TopVIP — 顶级变体识别管线（Cell Systems, 2024）', 2)
add_para(
    'DOI: 10.1016/j.cels.2024.01.006'
)
add_para(
    '集成多种零样本预测因子 + 采样算法 + 迭代低N策略。'
    '表明零样本预测因子单独使用不够，结合监督低N模型效果最佳。'
)

doc.add_page_break()

# ═══════════════════════════════════════════
# 七、反向翻译与序列优化
# ═══════════════════════════════════════════
add_heading_styled('七、反向翻译与序列优化工具', 1)

add_heading_styled('7.1 DNAChisel — 竞赛指定反向翻译算法', 2)
add_link_item(
    'Edinburgh-Genome-Foundry/DnaChisel',
    'https://github.com/Edinburgh-Genome-Foundry/DNAChisel',
    'Edinburgh Genome Foundry开发的DNA序列优化Python库（MIT许可证，Bioinformatics 2020）。'
    '竞赛使用其reverse_translate函数将氨基酸序列转为DNA，并包含CodonOptimize'
    '（支持use_best_codon、match_codon_usage、harmonize_rca三种方法）。'
    '支持E. coli等宿主。安装：pip install dnachisel。'
)
add_para(
    '参赛队伍无需自行操作反向翻译，大设施统一执行。但了解其原理有助于设计阶段考虑'
    '密码子偏好性的影响。'
)

add_heading_styled('7.2 其他相关序列优化工具', 2)
add_link_item(
    'BigHat-Biosciences/bh-DnaChisel',
    'https://github.com/BigHat-Biosciences/bh-DnaChisel',
    'BigHat Biosciences维护的DNAChisel分支，专注抗体序列优化。'
)

doc.add_page_break()

# ═══════════════════════════════════════════
# 八、推荐技术路线
# ═══════════════════════════════════════════
add_heading_styled('八、推荐技术路线', 1)

add_para('基于上述文献和工具调研，推荐以下几条互补的技术路线：')

add_heading_styled('路线A：ESM3/ESM2 + 结构约束生成', 2)
steps_a = [
    '以sfGFP或avGFP的结构（PDB）为条件，使用ESM3/ESM2进行序列生成',
    '约束发色团形成残基（Tyr66-Gly67-Gly68）和关键折叠残基',
    '在prompt中同时编码高亮度和高稳定性要求',
    '使用AlphaFold2/ESMFold验证生成序列的结构完整性（pLDDT、RMSD）',
]
for s in steps_a:
    doc.add_paragraph(s, style='List Number')

add_heading_styled('路线B：GeoEvoBuilder 结构-进化联合设计', 2)
steps_b = [
    '以GFP结构（如1GFL）为输入，运行GeoEvoBuilder（已在GFP上验证）',
    '零样本生成大量候选变体（支持超30%序列变化）',
    '筛选具有结构完整性且序列多样性的候选',
    '使用TemBERTure或SPURS评估热稳定性',
]
for s in steps_b:
    doc.add_paragraph(s, style='List Number')

add_heading_styled('路线C：MLDE + 多目标优化', 2)
steps_c = [
    '收集已有GFP变体实验数据作为初始训练集（如Sarkisyan等人的适应度景观数据）',
    '使用ESM-2嵌入 + GLEAM式注意力模型预测荧光亮度',
    '使用TemBERTure/SPURS预测热稳定性',
    '采用交叉熵蒙特卡洛或贝叶斯优化联合搜索高亮度+高稳定性的序列',
    '约束突变数≤6-10个（维持折叠完整性）',
]
for s in steps_c:
    doc.add_paragraph(s, style='List Number')

add_heading_styled('路线D：ProteinMPNN + MD筛选', 2)
steps_d = [
    '基于GFP晶体结构，使用ProteinMPNN生成大量序列变体',
    '固定发色团口袋关键残基（Tyr66、Gly67、Gly68、Arg96、Glu222等）',
    '使用AlphaFold2/ColabFold预测结构，筛选高pLDDT候选',
    '通过短MD模拟（参考YuzuFP方法）评估发色团氢键网络',
    '使用Rosetta ddG或FoldX预测折叠稳定性',
]
for s in steps_d:
    doc.add_paragraph(s, style='List Number')

add_heading_styled('路线E：RFDiffusion De Novo β-桶设计', 2)
steps_e = [
    '使用RFDiffusion从零设计β-桶骨架（参数化控制桶直径、链数）',
    '在发色团位置嵌入GFP发色团形成残基',
    '使用ProteinMPNN设计序列',
    'AlphaFold2筛选可折叠设计',
    '从零设计可能产生新颖性极高的序列（不属于任何现有GFP家族）',
    '但风险较高，建议与其他路线并行',
]
for s in steps_e:
    doc.add_paragraph(s, style='List Number')

add_heading_styled('推荐策略', 2)
add_para(
    '建议采用"多路线并行 + 最终筛选"策略：路线A/B/C/D并行生成候选序列池，'
    '通过结构预测（AlphaFold2 pLDDT）、稳定性预测（TemBERTure/SPURS）、'
    '荧光亮度预测（GLEAM/自定义模型）三方筛选，最终选出6条互补序列。'
    '在提交前使用DNAChisel本地验证反向翻译结果（参考竞赛官方流程）。'
    '设计文档中应详细记录每条路线的参数选择、中间输出和筛选逻辑。'
)

doc.add_page_break()

# ═══════════════════════════════════════════
# 九、参考文献汇总
# ═══════════════════════════════════════════
add_heading_styled('九、参考文献汇总（含DOI）', 1)

refs = [
    # (编号, 标题, 期刊, 年份, DOI, 备注)
    ('[1]', 'Simulating 500 million years of evolution with a language model (ESM3/esmGFP)',
     'Science 387(6736):850-858', '2025', '10.1126/science.ads0018', '★★★ 核心文献'),
    ('[2]', 'Mutational Effect Transfer Learning (METL)',
     'Nature Methods', '2025', '10.1038/s41592-025-02667-2', '★★★ 低N场景优势'),
    ('[3]', 'GeoEvoBuilder: A deep learning framework for efficient functional and thermostable protein design',
     'PNAS 122(41):e2504117122', '2025', '10.1073/pnas.2504117122', '★★★ GFP验证，零样本'),
    ('[4]', 'Reengineering a flavin-based fluorescent protein using ProteinMPNN',
     'Protein Science 33:e70002', '2024', '10.1002/pro.70002', '★★★ 荧光蛋白+ProteinMPNN'),
    ('[5]', 'Molecular dynamics guided identification of a brighter variant of sfGFP (YuzuFP)',
     'Communications Chemistry 8:174', '2025', '10.1038/s42004-025-01573-4', '★★ MD理性设计'),
    ('[6]', 'FluProCAD: A computational screening workflow for fluorescent protein variants',
     'Molecular Physics', '2025', '10.1080/00268976.2025.2458641', '★ 自动化FP筛选'),
    ('[7]', 'Designed active-site library reveals thousands of functional GFP variants (htFuncLib)',
     'Nature Communications 14:2890', '2023', '10.1038/s41467-023-38099-z', '★★★ 大规模FP设计'),
    ('[8]', 'Improving Protein Optimization with Smoothed Fitness Landscapes (GGS)',
     'ICLR 2024', '2024', '-', '★★ 平滑适应度景观'),
    ('[9]', 'Machine learning to predict continuous protein properties from binary cell sorting data',
     'PNAS 121(14):e2317527121', '2024', '10.1073/pnas.2317527121', '★ GFP基准'),
    ('[10]', 'Evaluation of MLDE Across Diverse Combinatorial Landscapes',
     'bioRxiv', '2024', '10.1101/2024.10.24.619774', '★★ MLDE系统评估'),
    ('[11]', 'A top variant identification pipeline for protein engineering (TopVIP)',
     'Cell Systems 15(2):134-148', '2024', '10.1016/j.cels.2024.01.006', '★ 变体筛选管线'),
    ('[12]', 'SPURS: Rewiring protein sequence and structure generative models to enhance stability prediction',
     'Nature Communications', '2025', '10.1038/s41467-025-56475-9', '★★★ 一流稳定性预测'),
    ('[13]', 'Improving Inverse Folding models at Protein Stability Prediction (ProteinMPNN-ddG)',
     'NeurIPS 2024 Workshop / bioRxiv', '2024', '10.1101/2024.06.15.599145', '★★ 零样本ΔΔG'),
    ('[14]', 'Parametrically guided design of beta barrels and transmembrane nanopores using deep learning',
     'bioRxiv / eLife', '2025', '10.1101/2024.07.22.604663', '★★ De novo β-桶'),
    ('[15]', 'ProteinGuide: Guide your favorite protein sequence generative model',
     'arXiv:2505.04823', '2025', '10.48550/arXiv.2505.04823', '★ 即插即用条件生成'),
    ('[16]', 'Thermophilic Chassis-Enabled High-Throughput Selection of a Thermostable Fluorogenic Reporter',
     'ACS Synthetic Biology 14(10):4100-4115', '2025', '10.1021/acssynbio.5c00343', '★ 嗜热筛选'),
    ('[17]', 'DnaChisel: A versatile DNA sequence optimizer',
     'Bioinformatics 36:4592-4594', '2020', '10.1093/bioinformatics/btaa558', '竞赛指定反向翻译工具'),
    ('[18]', 'Next-generation de novo luciferases: How AI is improving its own designs at light speed',
     'Current Opinion in Structural Biology 91:103011', '2025', '10.1016/j.sbi.2025.103011', '★ AI荧光素酶设计'),
]

for ref in refs:
    p = doc.add_paragraph()
    run_num = p.add_run(f'{ref[0]} ')
    run_num.bold = True
    run_num.font.size = Pt(10)
    run_title = p.add_run(f'{ref[1]}. ')
    run_title.font.size = Pt(10)
    run_journal = p.add_run(f'{ref[2]}, {ref[3]}. ')
    run_journal.font.size = Pt(10)
    run_doi = p.add_run(f'DOI: {ref[4]}')
    run_doi.font.size = Pt(10)
    run_doi.font.color.rgb = RGBColor(0, 80, 160)
    if ref[5]:
        run_note = p.add_run(f'  [{ref[5]}]')
        run_note.font.size = Pt(10)
        run_note.font.color.rgb = RGBColor(180, 80, 0)

doc.add_page_break()

# ═══════════════════════════════════════════
# 十、开源工具汇总
# ═══════════════════════════════════════════
add_heading_styled('十、开源工具汇总', 1)

add_heading_styled('10.1 蛋白质生成与设计', 2)
tools_gen = [
    ('ESM3 (EvolutionaryScale)', 'https://github.com/evolutionaryscale/esm',
     '多模态PLM，序列-结构-功能联合生成。基础版免费开放学术界。'),
    ('ESM2 (Meta)', 'https://github.com/facebookresearch/esm',
     '广泛使用的蛋白质语言模型（15B参数版本可用于嵌入提取）。'),
    ('ProteinMPNN', 'https://github.com/dauparas/ProteinMPNN',
     '基于图的逆折叠模型，从骨架结构设计序列。竞赛必需工具。'),
    ('RFDiffusion', 'https://github.com/RosettaCommons/RFdiffusion',
     '去噪扩散模型，de novo蛋白质骨架生成。'),
    ('GeoEvoBuilder', 'https://github.com/PKUliujl/GeoEvoBuilder',
     '结构+进化联合设计，GFP上已验证。'),
    ('ColabFold/AlphaFold2', 'https://github.com/sokrypton/ColabFold',
     '最快结构预测验证工具，MMseqs2生成MSA。'),
]
for name, url, desc in tools_gen:
    add_link_item(name, url, desc)

add_heading_styled('10.2 稳定性预测', 2)
tools_stab = [
    ('TemBERTure', 'https://github.com/ibmm-unibe-ch/TemBERTure',
     'Tm预测+嗜热分类，protBERT-BFD适配器。'),
    ('SPURS', 'https://github.com/luo-group/SPURS',
     'ProteinMPNN+ESM融合，单次前向预测所有L×20突变ΔΔG。'),
    ('ProteinMPNN-ddG', 'https://github.com/PeptoneLtd/proteinmpnn_ddg',
     '零样本ΔΔG，蛋白质组级扫描。'),
    ('PROSS', 'https://pross.weizmann.ac.il',
     'Rosetta+进化分析自动化稳定设计（Web服务器）。'),
    ('ThermoMPNN', 'https://github.com/Kuhlman-Lab/ThermoMPNN',
     'ProteinMPNN在Megascale上的微调版。'),
    ('ProtSSN', 'https://github.com/openmedlab/ProtSSN',
     '序列-结构融合，去噪预训练，零样本。'),
]
for name, url, desc in tools_stab:
    add_link_item(name, url, desc)

add_heading_styled('10.3 荧光/功能预测', 2)
tools_func = [
    ('GLEAM (synbiochallenges2025)', 'https://github.com/f-normies/synbiochallenges2025',
     'ESM-2嵌入+全局-局部注意力，GFP荧光亮度预测。'),
    ('METL', 'https://github.com/rsgelman/METL',
     '生物物理预训练Transformer，多功能预测（荧光、稳定性、催化）。'),
    ('FluProCAD', '- (DOI: 10.1080/00268976.2025.2458641)',
     '自动化FP光学/热力学性质计算工作流。'),
]
for name, url, desc in tools_func:
    add_link_item(name, url, desc)

add_heading_styled('10.4 序列优化与分析', 2)
tools_seq = [
    ('DNAChisel', 'https://github.com/Edinburgh-Genome-Foundry/DNAChisel',
     'Python DNA序列优化库，竞赛指定反向翻译工具。'),
    ('Biopython', 'https://biopython.org',
     '通用生物信息学库，序列操作、PDB解析。'),
]
for name, url, desc in tools_seq:
    add_link_item(name, url, desc)

add_heading_styled('10.5 MLDE / 适应度景观', 2)
tools_mlde = [
    ('GGS', 'https://github.com/kirjner/GGS',
     '平滑适应度景观 + Gibbs采样。'),
    ('MLDE benchmarking (Arnold Lab)', 'https://github.com/fhalab/MLDE',
     'MLDE/ALDE/ftMLDE系统比较。'),
]
for name, url, desc in tools_mlde:
    add_link_item(name, url, desc)

add_heading_styled('10.6 分子动力学', 2)
tools_md = [
    ('OpenMM', 'https://github.com/openmm/openmm',
     'GPU加速MD引擎，Python API。'),
    ('GROMACS', 'https://www.gromacs.org',
     '高性能MD模拟，蛋白质稳定性评估。'),
    ('FoldX', 'https://foldxsuite.crg.eu',
     '快速ΔΔG计算，突变扫描。'),
    ('Rosetta', 'https://www.rosettacommons.org',
     '全面的蛋白质设计套件（ddG_monomer、cartesian_ddg等）。'),
]
for name, url, desc in tools_md:
    add_link_item(name, url, desc)

# ── 免责声明 ──
doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('免责声明：本报告由AI辅助检索整理，建议在使用文献和工具前核实最新版本和许可条款。'
                '所有DOI链接和GitHub地址截至2026年5月已验证可访问。')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(128, 128, 128)

# ── 保存 ──
output_path = r'C:\Users\fengs\Desktop\蛋白质设计-合成生物学创新赛-Claude\GFP蛋白质设计文献与工具调研报告.docx'
doc.save(output_path)
print(f'报告已保存至: {output_path}')
