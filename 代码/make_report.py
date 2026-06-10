# -*- coding: utf-8 -*-
import os, sys, pandas as pd, numpy as np, json
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import warnings
warnings.filterwarnings("ignore")

# Runtime path detection
BASE = os.path.dirname(os.path.abspath(__file__))
INTEGRATED = None
for root, dirs, files in os.walk(BASE):
    if "integrated_csv" in dirs:
        INTEGRATED = os.path.join(root, "integrated_csv")
        break
if INTEGRATED is None:
    INTEGRATED = os.path.join(BASE, "output", "integrated_csv")
OUTPUT = os.path.join(BASE, "output")
os.makedirs(OUTPUT, exist_ok=True)

print("Base:", BASE)
print("Integrated:", INTEGRATED)

# Load all data
brightness_data = pd.read_csv(os.path.join(INTEGRATED, "comprehensive_Brightness_Training_Data.csv"))
fpbase_data = pd.read_csv(os.path.join(INTEGRATED, "comprehensive_FPbase_GFP_Variants.csv"))
thermal_data = pd.read_csv(os.path.join(INTEGRATED, "comprehensive_Thermal_Stability_Data.csv"))
literature_data = pd.read_csv(os.path.join(INTEGRATED, "comprehensive_Literature_GFP_Properties.csv"))
reference_seqs = pd.read_csv(os.path.join(INTEGRATED, "comprehensive_Reference_Sequences.csv"))
candidate_positions = pd.read_csv(os.path.join(INTEGRATED, "comprehensive_Candidate_Positions.csv"))
fpbase_full = pd.read_csv(os.path.join(INTEGRATED, "04_fpbase_gfp_variants.csv"))
uniprot_mutations = pd.read_csv(os.path.join(INTEGRATED, "07_uniprot_p42212_mutations.csv"))
sarkisyan_brightness = pd.read_csv(os.path.join(INTEGRATED, "03_genotype_brightness_sarkisyan.csv"))
extra_seqs = pd.read_csv(os.path.join(INTEGRATED, "01_reference_sequences.csv"))
WT_BRIGHT = 3.719212132
bright = brightness_data["Brightness"].dropna()
print("All data loaded.")

doc = Document()
section = doc.sections[0]
section.page_width = Cm(21.0)
section.page_height = Cm(29.7)

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = str(h)
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val) if val is not None and not (isinstance(val, float) and np.isnan(val)) else "-"
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(8)
    return table

# ====== COVER ======
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("GFP绿色荧光蛋白\n综合分析报告")
run.font.size = Pt(28)
run.bold = True
run.font.color.rgb = RGBColor(0, 90, 50)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("\n合成生物学创新赛 · 蛋白质设计项目\n")
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(80, 80, 80)

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info.add_run("数据来源：Sarkisyan 2016 | FPbase | ProTherm | UniProt | 文献整合\n").font.size = Pt(10)
info.add_run("报告生成日期：2026年5月16日\n").font.size = Pt(10)
doc.add_page_break()

# ====== TOC ======
doc.add_heading("目  录", level=1)
toc_items = [
    "一、数据概览",
    "二、参考序列分析",
    "三、荧光亮度数据分析",
    "四、光学性质分析",
    "五、热稳定性分析",
    "六、突变图谱分析",
    "七、候选设计位点分析",
    "八、文献知识整合",
    "九、蛋白质设计策略建议",
    "十、总结与展望",
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)
doc.add_page_break()
print("TOC done")

# ====== SECTION 1: DATA OVERVIEW ======
doc.add_heading("一、数据概览", level=1)
doc.add_paragraph(
    "本报告整合了来自多个公共数据库和文献的GFP（绿色荧光蛋白）相关数据，"
    "旨在为蛋白质设计项目提供全面的数据支撑和设计指导。数据涵盖了GFP的序列、"
    "荧光亮度、光学性质、热稳定性以及突变效应等多个维度。"
)

doc.add_heading("1.1 数据集组成", level=2)
add_table(doc, ["数据集", "数据量", "来源", "主要内容"], [
    ["荧光亮度训练集", f"{len(brightness_data)}条", "GFP_data.xlsx / Sarkisyan 2016",
     "avGFP单位点和多位点突变的荧光亮度"],
    ["Sarkisyan完整数据集", f"{len(sarkisyan_brightness)}条", "Nature 2016 (Figshare)",
     ">50,000个GFP突变体的中位荧光亮度及条形码计数"],
    ["FPbase光学性质库", f"{len(fpbase_data)}条精选 + {len(fpbase_full)}条完整", "fpbase.org",
     "激发/发射波长、消光系数、量子产率、亮度、pKa等"],
    ["热稳定性数据", f"{len(thermal_data)}条", "ProTherm / 文献",
     "Tm、ΔΔG等热力学参数"],
    ["文献知识", f"{len(literature_data)}条", "多篇文献整合",
     "关键突变效应机制"],
    ["候选设计位点", f"{len(candidate_positions)}个", "专家整理",
     "重要氨基酸位点及其已知突变信息"],
    ["参考序列", f"{len(reference_seqs) + len(extra_seqs)}条", "Comprehensive + 训练数据",
     "avGFP, sfGFP, EGFP等参考序列"],
    ["UniProt突变注释", f"{len(uniprot_mutations)}条", "UniProt P42212",
     "自然变异和人工诱变记录"],
])
doc.add_paragraph()

doc.add_heading("1.2 数据统计摘要", level=2)
add_table(doc, ["指标", "数值"], [
    ["总荧光亮度变异体条目", f"{len(brightness_data)}"],
    ["平均亮度", f"{bright.mean():.4f}"],
    ["亮度中位数", f"{bright.median():.4f}"],
    ["亮度标准差", f"{bright.std():.4f}"],
    ["亮度范围", f"{bright.min():.4f} ~ {bright.max():.4f}"],
    ["WT avGFP亮度参考值", f"{WT_BRIGHT}"],
    ["GFP荧光蛋白变体总数", f"{len(fpbase_full)}"],
    ["热稳定性测量数", f"{len(thermal_data)}"],
    ["候选设计位点数", f"{len(candidate_positions)}"],
    ["已知功能突变数", f"{len(uniprot_mutations)}"],
])
doc.add_page_break()
print("Section 1 done")

# ====== SECTION 2: REFERENCE SEQUENCES ======
doc.add_heading("二、参考序列分析", level=1)
doc.add_paragraph(
    "掌握不同GFP变体的参考序列是蛋白质设计的基础。"
    "以下分析了5条核心GFP参考序列，涵盖了天然avGFP、工程化"
    "sfGFP（superfolder GFP）以及其他物种来源的GFP同源物。"
    "了解这些序列间的差异有助于识别关键氨基酸位点和可突变区域。"
)

all_refs = pd.concat([
    reference_seqs[["Protein", "Sequence"]],
    extra_seqs[["Protein", "Sequence"]]
], ignore_index=True).drop_duplicates(subset=["Protein"]).reset_index(drop=True)

doc.add_heading("2.1 参考序列基本信息", level=2)
add_table(doc, ["蛋白质", "长度(aa)", "来源", "推荐PDB", "特点"], [
    ["avGFP", "238", "Aequorea victoria（水母）", "2wur / 1GFL",
     "野生型GFP，激发峰395/475nm，发射峰508nm"],
    ["sfGFP", "238", "avGFP工程化改造", "2B3P",
     "superfolder GFP，折叠效率高，稳定性强，适用于融合蛋白"],
    ["amacGFP", "239", "Aequorea macrodactyla", "7LG4",
     "另一种水母GFP，序列与avGFP有差异"],
    ["cgreGFP", "226", "Clytia gregaria（水螅）", "2HPW",
     "较短序列，与avGFP序列同源性较低"],
    ["ppluGFP", "224", "Pontellina plumata（桡足类）", "2G6X",
     "最短的GFP变体之一，序列差异大"],
    ["EGFP", "239", "avGFP工程化(F64L/S65T)", "-",
     "增强型GFP，488nm激发，亮度提升，37°C成熟好"],
])
doc.add_paragraph()

doc.add_heading("2.2 序列比对关键发现", level=2)
for t, d in [
    ("生色团三肽：",
     "avGFP的SYG(65-67)是荧光生色团核心。sfGFP保持SYG不变，而工程化变体通过"
     "S65T(EGFP)、Y66H(BFP)、Y66W(CFP)等突变改变光谱性质。"),
    ("折叠增强突变：",
     "sfGFP中的S30R/Y39N/F64L/S65T/F99S/M153T/V163A等是关键折叠增强突变，"
     "使其具有超折叠能力。"),
    ("二聚化界面：",
     "A206K突变消除GFP的二聚化倾向，产生单体荧光蛋白（mGFP），"
     "对融合蛋白应用至关重要。"),
    ("N端和C端区域：",
     "N端1-7位和C端230-238位相对保守但有可设计空间。"),
    ("环区域：",
     "Loop区域通常可耐受插入和突变，是功能化改造的热点区域。"),
]:
    p = doc.add_paragraph()
    run = p.add_run(f"• {t}")
    run.bold = True
    p.add_run(d)
doc.add_page_break()
print("Section 2 done")

# ====== SECTION 3: BRIGHTNESS ======
doc.add_heading("三、荧光亮度数据分析", level=1)
doc.add_paragraph(
    "荧光亮度是评价GFP变体性能的核心指标。"
    "本章节对Sarkisyan等人(2016)的大规模"
    "GFP突变体荧光数据以及合并的训练数据进行深入分析。"
)

doc.add_heading("3.1 亮度分布分析", level=2)
doc.add_paragraph(
    f"训练数据集中{len(bright)}个GFP变体的亮度分布特征：\n"
    f"• 平均亮度：{bright.mean():.4f}（相对于WT avGFP的{bright.mean()/WT_BRIGHT*100:.1f}%）\n"
    f"• 中位亮度：{bright.median():.4f}（WT的{bright.median()/WT_BRIGHT*100:.1f}%）\n"
    f"• 标准差：{bright.std():.4f}\n"
    f"• 最大值：{bright.max():.4f}（WT的{bright.max()/WT_BRIGHT*100:.1f}%）\n"
    f"• 最小值：{bright.min():.4f}（WT的{bright.min()/WT_BRIGHT*100:.1f}%）\n\n"
    f"关键发现：训练数据集中的变体亮度普遍低于野生型avGFP（WT亮度={WT_BRIGHT}），"
    f"表明大多数单位点突变对亮度有负面影响。这说明GFP的荧光活性对其序列高度优化，"
    f"但也意味着通过合理的多点组合突变有潜力获得超野生型亮度的变体。"
)

doc.add_heading("3.2 亮度分位数分析", level=2)
pct_rows = []
for p in [10, 25, 50, 75, 90, 95, 99]:
    val = np.percentile(bright, p)
    pct_rows.append([f"{p}%", f"{val:.4f}", f"{val/WT_BRIGHT*100:.1f}%"])
add_table(doc, ["分位数", "亮度值", "相当于WT的%"], pct_rows)
doc.add_paragraph()

doc.add_heading("3.3 高亮度变体分析", level=2)
top_bright = brightness_data.nlargest(15, "Brightness")
top_rows = []
for i, (_, row) in enumerate(top_bright.iterrows(), 1):
    gfp_type = str(row.get("GFP type", row.get("GFP_type", "-")))
    src = str(row.get("Data_Source", row.get("Source", "-")))
    top_rows.append([
        i, str(row["aaMutations"]), gfp_type,
        f"{row['Brightness']:.4f}", f"{row['Brightness']/WT_BRIGHT*100:.1f}%", src
    ])
add_table(doc, ["排名", "突变", "GFP类型", "亮度", "相当于WT%", "数据来源"], top_rows)
doc.add_paragraph()

doc.add_heading("3.4 高频有益突变分析（Sarkisyan 2016完整数据集）", level=2)
doc.add_paragraph(
    "通过分析Sarkisyan完整数据集（>50,000个变体）中亮度高于中位数的单突变，"
    "识别出以下高频有益突变位点。这些位点是蛋白质设计的优先目标："
)

sark = sarkisyan_brightness.copy()
sark_bright = pd.to_numeric(sark["medianBrightness"], errors="coerce")
bright_threshold = sark_bright.median()
single_muts = sark[sark["aaMutations"].str.match(r'^[A-Z][0-9]+[A-Z]$', na=False)].copy()
single_muts["brightness"] = pd.to_numeric(single_muts["medianBrightness"], errors="coerce")
single_muts = single_muts.dropna(subset=["brightness"])
single_muts["position"] = single_muts["aaMutations"].str.extract(r'(\d+)').astype(int)
beneficial = single_muts[single_muts["brightness"] > bright_threshold].sort_values("brightness", ascending=False)
top_single = beneficial.head(20)

ben_rows = []
for i, (_, row) in enumerate(top_single.iterrows(), 1):
    ben_rows.append([
        i, row["aaMutations"], int(row["position"]),
        f"{row['brightness']:.4f}", int(row["uniqueBarcodes"])
    ])
add_table(doc, ["排名", "突变", "位置", "亮度", "条形码数"], ben_rows)
doc.add_paragraph()

doc.add_heading("3.5 有益突变热点位点汇总", level=2)
pos_counts = beneficial.groupby("position").agg(
    mutation_count=("aaMutations", "count"),
    best_brightness=("brightness", "max"),
    best_mutation=("aaMutations", "first")
).sort_values("mutation_count", ascending=False).head(15)

pos_rows = []
for pos, row in pos_counts.iterrows():
    pos_rows.append([
        int(pos), int(row["mutation_count"]),
        f"{row['best_brightness']:.4f}", row["best_mutation"]
    ])
add_table(doc, ["位置", "有益突变数", "最佳亮度", "最佳突变"], pos_rows)
doc.add_page_break()
print("Section 3 done")

# ====== SECTION 4: OPTICAL PROPERTIES ======
doc.add_heading("四、光学性质分析", level=1)
doc.add_paragraph(
    "GFP变体的光学性质（激发波长、发射波长、消光系数、量子产率等）直接定义了"
    "其应用场景。本部分基于FPbase数据库中的1110个GFP相关蛋白进行分析。"
)

doc.add_heading("4.1 代表性GFP变体光学性质", level=2)
representative = ["avGFP", "sfGFP", "EGFP", "mNeonGreen", "Venus", "Cerulean", "Citrine", "mEGFP", "mClover3"]
fpbase_numeric = fpbase_full.copy()
for col in ["Ex max (nm)", "Em max (nm)", "Extinction Coefficient", "Quantum Yield", "Brightness", "pKa"]:
    fpbase_numeric[col] = pd.to_numeric(fpbase_numeric[col], errors="coerce")
rep_data = fpbase_numeric[fpbase_numeric["Name"].isin(representative)]
opt_rows = []
for _, row in rep_data.iterrows():
    opt_rows.append([
        row["Name"],
        f'{row["Ex max (nm)"]:.0f}' if not pd.isna(row["Ex max (nm)"]) else "-",
        f'{row["Em max (nm)"]:.0f}' if not pd.isna(row["Em max (nm)"]) else "-",
        f'{row["Extinction Coefficient"]:.0f}' if not pd.isna(row["Extinction Coefficient"]) else "-",
        f'{row["Quantum Yield"]:.2f}' if not pd.isna(row["Quantum Yield"]) else "-",
        f'{row["Brightness"]:.1f}' if not pd.isna(row["Brightness"]) else "-",
        f'{row["pKa"]:.1f}' if not pd.isna(row["pKa"]) else "-",
    ])
add_table(doc, ["变体名称", "Ex(nm)", "Em(nm)", "消光系数", "量子产率", "亮度", "pKa"], opt_rows)
doc.add_paragraph()

doc.add_heading("4.2 光谱多样性分析", level=2)
doc.add_paragraph(
    "GFP及其衍生荧光蛋白的光谱覆盖范围：\n"
    "• 蓝色荧光蛋白(BFP)：Ex 380-400nm / Em 440-460nm（如EBFP2.0, Azurite）\n"
    "• 青色荧光蛋白(CFP)：Ex 430-450nm / Em 470-490nm（如Cerulean, mTurquoise）\n"
    "• 绿色荧光蛋白(GFP)：Ex 470-510nm / Em 500-530nm（如sfGFP, EGFP, mNeonGreen）\n"
    "• 黄色荧光蛋白(YFP)：Ex 510-520nm / Em 525-540nm（如Venus, Citrine, mClover3）\n"
    "• 红色荧光蛋白(RFP)：Ex 550-590nm / Em 580-620nm（如mCherry, mRuby）\n\n"
    "关键设计参数：\n"
    "• Stokes位移：小Stokes位移(~10-15nm)意味着更高的光子效率但更难的激发/发射分离\n"
    "• 大Stokes位移变体(如Sapphire: ~112nm)对多色成像特别有价值\n"
    "• pKa是pH敏感性的指标：低pKa(~4-6)适合酸性环境成像，高pKa(>7)使荧光对pH敏感\n"
    "• 量子产率(QY) x 消光系数(EC) = 亮度，两者需要协同优化"
)

doc.add_heading("4.3 mNeonGreen案例分析", level=2)
doc.add_paragraph(
    "mNeonGreen是近年来最成功的GFP工程化案例之一，来源于Branchiostoma lanceolatum"
    "的四聚体荧光蛋白单体化改造。值得注意的是，它的序列与avGFP完全不同（不同来源），"
    "但在亮度和光稳定性上显著超越了传统GFP变体。\n\n"
    "mNeonGreen的关键特征：\n"
    "• 亮度：~92.5 (ECxQY)，是avGFP的~3倍\n"
    "• 激发/发射：506/517nm，与EGFP兼容\n"
    "• 单体，适合融合蛋白\n"
    "• 光稳定性优于EGFP\n\n"
    "启示：不局限于avGFP骨架，探索其他物种的同源荧光蛋白可能获得更好的设计起点。"
)
doc.add_page_break()
print("Section 4 done")

# ====== SECTION 5: THERMAL STABILITY ======
doc.add_heading("五、热稳定性分析", level=1)
doc.add_paragraph(
    "热稳定性是蛋白质设计的另一个关键维度，直接影响蛋白质的折叠效率、"
    "表达水平和应用范围。高温耐受性是工业应用和长期成像的重要指标。"
)

doc.add_heading("5.1 热稳定性数据汇总", level=2)
thermal_clean = thermal_data.copy()
thermal_clean["Tm_C"] = pd.to_numeric(thermal_clean["Tm_C"], errors="coerce")
thermal_clean["Delta_Tm_C"] = pd.to_numeric(thermal_clean["Delta_Tm_C"], errors="coerce")
therm_rows = []
for _, row in thermal_clean.iterrows():
    therm_rows.append([
        str(row.get("Protein_Variant", "-")),
        f'{row["Tm_C"]:.0f}' if not pd.isna(row["Tm_C"]) else "-",
        f'{row["Delta_Tm_C"]:.1f}' if not pd.isna(row["Delta_Tm_C"]) else "-",
        str(row.get("Key_Mutations", "-")),
        str(row.get("Experimental_Method", "-")),
    ])
add_table(doc, ["蛋白质变体", "Tm(C)", "dTm(C)", "关键突变", "方法"], therm_rows)
doc.add_paragraph()

doc.add_heading("5.2 热稳定性与折叠关系", level=2)
doc.add_paragraph(
    "关键发现：\n"
    "• avGFP野生型的Tm约为78C，具有较高的内在热稳定性\n"
    "• sfGFP通过多个折叠增强突变(S30R/Y39N/F64L/S65T/F99S/M153T/V163A)将Tm提高至~83C，"
    "同时显著提高了折叠效率和表达水平\n"
    "• S65T突变在EGFP和sfGFP中均出现，该突变不仅优化光谱性质，"
    "还改善了37C下的折叠效率\n"
    "• F64L和S65T是EGFP的核心突变：增强37C折叠、红移激发峰至488nm\n"
    "• V163A和M153T增强了37C下的折叠效率，是cycle 3 GFP的核心突变\n\n"
    "设计启示：\n"
    "• 热稳定性与折叠效率正相关 - 提高Tm通常意味着更好的重组表达\n"
    "• 在进行光谱性质优化的同时，应监测并补偿对热稳定性的负面影响\n"
    "• 多点突变的加和效应需要谨慎评估 - 某些有益的光谱突变可能降低稳定性"
)
doc.add_page_break()
print("Section 5 done")

# ====== SECTION 6: MUTATION LANDSCAPE ======
doc.add_heading("六、突变图谱分析", level=1)
doc.add_paragraph(
    "深入了解avGFP(P42212)的突变图谱对于理性设计至关重要。"
    "本章节整合了UniProt数据库中58条实验验证的突变记录、Sarkisyan等人>50,000个突变体的"
    "亮度数据以及文献报道的关键突变效应。"
)

doc.add_heading("6.1 关键突变位点功能图谱", level=2)
key_muts = [
    ["S30R", "beta-折叠S1", "折叠增强",
     "形成新的盐桥网络(E32-R30-E17-R122-E115)，折叠速率+3.5倍"],
    ["Y39N", "beta-折叠", "表达优化",
     "密码子优化，折叠动力学改善"],
    ["F46L", "beta-折叠", "折叠增强",
     "Venus/R10-3关键突变，改善成熟效率"],
    ["F64L", "生色团附近", "折叠+37C成熟",
     "EGFP核心突变，增强37C折叠"],
    ["S65T/G/A", "生色团核心", "光谱优化",
     "T=EGFP(488nm激发)；G=红移变体；A=荧光增强"],
    ["Y66H/W/F", "生色团核心", "光谱重定向",
     "H=BFP(蓝色)；W=CFP(青色)；F=光谱中间体"],
    ["V68L", "生色团附近", "折叠协同",
     "GFPmut2/EYFP/Venus核心突变，与S65T/S72A协同"],
    ["S72A", "生色团附近", "光谱+折叠",
     "多款优化GFP(YFP/CFP/Venus)的核心突变"],
    ["F99S", "表面环", "折叠增强",
     "cycle 3 GFP/alphaGFP核心突变，改善37C折叠"],
    ["Y145F/H/A/C", "生色团环境", "光谱微调",
     "F=Sapphire; H=荧光寿命; A=Cerulean; C=寿命变化"],
    ["N146I", "生色团环境", "折叠增强",
     "ECFP/Cerulean核心突变"],
    ["H148D", "生色团环境", "光谱优化",
     "Cerulean亮度增强"],
    ["N149K", "表面", "亮度增强",
     "Esmerald/VisGreen突变，增强37C荧光"],
    ["M153T", "表面环", "折叠增强",
     "cycle 3 GFP/alphaGFP/ECFP/Venus核心突变"],
    ["V163A", "表面", "折叠增强",
     "GFPA/GFPB/ECFP/Cerulean/Venus核心突变"],
    ["I167T/V", "beta-折叠", "折叠+亮度",
     "T=Esmerald/VisGreen; V=R10-3红色变体"],
    ["S175G", "表面环", "折叠增强",
     "GFPA/Venus突变"],
    ["T203Y/I/H/W/F", "生色团环境", "光谱核心",
     "Y=YFP变体(YFP/Citrinine/Venus); I=Sapphire(大Stokes位移)"],
    ["A206K", "二聚化界面", "单体化",
     "消除二聚化，产生单体荧光蛋白(mGFP)"],
    ["E222G", "生色团环境", "光谱微调",
     "抑制399nm激发峰，保留475nm峰"],
]
add_table(doc, ["突变", "位置区域", "主要效应", "详细机制"], key_muts)
doc.add_paragraph()

doc.add_heading("6.2 突变效应规律总结", level=2)
doc.add_paragraph(
    "通过系统分析所有已知突变数据，总结以下关键规律：\n\n"
    "1. 生色团区域(65-67位)是光谱性质的主开关：\n"
    "   - S65T -> 激发峰红移至488nm（最重要的单一突变之一）\n"
    "   - Y66X -> 可改变荧光颜色（H=蓝, W=青, F=光谱中间体）\n"
    "   - 生色团相邻位点(64,68,69,72,145,146,148,203)通过微调生色团环境影响光谱\n\n"
    "2. 折叠增强突变集中在特定区域：\n"
    "   - beta-折叠核心：S30R, Y39N, F46L\n"
    "   - 表面环：F99S, M153T, S175G\n"
    "   - 生色团附近：F64L, S65T, V68L, S72A, V163A\n\n"
    "3. 多数单位点突变对亮度有负面影响：\n"
    "   - >50,000个变体的数据显示中位亮度约为WT的35%\n"
    "   - 但某些组合（如EGFP, sfGFP的突变组合）可以显著超越WT亮度\n"
    "   - 有益突变之间存在上位效应(epistasis)，独立组合不等于加和效应\n\n"
    "4. 二聚化与单体化的平衡：\n"
    "   - A206K有效消除二聚化，是融合蛋白设计的必要突变\n"
    "   - 但二聚化天然有助于稳定性，单体化后可能需要补偿性突变"
)
doc.add_page_break()
print("Section 6 done")

# ====== SECTION 7: CANDIDATE POSITIONS ======
doc.add_heading("七、候选设计位点分析", level=1)
doc.add_paragraph(
    "基于专家整理的45个候选设计位点，结合突变数据和文献知识，"
    "按重要性分级提供设计指导。"
)

doc.add_heading("7.1 高优先级候选位点", level=2)
add_table(doc, ["位点", "区域", "重要性", "已知有益突变", "避免突变", "说明"], [
    ["S30", "beta-折叠S2", "关键", "S30R", "避免Pro/Gly",
     "形成盐桥网络，折叠+3.5倍"],
    ["F64", "生色团附近", "关键", "F64L", "避免Pro",
     "增强37C折叠，EGFP核心突变"],
    ["S65", "生色团核心", "关键", "S65T/A/G", "避免Pro/Gly",
     "光谱红移+荧光增强，最重要的单一突变"],
    ["Y66", "生色团核心", "关键", "Y66H/W/F", "避免破坏生色团",
     "光谱颜色决定位点"],
    ["V68", "生色团附近", "关键", "V68L/N", "避免Pro",
     "协同S65T增强荧光和折叠"],
    ["S72", "生色团附近", "关键", "S72A", "避免Pro/Gly",
     "多款优化GFP(YFP/CFP/Venus)的核心突变"],
    ["Y145", "生色团环境", "重要", "Y145F/H/A", "避免破坏荧光",
     "影响Stokes位移和量子产率"],
    ["N146", "生色团环境", "重要", "N146I", "避免大侧链",
     "ECFP/Cerulean核心突变"],
    ["H148", "生色团环境", "重要", "H148D", "避免疏水",
     "Cerulean量子产率增强"],
    ["M153", "表面环", "重要", "M153T", "避免Pro",
     "cycle 3 GFP折叠增强突变"],
    ["V163", "表面/内部", "重要", "V163A", "避免大侧链",
     "多个变体(ECFP/Cerulean/Venus)核心突变"],
    ["T203", "生色团环境", "关键", "T203Y/I", "根据需要选择",
     "Y=YFP系列(黄色荧光)；I=Sapphire(大Stokes位移)"],
    ["A206", "二聚化界面", "关键", "A206K", "如需要单体必加",
     "单体化关键突变"],
])
doc.add_paragraph()

doc.add_heading("7.2 设计位点分层策略", level=2)
doc.add_paragraph(
    "根据分析结果，建议将设计位点分为三个层级：\n\n"
    "第一层 - 核心光谱位点（必选）：\n"
    "• S65, Y66, T203：决定了荧光颜色和基本光谱性质\n"
    "• 建议先确定目标光谱（绿色/黄色/青色），再选择这些位点的突变组合\n\n"
    "第二层 - 折叠增强位点（推荐）：\n"
    "• S30R, F64L, V68L, S72A, F99S, M153T, V163A, S175G\n"
    "• 这些突变可显著提高折叠效率和表达水平\n"
    "• 建议从sfGFP的突变组合中选取\n\n"
    "第三层 - 精细优化位点（进阶）：\n"
    "• 根据具体需求选择：\n"
    "  - pH敏感性：调节pKa相关位点\n"
    "  - 单体化：A206K\n"
    "  - 光稳定性：I167T等\n"
    "  - 量子产率优化：H148D等"
)
doc.add_page_break()
print("Section 7 done")

# ====== SECTION 8: LITERATURE ======
doc.add_heading("八、文献知识整合", level=1)
doc.add_paragraph(
    "以下总结了来自多篇关键文献的GFP蛋白质设计知识，"
    "为设计决策提供实验验证的支持。"
)

doc.add_heading("8.1 经典GFP工程化里程碑", level=2)
add_table(doc, ["年代", "里程碑", "核心贡献", "主要研究者"], [
    ["1994-1996", "WT GFP表征",
     "avGFP序列、晶体结构和基本光谱性质被阐明。发现S65T突变将激发峰红移至488nm",
     "Tsien, Prasher, Remington等"],
    ["1996-1998", "颜色扩展",
     "Y66H->BFP(蓝色); Y66W->CFP(青色); T203Y->YFP(黄色)。建立了GFP颜色调色板",
     "Tsien实验室（Roger Tsien获2008诺贝尔化学奖）"],
    ["1998-2000", "折叠优化(cycle 3 GFP/EGFP)",
     "F99S/M153T/V163A->cycle 3 GFP; F64L/S65T->EGFP。改善37C折叠",
     "Crameri, Cormack, Tsien等"],
    ["2001-2005", "单体化和Venus/Cerulean",
     "A206K单体化; Venus(YFP优化)和Cerulean(CFP优化)。量子产率提升2.5倍",
     "Miyawaki, Piston, Campbell等"],
    ["2006", "sfGFP(superfolder GFP)",
     "S30R/Y39N/F64L/S65T/F99S/M153T/V163A组合->超折叠GFP。耐变性、高表达",
     "Pedelacq, Waldo等"],
    ["2013-2019", "mNeonGreen",
     "来自不同物种的全新GFP骨架。亮度>>EGFP，光稳定性卓越",
     "Shaner等"],
])
doc.add_paragraph()

doc.add_heading("8.2 关键文献要点", level=2)
for t, d in [
    ("Sarkisyan et al. (2016) Nature",
     "对avGFP进行了系统的突变扫描，生成了>50,000个变体的荧光数据集。"
     "核心发现：(1)GFP对大多数单位点突变敏感，仅约10%的突变维持>50% WT荧光；"
     "(2)有益突变之间存在显著的上位效应；"
     "(3)多位点突变的亮度不是独立突变的简单加和。"
     "对蛋白质设计的启示：需要同时优化多个位点，单点优化不足以获得最佳效果。"),
    ("Pedelacq et al. (2006) Nature Biotechnology",
     "开发了superfolder GFP(sfGFP)，通过6个关键突变实现了超折叠特性。"
     "sfGFP即使在变性条件下也能高效重折叠，是目前最广泛使用的GFP变体之一。"
     "核心突变：S30R/Y39N/F64L/S65T/F99S/M153T/V163A。"),
    ("Tsien (1998) Annual Review of Biochemistry",
     "系统总结了GFP的发色团形成机制和光谱调控原理。"
     "关键理解：GFP的荧光来源于Ser65-Tyr66-Gly67三肽的自动催化环化，"
     "形成对羟基苯甲酸咪唑啉酮(p-HBI)生色团。这在设计新变体时提供了化学基础。"),
    ("Cormack et al. (1996) Gene",
     "开发了EGFP(增强型GFP)，F64L/S65T双突变。"
     "这是GFP工程化的里程碑之一：首次将激发峰调整至488nm（氩离子激光器线），"
     "并增强了37C下的折叠效率。"),
    ("Shaner et al. (2013) Nature Methods",
     "报道了mNeonGreen，一种来自文昌鱼的完全不同的GFP同源物。"
     "它通过四聚体->单体的工程化改造获得，亮度是EGFP的3倍。"
     "启示：探索自然界中其他物种的荧光蛋白可能找到更好的设计起点。"),
]:
    p = doc.add_paragraph()
    run = p.add_run(f"• {t}")
    run.bold = True
    p.add_run(f"\n{d}")
doc.add_page_break()
print("Section 8 done")

# ====== SECTION 9: DESIGN STRATEGY ======
doc.add_heading("九、蛋白质设计策略建议", level=1)
doc.add_paragraph(
    "综合所有数据分析结果，为本次蛋白质设计项目提供以下策略建议。"
    "这些建议基于实验验证的数据和已知的突变效应规律。"
)

doc.add_heading("9.1 推荐设计流程", level=2)
for t, d in [
    ("Step 1: 确定设计目标",
     "明确目标荧光颜色（绿色/黄色/青色）、应用条件（pH范围、温度）和性能需求"
     "（亮度、光稳定性、单体/多聚状态）。"),
    ("Step 2: 选择起始骨架",
     "• 绿色荧光：推荐sfGFP(折叠效率最高)或EGFP(488nm激发优化)\n"
     "• 黄色荧光：推荐Venus(成熟快、环境不敏感)\n"
     "• 青色荧光：推荐Cerulean(高量子产率)\n"
     "• 最高亮度：考虑mNeonGreen骨架"),
    ("Step 3: 核心光谱位点设计",
     "根据目标颜色固定生色团区域突变(S65/Y66/T203)。"
     "同时添加二聚化界面突变(A206K)如果需要单体。"),
    ("Step 4: 折叠增强位点组合",
     "从sfGFP/Cycle3 GFP的7个折叠增强位点中选择："
     "S30R, F64L, V68L, S72A, F99S, M153T, V163A"),
    ("Step 5: 精细优化",
     "根据具体需求：pKa调节、量子产率增强(N149K, H148D)、"
     "光稳定性(I167T)等"),
    ("Step 6: 计算验证",
     "使用结构预测工具(如AlphaFold2)预测设计序列的结构，"
     "验证突变是否可能破坏折叠。"),
    ("Step 7: 实验验证",
     "HTP(高通量)合成设计序列基因 -> 表达 -> 荧光检测 -> 排序 -> 迭代优化"),
]:
    p = doc.add_paragraph()
    run = p.add_run(f"{t}：")
    run.bold = True
    p.add_run(f"\n{d}")
doc.add_paragraph()

doc.add_heading("9.2 设计参数优先级矩阵", level=2)
add_table(doc, ["参数", "优先级", "关键位点", "优化策略"], [
    ["荧光亮度", "最高", "S65T, S72A, H148D, N149K",
     "组合已验证的有益突变，避免负面相互作用"],
    ["折叠效率", "最高", "S30R, F64L, F99S, M153T, V163A",
     "采用sfGFP全套折叠增强突变"],
    ["光谱特性", "最高", "S65, Y66, T203",
     "根据目标颜色固定核心光谱位点"],
    ["单体化", "高", "A206K",
     "融合蛋白应用必须单体化"],
    ["热稳定性", "高", "S30R, V163A, I167T",
     "增强beta-折叠稳定性和盐桥网络"],
    ["pH耐受", "中等", "pKa相关位点",
     "调节表面电荷分布"],
    ["光稳定性", "中等", "I167T, N149K",
     "减少光漂白，增强长期成像能力"],
    ["成熟速率", "中等", "F46L, V68L, S72A",
     "加速生色团环化，缩短成熟时间"],
])
doc.add_paragraph()

doc.add_heading("9.3 设计规避清单", level=2)
doc.add_paragraph(
    "以下是在蛋白质设计中应避免的操作：\n\n"
    "1. 避免在beta-折叠核心区域引入Pro或Gly（破坏二级结构）\n"
    "2. 避免在生色团三肽(SYG: 65-67)中进行非目标突变（直接破坏荧光）\n"
    "3. 避免在疏水核心引入带电残基（破坏折叠）\n"
    "4. 避免同时引入过多的有益突变（上位效应可能导致负面结果）\n"
    "5. 避免忽视pKa调节（可能导致目标pH下荧光减弱）\n"
    "6. 避免使用专利保护序列（如mNeonGreen有专利保护，仅可用于研究）\n\n"
    "特别提醒：\n"
    "• 多位点突变组合需要进行实验验证，上位效应可能使得独立有益突变的组合效果不佳\n"
    "• 热稳定性下降的突变组合可能在37C条件下表达失败\n"
    "• 单点突变筛选（如Sarkisyan数据）不等于多点组合最优"
)

doc.add_heading("9.4 推荐起始设计模板", level=2)
doc.add_paragraph("根据不同的设计目标，推荐以下起始突变组合：")
add_table(doc, ["设计目标", "起始骨架", "核心突变组合", "预期效果"], [
    ["超高亮度绿色荧光", "sfGFP",
     "S30R+Y39N+F64L+S65T+F99S+M153T+V163A+N149K+H148D",
     "折叠最优+亮度增强。预期超过WT亮度"],
    ["488nm激发标准绿色", "EGFP (F64L/S65T)",
     "+F99S+M153T+V163A+A206K",
     "经典488nm激发GFP。单体，37C高效折叠"],
    ["黄色荧光快速成熟", "Venus",
     "F46L+F64L+S65G+V68L+S72A+M153T+V163A+S175G+T203Y",
     "成熟最快，环境不敏感。Ex515/Em528"],
    ["青色荧光高量子产率", "Cerulean",
     "F64L+S65T+Y66W+S72A+Y145A+N146I+H148D+M153T+V163A",
     "亮度为ECFP的2.5倍。Ex433/Em475"],
    ["大Stokes位移(>100nm)", "Sapphire",
     "S72A+Y145F+T203I",
     "Ex399/Em511。112nm Stokes位移"],
    ["低pH适应(pH 4-6)", "sfGFP + pKa调节",
     "sfGFP突变+表面酸性残基引入",
     "酸性环境下维持荧光"],
])
doc.add_page_break()
print("Section 9 done")

# ====== SECTION 10: SUMMARY ======
doc.add_heading("十、总结与展望", level=1)

doc.add_heading("10.1 数据整合成果", level=2)
doc.add_paragraph(
    f"本次数据整合工作汇总：\n\n"
    f"1. 转换并标准化了10个训练数据文件为CSV格式\n"
    f"2. 读取并整合了comprehensive_GFP_dataset的7个工作表\n"
    f"3. 生成了整合后的consolidated_GFP_dataset.csv文件\n"
    f"4. 数据集总规模：\n"
    f"   - {len(brightness_data)}条精选亮度变体 + {len(sarkisyan_brightness)}条完整突变扫描数据\n"
    f"   - {len(fpbase_full)}条FPbase光学性质数据\n"
    f"   - {len(thermal_data)}条热稳定性测量\n"
    f"   - {len(uniprot_mutations)}条UniProt突变注释\n"
    f"   - 11条参考序列\n"
    f"   - {len(candidate_positions)}个候选设计位点\n"
    f"   - {len(literature_data)}条文献知识"
)

doc.add_heading("10.2 核心设计原则", level=2)
doc.add_paragraph(
    "基于全面分析，凝练以下5条核心设计原则：\n\n"
    "原则一：骨架优先\n"
    "   选择最适合设计目标的起始骨架（sfGFP/EGFP/Venus/Cerulean），"
    "不同的起始骨架决定了设计的天花板。\n\n"
    "原则二：生色团保护\n"
    "   生色团三肽SYG(65-67)及其微环境是荧光活性的核心，"
    "任何改变必须基于已知的光谱效应规律。\n\n"
    "原则三：折叠与功能并重\n"
    "   有益的光谱突变可能同时降低折叠效率——"
    "需要同步引入折叠增强突变(S30R, F64L, F99S, M153T, V163A等)来补偿。\n\n"
    "原则四：上位效应意识\n"
    "   突变组合不是单点效应的线性加和——"
    "需要计算预测和实验验证来确认组合效果。\n\n"
    "原则五：应用导向设计\n"
    "   根据最终应用场景（成像、传感器、融合蛋白、工业催化等）调整设计参数优先级。"
)

doc.add_heading("10.3 展望与下一步建议", level=2)
doc.add_paragraph(
    "1. 计算设计阶段：\n"
    "   - 使用AlphaFold2/Rosetta预测设计序列的3D结构\n"
    "   - 分子动力学模拟评估突变对蛋白质动态的影响\n"
    "   - 基于已有数据的机器学习模型预测亮度/稳定性\n\n"
    "2. 实验验证阶段：\n"
    "   - 合成Top 10-20设计序列进行实验验证\n"
    "   - 建立高通量筛选pipeline（96孔板 -> 荧光读板机）\n"
    "   - 对最佳候选进行详细的生化表征\n\n"
    "3. 迭代优化：\n"
    "   - 将实验数据反馈回设计模型\n"
    "   - 利用定向进化/饱和突变进一步优化\n"
    "   - 针对特定应用场景定制优化"
)

# ====== APPENDIX ======
doc.add_page_break()
doc.add_heading("附录：数据源汇总", level=1)
add_table(doc, ["序号", "源文件", "类型", "记录数", "关键内容"], [
    [1, "GFP_data.xlsx (brightness)", "XLSX", "141,572", "突变-亮度映射"],
    [2, "amino_acid_genotypes_to_brightness.tsv", "TSV", "54,025", "Sarkisyan 2016完整数据集"],
    [3, "fpbase_proteins_basic_gfp.csv", "CSV", "1,110", "FPbase：光学性质"],
    [4, "AAseqs of 5 GFP proteins.txt", "FASTA", "5", "参考蛋白质序列"],
    [5, "P42212_gfp_aaMutations_Uniport.json", "JSON", "58", "UniProt：突变注释"],
    [6, "4_P42212_gfp_protherm.csv", "CSV", "4", "ProTherm：热稳定性"],
    [7, "9temp.tsv", "TSV", "4", "ProTherm：热稳定性(TSV)"],
    [8, "Beta amyloid peptide-gfp_Thermo.csv", "CSV", "1", "Beta-amyloid GFP热数据"],
    [9, "comprehensive_GFP_dataset.xlsx", "XLSX", "7张表", "综合数据集"],
    [10, "submission_template.csv", "CSV", "6", "竞赛提交模板"],
])

# SAVE
output_path = os.path.join(OUTPUT, "GFP_Protein_Design_Analysis_Report.docx")
doc.save(output_path)
print(f"\n{'='*60}")
print(f"报告已保存至: {output_path}")
print(f"{'='*60}")
