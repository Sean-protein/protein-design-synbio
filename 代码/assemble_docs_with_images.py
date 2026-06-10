# -*- coding: utf-8 -*-
"""
组装学习文档——将生成的配图插入四个专题文档
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

base_dir = r'C:\Users\fengs\Desktop\蛋白质设计-合成生物学创新赛-Claude'
img_dir = os.path.join(base_dir, 'images')

def add_image_to_doc(doc, img_name, width=5.5, caption=""):
    """Add an image with centered caption to the document."""
    img_path = os.path.join(img_dir, img_name)
    if os.path.exists(img_path):
        doc.add_paragraph()  # spacer
        last_para = doc.paragraphs[-1]
        run = last_para.add_run()
        run.add_picture(img_path, width=Inches(width))
        last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if caption:
            cap = doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cap.add_run(caption)
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(100, 100, 100)
        return True
    return False


def h(doc, text, level=1):
    """Global helper: add a colored heading."""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0, 51, 102)


def add_p(doc, text, bold=False, size=11):
    """Global helper: add a paragraph."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)


def bullet(doc, text):
    """Global helper: add a bullet point."""
    doc.add_paragraph(text, style='List Bullet')


def rebuild_doc1_with_images():
    """专题一：总体学习路径+路线图+时间线"""
    doc = Document()

    # Style
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(11)

    # Cover
    for _ in range(4):
        doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('蛋白质设计入门学习与实践\n——专题一：总体学习路径与路线图')
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0, 51, 102)
    doc.add_paragraph()
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run('从零基础到合成生物学创新赛蛋白质设计赛道')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0, 80, 130)
    doc.add_page_break()

    # Section 1: 学习路线图
    h(doc, '一、学习路线图总览', 1)
    add_p(doc, '以下路线图直观展示了从零基础到竞赛准备的完整学习路径。四个阶段递进式推进，建议按顺序学习。')
    add_image_to_doc(doc, 'learning_roadmap.png', 5.5, '图1: GFP蛋白质设计学习路线图——四阶段递进式知识构建')

    add_image_to_doc(doc, 'timeline.png', 5.5, '图2: 两周学习时间线推荐——每日学习主题安排')

    h(doc, '1.1 第一阶段：GFP基础理论（2-3天）', 2)
    items = [
        'GFP的结构：β-桶、中心α螺旋、发色团的三维位置关系',
        '发色团自催化成熟三步骤：环化→脱水→氧化，关键催化残基',
        '荧光光物理：吸收、激发态质子转移（ESPT）、发射的完整过程',
        '发色团质子化状态：A态（中性, 395nm）vs B态（阴离子, 488nm），pKa的概念',
        '关键残基功能：Tyr66, Gly67, Arg96, Glu222, Ser205, His148',
        '用手在Mol*中查看关键PDB结构：1GFL, 1EMA, 2B3P, 8q79',
    ]
    for item in items:
        bullet(doc, item)

    h(doc, '1.2 第二阶段：AI蛋白质设计工具入门（3-5天）', 2)
    items = [
        'ESM-2/ESM3：蛋白质语言模型——序列嵌入提取、零样本突变效应预测',
        'ProteinMPNN：逆折叠模型——给定结构设计氨基酸序列',
        'AlphaFold2/ESMFold：结构预测——验证设计序列是否可折叠（pLDDT, PAE, RMSD）',
        'RFdiffusion（进阶）：扩散模型——从零生成蛋白质骨架结构',
        '所有工具均可在Google Colab免费GPU上运行——无需本地硬件',
    ]
    for item in items:
        bullet(doc, item)

    h(doc, '1.3 第三阶段：稳定性与功能预测（2-3天）', 2)
    items = [
        '理解GFP适应度景观：75%单突变降低荧光，30%多突变有负上位效应',
        'Sigmoid截断函数模型：为什么多个中性突变组合会导致荧光灾难性丧失',
        '热稳定性评估工具：TemBERTure（Tm预测）、SPURS（ΔΔG预测）、ProteinMPNN-ddG',
        'Rosetta/FoldX ΔΔG计算：精确预测单突变对稳定性的影响',
        'PROSS Web服务器：自动化蛋白质稳定性设计（零编程，在线使用）',
    ]
    for item in items:
        bullet(doc, item)

    h(doc, '1.4 第四阶段：竞赛设计策略与实践（3-5天）', 2)
    items = [
        '深入理解评分规则：双能乘积制 = (F_initial/F_WT) × (F_final/F_initial)',
        '基准蛋白sfGFP的6个关键突变及其协同稳定化机制',
        '设计策略组合：表面工程(TGP风格) + 核心堆积优化 + 静电网络增强 + 发色团环境加固',
        '管线搭建：ESM嵌入 → ML模型训练 → 候选序列生成 → 稳定性筛选 → 结构验证',
        '文档撰写要点：完整算法管线描述、Agent逻辑树、关键执行日志、复现说明',
    ]
    for item in items:
        bullet(doc, item)

    doc.add_page_break()

    # Section 2: 学习资源导航
    h(doc, '二、学习资源导航（所有链接可追溯）', 1)
    add_p(doc, '以下资源已经过验证，可在浏览器中直接打开：')

    h(doc, '在线课程与视频', 2)
    resources = [
        ('国家智慧教育平台：AI驱动的蛋白质智能设计',
         'https://cetc.nwu.edu.cn/info/1151/2836.htm',
         '浙江大学于浩然研究员, 涵盖MLDE、ProteinMPNN、GFP案例'),
        ('B站：AI驱动蛋白质设计全流程 (BindCraft/RFdiffusion/ProteinMPNN/AF2)',
         'https://www.bilibili.com/video/BV1RVopBwEsL/', '中文全流程实操'),
        ('B站：ProteinMPNN定向进化改造蛋白质实操',
         'https://www.bilibili.com/video/BV1hN6HYnEPR/', '使用ProteinMPNN进行定向进化的完整操作'),
        ('AAAI 2025 Tutorial: AI for Protein Design (英文)',
         'https://deepgraphlearning.github.io/ProteinTutorial_AAAI2025/', '学术级教程，涵盖全部主要工具'),
        ('Rosetta Commons YouTube: MPNN讲座', 'https://www.youtube.com/watch?v=6z4XmUAwdNA', '理论深度讲解'),
        ('Rosetta Commons YouTube: RFdiffusion讲座', 'https://www.youtube.com/watch?v=OEnY2yA3jy8', '扩散模型原理与实操'),
        ('Rosetta Commons YouTube: AlphaFold讲座', 'https://www.youtube.com/watch?v=SVrn8_8aKO8', '结构预测用于设计'),
    ]
    for title, url, desc in resources:
        p = doc.add_paragraph()
        run = p.add_run(f"■ {title}")
        run.bold = True
        run.font.size = Pt(10)
        run = p.add_run(f"\n  {url}")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0, 80, 160)
        run = p.add_run(f"\n  {desc}")
        run.font.size = Pt(9)

    h(doc, 'Colab Notebook（零安装，浏览器运行）', 2)
    colab_resources = [
        ('ColabDesign (sokrypton)', 'https://github.com/sokrypton/ColabDesign',
         'ProteinMPNN + RFdiffusion + AfDesign一站式Colab，公认最佳入门工具'),
        ('ProteinMPNN官方QuickDemo', 'https://github.com/dauparas/ProteinMPNN',
         'colab_notebooks/quickdemo.ipynb'),
        ('DL4Proteins Notebooks (Gray Lab)', 'https://github.com/Graylab/DL4Proteins-notebooks',
         '完整教学系列，第9章：RFdiffusion→ProteinMPNN→AlphaFold全管线'),
        ('ProteinMPNN-ddG Colab', 'https://github.com/PeptoneLtd/proteinmpnn_ddg',
         '零样本ΔΔG预测，用于快速评估突变稳定性'),
    ]
    for title, url, desc in colab_resources:
        p = doc.add_paragraph()
        run = p.add_run(f"■ {title}")
        run.bold = True
        run.font.size = Pt(10)
        run = p.add_run(f"\n  {url}")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0, 80, 160)
        run = p.add_run(f"\n  {desc}")
        run.font.size = Pt(9)

    h(doc, '关键文献（必读10篇）', 2)
    papers = [
        ('Tsien RY (1998) "The Green Fluorescent Protein" — Annu Rev Biochem', 'DOI: 10.1146/annurev.biochem.67.1.509', 'GFP领域奠基性综述'),
        ('Pedelacq JD et al. (2006) "Superfolder GFP" — Nat Biotechnol', 'DOI: 10.1038/nbt1172', 'sfGFP原始文献（竞赛基准蛋白）'),
        ('Sarkisyan KS et al. (2016) "Local fitness landscape of GFP" — Nature', 'DOI: 10.1038/nature17995', 'GFP突变耐受和上位效应核心文献'),
        ('Hirano M et al. (2022) "StayGold" — Nat Biotechnol', 'DOI: 10.1038/s41587-022-01278-2', '超高光稳定性GFP'),
        ('Zhang H et al. (2024) "mBaoJin" — Nat Methods', 'DOI: 10.1038/s41592-024-02203-y', '最新单体超亮GFP'),
        ('Rives A et al. (2025) "ESM3/esmGFP" — Science', 'DOI: 10.1126/science.ads0018', 'AI生成全新GFP'),
        ('Close DW et al. (2015) "TGP" — Proteins', 'DOI: 10.1002/prot.24699', '结构导向表面工程创造热稳定性'),
        ('Gonzalez Somermeyer et al. (2022) — eLife', 'DOI: 10.7554/eLife.75842', '多同源GFP适应度景观比较'),
        ('Church GM et al. (2025) — Nat Rev Bioeng', 'DOI: 10.1038/s44222-025-00349-8', 'AI蛋白质设计路线图综述'),
        ('Nikolaev G et al. (2024) — Protein Sci', 'DOI: 10.1002/pro.70002', 'ProteinMPNN设计荧光蛋白验证'),
    ]
    for title, doi, desc in papers:
        p = doc.add_paragraph()
        run = p.add_run(f"■ {title}")
        run.bold = True
        run.font.size = Pt(10)
        run = p.add_run(f"\n  {doi}  [{desc}]")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(128, 128, 128)

    doc.add_page_break()

    # Section 3: Environment Setup
    h(doc, '三、软件环境配置指南', 1)
    add_p(doc, '方案A（推荐初学者）：纯云端', bold=True)
    bullet(doc, 'Google Colab：运行所有蛋白质设计Notebook（免费GPU）')
    bullet(doc, 'HuggingFace Spaces：在线运行模型Demo')
    bullet(doc, 'PROSS Web Server (https://pross.weizmann.ac.il)：在线稳定性设计')
    bullet(doc, 'AlphaFold3 Server (https://alphafoldserver.com)：在线结构预测')
    bullet(doc, 'RCSB PDB Mol* Viewer (https://www.rcsb.org/3d-view)：在线三维结构查看')

    add_p(doc, '方案B（有编程基础）：本地轻量环境', bold=True)
    bullet(doc, 'Python 3.9+, pip install fair-esm transformers biopython dnachisel')
    bullet(doc, 'Jupyter Lab：本地运行Notebook')
    bullet(doc, 'VS Code + Copilot：AI辅助编程')

    add_p(doc, '方案C（竞赛级）：GPU加速环境', bold=True)
    bullet(doc, 'CUDA 12.x + PyTorch 2.x + NVIDIA GPU (RTX 3090/A100)')
    bullet(doc, 'ESM-2 3B嵌入提取（本地GPU加速）')
    bullet(doc, 'ColabFold本地化部署（大规模结构验证）')

    doc.add_page_break()

    # Section 4: Quick Reference
    h(doc, '四、快速参考卡', 1)
    add_p(doc, 'GFP核心参数速记：', bold=True)
    bullet(doc, '序列长度: 238 aa (avGFP/wtGFP), 分子量: ~27 kDa')
    bullet(doc, '发色团三肽: Ser/Thr65-Tyr66-Gly67 (XZG基序)')
    bullet(doc, 'AvGFP: 激发395/475nm, 发射508nm; EGFP: 激发488nm, 发射508nm')
    bullet(doc, '亮度 = ε × Φ (消光系数 × 量子产率)')
    bullet(doc, 'sfGFP: ε≈58,500, Φ≈0.65, pKa≈5.4')
    bullet(doc, 'mBaoJin: ε≈90,000, Φ≈0.75, pKa≈4.37 (比EGFP光稳定15倍)')

    add_p(doc, '关键PDB ID速记：', bold=True)
    bullet(doc, '1GFL — wtGFP (1.9Å) | 1EMA — EGFP | 2B3P — sfGFP | 8q79 — mBaoJin')

    add_p(doc, '竞赛提交检查清单：', bold=True)
    bullet(doc, '[ ] 6条序列, 220-250 aa, 以M开头, 仅20种标准氨基酸')
    bullet(doc, '[ ] 与Exclusion_List.csv逐一比对, 无完全一致匹配')
    bullet(doc, '[ ] CSV文件: Team_Name, Seq_ID, Sequence 三列')
    bullet(doc, '[ ] PDF设计文档: 完整管线+Agent逻辑树+执行日志')
    bullet(doc, '[ ] GitHub/开源仓库URL: 含README.md')
    bullet(doc, '[ ] 每提交一条序列前再次确认格式正确')

    # Save
    output_path = os.path.join(base_dir, '学习文档_专题一_总体学习路径与路线图（含配图）.docx')
    doc.save(output_path)
    print(f'  保存: {output_path}')


def rebuild_doc2_with_images():
    """专题二：GFP基础 + 发色团互作网络图"""
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(11)

    # (Simple version - key content + images)
    # ... content similar to build_doc2_gfp_basics but with images from existing mechanism report

    # Since build_doc2 already has extensive content, let's just note that images
    # from the mechanism report (fig1-fig6) should be viewed alongside this doc
    # and embed the new chromophore network image

    pass  # Content already rich from build_doc2, images from mechanism report available


def rebuild_doc3_with_images():
    """专题三：AI工具 + 管线流程图"""
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(11)

    # Cover
    for _ in range(4):
        doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('专题三：AI蛋白质设计工具与实战\n——ESM、ProteinMPNN、AlphaFold等工具的实操指南')
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0, 51, 102)
    doc.add_paragraph()
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run('从零开始使用AI工具进行蛋白质序列设计与评估')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0, 80, 130)
    doc.add_page_break()

    h(doc, '一、AI蛋白质设计工具生态系统', 1)
    add_p(doc, '2024-2025年，AI蛋白质设计领域呈现"四大类工具协同工作"的格局。下图展示了标准设计管线的完整流程：')
    add_image_to_doc(doc, 'ai_pipeline.png', 5.5, '图1: AI蛋白质设计标准管线——从PDB结构到Top 6候选序列的完整工作流')

    add_p(doc, '工具分类：')
    bullet(doc, '蛋白质语言模型（PLM）：ESM-2/ESM3, ProtBERT, SaProt —— 序列特征提取、零样本突变预测、序列生成')
    bullet(doc, '逆折叠模型：ProteinMPNN, ESM-IF —— 给定骨架结构，设计氨基酸序列')
    bullet(doc, '结构预测模型：AlphaFold2/3, ESMFold, Chai-1 —— 验证设计序列是否折叠为目标结构')
    bullet(doc, '骨架生成模型（扩散）：RFdiffusion, FrameBuilder —— 从头生成蛋白质骨架')

    h(doc, '二、ProteinMPNN快速上手（5分钟）', 1)
    add_p(doc, '最快上手方式——无需安装任何软件：', bold=True)
    bullet(doc, '1. 打开 https://github.com/dauparas/ProteinMPNN')
    bullet(doc, '2. 找到 colab_notebooks/quickdemo.ipynb → "Open in Colab"')
    bullet(doc, '3. 上传PDB文件（如2B3P——sfGFP的结构）')
    bullet(doc, '4. 运行所有单元格 → 获得设计序列')
    bullet(doc, '5. 尝试修改温度参数(T=0.1保守→T=1.0多样)观察序列多样性变化')

    add_p(doc, '设计GFP变体的关键约束（确保荧光功能）：', bold=True)
    bullet(doc, '必须固定发色团三肽（Tyr66-Gly67-Gly68 / sfGFP编号）——不可替换')
    bullet(doc, '强烈建议固定Arg96（催化环化）和Glu222（质子转移）')
    bullet(doc, '考虑固定His148和Thr203（桶盖残基，保护发色团不受溶剂淬灭）')
    bullet(doc, '允许其余β-桶表面和核心残基自由设计 → 探索序列空间')

    add_p(doc, 'ProteinMPNN设计GFP的Python代码框架：', bold=True)
    code = """# 在Colab中运行
import sys; sys.path.append('/content/ProteinMPNN')
from protein_mpnn_utils import *

# 固定发色团和催化残基
fixed_positions = {"A": [65, 66, 67, 96, 222]}  # 1-based residue numbers
temperature = 0.3  # 较低温度→保守设计
num_sequences = 100  # 生成候选数

# 运行ProteinMPNN生成序列
# (完整代码参见官方quickdemo.ipynb)"""
    cp = doc.add_paragraph()
    run = cp.add_run(code)
    run.font.size = Pt(9)
    run.font.name = 'Courier New'

    h(doc, '三、ColabFold结构验证', 1)
    bullet(doc, 'GitHub: https://github.com/sokrypton/ColabFold')
    bullet(doc, '快速模式 (colabfold_batch)：适合批量验证设计序列')
    bullet(doc, '质量判断标准：pLDDT > 80 (高置信度), pTM > 0.7 (整体折叠好), Cα RMSD (与2B3P) < 3Å')
    bullet(doc, 'pLDDT（predicted LDDT）：0-100，>90极高置信度，<50可能无序/错误折叠')
    bullet(doc, 'PAE（Predicted Aligned Error）：低值表示残基间距离预测准确')

    h(doc, '四、ESM-2/ESM3序列特征提取', 1)
    bullet(doc, '提取序列嵌入作为ML模型输入特征 → 用于亮度/稳定性回归预测')
    bullet(doc, 'ESM-1v零样本突变效应预测：计算突变前后的log-likelihood变化')
    bullet(doc, 'ESM-2嵌入可用于t-SNE/UMAP可视化序列空间，发现高亮度"簇"')
    bullet(doc, 'ESM3 (2025)：支持多模态条件生成（序列+结构+功能prompt联合引导）')

    h(doc, '五、在线工具和数据库', 1)
    resources = [
        ('ColabDesign', 'https://github.com/sokrypton/ColabDesign', 'ProteinMPNN/RFdiffusion Colab一站式'),
        ('ESM GitHub', 'https://github.com/facebookresearch/esm', 'ESM-2预训练模型'),
        ('ESM3 GitHub', 'https://github.com/evolutionaryscale/esm', 'ESM3多模态模型'),
        ('RCSB PDB for 2B3P', 'https://www.rcsb.org/structure/2B3P', 'sfGFP三维结构'),
        ('FPbase', 'https://www.fpbase.org/', '荧光蛋白数据库'),
    ]
    for title, url, desc in resources:
        p = doc.add_paragraph()
        run = p.add_run(f"■ {title}")
        run.bold = True
        run.font.size = Pt(10)
        run = p.add_run(f"\n  {url}")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0, 80, 160)
        run = p.add_run(f"\n  {desc}")
        run.font.size = Pt(9)

    output_path = os.path.join(base_dir, '学习文档_专题三_AI蛋白质设计工具与实战（含配图）.docx')
    doc.save(output_path)
    print(f'  保存: {output_path}')


def rebuild_doc4_with_images():
    """专题四：稳定性 + 策略图 + 适应度景观"""
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(11)

    # Cover
    for _ in range(4):
        doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('专题四：蛋白质稳定性与热稳定性工程\n——ΔΔG预测、表面工程与稳定性设计策略')
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0, 51, 102)
    doc.add_paragraph()
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run('掌握提升GFP热稳定性的六大工程策略')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0, 80, 130)
    doc.add_page_break()

    h(doc, '一、GFP热稳定性提升——六大工程策略', 1)
    add_p(doc, '下图为GFP热稳定性提升的六大互补工程策略总览：')
    add_image_to_doc(doc, 'stability_strategies.png', 5.5, '图1: GFP热稳定性提升六大工程策略——从核心堆积到表面电荷的全方位设计方法')

    h(doc, '二、理解GFP适应度景观——避免设计中的"陷阱"', 1)
    add_p(doc, 'Sarkisyan等人(Nature 2016)系统揭示了GFP的突变耐受规律。这是所有参赛者必须理解的核心概念：')
    add_image_to_doc(doc, 'fitness_landscape.png', 5.5, '图2: GFP适应度景观——Sigmoid截断模型与多突变荧光丧失比例 (数据来源: Sarkisyan et al., Nature 2016)')

    add_p(doc, '核心发现与直接启示：', bold=True)
    bullet(doc, '75%单突变降低荧光 → GFP对突变高度敏感 → 设计时需要全局稳定性评估')
    bullet(doc, '30%多突变有负上位效应 → 不能简单叠加"好的"单点突变')
    bullet(doc, 'Sigmoid截断函数模型(R²=0.85) → 稳定性低于阈值则荧光急剧丧失')
    bullet(doc, '策略：使用超稳定骨架(sfGFP)作为"突变缓冲器" + 迭代验证而非一次性引入所有突变')

    h(doc, '三、发色团微环境关键残基互作网络', 1)
    add_p(doc, '了解发色团周围的关键残基网络是进行精细设计和突变筛选的前提：')
    add_image_to_doc(doc, 'chromophore_network.png', 5.0, '图3: GFP发色团微环境关键残基网络——固定残基(红色标注) + 可优化残基(蓝色标注)')

    add_p(doc, '设计原则：', bold=True)
    bullet(doc, '红色标注残基 → 设计时必须固定，不可突变（T66-G67-G68, R96, E222）')
    bullet(doc, '蓝色标注残基 → 可以优化，是提升稳定性和亮度的主要操作空间')
    bullet(doc, "W1保守水分子 → 参与关键H-bond网络，保持其结合环境")
    bullet(doc, 'Cl⁻ 结合口袋(mBaoJin) → 氯离子结合提升光稳定性和亮度')

    h(doc, '四、已知的GFP稳定化突变汇总', 1)
    add_p(doc, '以下突变已在实验中被验证可提升GFP的稳定性或亮度：')

    mut_data = [
        ('S30R', 'sfGFP', '五元静电网络 (E32-R30-E17-R122-E115)', '折叠+稳定性显著提升', '不可逆改'),
        ('Y145F', 'sfGFP', '移除OH改善疏水核心堆积', '+3-4°C Tm', '可独立迁移'),
        ('H193Y', 'TGP', 'π-堆积稳定发色团', '99°C仍维持荧光', '强烈推荐'),
        ('H148S', 'YuzuFP', '改善与发色团和W1的H-bond', '亮度↑1.5倍+光稳定性↑3倍', 'MD引导设计'),
        ('Q66E', 'TGP', '改善发色团环境', '化学稳定性和pH稳定性', '改善耐酸性'),
        ('C165Y', 'mBaoJin', '单体化+局部结构稳定', '光稳定性+单体性', '减少二聚化'),
        ('K/R→E', 'TGP', '表面正电荷→谷氨酸', '90°C半衰期2.2倍(175→380min)', '通用表面策略'),
    ]
    table = doc.add_table(rows=len(mut_data)+1, cols=5, style='Light Grid Accent 1')
    for i, h_text in enumerate(['突变', '来源', '机制', '效果', '可迁移性']):
        table.rows[0].cells[i].text = h_text
        for ph in table.rows[0].cells[i].paragraphs:
            for r in ph.runs:
                r.bold = True
                r.font.size = Pt(8)
    for i, row in enumerate(mut_data):
        for j, val in enumerate(row):
            table.rows[i+1].cells[j].text = val
            for ph in table.rows[i+1].cells[j].paragraphs:
                for r in ph.runs:
                    r.font.size = Pt(8)

    h(doc, '五、ΔΔG预测的计算方法速查', 1)
    add_p(doc, '四种主要方法及其适用场景：')
    ddg_data = [
        ('Rosetta cartesian_ddg', '力场+骨架柔性', '精确(ρ~0.69)', '中(30-60min/变体)', '结构导向精确设计'),
        ('FoldX', '经验力场', '中等(ρ~0.5-0.7)', '快(1-5min/变体)', '快速扫描大量突变'),
        ('SPURS', 'ProteinMPNN+ESM', '最高(ρ~0.83)', '快(GPU)', '对全部L×20单突变同时预测'),
        ('ProteinMPNN-ddG', '零样本ML', '高(ρ~0.77)', '快(GPU)', '高通量ΔΔG预筛选'),
    ]
    table2 = doc.add_table(rows=len(ddg_data)+1, cols=5, style='Light Grid Accent 1')
    for i, h_text in enumerate(['方法', '原理', '精度', '速度', '场景']):
        table2.rows[0].cells[i].text = h_text
        for ph in table2.rows[0].cells[i].paragraphs:
            for r in ph.runs:
                r.bold = True
                r.font.size = Pt(9)
    for i, row in enumerate(ddg_data):
        for j, val in enumerate(row):
            table2.rows[i+1].cells[j].text = val
            for ph in table2.rows[i+1].cells[j].paragraphs:
                for r in ph.runs:
                    r.font.size = Pt(9)

    h(doc, '六、竞赛设计策略速查', 1)
    add_p(doc, '保守路径（推荐初学者）：', bold=True)
    bullet(doc, '以sfGFP为模板，保留6个关键突变+引入8-12个优化突变')
    bullet(doc, '使用PROSS在线服务器生成初始设计 → FoldX验证ΔΔG → ColabFold验证结构')
    bullet(doc, '最终突变控制在8-12个，降低上位效应风险')

    add_p(doc, '数据驱动路径（中等难度）：', bold=True)
    bullet(doc, '利用竞赛GFP_data.xlsx训练ML模型 → 多目标优化 → 交叉熵蒙特卡洛搜索')
    bullet(doc, '使用TemBERTure+SPURS联合筛选 → 选出高亮度+高稳定性候选')
    bullet(doc, '管线和Agent逻辑树必须详细记录——评审重点考察')

    add_p(doc, '生成式路径（高难度·高回报）：', bold=True)
    bullet(doc, 'ESM3/GeoEvoBuilder → 零样本生成大量候选 → 三轮迭代筛选')
    bullet(doc, 'ProteinMPNN+固定发色团约束 → 探索超越现有GFP家族的序列空间')
    bullet(doc, '提交6条具有互补性的序列——最大化至少一条高分的概率')

    # Save
    output_path = os.path.join(base_dir, '学习文档_专题四_蛋白质稳定性与热稳定性工程（含配图）.docx')
    doc.save(output_path)
    print(f'  保存: {output_path}')


# ===================================================================
# Main
# ===================================================================
if __name__ == '__main__':
    print("生成含配图的专题文档...")
    rebuild_doc1_with_images()
    rebuild_doc3_with_images()
    rebuild_doc4_with_images()
    print("\n含配图的文档生成完成！")
    print("注意：专题二的原始文档已包含丰富内容，机制报告的fig1-fig6可作为参考配图。")
