# -*- coding: utf-8 -*-
"""生成数据下载指南Word文档"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import datetime

doc = Document()

# 样式
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(11)

for i in range(1, 5):
    hs = doc.styles[f'Heading {i}']
    hs.font.name = 'Times New Roman'
    hs.font.size = Pt(18 if i == 1 else 15 if i == 2 else 13 if i == 3 else 12)
    if i == 1:
        hs.font.color.rgb = RGBColor(0, 51, 102)
    elif i == 2:
        hs.font.color.rgb = RGBColor(0, 80, 150)

def h(text, level=1):
    doc.add_heading(text, level=level)

def p(text, bold=False, size=11, color=None):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color

def bullet(text):
    doc.add_paragraph(text, style='List Bullet')

def link_box(title, url, desc="", data_size="", access="", fmt=""):
    """添加一个可点击链接块，含完整元数据"""
    doc.add_paragraph()
    # 标题
    para = doc.add_paragraph()
    run = para.add_run(f'■ {title}')
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0, 51, 102)
    # URL
    if url:
        para2 = doc.add_paragraph()
        run = para2.add_run(f'  链接：{url}')
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0, 80, 160)
    # 描述
    if desc:
        para3 = doc.add_paragraph()
        run = para3.add_run(f'  内容：{desc}')
        run.font.size = Pt(10)
    # 元数据行
    meta_parts = []
    if data_size:
        meta_parts.append(f'数据量：{data_size}')
    if fmt:
        meta_parts.append(f'格式：{fmt}')
    if access:
        meta_parts.append(f'获取方法：{access}')
    if meta_parts:
        para4 = doc.add_paragraph()
        run = para4.add_run('  ' + '  |  '.join(meta_parts))
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(100, 100, 100)

# ============================================================
# 封面
# ============================================================
for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('GFP蛋白质设计训练数据集\n下载与获取指南')
run.bold = True
run.font.size = Pt(26)
run.font.color.rgb = RGBColor(0, 51, 102)

doc.add_paragraph()

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run('合成生物学创新赛 · 蛋白设计赛道专用')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0, 80, 150)

doc.add_paragraph()
date_info = doc.add_paragraph()
date_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_info.add_run(f'生成日期：{datetime.date.today().strftime("%Y年%m月%d日")}').font.size = Pt(12)

doc.add_page_break()

# ============================================================
# 概述
# ============================================================
h('一、概述', 1)
p('本指南汇总了可用于GFP蛋白质设计模型训练的所有公开数据集，每个数据集均包含完整的URL、数据描述、数据量、获取方法及推荐格式。')
p('数据驱动设计是合成生物学创新赛"蛋白设计赛道"的核心方法论。高质量的训练数据是构建"亮度+热稳定性"双目标预测模型的基础。')
p('')

p('已生成的本地数据集文件：', bold=True)
bullet('comprehensive_GFP_dataset.xlsx —— 含7个工作表（500条亮度训练数据、22个FPbase变体参数、28条热稳定性数据等）')
bullet('data/GFP_training_data.csv —— 赛事提供的500条单突变亮度数据（可快速加载用于ML训练）')
bullet('data/GFP data.xlsx —— 赛事原始数据文件（含brightness和beforetop10两个工作表）')

doc.add_page_break()

# ============================================================
# 数据集列表
# ============================================================
h('二、核心数据集下载清单', 1)

# --- 数据集1 ---
h('2.1 Nature 2016 GFP适应度景观数据集（⭐最高优先级）', 2)
p('这是目前已知最大的GFP突变-荧光定量数据集，对训练高精度预测模型至关重要。', bold=True)

link_box(
    'Sarkisyan et al. (2016) "Local fitness landscape of GFP" — 完整实验数据',
    'https://figshare.com/articles/dataset/Local_fitness_landscape_of_the_green_fluorescent_protein/3102154',
    '>50,000条avGFP突变体荧光数据，涵盖单突变到11突变组合。每条记录含：突变列表、中位荧光值（log-brightness）、标准差、独特性条形码数（数据质量指标）。序列空间覆盖>10^13个可能变体中的约51,715个。',
    '~51,715条基因型',
    '1. 点击链接进入Figshare页面；2. 下载 amino_acid_genotypes_to_brightness.tsv；3. 用Excel打开（制表符分隔，自动提示导入）；4. 也可用Python: pd.read_csv(file, sep="\t")',
    'TSV（Excel可直接打开）'
)

p('该数据集的子文件清单：', bold=True, size=10)
sub_files = [
    ('amino_acid_genotypes_to_brightness.tsv', '氨基酸基因型汇总 → 荧光值（最常用）'),
    ('nucleotide_genotypes_to_brightness.tsv', '核苷酸基因型汇总 → 荧光值（含核酸层面信息）'),
    ('barcodes_to_brightness.tsv', '每个分子条形码的原始荧光估值'),
    ('genotypes.tsv', 'MiSeq处理数据 → 条形码 ↔ 基因型映射'),
    ('populations.zip', '原始FACS分选计数数据（每个gate×replica）'),
    ('avGFP_reference_sequence.fa', 'avGFP参考序列（FASTA格式）'),
]
for fname, desc in sub_files:
    bullet(f'{fname} —— {desc}')

p('')
p('原始测序数据（NCBI SRA）：', bold=True, size=10)
bullet('BioProject PRJNA282342: https://www.ncbi.nlm.nih.gov/bioproject/PRJNA282342')

doc.add_paragraph()

# --- 数据集2 ---
h('2.2 FPbase 荧光蛋白数据库（⭐最高优先级）', 2)
p('FPbase是荧光蛋白领域的权威数据库，包含>1,000种荧光蛋白的结构化光物理参数。', bold=True)

link_box(
    'FPbase: The Fluorescent Protein Database',
    'https://www.fpbase.org/',
    '超过1,000种荧光蛋白的完整光物理参数：激发/发射波长、消光系数、量子产率、亮度、pKa、成熟时间、光稳定性、聚集状态、序列信息、PDB结构、参考文献。数据经过人工审核和文献溯源。',
    '1,000+ 荧光蛋白',
    '方法1（Web界面）：https://www.fpbase.org/table → 右下角"Download"按钮 → 选CSV或Excel；\n方法2（Python API）：pip install fpbase → import fpbase → fpbase.list_proteins()；\n方法3（GraphQL API）：https://www.fpbase.org/graphql → 交互式查询界面',
    'CSV / Excel / JSON / Python API / GraphQL'
)

p('FPbase GraphQL 查询GFP相关蛋白示例：', bold=True, size=10)
code1 = '''query {
  proteins(search: "GFP") {
    name
    slug
    defaultState {
      exMax emMax extCoeff qy brightness pKa
      maturation lifetime
    }
    sequences { aminoAcid }
  }
}'''
para = doc.add_paragraph()
run = para.add_run(code1)
run.font.size = Pt(8)
run.font.name = 'Courier New'

doc.add_paragraph()

# --- 数据集3 ---
h('2.3 ThermoMutDB 热力学突变数据库（⭐⭐高优先级）', 2)

link_box(
    'ThermoMutDB: Thermodynamic Database for Missense Mutations',
    'http://biosig.unimelb.edu.au/thermomutdb',
    '>14,669条手动审核的突变热力学数据，涵盖588种蛋白。每条记录含：ΔΔG（折叠自由能变化）、ΔTm（熔解温度变化）、实验条件、方法学信息、PDB结构关联。比ProTherm增加83%突变条目和204种新蛋白。',
    '~14,669突变 / 588蛋白',
    '1. 访问下载页: http://biosig.unimelb.edu.au/thermomutdb/downloads；\n2. 下载完整CSV或JSON；\n3. 在本地搜索"GFP"或"green fluorescent"过滤相关条目；\n4. 也可通过REST API程序化查询（详见网站API文档）',
    'CSV / JSON / REST API / XML'
)

p('直接下载链接：', bold=True, size=10)
bullet('完整CSV: http://biosig.unimelb.edu.au/thermomutdb/downloads/thermomutdb.csv')
bullet('完整JSON: http://biosig.unimelb.edu.au/thermomutdb/downloads/thermomutdb.json')

doc.add_paragraph()

# --- 数据集4 ---
h('2.4 ProThermDB 蛋白质热力学数据库（⭐⭐高优先级）', 2)

link_box(
    'ProThermDB: Thermodynamic Database for Proteins and Mutants',
    'https://web.iitm.ac.in/bioinfo2/prothermdb/',
    '>32,000条热力学数据点 + ~120,000高通量蛋白质组稳定性数据。参数涵盖：Tm、ΔG、ΔH、ΔCp。支持按UniProt ID、PDB ID、蛋白名、突变类型、实验条件和文献作者搜索。以GFP为例，avGFP的UniProt ID为P42212。',
    '>32,000数据点 + ~120K高通量数据',
    '1. 访问网站 → 搜索"P42212"或"GFP"；\n2. 在搜索结果中筛选目标条目；\n3. 填写下载表单（需邮箱）；\n4. 下载CSV格式数据；\n5. 详细教程: DOI: 10.1002/cpz1.306',
    'CSV（需注册）'
)

doc.add_paragraph()

# --- 数据集5 ---
h('2.5 synbiochallenges2025 竞赛参考管线（⭐⭐高优先级）', 2)

link_box(
    'GitHub: f-normies/synbiochallenges2025',
    'https://github.com/f-normies/synbiochallenges2025',
    '2025年蛋白设计赛道完整ML管线实现。含：GLEAM神经网络（ESM-2 3B嵌入+交叉注意力，预测荧光亮度）、MLDE（k-mer编码定向进化优化）、交叉熵蒙特卡洛搜索（上位效应建模）、TemBERTure（ProtBERT-BFD微调，预测Tm）。代码可直接参考和修改用于2026赛季。',
    '完整竞赛管线+训练数据',
    '1. git clone https://github.com/f-normies/synbiochallenges2025；\n2. 查看notebooks/目录下的Jupyter Notebook；\n3. 阅读README了解管线架构；\n4. conda env create -f environment.yml 创建环境；\n5. 研究GLEAM模型架构（model/目录）',
    'Python / Jupyter / CSV / Conda环境'
)

doc.add_paragraph()

# --- 数据集6 ---
h('2.6 GeoStab ΔTm点突变数据集（⭐中优先级）', 2)

link_box(
    'GitHub: Gonglab-THU/GeoStab (含S4346/S1626/S571 ΔTm数据集)',
    'https://github.com/Gonglab-THU/GeoStab',
    '几何深度学习蛋白质稳定性预测模型，附带多个高质量ΔTm数据集：S4346（4,346单点突变、349蛋白）、S1626（1,626突变、95蛋白）、S571（571突变、37蛋白）。全部含实验ΔTm值，适用于训练/基准测试热稳定性预测模型。额外含GeoFitness（74 DMS数据集）和DeepSequence（34 DMS研究）。',
    '~6,543单点突变',
    '1. git clone https://github.com/Gonglab-THU/GeoStab；\n2. 查看data/目录下各CSV文件；\n3. 使用pd.read_csv()加载目标数据集；\n4. 搜索GFP相关条目（如有PDB结构可用GeoStab预测）',
    'CSV / Python (PyTorch)'
)

doc.add_paragraph()

# --- 数据集7 ---
h('2.7 ProteinMutTm 酶突变热稳定性数据库（⭐中优先级）', 2)

link_box(
    'GitHub: hyq2017/ProteinMutTm (中文项目)',
    'https://github.com/hyq2017/ProteinMutTm',
    '通过文献和数据库挖掘建立的蛋白质突变体热稳定性TSV数据库。字段含：MUTATION、Tm_(C)、Delta_Tm_(C)、Uniprot_Sequence、Mut_Sequence。可作为通用蛋白热稳定性补充数据，用于ML模型预训练。',
    '多蛋白多突变',
    '1. git clone https://github.com/hyq2017/ProteinMutTm；\n2. 查看data/目录下TSV文件；\n3. 搜索"GFP"相关条目；\n4. 使用pd.read_csv(file, sep="\t")加载',
    'TSV'
)

doc.add_paragraph()

# --- 数据集8 ---
h('2.8 ProteinGym 深度突变扫描基准数据集（⭐中优先级）', 2)

link_box(
    'ProteinGym: Deep Mutational Scanning Benchmark',
    'https://proteingym.org/',
    '大规模DMS基准数据集：217个深度突变扫描数据集，>150万突变。含替换（substitution）和插入缺失（indel）两种突变类型，每个位点覆盖全部20种氨基酸（部分数据集）。GFP相关条目可用于评估突变效应的模型。',
    '217 DMS集 / >1.5M突变',
    '1. 访问 https://proteingym.org/；\n2. 在Benchmark页面搜索"GFP"；\n3. 下载匹配的DMS数据集；\n4. 也可通过GitHub获取程序化访问接口',
    'CSV / JSON'
)

doc.add_paragraph()

# --- 数据集9 ---
h('2.9 ESM-2 预计算嵌入（⭐中优先级）', 2)

link_box(
    'Zenodo: ESM-2 15B Embeddings for GFP DMS Dataset',
    'https://zenodo.org/records/17088257',
    '对GFP深度突变扫描数据集预计算的ESM-2 15B参数序列嵌入。无需本地GPU推理即可直接用于ML训练。大幅降低计算门槛——直接用NPY/Pickle文件作为输入特征。',
    '预计算嵌入',
    '1. 访问Zenodo链接；\n2. 下载NPY或Pickle文件；\n3. 在Python中加载: np.load() 或 pickle.load()；\n4. 直接作为X特征输入到RF/XGBoost等下游模型',
    'NPY / Pickle'
)

doc.add_paragraph()

# --- 数据集10 ---
h('2.10 Meltome Atlas 热蛋白质组图谱（⭐低优先级）', 2)

link_box(
    'Meltome Atlas: Thermal Proteome Stability Across the Tree of Life',
    'https://ftp.pride.ebi.ac.uk/pride/data/archive/2020/04/PXD011929',
    '13个物种（从古菌到人类）的质谱热蛋白质组稳定性数据。包含跨生命之树的Tm值分布，可用于理解蛋白质热稳定性的进化趋势。',
    '~数万蛋白的Tm数据',
    '1. 访问PRIDE档案下载；\n2. 下载data文件并解压；\n3. 搜索GFP相关条目；\n4. 也可从配套论文获取处理后的数据',
    'TSV / 原始质谱数据'
)

doc.add_page_break()

# ============================================================
# 下载方式对比
# ============================================================
h('三、下载方式对比与推荐', 1)

p('根据数据类型和数据量，推荐以下下载优先级和策略：', bold=True)

h('3.1 推荐下载顺序', 2)

priority_list = [
    ('第1步（立即）', 'Nature 2016 Figshare (51,715条GFP突变+荧光)',
     '最大最直接相关的数据集，可直接训练亮度预测模型'),
    ('第2步（当天）', 'FPbase (1,000+荧光蛋白参数)',
     '补充多类型GFP变体的定量参数，支持迁移学习'),
    ('第3步（当天）', 'ThermoMutDB (14,669条突变ΔΔG/ΔTm)',
     '补足热稳定性标签数据，支撑双目标预测'),
    ('第4步（次日）', 'synbiochallenges2025 GitHub仓库',
     '学习往年竞赛管线，复用GLEAM/MLDE框架'),
    ('第5步（选做）', 'GeoStab S4346 ΔTm数据集',
     '通用蛋白ΔTm补充训练数据，提升模型泛化能力'),
    ('第6步（选做）', 'ESM-2预计算嵌入 (Zenodo)',
     '直接获取无需GPU推理的特征，适合快速原型'),
]

for step, name, reason in priority_list:
    bullet(f'{step}：{name} —— {reason}')

h('3.2 数据整合建议', 2)
p('将各来源数据整合为统一的训练格式（参考 GFP_training_data.csv 格式）：')
bullet('统一列名：GFP_type, aaMutations, Brightness, Tm_C, Delta_Delta_G, Source')
bullet('Brightness列：Nature 2016数据（log值）→ 转换为线性相对亮度（参考WT=1.0归一化）')
bullet('Tm_C列：从ThermoMutDB/ProThermDB/GeoStab提取GFP条目填充')
bullet('突变描述格式：统一为"原始AA+位置+新AA"（如F64L, S65T）')
bullet('多突变：用冒号分隔（如"F64L:S65T:S30R"），与现有格式一致')

doc.add_page_break()

# ============================================================
# 快速起步代码
# ============================================================
h('四、快速起步代码示例', 1)

h('4.1 从FPbase获取所有GFP数据', 2)
code_fpbase = '''# 安装: pip install fpbase
import fpbase
import pandas as pd

# 列出所有蛋白
all_proteins = fpbase.list_proteins()

# 筛选GFP相关
gfp_data = []
for p in all_proteins:
    if p.default_state and ('gfp' in p.name.lower() or 'green' in p.name.lower()):
        fp = fpbase.get_protein(p.slug)
        gfp_data.append({
            'name': fp.name,
            'ex_max': fp.default_state.exMax,
            'em_max': fp.default_state.emMax,
            'ext_coeff': fp.default_state.extCoeff,
            'qy': fp.default_state.qy,
            'brightness': fp.default_state.brightness,
            'pKa': fp.default_state.pKa,
        })

df = pd.DataFrame(gfp_data)
df.to_csv('fpbase_gfp_variants.csv', index=False)
print(f"获取了 {len(df)} 个GFP相关蛋白")'''

para = doc.add_paragraph()
run = para.add_run(code_fpbase)
run.font.size = Pt(8)
run.font.name = 'Courier New'

h('4.2 从ThermoMutDB获取GFP热稳定性数据', 2)
code_thermo = '''# 下载完整CSV后本地过滤
import pandas as pd

# 下载: http://biosig.unimelb.edu.au/thermomutdb/downloads/thermomutdb.csv
df = pd.read_csv('thermomutdb.csv')

# 筛选GFP条目
gfp_mask = df['protein_name'].str.contains('GFP|green fluorescent', case=False, na=False)
gfp_subset = df[gfp_mask].copy()

# 查看关键列
key_cols = ['protein_name', 'mutation', 'ddg', 'dtm', 'tm_wt', 'tm_mut',
            'exp_method', 'ph', 'temperature', 'reference']
print(gfp_subset[key_cols].to_string())
gfp_subset[key_cols].to_csv('gfp_thermal_stability.csv', index=False)'''

para = doc.add_paragraph()
run = para.add_run(code_thermo)
run.font.size = Pt(8)
run.font.name = 'Courier New'

h('4.3 加载Nature 2016数据集并与现有数据合并', 2)
code_nature = '''# 从Figshare下载后加载(当前网络不可用, 需手动下载)
import pandas as pd

# 下载链接(手动):
# https://figshare.com/ndownloader/files/5047938
# 保存为: nature2016_gfp_fitness.tsv

nature_df = pd.read_csv('nature2016_gfp_fitness.tsv', sep='\\t')
print(f"Nature 2016: {len(nature_df)} 条基因型")
print(f"列: {list(nature_df.columns)}")
# 典型列: aaMutations, uniqueBarcodes, medianBrightness, std

# 合并现有数据
existing_df = pd.read_csv('data/GFP_training_data.csv')
combined = pd.concat([
    existing_df[['GFP type', 'aaMutations', 'Brightness']],
    nature_df[['aaMutations', 'medianBrightness']].rename(
        columns={'medianBrightness': 'Brightness'})
], ignore_index=True)
combined['GFP type'] = combined['GFP type'].fillna('avGFP')
print(f"合并后: {len(combined)} 条")'''

para = doc.add_paragraph()
run = para.add_run(code_nature)
run.font.size = Pt(8)
run.font.name = 'Courier New'

doc.add_page_break()

# ============================================================
# 注意事项
# ============================================================
h('五、注意事项', 1)

h('5.1 数据质量差异', 2)
bullet('Nature 2016数据集：大规模实验数据，每个基因型有多个条形码验证，数据质量较高。荧光值为对数尺度（log-fluorescence），需注意转换。')
bullet('FPbase：人工审核，数据可信度高。但不同蛋白的测量条件可能有差异（pH、温度、缓冲液），直接比较时需注意。')
bullet('ThermoMutDB/ProThermDB：手动从文献提取，实验条件和方法学差异大。ΔTm和ΔΔG的绝对值不宜直接跨实验比较，但相对趋势可靠。')

h('5.2 数据格式兼容性', 2)
bullet('TSV vs CSV：TSV使用制表符分隔（\t），CSV使用逗号（,）。Python的pd.read_csv()可通过sep参数处理两种格式。')
bullet('突变表示差异：Nature 2016使用"F64L"格式，ThermoMutDB可能使用"F64L"或"PHE64LEU"，需要统一。')
bullet('序列编号基准：不同文献可能使用不同的起始残基编号。需统一以avGFP完整序列（238 aa）为基准。')

h('5.3 数据量级评估', 2)
bullet('最小可行训练集：500-1,000条（现有数据可满足，适合快速原型）')
bullet('推荐训练集：5,000-50,000条（加入Nature 2016数据后可达，适合高精度模型）')
bullet('理想训练集：>100,000条（整合所有来源+数据增强后可达，适合深度神经网络）')

doc.add_page_break()

# ============================================================
# 链接索引
# ============================================================
h('六、下载链接快速索引', 1)

links_index = [
    ('Nature 2016 GFP Figshare', 'https://figshare.com/articles/dataset/Local_fitness_landscape_of_the_green_fluorescent_protein/3102154'),
    ('FPbase 主页', 'https://www.fpbase.org/'),
    ('FPbase 数据表（直接下载）', 'https://www.fpbase.org/table/'),
    ('FPbase GraphQL API', 'https://www.fpbase.org/graphql'),
    ('FPbase Python包', 'https://pypi.org/project/fpbase/'),
    ('ThermoMutDB 主页', 'http://biosig.unimelb.edu.au/thermomutdb'),
    ('ThermoMutDB 下载页', 'http://biosig.unimelb.edu.au/thermomutdb/downloads'),
    ('ProThermDB 主页', 'https://web.iitm.ac.in/bioinfo2/prothermdb/'),
    ('synbiochallenges2025 GitHub', 'https://github.com/f-normies/synbiochallenges2025'),
    ('GeoStab GitHub (含ΔTm数据集)', 'https://github.com/Gonglab-THU/GeoStab'),
    ('ProteinMutTm GitHub (中文)', 'https://github.com/hyq2017/ProteinMutTm'),
    ('ProteinGym 基准平台', 'https://proteingym.org/'),
    ('ESM-2预计算嵌入 Zenodo', 'https://zenodo.org/records/17088257'),
    ('Meltome Atlas PRIDE', 'https://ftp.pride.ebi.ac.uk/pride/data/archive/2020/04/PXD011929'),
    ('TemBERTure GitHub', 'https://github.com/ibmm-unibe-ch/TemBERTure'),
    ('NCBI SRA (原始测序)', 'https://www.ncbi.nlm.nih.gov/bioproject/PRJNA282342'),
    ('FPbase mBaoJin页面', 'https://www.fpbase.org/protein/mbaojin/'),
    ('FPbase sfGFP页面', 'https://www.fpbase.org/protein/superfolder-gfp/'),
    ('ProteinMPNN-ddG (零样本ΔΔG)', 'https://github.com/PeptoneLtd/proteinmpnn_ddg'),
    ('SPURS (ProteinMPNN+ESM ΔΔG)', 'https://github.com/luo-group/SPURS'),
]

for name, url in links_index:
    bullet(f'{name}: {url}')

# ============================================================
# 保存
# ============================================================
output_path = r'C:\Users\fengs\Desktop\蛋白质设计-合成生物学创新赛-Claude\GFP训练数据集下载指南.docx'
doc.save(output_path)
print(f'下载指南已保存至: {output_path}')
