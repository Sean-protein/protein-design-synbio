
import os, sys, pandas as pd, numpy as np, json
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import warnings
warnings.filterwarnings("ignore")

BASE = r"C:\Users\fengs\Desktop\蛋白质设计-合成生物学创新赛-Claude"
INTEGRATED = None
for root, dirs, files in os.walk(BASE):
    if "integrated_csv" in dirs:
        INTEGRATED = os.path.join(root, "integrated_csv")
        break
OUTPUT = os.path.join(BASE, "output")
os.makedirs(OUTPUT, exist_ok=True)

brightness_data = pd.read_csv(os.path.join(INTEGRATED, "comprehensive_Brightness_Training_Data.csv"))
fpbase_data = pd.read_csv(os.path.join(INTEGRATED, "comprehensive_FPbase_GFP_Variants.csv"))
thermal_data = pd.read_csv(os.path.join(INTEGRATED, "comprehensive_Thermal_Stability_Data.csv"))
literature_data = pd.read_csv(os.path.join(INTEGRATED, "comprehensive_Literature_GFP_Properties.csv"))
reference_seqs = pd.read_csv(os.path.join(INTEGRATED, "comprehensive_Reference_Sequences.csv"))
candidate_positions = pd.read_csv(os.path.join(INTEGRATED, "comprehensive_Candidate_Positions.csv"))
fpbase_full = pd.read_csv(os.path.join(INTEGRATED, "04_fpbase_gfp_variants.csv"))
uniprot_mutations = pd.read_csv(os.path.join(INTEGRATED, "07_uniprot_p42212_mutations.csv"))
sarkisyan_brightness = pd.read_csv(os.path.join(INTEGRATED, "03_genotype_brightness_sarkisyan.csv"))
extra_seqs = pd.read_csv(os.path.join(INTEGRATED, "01_reference_sequences.csv"))
WT_BRIGHT = 3.719212132
bright = brightness_data["Brightness"].dropna()

doc = Document()
section = doc.sections[0]
section.page_width = Cm(21.0)
section.page_height = Cm(29.7)

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = str(h)
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val) if val is not None and not (isinstance(val, float) and np.isnan(val)) else "-"
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(8)
    return table

# ====== COVER ======
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("GFP绿色荧光蛋白\n综合分析报告")
run.font.size = Pt(28)
run.bold = True
run.font.color.rgb = RGBColor(0, 90, 50)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("\n合成生物学创新赛 · 蛋白质设计项目\n")
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(80, 80, 80)

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info.add_run("数据来源：Sarkisyan 2016 | FPbase | ProTherm | UniProt | 文献整合\n").font.size = Pt(10)
info.add_run("报告生成日期：2026年5月16日\n").font.size = Pt(10)
doc.add_page_break()
print("Cover done")
