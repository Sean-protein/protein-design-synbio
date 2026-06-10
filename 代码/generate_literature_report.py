# -*- coding: utf-8 -*-
"""
生成2026蛋白质设计合成生物学挑战赛文献调研报告Word文档
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
import os

doc = Document()

# ========== 样式设置 ==========
style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# 标题样式
for i in range(1, 4):
    heading_style = doc.styles[f'Heading {i}']
    heading_style.font.name = '微软雅黑'
    heading_style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    if i == 1:
        heading_style.font.size = Pt(22)
        heading_style.font.color.rgb = RGBColor(0, 51, 102)
    elif i == 2:
        heading_style.font.size = Pt(16)
        heading_style.font.color.rgb = RGBColor(0, 80, 150)
    elif i == 3:
        heading_style.font.size = Pt(13)
        heading_style.font.color.rgb = RGBColor(0, 100, 180)

# ========== 封面 ==========
for _ in range(4):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('2026 蛋白质设计-合成生物学创新赛\n文献调研报告')
run.font.size = Pt(28)
run.font.bold = True
run.font.color.rgb = RGBColor(0, 51, 102)

doc.add_paragraph()

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('—— 绿色荧光蛋白（GFP）设计前沿文献综述')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(100, 100, 100)

doc.add_paragraph()
doc.add_paragraph()

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run('报告日期：2026年5月14日')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(120, 120, 120)

doc.add_page_break()

# ========== 目录页 ==========
doc.add_heading('目录', level=1)
doc.add_paragraph('一、2026年赛事概况与规则解读')
doc.add_paragraph('二、核心参考文献综述')
doc.add_paragraph('    2.1 StayGold / mStayGold —— 超高光稳定性荧光蛋白的单体化突破')
doc.add_paragraph('    2.2 mBaoJin —— 新一代单体超亮绿色荧光蛋白')
doc.add_paragraph('    2.3 TGP —— 结构导向表面工程创造极端热稳定性')
doc.add_paragraph('    2.4 Superfolder GFP —— 超折叠GFP奠基性工作')
doc.add_paragraph('    2.5 Nature GFP适应度景观 —— 理解GFP突变耐受的底层逻辑')
doc.add_paragraph('三、文献核心发现对比与赛事设计启示')
doc.add_paragraph('四、参考文献')

doc.add_page_break()

# ========== 第一部分：赛事概况 ==========
doc.add_heading('一、2026年赛事概况与规则解读', level=1)

doc.add_heading('1.1 赛事目标', level=2)
doc.add_paragraph(
    '2026年合成生物学创新赛（SynBio Challenges）蛋白质设计赛道的核心目标是：'
    '在"全序列开放设计"的条件下，参赛者需运用计算生物学、生物信息学以及可编排的Agent工作流，'
    '设计出兼具极高初始亮度与极限热稳定性的新型绿色荧光蛋白（GFP），打造所谓的"六边形战士"。'
)
doc.add_paragraph(
    '本年度基准对照蛋白从以往的普通GFP升级为超折叠GFP（sfGFP），这对参赛者提出了更高的要求。'
)

doc.add_heading('1.2 提交要求', level=2)
doc.add_paragraph('每队最多提交6条自行设计的蛋白序列，序列长度严格限制在220-250个氨基酸之间。')
doc.add_paragraph('序列必须以甲硫氨酸（M）开头，仅允许使用20种标准氨基酸的大写字母，严禁包含终止密码子或标点符号。')
doc.add_paragraph('提交内容包括三部分：')
doc.add_paragraph('  (1) 设计序列文件（.csv）：包含Team_Name、Seq_ID、Sequence三列', style='List Bullet')
doc.add_paragraph('  (2) 设计思路与说明文档（.pdf）：展示完整算法管线及Agent逻辑树', style='List Bullet')
doc.add_paragraph('  (3) 开源项目链接（URL）：GitHub/GitLab/HuggingFace等公开仓库', style='List Bullet')

doc.add_heading('1.3 实验测试流程（大设施自动执行）', level=2)
doc.add_paragraph(
    '所有序列由大设施统一执行NEBExpress Cell-Free E. coli Protein Synthesis System进行无细胞表达。测试流程分两个步骤：'
)
doc.add_paragraph('步骤一——初始亮度检测：在96孔黑色平底板中加入Cell-Free反应液与标准化DNA模板，30°C孵育3小时，25°C静置30分钟确保发色团完全成熟，使用酶标仪读取初始荧光强度（Ex: 485nm / Em: 515nm）。')
doc.add_paragraph('步骤二——极端热稳定性检测：自动化工作站将表达产物转移至PCR板并封膜，在PCR仪中执行72°C高温加热10分钟，然后降温至25°C保持5分钟复性（消除温度物理淬灭效应），读取残余荧光强度。')

doc.add_heading('1.4 评分规则——双能乘积制', level=2)
doc.add_paragraph('本次评分采用"双能乘积制"：')
doc.add_paragraph('项一：相对亮度得分 —— 评估蛋白在Cell-Free体系中的折叠效率与消光系数', style='List Bullet')
doc.add_paragraph('项二：热稳定性保留率 —— 评估蛋白抵抗72°C高温变性的能力', style='List Bullet')
doc.add_paragraph(
    '极低亮度阈值淘汰：若某条提交序列的初始亮度不到同批次野生型sfGFP的30%，'
    '说明该序列基础结构功能遭到严重破坏，该序列直接淘汰出局（两项及总成绩均记为0分）。'
)
doc.add_paragraph(
    '排名与奖项：以6条序列中得分最高的一条作为本队最终打榜成绩。前30%金奖、接下来若干比例银奖。'
    '此外设立"最佳亮度奖"和"最佳热稳定奖"作为独立荣誉。'
)

doc.add_heading('1.5 赛事提供的资源数据', level=2)
doc.add_paragraph('赛事为选手提供了以下辅助资源：')
doc.add_paragraph('GFP突变与亮度数据集（GFP_data.xlsx）：涵盖多品种GFP蛋白的历史数据，可用于大模型微调训练', style='List Bullet')
doc.add_paragraph('历年优秀序列（Top 20）：2024和2025年各10条优秀序列及对应年份字段', style='List Bullet')
doc.add_paragraph('PDB参考序列：sfGFP、avGFP、amacGFP、cgreGFP、ppluGFP共5条参考氨基酸序列及推荐PDB（如2B3P）', style='List Bullet')
doc.add_paragraph('排除序列列表（Exclusion_List.csv）：包含历史测试序列、基于FPbase的已知天然/工程变体等，任何完全一致的提交视为无效', style='List Bullet')

doc.add_page_break()

# ========== 第二部分：核心文献综述 ==========
doc.add_heading('二、核心参考文献综述', level=1)
doc.add_paragraph(
    '以下5篇文献是赛事方提供的核心参考资料，涵盖了GFP从基础折叠机制到前沿工程设计的全方位知识。'
    '每篇文献均对GFP蛋白质设计有直接指导意义。'
)

# --- 2.1 StayGold ---
doc.add_heading('2.1 StayGold / mStayGold —— 超高光稳定性荧光蛋白的单体化突破', level=2)
doc.add_paragraph(
    '文献：Ivorra-Molla E, et al. "A monomeric StayGold fluorescent protein." '
    'Nature Biotechnology, 2023.'
)

doc.add_heading('研究背景', level=3)
doc.add_paragraph(
    'StayGold是一种从Cytaeis uchidae（水母）GFP中工程改造而来的荧光蛋白，以其极高的亮度和卓越的抗光漂白性能著称。'
    '然而，StayGold在溶液中形成天然二聚体，这严重限制了其作为荧光标签的应用——二聚化可能导致融合蛋白定位错误和聚集。'
)

doc.add_heading('核心发现', level=3)
doc.add_paragraph('(1) 结构解析：通过X射线晶体学解析了StayGold的1.6Å高分辨率结构。StayGold单体由11条链组成的β-桶状结构构成，与经典的avGFP结构几乎相同。发色团由Gly57-Tyr58-Gly59形成。')
doc.add_paragraph('(2) 关键结构特征：StayGold发色团旁边存在一个氯离子（Cl⁻），位于与发色团同一平面内，通过与Lys61、Lys192及Arg86的静电相互作用保持稳定。这一特征在mNeonGreen中也存在。')
doc.add_paragraph('(3) 单体化突变：通过E138D单点突变成功将StayGold转变为完全单体化的mStayGold，且不损失其亮度和光稳定性。mStayGold在活细胞和固定细胞中均表现出与原始蛋白相当的优越光稳定性。')
doc.add_paragraph('(4) 性能对比：在激光照射下，mStayGold E138D的半衰期约为70秒，而sfGFP仅为11秒——光稳定性提升约6.4倍。')

doc.add_heading('对赛事的启示', level=3)
doc.add_paragraph(
    'StayGold/mStayGold代表了当前GFP光稳定性的"天花板"。其结构特征——特别是发色团附近的氯离子结合位点、'
    '以及维持发色团平面性和刚性的关键氨基酸网络——直接启发了后续mBaoJin等变体的设计。'
    '参赛者在追求稳定性的同时，需要确保单体性，E138D是一个可靠的参考位点。'
    '结构-功能关系表明：发色团微环境的刚性对光稳定性至关重要。'
)

# --- 2.2 mBaoJin ---
doc.add_heading('2.2 mBaoJin —— 新一代单体超亮绿色荧光蛋白', level=2)
doc.add_paragraph(
    '文献：Zhang H, Lesnov GD, et al. "Bright and stable monomeric green fluorescent protein derived from StayGold." '
    'Nature Methods, 2024.'
)

doc.add_heading('研究背景', level=3)
doc.add_paragraph(
    '在StayGold被报道后，多个研究组竞相开发其单体化版本。mBaoJin的研究团队采用了不同于E138D结构导向方法的策略——'
    '他们开发了一种巧妙的细菌筛选系统，无需预先了解二聚化结构即可筛选单体化变体。'
)

doc.add_heading('核心发现', level=3)
doc.add_paragraph('(1) 创新的筛选系统：利用AraC蛋白的DNA结合域——当AraC DNA通过二聚化的FP连接时会激活报告基因（mTagBFP）的转录。二聚化越弱的FP变体产生的蓝色荧光越少，从而实现了高通量筛选。')

doc.add_paragraph('(2) 8轮定向进化筛选：经过8轮随机诱变和筛选，获得了关键突变：S55T、H77R、E80G、Q140P、H141Q、C165Y、N171Y、T201A。最终变体命名为mBaoJin。')

doc.add_paragraph('(3) 单体性验证：在56μM高浓度下，mBaoJin约99%为单体——甚至高于mNeonGreen的95%。OSER实验评分为93.7%，达到单体性阈值。')

doc.add_paragraph('(4) 卓越性能特征：')
doc.add_paragraph('    - 亮度：细胞内亮度比EGFP高70-140%', style='List Bullet')
doc.add_paragraph('    - 光稳定性：在活细胞成像中远超mNeonGreen', style='List Bullet')
doc.add_paragraph('    - 成熟速度极快：比StayGold和mNeonGreen分别快1.4倍和1.8倍', style='List Bullet')
doc.add_paragraph('    - pH稳定性：pKa值4.37（极低，适合酸性环境）', style='List Bullet')
doc.add_paragraph('    - 化学稳定性：在6M盐酸胍中孵育24小时后荧光反而增加18%，而mNeonGreen完全淬灭', style='List Bullet')

doc.add_paragraph('(5) 结构分析：解析了mBaoJin在pH 4.6、6.5、8.5下的多个X射线晶体结构。揭示了发色团周围的致密氢键网络和疏水接触是其高光稳定性的结构基础。N137K、S134P、V152E突变分别使光漂白加速1.5、6.5和4.5倍。')

doc.add_paragraph('(6) 应用验证：成功用于超分辨率成像（SIM，约65nm分辨率）、扩展显微镜（ExM）、秀丽隐杆线虫神经元标记和小鼠脑组织成像。')

doc.add_heading('对赛事的启示', level=3)
doc.add_paragraph(
    'mBaoJin的开发范式展示了"高通量筛选+结构验证"的双轮驱动策略。其对发色团微环境（N137、S134、V152位点）'
    '的分析直接揭示了影响光稳定性的关键残基。注意mBaoJin的N端和C端来自mNeonGreen以确保哺乳动物细胞稳定性——'
    '这说明末端序列对蛋白在异源系统中的表达和折叠至关重要。'
)

# --- 2.3 TGP ---
doc.add_heading('2.3 TGP —— 结构导向表面工程创造极端热稳定性', level=2)
doc.add_paragraph(
    '文献：Close DW, et al. "TGP, an extremely stable, non-aggregating fluorescent protein created '
    'by structure-guided surface engineering." Proteins, 2015.'
)

doc.add_heading('研究背景', level=3)
doc.add_paragraph(
    'eCGP123是一种高度热稳定的荧光蛋白，但存在严重的聚集问题。TGP研究的目标是：'
    '在不损害热稳定性和荧光特性的前提下，通过结构导向的表面工程改善其溶解性。'
)

doc.add_heading('核心发现', level=3)
doc.add_paragraph('(1) eCGP123晶体结构：解析了eCGP123的2.1Å晶体结构（P1空间群，8个分子在不对称单元中）。其整体折叠为经典11链β-桶状结构，发色团为QYG三联体（Gln62-Tyr63-Gly64），His193通过π-π堆积与发色团的4-羟基苄基环相互作用。')

doc.add_paragraph('(2) 表面工程设计策略——两大原则：')
doc.add_paragraph('    原则一：破坏晶体晶格接触——识别并突变在晶体中形成分子间界面的残基，消除引起聚集的分子间相互作用', style='List Bullet')
doc.add_paragraph('    原则二：增加表面净电荷——将表面碱性残基（K、R）突变为谷氨酸（E），通过电荷排斥减少聚集', style='List Bullet')

doc.add_paragraph('(3) 具体突变：K45E、K73E、K117E、R149E、N158E（表面正/中性残基→谷氨酸）。C末端219-225（MLPSQAK）替换为GGGSGGG柔性连接序列。')

doc.add_paragraph('(4) TGP的性能提升：')
doc.add_paragraph('    - 90°C下的荧光半衰期：TGP约380分钟 vs eCGP123约175分钟（提升约2.2倍）', style='List Bullet')
doc.add_paragraph('    - 溶解性显著提高，不再聚集', style='List Bullet')
doc.add_paragraph('    - 光物理特性（量子产率0.66、荧光寿命3.0-3.2ns）基本保持不变', style='List Bullet')

doc.add_paragraph('(5) 反直觉发现：尽管引入了高熵谷氨酸侧链，TGP反而更容易结晶（在多种条件下均可结晶，并获得更高分辨率1.9Å的结构），这与"表面熵降低促进结晶"的传统策略看似矛盾。作者认为：溶解性改善允许在更高浓度下进行结晶试验，从而补偿了表面熵的增加。')

doc.add_heading('对赛事的启示', level=3)
doc.add_paragraph(
    'TGP的工作对本次赛事最具直接指导意义——它直接演示了如何通过结构导向的表面工程提升热稳定性。'
    '核心策略可概括为：①解析高分结构→②识别分子间晶格接触→③破坏聚集界面+增加表面负电荷→④验证热稳定性提升。'
    '2.2倍的稳定性提升（90°C下175→380分钟）虽然卓越，但赛事要求72°C 10分钟后的保留率，'
    '意味着在更温和的变性条件下，需要探索更精细的稳定性增强策略。'
)

doc.add_page_break()

# --- 2.4 Superfolder GFP ---
doc.add_heading('2.4 Superfolder GFP —— 超折叠GFP的奠基性工作', level=2)
doc.add_paragraph(
    '文献：Pédelacq JD, et al. "Engineering and characterization of a superfolder green fluorescent protein." '
    'Nature Biotechnology, 2006.'
)

doc.add_heading('研究背景', level=3)
doc.add_paragraph(
    '普通GFP在作为融合标签时会因为融合伴侣蛋白的错误折叠而自身折叠失败、丧失荧光。'
    '该研究的目标是创造一种即使在融合了错误折叠多肽后仍能正确折叠的"超折叠"GFP变体。'
    '这也是本次赛事基准蛋白sfGFP的原始文献。'
)

doc.add_heading('核心发现', level=3)
doc.add_paragraph('(1) 创新的筛选策略：将GFP变体库与N端"诱饵"多肽（铁蛋白ferritin——一种强烈错误折叠的蛋白）融合表达，筛选仍能产生荧光的克隆。这套策略的压力远大于传统的单独表达筛选。')

doc.add_paragraph('(2) 超折叠GFP的6个关键突变：')
doc.add_paragraph('    - S30R：最重要的突变，贡献最大', style='List Bullet')
doc.add_paragraph('    - Y39N', style='List Bullet')
doc.add_paragraph('    - N105T', style='List Bullet')
doc.add_paragraph('    - Y145F', style='List Bullet')
doc.add_paragraph('    - I171V', style='List Bullet')
doc.add_paragraph('    - A206V：减少二聚化倾向', style='List Bullet')

doc.add_paragraph('(3) S30R的分子机制（1.4Å高分辨率结构揭示）：S30位于β-链S2中间，S30R突变引发了局部构象变化，在相邻的β-链间形成了五元静电网络：E32—R30—E17—R122—E115（涉及S1、S2、S5、S6四条相邻β-链）。这一离子网络显著增强了β-桶结构的全局稳定性。')

doc.add_paragraph('(4) 性能提升：')
doc.add_paragraph('    - 折叠动力学：初始再折叠速率提升约3.5倍', style='List Bullet')
doc.add_paragraph('    - 化学稳定性：尿素耐受性显著增强', style='List Bullet')
doc.add_paragraph('    - 环状排列耐受性：几乎所有14种环状排列变体均有荧光（普通GFP仅3种）', style='List Bullet')
doc.add_paragraph('    - 随机突变耐受性：随机突变库中含有更少的弱荧光/非荧光克隆', style='List Bullet')

doc.add_paragraph('(5) 与融合蛋白的表达：当与18种来自Pyrobaculum aerophilum的测试蛋白融合表达时，超折叠GFP的全细胞荧光与融合表达总量成正比，而传统GFP的荧光则与融合伴侣的折叠产量强相关——说明sfGFP能够独立于融合伴侣的折叠状态自行折叠。')

doc.add_heading('对赛事的启示', level=3)
doc.add_paragraph(
    'sfGFP是本赛事的基准对照蛋白，深刻理解其结构和功能至关重要。S30R形成的跨β-链静电网络'
    '提供了一个经典案例：少数关键位置的精确突变就能大幅提升全局稳定性。'
    '在sfGFP基础上进一步提升热稳定性，需要考虑在维持或增强此类"稳定化网络"的前提下，'
    '引入额外的结构加固。sfGFP的6个突变位点是设计起点模板的核心参考。'
)

# --- 2.5 Nature GFP fitness landscape ---
doc.add_heading('2.5 Nature GFP适应度景观 —— 理解GFP突变耐受的底层逻辑', level=2)
doc.add_paragraph(
    '文献：Sarkisyan KS, et al. "Local fitness landscape of the green fluorescent protein." '
    'Nature, 2016, 533: 397-401.'
)

doc.add_heading('研究背景', level=3)
doc.add_paragraph(
    '理解蛋白质的"适应度景观"（fitness landscape）——即基因型与表型之间的映射关系——'
    '对于蛋白质设计至关重要。该研究首次在蛋白质全序列尺度上系统描绘了GFP的局部适应度景观。'
)

doc.add_heading('核心发现', level=3)
doc.add_paragraph('(1) 实验规模：通过对avGFP进行随机诱变和荧光分选，测量了数万种衍生基因型的荧光水平。对单突变、双突变及多突变组合进行了系统评估。')

doc.add_paragraph('(2) GFP适应度景观极为狭窄：')
doc.add_paragraph('    - 75%的单突变导致荧光降低（仅约23%为中性突变）', style='List Bullet')
doc.add_paragraph('    - 50%的含4个突变的基因型完全丧失荧光', style='List Bullet')
doc.add_paragraph('    - 绝大多数多突变组合具有比预期更低的荧光（负上位效应）', style='List Bullet')

doc.add_paragraph('(3) 上位效应（Epistasis）广泛存在：')
doc.add_paragraph('    - 高达30%的多突变基因型表现出显著的上位效应（主要是负上位效应）', style='List Bullet')
doc.add_paragraph('    - 负上位效应主要由"弱有害突变的累积效应"驱动——单个看似中性的突变组合在一起会突破蛋白质稳定性的阈值，导致折叠失败和荧光丧失', style='List Bullet')
doc.add_paragraph('    - 正上位效应极为罕见', style='List Bullet')

doc.add_paragraph('(4) 适应度景观可描述为Sigmoid截断函数：简单的线性非上位模型仅解释70%的方差；引入Sigmoid函数形式的"适应度势能"后解释率达到85%。这意味着蛋白质荧光与底层稳定性之间存在非线性阈值关系——稳定性降低到某个临界点后，荧光急剧丧失。')

doc.add_paragraph('(5) 进化尺度的一致性：通过分析跨越数亿年的GFP直系同源序列，发现上位效应的演化预测与实验数据高度吻合——局部适应度景观的特征与全局进化中的特征具有一致性。')

doc.add_heading('对赛事的启示', level=3)
doc.add_paragraph(
    '这是对参赛者最为重要的一篇警示性文献。其核心教训是：'
    '①不能将单个突变的效应简单加和——突变之间存在复杂的非线性相互作用（上位效应）；'
    '②多个看似无害的"中性"突变组合起来可能导致蛋白质灾难性失稳——因此需要在设计过程中'
    '持续评估全局稳定性；③实验验证不可替代——计算预测无法完全捕捉所有上位效应；'
    '④设计策略应考虑突变缓冲（mutational buffering）——利用如sfGFP这类超稳定骨架来吸收突变的不利影响。'
)

doc.add_page_break()

# ========== 第三部分：对比与启示 ==========
doc.add_heading('三、文献核心发现对比与赛事设计启示', level=1)

doc.add_heading('3.1 各文献核心性能参数对比', level=2)

# 创建对比表格
table = doc.add_table(rows=6, cols=5, style='Light Grid Accent 1')
table.autofit = True

# 表头
headers = ['蛋白/文献', '亮度/量子产率', '热稳定性', '光稳定性', '关键创新']
for i, header in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = header
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.bold = True
            run.font.size = Pt(9)

# 数据行
data = [
    ['sfGFP\n(2006, Nat Biotechnol)', 'QY=0.65\nε=83,300', '尿素耐受性增强\n（vs折叠报告GFP）', '略快于折叠报告GFP', 'S30R五元静电网络\n环状排列耐受\n融合蛋白折叠独立'],
    ['StayGold/mStayGold\n(2023, Nat Biotechnol)', '极高亮度\n（具体QY见原文）', '—', '半衰期~70s(vs sfGFP~11s)\n6.4倍光稳定性', 'E138D单体化\n氯离子配位发色团\n1.6Å结构'],
    ['mBaoJin\n(2024, Nat Methods)', '比EGFP高70-140%\n细胞亮度', '6M GdnHCl 24h\n荧光增加18%', '活细胞超强抗漂白\nSIM成像>60min不漂白', 'AraC筛选系统\n99%单体@56μM\npKa=4.37'],
    ['TGP\n(2015, Proteins)', 'QY=0.66\nε未显著改变', '90°C半衰期~380min\n(vs eCGP123~175min)', '—（原文侧重热稳定性）', '结构导向表面工程\n破坏晶格接触\n增加表面负电荷'],
    ['avGFP适应度景观\n(2016, Nature)', '—', '75%单突变降低荧光\n上位效应~30%\nSigmoid截断模型', '—', '全序列适应度景观\n上位效应量化\n进化-实验一致性'],
]

for row_idx, row_data in enumerate(data):
    for col_idx, cell_text in enumerate(row_data):
        table.rows[row_idx + 1].cells[col_idx].text = cell_text
        for paragraph in table.rows[row_idx + 1].cells[col_idx].paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(8)

doc.add_heading('3.2 对赛事设计策略的综合启示', level=2)

doc.add_paragraph(
    '【策略一：以超稳定骨架为基础】sfGFP已经具有显著优于普通GFP的折叠稳健性。'
    '在sfGFP骨架基础上进行进一步优化，可以借助其"突变缓冲"能力吸收更多的工程突变。'
    '参赛者应从sfGFP的6个关键突变出发，特别是S30R形成的静电网络，作为不可触碰的核心稳定化元件。'
)

doc.add_paragraph(
    '【策略二：借鉴TGP的表面工程逻辑】TGP证明：通过破坏分子间聚集界面并增加表面负电荷，'
    '可以显著提升热稳定性（90°C下半衰期翻倍）。结合sfGFP的结构，识别其β-桶表面的潜在聚集热点'
    '并引入电荷修饰，是提升热稳定性保留率的可行路径。注意要避免破坏关键的结构稳定化网络。'
)

doc.add_paragraph(
    '【策略三：发色团微环境的精细化优化】mBaoJin和StayGold的结构分析揭示了：'
    '发色团周围的氢键网络密度、疏水接触的完整性、以及氯离子结合位点的保留，'
    '对光稳定性和亮度至关重要。N137、S134、V152（mBaoJin编号）等位点的优化可直接影响光稳定性。'
    '对应到sfGFP结构（PDB: 2B3P），参赛者需要识别同源的发色团微环境残基并进行精密优化。'
)

doc.add_paragraph(
    '【策略四：警惕上位效应的陷阱】Nature GFP适应度景观研究明确指出：'
    '多个看似中性的突变组合在一起可能导致灾难性的稳定性丧失。这意味着：'
    '参赛者不能简单地组合多个"好的"单点突变——必须在设计的每一步评估全局稳定性。'
    '建议采用"迭代设计-筛选"策略，或使用能预测上位效应的计算工具（如Rosetta ddG预测、'
    'ESM系列模型等），在突变引入后进行全局结构评估。'
)

doc.add_paragraph(
    '【策略五：平衡双目标——拒绝"偏科"】评分规则采用乘积制，意味着亮度或热稳定性任一项偏低'
    '都会导致总分大幅缩水。理想的设计应同时兼顾：'
    '①高效折叠（通过亲水表面、优化折叠动力学）→ 高初始亮度；'
    '②高结构刚性（通过核心疏水区优化、静电网络增强）→ 高热稳定性保留率。'
    '这两方面可能存在权衡——例如，过多的表面电荷可能影响折叠效率。'
    '需要找到"帕累托最优"区域。'
)

doc.add_heading('3.3 可利用的计算工具和方法', level=2)
doc.add_paragraph('基于各文献的方法论，以下工具和策略可直接用于本次设计：')
doc.add_paragraph('蛋白质结构预测与设计：AlphaFold2/3、ESMFold（结构预测）；ProteinMPNN（逆向折叠设计）；Rosetta（能量计算、ddG预测、Pack设计）', style='List Bullet')
doc.add_paragraph('序列-功能关系建模：ESM-1v/ESM-2/ESM-3（蛋白质语言模型嵌入）；SaProt（结构感知蛋白语言模型）；Tranception（进化规模的蛋白语言模型）', style='List Bullet')
doc.add_paragraph('上位效应评估：Rosetta ΔΔG计算、FoldX稳定性预测、基于MSA的耦合残基分析（如EVcouplings）', style='List Bullet')
doc.add_paragraph('高通量虚拟筛选：多样本Rosetta设计+能量排序、ML模型预测荧光/稳定性、多级级联漏斗筛选', style='List Bullet')

doc.add_page_break()

# ========== 第四部分：参考文献 ==========
doc.add_heading('四、参考文献', level=1)

references = [
    '1. Ivorra-Molla E, Akhuli D, McAndrew MBL, et al. A monomeric StayGold fluorescent protein. Nature Biotechnology, 2023. https://doi.org/10.1038/s41587-023-02018-w',
    '2. Zhang H, Lesnov GD, Subach OM, et al. Bright and stable monomeric green fluorescent protein derived from StayGold. Nature Methods, 2024.',
    '3. Close DW, Don Paul C, Langan PS, et al. TGP, an extremely stable, non-aggregating fluorescent protein created by structure-guided surface engineering. Proteins, 2015, 83(7): 1225-1237.',
    '4. Pédelacq JD, Cabantous S, Tran T, et al. Engineering and characterization of a superfolder green fluorescent protein. Nature Biotechnology, 2006, 24(1): 79-88.',
    '5. Sarkisyan KS, Bolotin DA, Meer MV, et al. Local fitness landscape of the green fluorescent protein. Nature, 2016, 533: 397-401.',
    '6. 2026 Protein Design in SynBio Challenges — 官方赛题文档, 2026.',
]

for ref in references:
    doc.add_paragraph(ref, style='List Number')

# ========== 保存文件 ==========
output_path = 'C:/Users/fengs/Desktop/蛋白质设计-合成生物学创新赛-Claude/2026ProteinDesign_文献调研报告.docx'
doc.save(output_path)
print(f'报告已保存至: {output_path}')
print('完成！')
