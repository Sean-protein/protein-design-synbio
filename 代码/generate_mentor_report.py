# -*- coding: utf-8 -*-
"""生成河南大学导师匹配分析报告 Word 文档"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

def scf(cell, text, bold=False, size=Pt(10), color=None, align='left'):
    """设置表格单元格文字"""
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = {'left': WD_ALIGN_PARAGRAPH.LEFT, 'center': WD_ALIGN_PARAGRAPH.CENTER}[align]
    run = p.add_run(text)
    run.font.size = size
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.bold = bold
    if color:
        run.font.color.rgb = color

def shade_row(row, color="1F497D"):
    for cell in row.cells:
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), color)
        shading.set(qn('w:val'), 'clear')
        cell._tc.get_or_add_tcPr().append(shading)
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

def heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return h

def para(doc, text, bold=False, size=Pt(11), indent=True):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Pt(22)
    run = p.add_run(text)
    run.font.size = size
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.bold = bold
    return p

def bullet(doc, text, size=Pt(10.5)):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run('• ' + text)
    run.font.size = size
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def numbered(doc, num, text, size=Pt(10.5)):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(f'{num}. {text}')
    run.font.size = size
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def make_info_table(doc, data):
    """data: list of (key, value) tuples"""
    n = len(data)
    table = doc.add_table(rows=n, cols=2, style='Table Grid')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (k, v) in enumerate(data):
        scf(table.cell(i, 0), k, bold=True, size=Pt(10))
        scf(table.cell(i, 1), v, bold=False, size=Pt(10))
        if i == 0:
            shade_row(table.rows[i])
    for row in table.rows:
        row.cells[0].width = Cm(3)
        row.cells[1].width = Cm(13)
    doc.add_paragraph()
    return table

# ═══════════════════════ 封面 ═══════════════════════
doc.add_paragraph()
doc.add_paragraph()
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('2026蛋白质设计-合成生物学创新赛\n河南大学导师资源匹配分析报告')
r.font.size = Pt(22); r.font.name = '微软雅黑'
r._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
r.bold = True; r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

doc.add_paragraph()
st = doc.add_paragraph(); st.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = st.add_run('——基于AI的GFP蛋白质亮度预测与设计项目')
r.font.size = Pt(14); r.font.name = '微软雅黑'
r._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
r.font.color.rgb = RGBColor(0x59, 0x56, 0x59)

doc.add_paragraph(); doc.add_paragraph()
info_p = doc.add_paragraph(); info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = info_p.add_run(f'生成日期：{datetime.date.today().strftime("%Y年%m月%d日")}\n数据来源：河南大学各学院/实验室官方网站 + 网络学术数据库')
r.font.size = Pt(11); r.font.name = '宋体'
r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

doc.add_page_break()

# ═══════════════════════ 目录 ═══════════════════════
heading(doc, '目  录', 1)
for item in [
    '一、项目背景与需求分析',
    '二、导师筛选方法与范围',
    '三、重点推荐导师（按优先级排序）',
    '    3.1 第一梯队：AI+蛋白质方向（最高优先级）',
    '    3.2 第二梯队：蛋白质工程与设计方向（高优先级）',
    '    3.3 第三梯队：合成生物学与计算模拟方向（中优先级）',
    '四、各学院/实验室完整导师名录表',
    '五、导师匹配度综合评估矩阵',
    '六、联系方案与话术模板',
    '七、附录：重要平台与资源',
]:
    para(doc, item, size=Pt(11), indent=False)

doc.add_page_break()

# ═══════════════════════ 一、项目背景 ═══════════════════════
heading(doc, '一、项目背景与需求分析', 1)

heading(doc, '1.1 项目概述', 2)
para(doc, '本项目为2026年合成生物学创新赛参赛项目，核心目标是通过AI计算方法（蛋白质语言模型+机器学习），预测和设计高亮度、高热稳定性的绿色荧光蛋白（GFP）变体，并筛选出最优候选序列。项目以5种GFP蛋白质（sfGFP、avGFP、amacGFP、cgreGFP、ppluGFP）为研究对象，重点针对avGFP进行突变设计和亮度预测。')

heading(doc, '1.2 技术路线', 2)
para(doc, '项目已实现完整的技术Pipeline：', indent=True)
for item in [
    '蛋白质语言模型嵌入：使用Meta ESM-2（esm2_t12_35M_UR50D）或SaProt（SaProt_35M_AF2）对GFP氨基酸序列进行特征提取，生成序列嵌入向量；',
    '机器学习预测：基于Random Forest Regressor，在~5000条avGFP突变体亮度数据上训练预测模型；',
    '候选序列生成：在Superfolder GFP的24个关键位点池中随机组合≤6个突变，生成候选序列；',
    '筛选与排序：排除已知序列后，按预测亮度降序输出Top 6候选序列。',
]:
    bullet(doc, item)

para(doc, '技术栈：Python, PyTorch, fair-esm, HuggingFace Transformers, scikit-learn, pandas, numpy')

heading(doc, '1.3 当前进度', 2)
for item in [
    '已完成：数据收集与整合（comprehensive_GFP_dataset.xlsx，含亮度数据）；',
    '已完成：ESM/SaProt蛋白质语言模型嵌入代码开发与调试；',
    '已完成：随机森林回归模型训练与验证（R²评估）；',
    '已完成：候选序列自动生成与预测Pipeline（gfp_design.py）；',
    '已完成：GFP理论基础、AI蛋白质设计工具、蛋白质稳定性工程等系统学习文档（12份）；',
    '待完成：候选序列的实验验证（湿实验——表达、纯化、荧光/稳定性测定）；',
    '待优化：模型精度提升、更多蛋白质特征的引入（如结构特征）、多目标优化（亮度+稳定性）；',
    '待探索：深度学习方法（CNN/Transformer/GNN）直接预测功能。',
]:
    bullet(doc, item)

heading(doc, '1.4 导师需求分析', 2)
para(doc, '基于项目当前阶段，需要在以下方面获得导师指导：')
for item in [
    'AI方法学指导：蛋白质语言模型的最优使用策略、深度学习模型的引入、模型可解释性分析；',
    '蛋白质设计经验：GFP关键功能位点的选择建议、突变策略优化、多目标优化平衡；',
    '实验验证资源：候选序列的湿实验验证，反馈实验数据用于模型迭代；',
    '竞赛策略建议：创新赛评审标准解读、项目展示优化、学术报告写作指导；',
    '学术发展指导：未来在AI蛋白质设计方向深造或发展的建议。',
]:
    bullet(doc, item)

doc.add_page_break()

# ═══════════════════════ 二、筛选方法 ═══════════════════════
heading(doc, '二、导师筛选方法与范围', 1)

heading(doc, '2.1 检索范围', 2)
for item in [
    '河南大学 生命科学学院 (bio.henu.edu.cn)',
    '河南大学 化学与分子科学学院 (ccce.henu.edu.cn)',
    '河南大学 人工智能学院 (ai.henu.edu.cn)',
    '河南大学 计算机与信息工程学院 (cs.henu.edu.cn)',
    '河南大学 未来技术学院 (future.henu.edu.cn) —— 确认为量子信息方向，不相关',
    '省部共建作物逆境适应与改良国家重点实验室/棉花楼 (csai.henu.edu.cn)',
    '植物逆境生物学重点实验室 (cps.henu.edu.cn)',
    '河南大学 农学院 (nxy.henu.edu.cn)',
    '河南大学 医学院 (med.henu.edu.cn)',
    '国际生物医学联合创新中心 JCBI (bs.henu.edu.cn)',
    '中州实验室 (zzlab.henu.edu.cn)',
    '河南大学 数学与统计学院',
]:
    bullet(doc, item)

heading(doc, '2.2 筛选标准', 2)
for i, item in enumerate([
    '近3年研究方向与蛋白质设计、AI蛋白质预测、蛋白质工程直接相关；',
    '具有蛋白质计算模拟（分子动力学、分子对接）或生物信息学背景；',
    '从事合成生物学、蛋白质组学相关研究；',
    '具备指导本科生/研究生竞赛的经验或意愿；',
    '有相关国家级/省部级科研项目支撑，能够提供实验或计算资源。',
]):
    numbered(doc, i+1, item)

doc.add_page_break()

# ═══════════════════════ 三、重点推荐 ═══════════════════════
heading(doc, '三、重点推荐导师（按优先级排序）', 1)
para(doc, '根据项目核心技术需求（AI蛋白质预测 + 蛋白质设计 + 合成生物学），将检索到的导师分为三个优先级梯队。第一梯队为研究方向与项目高度匹配的导师，强烈建议优先联系。')

# ── 3.1 ──
heading(doc, '3.1 第一梯队：AI+蛋白质方向（最高优先级 ★★★★★）', 2)
para(doc, '本梯队导师的研究方向与项目核心技术（AI驱动的蛋白质预测与设计）高度吻合，能够提供最直接的学术指导和技术支持。')

heading(doc, '【1】杨伟 — 蛋白质二级结构预测 + 深度学习（最匹配·首选）', 3)
make_info_table(doc, [
    ('姓名', '杨伟'),
    ('职称/身份', '副教授、硕士生导师'),
    ('所在学院', '计算机与信息工程学院（数据科学系）'),
    ('研究方向', '深度学习、度量学习、计算生物学、蛋白质二级结构预测'),
    ('核心成果', '① 2022年在Knowledge-Based Systems（SCI一区，IF=8.8）发表基于深度度量学习的蛋白质二级结构预测方法（DML_SS），代码已开源；\n② 同年同刊发表轻量级卷积网络+标签分布感知损失方法（ShuffleNet_SS）；\n③ 使用预训练蛋白质语言模型ProtT5-XL-U50进行特征提取，与项目ESM思路高度一致；\n④ 拥有蛋白质二级结构预测相关国家发明专利。'),
    ('教育背景', '博士：哈尔滨工业大学（计算机应用技术，2011年）\n访学：香港理工大学'),
    ('在研项目', '国家自然科学基金项目、河南省科技攻关项目'),
    ('匹配理由', '1. 研究方向与项目技术栈高度吻合——使用预训练蛋白质语言模型进行特征提取后做预测；\n2. 具备深度学习+蛋白质的交叉背景，可为ESM/SaProt模型优化提供直接指导；\n3. 硕士生导师，有指导学生的经验；\n4. 代码开源（GitHub），工程能力强，能提供实战级别的技术建议。'),
    ('邮箱', 'yang0sun@gmail.com'),
    ('个人主页', 'http://cs.henu.edu.cn/info/1050/1887.htm'),
    ('所在校区', '开封·金明校区'),
    ('优先级', '★★★★★（最高优先级·首选）'),
    ('建议联系方式', '邮件联系，附项目介绍和代码仓库链接'),
])

heading(doc, '【2】刘翘铭 — 深度学习 + 生物信息学', 3)
make_info_table(doc, [
    ('姓名', '刘翘铭'),
    ('职称/身份', '河南大学"百人计划"B岗'),
    ('所在学院', '人工智能学院（郑州校区）'),
    ('研究方向', '深度学习、生物信息学（蛋白质相关）'),
    ('学术兼职', '中国计算机学会（CCF）生物信息学专业委员会委员'),
    ('核心成果', '以第一/通讯作者在Briefings in Bioinformatics、IEEE JHBI、IEEE TCBB、BMC Biology、IEEE BIBM等权威期刊/会议发表多篇论文。'),
    ('在研项目', '国家自然科学基金1项、博士后面上项目1项、河南省青年基金1项\n入选2023年度国家博士后资助计划'),
    ('匹配理由', '1. CCF生物信息学专委，方向完全对口的青年学者；\n2. 2025年3月刚入职，正在积极招研究生，联系成功率高；\n3. 深度学习+生物信息学双栖，可指导模型优化和生信分析；\n4. 郑州校区，未来可能在郑州发展。'),
    ('邮箱', 'cslqm@henu.edu.cn'),
    ('个人主页', 'http://ai.henu.edu.cn/old/info/1855/29472.htm'),
    ('所在校区', '郑州校区'),
    ('优先级', '★★★★★（最高优先级）'),
])

heading(doc, '【3】汪欣 — 计算模拟 + 机器学习 + 蛋白质', 3)
make_info_table(doc, [
    ('姓名', '汪欣'),
    ('职称/身份', '特聘教授（黄河学者）、硕士生导师'),
    ('所在学院/实验室', '国际生物医学联合创新中心（JCBI）'),
    ('研究方向', '多维度计算模拟、机器学习、蛋白质-小分子相互作用、药物设计'),
    ('技术手段', '虚拟筛选、分子对接、量子化学计算、QM/MM、分子动力学模拟、从头算分子动力学、机器学习'),
    ('核心成果', '以共同一作在Nature发表论文1篇；累计SCI论文30余篇。'),
    ('匹配理由', '1. 掌握分子动力学模拟、QM/MM等蛋白质结构模拟方法，可补充项目在结构层面的分析；\n2. 使用机器学习+多尺度模拟策略，与项目AI+蛋白质思路互补；\n3. Nature级别发表经验，学术水平有保障；\n4. 特聘教授，有国际视野和合作资源。'),
    ('邮箱', 'wx@henu.edu.cn'),
    ('个人主页', 'https://bs.henu.edu.cn/info/1030/2368.htm'),
    ('所在校区', '开封·金明校区'),
    ('优先级', '★★★★★（最高优先级）'),
])

doc.add_page_break()

# ── 3.2 ──
heading(doc, '3.2 第二梯队：蛋白质工程与设计方向（高优先级 ★★★★）', 2)
para(doc, '本梯队导师从事蛋白质理性设计、蛋白质工程和计算生物学研究，能为项目提供蛋白质改造的实验指导和设计策略。')

heading(doc, '【4】李华 — 蛋白质理性设计改造 + 定向进化', 3)
make_info_table(doc, [
    ('姓名', '李华'),
    ('职称/身份', '博士、副教授、硕士生导师\n生物工程系副主任、生物工程实验室主任'),
    ('所在学院', '生命科学学院'),
    ('研究方向', '蛋白质理性设计改造、定向进化、工业微生物理性改造、系统代谢工程'),
    ('核心成果', '在Frontiers in Microbiology等发表蛋白质定向进化研究成果；针对多肽类药物、高丝氨酸等生产瓶颈进行蛋白质理性设计。'),
    ('匹配理由', '1. 蛋白质理性设计改造是最直接的匹配关键词；\n2. 定向进化与项目中的突变体筛选直接相关；\n3. 有生物工程实验室，可提供实验验证条件；\n4. 硕士生导师，有指导学生经验。'),
    ('邮箱', 'lihua@henu.edu.cn'),
    ('个人主页', 'https://bio.henu.edu.cn/info/1007/4173.htm'),
    ('所在校区', '开封·金明校区'),
    ('优先级', '★★★★（高优先级）'),
])

heading(doc, '【5】张戈 — AI + 生物信息学', 3)
make_info_table(doc, [
    ('姓名', '张戈'),
    ('职称/身份', '博士、副教授、硕士生导师'),
    ('所在学院', '计算机与信息工程学院'),
    ('研究方向', '人工智能、生物信息学（药物协同预测、癌症诊断、多组学）'),
    ('核心成果', 'KGANSynergy（药物协同预测，Briefings in Bioinformatics, 2023）；MultiGATAE（癌症亚型识别）；拥有胃癌生物标志物识别国家发明专利。'),
    ('技术手段', '图神经网络（GNN）、注意力机制、特征选择、知识图谱'),
    ('匹配理由', '1. AI+生物信息学，可提供GNN等深度学习模型引入建议；\n2. 药物-蛋白质相互作用研究与蛋白质功能预测相通；\n3. 硕士生导师，有指导学生经验。'),
    ('邮箱', 'zhangge@henu.edu.cn'),
    ('个人主页', 'https://cs.henu.edu.cn/info/1267/5989.htm'),
    ('所在校区', '开封·金明校区'),
    ('优先级', '★★★★（高优先级）'),
])

heading(doc, '【6】陈欣 — 计算化学 + 生物信息学 + 蛋白质模拟', 3)
make_info_table(doc, [
    ('姓名', '陈欣'),
    ('职称/身份', '博士、副教授、硕士生导师'),
    ('所在学院', '化学与分子科学学院'),
    ('研究方向', '计算化学、生物信息学、蛋白质在材料表面的吸附模拟、拉伸分子动力学（SMD）'),
    ('核心成果', '使用SMD探索HIV-1蛋白酶动态结构；研究IGFBPs与IGF相互作用机理；讲授《分子模拟的理论与实践》课程。'),
    ('在研项目', '国家自然科学基金青年基金：糖尿病相关的IGFBPs与IGF相互作用机理研究'),
    ('匹配理由', '1. 蛋白质分子动力学模拟经验丰富，可辅助蛋白质结构稳定性分析；\n2. 生物信息学+计算化学双重背景；\n3. 讲授分子模拟课程，理论基础扎实。'),
    ('邮箱', 'xin_chen@henu.edu.cn'),
    ('个人主页', 'http://ccce.henu.edu.cn/info/1028/1057.htm'),
    ('所在校区', '开封·金明校区'),
    ('优先级', '★★★★（高优先级）'),
])

doc.add_page_break()

# ── 3.3 ──
heading(doc, '3.3 第三梯队：合成生物学与计算模拟方向（中优先级 ★★★）', 2)
para(doc, '本梯队导师从事合成生物学、蛋白质组学和大规模计算模拟研究，能为项目提供合成生物学竞赛策略指导和实验平台资源。')

tier3 = [
    ('【7】张立新', [
        ('职称/身份', '教授、博士生导师（国家杰青、中原学者）'),
        ('所在学院/实验室', '生命科学学院（院长）、河南省合成生物与生物制造重点实验室（主任）'),
        ('研究方向', '光合膜色素蛋白复合物生成与动态调节、光合作用光高效利用、合成生物学'),
        ('代表性成果', 'Cell, Nature Plants, PNAS, Plant Cell等发表60余篇；首次提出相分离驱动叶绿体内蛋白分选新机制（Cell）。'),
        ('匹配理由', '作为院长和合成生物重点实验室主任，可提供顶层学术资源和竞赛平台支持；但个人研究方向偏向植物光合作用，与AI蛋白质设计有一定距离。'),
        ('邮箱', '（建议通过学院办公室联系）'),
        ('个人主页', 'https://bio.henu.edu.cn/'),
    ]),
    ('【8】王强', [
        ('职称/身份', '教授'),
        ('所在学院/实验室', '作物逆境适应与改良国家重点实验室（CSAI）'),
        ('研究方向', '合成生物学、羧酶体蛋白质外壳、自组装纳米蛋白笼、纳米生物反应器'),
        ('代表性成果', 'ACS Nano (2024), Nature Communications (2020)；开发基于羧酶体外壳的可控自组装纳米蛋白笼工具包。'),
        ('匹配理由', '合成生物学+蛋白质工程的完美结合；蛋白质自组装方向与蛋白质设计互补；可提供合成生物学竞赛的整体策略指导。'),
        ('邮箱', '（可通过CSAI网站联系）'),
        ('个人主页', 'https://csai.henu.edu.cn/info/1043/2634.htm'),
    ]),
    ('【9】吕丹丹', [
        ('职称/身份', '教授、硕士生导师（河南大学特聘教授）'),
        ('所在学院/实验室', '作物逆境适应与改良国家重点实验室（CSAI）'),
        ('研究方向', '光能高效利用与合成生物学、蛋白质组学'),
        ('代表性成果', '主持国家重点研发计划"合成生物学"重点专项子课题（100万元，2020-2025）；德国马普所博士；Nature Communications, Proteomics等发表论文。'),
        ('匹配理由', '合成生物学+蛋白质组学双重背景；主持合成生物学专项，可提供竞赛策略指导。'),
        ('邮箱', 'ludandan@henu.edu.cn'),
        ('个人主页', 'http://csai.henu.edu.cn/info/1045/1985.htm'),
    ]),
    ('【10】白仲虎', [
        ('职称/身份', 'II级教授、博士生导师（英国皇家生物学会会士FRSB）'),
        ('所在学院', '生命科学学院'),
        ('研究方向', '生物过程工程和合成生物学、疫苗工程'),
        ('代表性成果', 'Nature Catalysis, Metabolic Engineering等顶刊50余篇SCI论文。'),
        ('匹配理由', '合成生物学领域资深专家；国际视野和产业经验丰富。'),
        ('邮箱', '（讲座教授，建议通过学院联系）'),
        ('个人主页', 'http://bio.henu.edu.cn/info/1033/10197.htm'),
    ]),
    ('【11】洪军', [
        ('职称/身份', '教授、河南省生物物理学会理事'),
        ('所在学院', '生命科学学院'),
        ('研究方向', '自组装人工酶构建与性能调控、蛋白质聚集与保护机制、纳米生物电化学'),
        ('代表性成果', '河南省教育厅学术技术带头人。'),
        ('匹配理由', '人工酶构建与蛋白质性能调控直接相关；蛋白质聚集研究可辅助稳定性工程。'),
        ('邮箱', 'hongjun@henu.edu.cn'),
        ('个人主页', 'https://bio.henu.edu.cn/info/1007/4816.htm'),
    ]),
    ('【12】赵媛', [
        ('职称/身份', '教授、博士生导师（河南省优青）'),
        ('所在学院', '医学院'),
        ('研究方向', '理论与计算药物化学/生物学、酶催化机制理论计算、分子动力学模拟、蛋白质修饰'),
        ('代表性成果', 'Chemical Science, ACS Catalysis等20余篇；主持国家自然科学基金面上项目。'),
        ('匹配理由', '分子动力学模拟+酶催化机制计算，可指导蛋白质稳定性计算分析。'),
        ('邮箱', 'zhaoyuan@henu.edu.cn'),
        ('个人主页', 'https://med.henu.edu.cn/info/2040/22288.htm'),
    ]),
]

for title, data in tier3:
    heading(doc, title, 3)
    make_info_table(doc, data)

# 第三梯队补充名单（简要）
heading(doc, '第三梯队补充名单（简要）', 3)
tier3_extra = [
    ('黄继红', '教授/博导/中原学者', '农学院', '合成生物学、无细胞蛋白质合成、体外定向进化', 'huangjih1216@126.com'),
    ('刘宇鹏', '教授/硕导', '生命科学学院', '蛋白质工程、生物催化', '（院办联系）'),
    ('卢福浩', '特聘教授/PI', 'CSAI国重实验室', '生物信息学、基因组学/转录组学', 'lufuhao@henu.edu.cn'),
    ('王沛', '教授/博导/省特聘', '数学与统计学院', '计算系统生物学、生物统计', 'wangpei@henu.edu.cn'),
    ('杨路', '黄河学者/校聘教授', '中州实验室', '蛋白质组学、代谢组学', 'yangluhnu@henu.edu.cn'),
    ('孙扬', '特聘教授/硕导', '生命科学学院', '光合作用合成生物学', 'sunyy@henu.edu.cn'),
    ('徐秀美', '教授（黄河学者）', 'CSAI国重实验室', '光合蛋白复合物、合成生物学', '—'),
    ('胡筑兵', '教授', 'CSAI国重实验室', '合成生物学、固碳模块', '—'),
    ('顾凯旋', '博士/硕导/中原青年拔尖', '化学与分子科学学院', '量子动力学模拟、催化机理', 'kxgu@henu.edu.cn'),
    ('张庆友', '教授/博导', '化学与分子科学学院', '化学信息学、机器学习', 'zhqingyou@henu.edu.cn'),
    ('时玉', '教授/博导/黄河学者', '生命科学学院', '微生物组大数据、生物信息学', '—'),
    ('徐小冬', '教授/博导', '生命科学学院', '生物钟、蛋白质组学', '—'),
]

t = doc.add_table(rows=len(tier3_extra)+1, cols=5, style='Table Grid')
t.alignment = WD_TABLE_ALIGNMENT.CENTER
for j, h in enumerate(['姓名', '职称', '学院/实验室', '研究方向', '联系方式']):
    scf(t.cell(0, j), h, bold=True, size=Pt(9), align='center')
shade_row(t.rows[0])
for i, row_data in enumerate(tier3_extra):
    for j, val in enumerate(row_data):
        scf(t.cell(i+1, j), val, size=Pt(8.5), align='center' if j in [0, 4] else 'left')

doc.add_page_break()

# ═══════════════════════ 四、完整名录 ═══════════════════════
heading(doc, '四、各学院/实验室完整导师名录表', 1)
para(doc, '以下按学院/实验室列出所有本次检索到的24位相关导师汇总信息。')

all_p = [
    ('杨伟', '副教授/硕导', '计算机与信息工程学院', '深度学习、蛋白质二级结构预测', '★★★★★', 'yang0sun@gmail.com'),
    ('刘翘铭', '百人计划B岗', '人工智能学院', '深度学习、生物信息学', '★★★★★', 'cslqm@henu.edu.cn'),
    ('汪欣', '特聘教授/硕导', 'JCBI', '计算模拟、ML、蛋白质-小分子', '★★★★★', 'wx@henu.edu.cn'),
    ('李华', '副教授/硕导', '生命科学学院', '蛋白质理性设计、定向进化', '★★★★', 'lihua@henu.edu.cn'),
    ('张戈', '副教授/硕导', '计算机与信息工程学院', 'AI、生物信息学', '★★★★', 'zhangge@henu.edu.cn'),
    ('陈欣', '副教授/硕导', '化学与分子科学学院', '计算化学、生物信息学、蛋白MD', '★★★★', 'xin_chen@henu.edu.cn'),
    ('张立新', '教授/博导/杰青', '生命科学学院', '光合膜蛋白、合成生物学', '★★★', '（院办联系）'),
    ('王强', '教授', 'CSAI国重实验室', '合成生物学、蛋白质纳米笼', '★★★', '（CSAI网站）'),
    ('吕丹丹', '教授/硕导', 'CSAI国重实验室', '合成生物学、蛋白质组学', '★★★', 'ludandan@henu.edu.cn'),
    ('白仲虎', 'II级教授/博导', '生命科学学院', '合成生物学、生物过程工程', '★★★', '（院办联系）'),
    ('洪军', '教授', '生命科学学院', '人工酶构建、蛋白质聚集', '★★★', 'hongjun@henu.edu.cn'),
    ('赵媛', '教授/博导', '医学院', '计算化学、MD、蛋白质修饰', '★★★', 'zhaoyuan@henu.edu.cn'),
    ('黄继红', '教授/博导/中原学者', '农学院', '合成生物学、无细胞蛋白合成', '★★', 'huangjih1216@126.com'),
    ('刘宇鹏', '教授/硕导', '生命科学学院', '蛋白质工程、生物催化', '★★', '（院办联系）'),
    ('卢福浩', '特聘教授/PI', 'CSAI国重实验室', '生物信息学、基因组学', '★★', 'lufuhao@henu.edu.cn'),
    ('王沛', '教授/博导/省特聘', '数学与统计学院', '计算系统生物学、生物统计', '★★', 'wangpei@henu.edu.cn'),
    ('杨路', '黄河学者/校聘教授', '中州实验室', '蛋白质组学、代谢组学', '★★', 'yangluhnu@henu.edu.cn'),
    ('孙扬', '特聘教授/硕导', '生命科学学院', '光合作用合成生物学', '★★', 'sunyy@henu.edu.cn'),
    ('徐秀美', '教授（黄河学者）', 'CSAI国重实验室', '光合蛋白复合物、合成生物学', '★★', '—'),
    ('胡筑兵', '教授', 'CSAI国重实验室', '合成生物学、固碳模块', '★★', '—'),
    ('顾凯旋', '博士/硕导', '化学与分子科学学院', '量子动力学模拟、催化机理', '★★', 'kxgu@henu.edu.cn'),
    ('张庆友', '教授/博导', '化学与分子科学学院', '化学信息学、机器学习', '★★', 'zhqingyou@henu.edu.cn'),
    ('时玉', '教授/博导/黄河学者', '生命科学学院', '微生物组、生物信息学', '★', '—'),
    ('徐小冬', '教授/博导', '生命科学学院', '生物钟、蛋白质组学', '★', '—'),
]

t2 = doc.add_table(rows=len(all_p)+1, cols=6, style='Table Grid')
t2.alignment = WD_TABLE_ALIGNMENT.CENTER
for j, h in enumerate(['姓名', '职称', '学院/实验室', '研究方向', '匹配度', '联系方式']):
    scf(t2.cell(0, j), h, bold=True, size=Pt(9), align='center')
shade_row(t2.rows[0])
for i, row_data in enumerate(all_p):
    for j, val in enumerate(row_data):
        scf(t2.cell(i+1, j), val, size=Pt(8.5), align='center' if j in [0, 4] else 'left')

doc.add_page_break()

# ═══════════════════════ 五、评估矩阵 ═══════════════════════
heading(doc, '五、导师匹配度综合评估矩阵', 1)
para(doc, '下表从5个维度对第一梯队和第二梯队的6位核心导师进行定量评估（每项1-5分，总分25分）。')

para(doc, '评估维度说明：', bold=True)
for d in [
    'AI技术匹配：导师研究中使用AI/ML方法与项目技术路线的契合程度；',
    '蛋白质领域匹配：导师在蛋白质结构/功能/设计方面的研究深度；',
    '指导可行性：导师指导学生竞赛/科研的经验、可用时间和意愿评估；',
    '资源与平台：导师能提供的计算资源、实验平台、项目经费等支持；',
    '学术影响力：导师的学术水平、发表论文质量、项目主持经验。',
]:
    bullet(doc, d, size=Pt(10))

matrix = [
    ('杨伟', '5', '5', '4', '3', '4', '21', '最优选择：技术栈完全匹配'),
    ('刘翘铭', '5', '4', '5', '3', '3', '20', '新入职，联系成功率高'),
    ('汪欣', '4', '5', '3', '4', '5', '21', '学术水平最高（Nature）'),
    ('李华', '2', '5', '4', '5', '3', '19', '实验资源最丰富'),
    ('张戈', '5', '3', '4', '3', '3', '18', 'AI方法学强'),
    ('陈欣', '3', '4', '3', '4', '3', '17', '计算模拟专长'),
]

t3 = doc.add_table(rows=len(matrix)+1, cols=8, style='Table Grid')
t3.alignment = WD_TABLE_ALIGNMENT.CENTER
for j, h in enumerate(['导师', 'AI技术\n(5)', '蛋白质\n(5)', '指导\n(5)', '资源\n(5)', '影响力\n(5)', '总分\n(25)', '备注']):
    scf(t3.cell(0, j), h, bold=True, size=Pt(8.5), align='center')
shade_row(t3.rows[0])
for i, row_data in enumerate(matrix):
    for j, val in enumerate(row_data):
        scf(t3.cell(i+1, j), val, bold=(j == 0), size=Pt(8.5), align='center' if j < 7 else 'left')

para(doc, '')
para(doc, '推荐策略：优先联系杨伟老师（技术栈最匹配），同时可联系刘翘铭和汪欣老师（不同方向的互补指导）。李华老师作为实验验证资源的重要补充。')

doc.add_page_break()

# ═══════════════════════ 六、联系方案 ═══════════════════════
heading(doc, '六、联系方案与话术模板', 1)

heading(doc, '6.1 联系策略与时机', 2)
para(doc, '建议按照"梯队优先序"分批次联系，每批次间隔3-5天，根据回复情况调整后续策略。')
for i, s in enumerate([
    '第一批次（Day 1-3）：同时联系第一梯队的杨伟、刘翘铭、汪欣三位导师。三位导师方向互补（蛋白质预测+生物信息学+计算模拟），可根据回复选择最合适的导师。',
    '第二批次（Day 5-7）：若第一梯队未获积极回复，联系李华、张戈、陈欣三位导师。',
    '第三批次（Day 8-10）：联系第三梯队的张立新院长、王强教授、吕丹丹教授等，寻求更广泛的指导资源。',
    '平行策略：也可同时联系不同梯队的导师（因指导内容不冲突），但需对每位导师保持信息透明。',
]):
    numbered(doc, i+1, s)

heading(doc, '6.2 邮件模板（首选·联系杨伟老师）', 2)

template1 = '''尊敬的杨老师：

您好！我是河南大学XX学院的本科生XX，目前正在参加2026年合成生物学创新赛，冒昧给您发邮件，希望能就我们的参赛项目得到您的指导。

我们的项目主题是"基于AI的GFP蛋白质亮度预测与设计"。技术路线：使用蛋白质语言模型（ESM-2/SaProt）对avGFP突变体序列进行嵌入编码，结合随机森林回归模型预测突变体的荧光亮度，并在此基础上生成和筛选高亮度候选序列。目前已完成了数据处理、模型训练和候选序列生成的完整Pipeline（代码见GitHub：XXX）。

我注意到您在Knowledge-Based Systems上发表的基于深度度量学习的蛋白质二级结构预测研究（DML_SS），其中使用ProtT5-XL-U50预训练蛋白质语言模型进行特征提取的思路，与我们的方法非常相似。您在该领域深厚的积累和工程经验，对我们优化模型、提升预测精度将具有极大的指导意义。

具体而言，我们希望在以下方面得到您的指导：
1. 蛋白质语言模型（ESM/SaProt）嵌入特征的最优利用策略；
2. 能否将深度学习方法直接引入预测Pipeline替代随机森林；
3. 模型可解释性分析——如何理解哪些序列特征决定了亮度变化。

如果您方便的话，我们非常希望能有机会与您进行一次简短的交流（线下或线上均可，约15-20分钟），听取您的宝贵建议。

附件中是我们项目的详细介绍和当前进展。非常感谢您在百忙之中阅读这封邮件！

祝您工作顺利！

学生：XXX
邮箱：XXX
日期：2026年X月X日'''

para(doc, template1, size=Pt(9.5))

heading(doc, '6.3 邮件模板（联系刘翘铭老师）', 2)

template2 = '''尊敬的刘老师：

您好！我是河南大学XX学院的本科生XX，正在参加2026年合成生物学创新赛。我关注到您今年刚加入河南大学人工智能学院，并且是CCF生物信息学专业委员会委员，研究方向是深度学习与生物信息学。

我们的参赛项目处于AI与蛋白质科学的交叉点——使用蛋白质语言模型（ESM-2/SaProt）进行GFP突变体序列的嵌入编码，通过机器学习预测荧光亮度，并生成高亮度候选序列。项目已完成完整的计算Pipeline，目前正在寻求AI方法学的优化指导和竞赛策略建议。

您在国际生物信息学权威期刊（Briefings in Bioinformatics等）上的深度学习研究成果，让我们非常期待能得到您的指导。特别希望向您请教深度学习模型在蛋白质序列分析中的最佳实践，以及如何更科学地设计计算实验来验证模型的有效性。

如果您有时间，我们非常希望有机会向您汇报项目进展并听取您的建议。感谢您！

祝您科研顺利！

学生：XXX
邮箱：XXX
日期：2026年X月X日'''

para(doc, template2, size=Pt(9.5))

heading(doc, '6.4 通用联系话术要点', 2)
for i, pt in enumerate([
    '开门见山：第一段直接说明来意（参赛+项目主题+为什么联系这位老师）；',
    '展示了解：引用导师的具体研究成果，证明你做过功课，不是群发邮件；',
    '明确需求：具体列出希望在哪些方面得到指导，让导师快速评估是否有能力帮助；',
    '降低门槛：建议"简短交流（15-20分钟）"而非"长期指导"，降低时间承诺压力；',
    '附带材料：准备好项目简介（1-2页PDF）、代码仓库链接、当前结果展示；',
    '跟进节奏：若5个工作日内未收到回复，可发送一次简短跟进邮件；若仍无回复则转向下一位导师；',
    '保持透明：如有多位导师同时指导（不同方向），应让每位导师知晓，避免误会。',
]):
    numbered(doc, i+1, pt)

heading(doc, '6.5 拜访前的准备工作清单', 2)
for i, pt in enumerate([
    '准备项目摘要（1-2页PDF），包含：项目背景、技术路线、当前进展、核心结果图表；',
    '准备3-5个具体的技术问题（展示思考的深度，而非泛泛而问）；',
    '准备代码演示（Jupyter Notebook或Python脚本，能现场运行展示）；',
    '了解导师近3年的代表性论文（至少读摘要和结论）；',
    '准备纸笔记录导师的建议；',
    '准备简短自我介绍（30秒版本和2分钟版本）；',
    '确认导师办公室位置和拜访时间，提前5分钟到达。',
]):
    numbered(doc, i+1, pt)

heading(doc, '6.6 多种场景的话术准备', 2)

scenarios = [
    ('场景一：导师表示有兴趣但时间有限',
     '非常感谢您愿意了解我们的项目！我们只需要15-20分钟简单汇报一下，如果您觉得有值得深入的点，我们可以后续再约时间。我们也可以先把项目摘要发给您看看。'),
    ('场景二：导师建议联系其他老师',
     '感谢您的建议！能否请您帮忙推荐一两位在这个方向上更合适的老师？或者方便的话，您能否帮忙引荐一下？'),
    ('场景三：导师犹豫是否能够提供足够帮助',
     '您太谦虚了！其实我们最需要的是您在[具体方向]方面的经验指导，哪怕只是一些方向性的建议对我们来说都非常宝贵。'),
    ('场景四：导师同意指导',
     '非常感谢您！我们会尽快整理当前项目的详细材料和具体问题清单发给您，之后根据您的时间安排定期汇报进展。请问您prefer什么沟通方式（邮件/微信/线下见面）？'),
]

for scenario_title, script in scenarios:
    para(doc, f'【{scenario_title}】', bold=True, indent=False)
    para(doc, script, size=Pt(10.5))

doc.add_page_break()

# ═══════════════════════ 七、附录 ═══════════════════════
heading(doc, '七、附录：重要平台与资源', 1)

heading(doc, '7.1 河南大学相关科研平台', 2)
for name, desc in [
    ('河南省合成生物与生物制造重点实验室', '张立新教授担任主任，合成生物学方向核心平台'),
    ('省部共建作物逆境适应与改良国家重点实验室（CSAI）', 'csai.henu.edu.cn，王强、吕丹丹等PI所在地，合成生物学专项依托平台'),
    ('棉花生物育种与综合利用全国重点实验室', '与CSAI联动，共享大型实验平台和计算资源'),
    ('植物逆境生物学重点实验室', 'cps.henu.edu.cn，卢福浩等生物信息学课题组所在地'),
    ('国际生物医学联合创新中心（JCBI）', 'bs.henu.edu.cn，汪欣教授所在地，计算模拟方向'),
    ('中州实验室', 'zzlab.henu.edu.cn，杨路等蛋白质组学方向平台'),
    ('河南省应用微生物工程研究中心', '刘宇鹏教授所在平台，蛋白质工程方向'),
    ('河南省人工智能理论及算法工程研究中心', '杨晓慧教授负责，AI算法方向'),
]:
    bullet(doc, f'{name}：{desc}')

heading(doc, '7.2 项目相关重要参考文献', 2)
for ref in [
    'Yang W, et al. "Deep metric learning for accurate protein secondary structure prediction." Knowledge-Based Systems, 2022. (杨伟老师代表作，与项目方法密切相关)',
    'Yang W, et al. "Protein secondary structure prediction using a lightweight convolutional network and label distribution aware margin loss." Knowledge-Based Systems, 2022.',
    'Wang Q, et al. "Reprogramming bacterial protein organelles as a nanoreactor for hydrogen production." Nature Communications, 2020. (王强教授代表作)',
    'Zhang G, et al. "KGANSynergy: knowledge graph attention network for drug synergy prediction." Briefings in Bioinformatics, 2023. (张戈老师代表作)',
    'Rives A, et al. "Biological structure and function emerge from scaling unsupervised learning to 250 million protein sequences." PNAS, 2021. (ESM模型原始论文)',
    'Su J, et al. "SaProt: Protein Language Modeling with Structure-aware Vocabulary." ICLR, 2024. (SaProt模型论文)',
    'Notin P, et al. "ProteinGym: Large-Scale Benchmarks for Protein Fitness Prediction and Design." NeurIPS, 2023.',
]:
    bullet(doc, ref, size=Pt(10))

heading(doc, '7.3 数据检索说明', 2)
para(doc, '本报告数据来源：')
for s in [
    '河南大学各学院/实验室官方网站（通过学校VPN访问）；',
    '网络学术搜索（Google Scholar、Web of Science、中国知网CNKI）；',
    '国家自然科学基金委员会项目数据库；',
    'CCF、CSIAM等学术组织成员数据库；',
    'ResearchGate、Google Scholar学者个人页面。',
]:
    bullet(doc, s)

para(doc, '')
para(doc, f'检索日期：{datetime.date.today().strftime("%Y年%m月%d日")}')
para(doc, '声明：本报告中的导师信息来源于公开网页和学术数据库，仅供内部参考。导师的招生状态和联系方式可能随时间变化，建议在联系前通过学院官网确认最新信息。部分导师的联系方式未能公开检索到，标注为"（院办联系）"或"（CSAI网站）"。')

# ── 保存 ──
output_path = r"C:\Users\fengs\Desktop\蛋白质设计-合成生物学创新赛-Claude\河南大学导师匹配分析报告_2026ProteinDesign.docx"
doc.save(output_path)
print(f"Done: {output_path}")
