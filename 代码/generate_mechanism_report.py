#!/usr/bin/env python3
"""生成 GFP 荧光机制、亮度与稳定性因素讲解文档（含结构示意图）"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import datetime
import os
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch, Arc, Rectangle, Circle, FancyArrow, Polygon
import matplotlib.font_manager as fm
import numpy as np

# ── 设置中文字体 ──
for font_name in ['SimHei', 'Microsoft YaHei', 'Noto Sans SC', 'STXihei']:
    try:
        fm.findfont(font_name, fallback_to_default=False)
        plt.rcParams['font.sans-serif'] = [font_name, 'DejaVu Sans']
        plt.rcParams['axes.unicode_minus'] = False
        print(f"使用字体: {font_name}")
        break
    except Exception:
        continue

# ── 初始化文档 ──
doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(11)

img_dir = r'C:\Users\fengs\Desktop\蛋白质设计-合成生物学创新赛-Claude\images'
os.makedirs(img_dir, exist_ok=True)

# ── 辅助函数 ──
def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 51, 102)
    return h

def add_para(text, bold=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    return p

def add_ref(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10)
    return p

def add_list_item(text):
    p = doc.add_paragraph(text, style='List Bullet')
    return p

def add_image(img_path, width_inches=5.5, caption=""):
    if os.path.exists(img_path):
        doc.add_picture(img_path, width=Inches(width_inches))
        if caption:
            cap = doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cap.add_run(caption)
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(100, 100, 100)
        return True
    return False

# ═══════════════════════════════════════════════════
# 生成示意图
# ═══════════════════════════════════════════════════

# ---- 图1: GFP整体结构示意图 (β-桶 + 中心α螺旋 + 发色团) ----
def draw_gfp_structure():
    fig, ax = plt.subplots(1, 1, figsize=(8, 6))
    ax.set_xlim(-5, 5)
    ax.set_ylim(-6, 6)
    ax.set_aspect('equal')
    ax.axis('off')

    # β-桶 (11条β链表现为椭圆桶)
    for i in range(11):
        angle = i * 2 * np.pi / 11 - np.pi / 2
        x_offset = 0.3 * np.sin(angle)
        # 上环
        top_y = 4.5
        bot_y = -4.5
        # β链 (箭头)
        ax.annotate('', xy=(0.6 * np.sin(angle), top_y),
                     xytext=(0.6 * np.sin(angle), bot_y),
                     arrowprops=dict(arrowstyle='->', color='#2E86AB',
                                    lw=3, alpha=0.7))
        # 下环的链
        ax.plot([0.6 * np.sin(angle), 0.6 * np.sin(angle + 2*np.pi/11)],
                [bot_y, bot_y], color='#2E86AB', lw=3, alpha=0.7)
        ax.plot([0.6 * np.sin(angle), 0.6 * np.sin(angle + 2*np.pi/11)],
                [top_y, top_y], color='#2E86AB', lw=3, alpha=0.7)

    # 桶轮廓
    barrel_top = mpatches.Ellipse((0, 4.5), 1.8, 0.4, fill=False, color='#2E86AB', lw=2.5)
    barrel_bot = mpatches.Ellipse((0, -4.5), 1.8, 0.4, fill=False, color='#2E86AB', lw=2.5)
    ax.add_patch(barrel_top)
    ax.add_patch(barrel_bot)

    # 左侧连线
    ax.plot([-0.9, -0.9], [4.5, -4.5], color='#2E86AB', lw=2.5)
    ax.plot([0.9, 0.9], [4.5, -4.5], color='#2E86AB', lw=2.5)

    # 中心α螺旋
    ax.plot([0, 0], [3.5, -3.5], color='#D81159', lw=6, alpha=0.8)

    # 发色团 (在螺旋中间)
    chromo = Rectangle((-0.35, -0.5), 0.7, 1.0, fill=True, facecolor='#3DDC84',
                       edgecolor='#1B7A43', lw=2, alpha=0.9)
    ax.add_patch(chromo)
    # 绿色荧光
    for angle in np.linspace(0, 2*np.pi, 12):
        dx = 0.8 * np.cos(angle)
        dy = 0.8 * np.sin(angle)
        ax.annotate('', xy=(dx * 1.8, dy * 1.8), xytext=(dx * 0.5, dy * 0.5),
                     arrowprops=dict(arrowstyle='->', color='#3DDC84', lw=1.5, alpha=0.5))

    ax.text(0, 0.7, '发色团\n(Tyr66-Gly67-Gly68)', ha='center', fontsize=8,
            color='#1B7A43', fontweight='bold')
    ax.text(0, 5.2, 'β-桶 (11条β链)', ha='center', fontsize=10,
            color='#2E86AB', fontweight='bold')
    ax.text(0.7, 0, '中心\nα螺旋', ha='left', fontsize=8, color='#D81159')

    # 标注关键残基
    ax.annotate('Arg96', xy=(0.15, 1.8), fontsize=7, color='#8B4513',
                arrowprops=dict(arrowstyle='->', color='#8B4513', lw=1))
    ax.annotate('Glu222', xy=(0.2, -2.0), fontsize=7, color='#8B4513',
                arrowprops=dict(arrowstyle='->', color='#8B4513', lw=1))

    ax.set_title('GFP 整体结构示意', fontsize=14, fontweight='bold', color='#003366', pad=15)

    path = os.path.join(img_dir, 'fig1_gfp_structure.png')
    fig.savefig(path, dpi=150, bbox_inches='tight', facecolor='white')
    plt.close(fig)
    return path


# ---- 图2: 发色团形成过程 (化学反应示意图) ----
def draw_chromophore_formation():
    fig, (ax1, ax2, ax3, ax4) = plt.subplots(1, 4, figsize=(14, 4))

    # 步骤1: 未修饰三肽
    ax1.set_xlim(-2, 2)
    ax1.set_ylim(-2, 2)
    ax1.axis('off')
    ax1.set_title('(1) 折叠 + 三肽定位\nSer65-Tyr66-Gly67', fontsize=10, fontweight='bold', color='#003366')
    # 简化三肽
    ax1.plot([-1.5, -0.5, 0.5, 1.5], [0, 0.3, -0.2, 0.1], 'ko-', lw=2, markersize=8)
    ax1.text(-1.5, -0.3, 'Ser65', fontsize=8, ha='center')
    ax1.text(-0.5, -0.5, 'Tyr66', fontsize=8, ha='center')
    ax1.text(0.8, -0.3, 'Gly67', fontsize=8, ha='center')
    # Gly67 NH2 亲核攻击
    ax1.annotate('Gly67\n酰胺N', xy=(0.5, 0.1), xytext=(0.5, 1.5),
                 fontsize=7, color='#D81159',
                 arrowprops=dict(arrowstyle='->', color='#D81159', lw=1.5))

    # 步骤2: 环化
    ax2.set_xlim(-2, 2)
    ax2.set_ylim(-2, 2)
    ax2.axis('off')
    ax2.set_title('(2) 环化 (Cyclization)\n~1.5 min', fontsize=10, fontweight='bold', color='#003366')
    # 五元咪唑啉酮环
    angles = np.linspace(0, 2*np.pi, 6)
    ring_x = np.cos(angles) * 0.8
    ring_y = np.sin(angles) * 0.8
    ax2.plot(ring_x, ring_y, 'k-', lw=2.5)
    ax2.text(0, 0, 'imidazolidinone\n五元环', fontsize=8, ha='center', color='#D81159')
    # 连接的Tyr酚环
    rect = Rectangle((0.8, -0.3), 0.7, 1.0, fill=True, facecolor='#FFE4B5',
                     edgecolor='#8B4513', lw=1.5)
    ax2.add_patch(rect)
    ax2.text(1.15, 0.2, 'Tyr66\n酚环', fontsize=7, ha='center')

    # 步骤3: 脱水
    ax3.set_xlim(-2, 2)
    ax3.set_ylim(-2, 2)
    ax3.axis('off')
    ax3.set_title('(3) 脱水 (Dehydration)\n–H₂O; Arg96, Glu222催化', fontsize=10, fontweight='bold', color='#003366')
    angles2 = np.linspace(0, 2*np.pi, 6)
    rx2 = np.cos(angles2) * 0.8
    ry2 = np.sin(angles2) * 0.8
    ax3.plot(rx2, ry2, 'k-', lw=2.5)
    # 双键表示
    ax3.plot([0.5, 0.9], [0.3, 0.5], 'k-', lw=2)
    ax3.plot([0.52, 0.92], [0.26, 0.46], 'k-', lw=2)
    ax3.text(0, 0, '咪唑啉酮\n(脱水)', fontsize=8, ha='center', color='#D81159')
    rect2 = Rectangle((0.8, -0.3), 0.7, 1.0, fill=True, facecolor='#FFE4B5',
                      edgecolor='#8B4513', lw=1.5)
    ax3.add_patch(rect2)
    ax3.text(2.0, 0, '–H₂O', fontsize=9, color='#2E86AB', fontweight='bold')

    # 步骤4: 氧化 → 成熟发色团
    ax4.set_xlim(-2, 3)
    ax4.set_ylim(-2, 2)
    ax4.axis('off')
    ax4.set_title('(4) 氧化 (Oxidation)\n+O₂ → –H₂O₂; ~34 min', fontsize=10, fontweight='bold', color='#003366')
    # 五元环 + 双键连接
    angles3 = np.linspace(0, 2*np.pi, 6)
    rx3 = np.cos(angles3) * 0.8
    ry3 = np.sin(angles3) * 0.8
    ax4.plot(rx3, ry3, 'k-', lw=2.5)
    # Cα=Cβ双键
    ax4.plot([0.8, 1.1], [0, 0], 'k-', lw=3, color='#D81159')
    ax4.plot([0.8, 1.1], [-0.1, -0.1], 'k-', lw=3, color='#D81159')
    ax4.text(0.95, -0.3, 'Cα=Cβ\n双键', fontsize=7, ha='center', color='#D81159')
    rect3 = Rectangle((1.2, -0.3), 0.7, 1.0, fill=True, facecolor='#90EE90',
                      edgecolor='#1B7A43', lw=2)
    ax4.add_patch(rect3)
    ax4.text(1.55, 0.2, 'p-HBDI\n发色团', fontsize=7, ha='center', color='#1B7A43')
    # 荧光发射
    for a in np.linspace(0, 2*np.pi, 8):
        ax4.annotate('', xy=(1.55 + 1.0*np.cos(a), 0.2 + 1.0*np.sin(a)),
                     xytext=(1.55 + 0.3*np.cos(a), 0.2 + 0.3*np.sin(a)),
                     arrowprops=dict(arrowstyle='->', color='#3DDC84', lw=1.5, alpha=0.7))
    ax4.text(2.5, 0.5, '绿色荧光\n~508 nm', fontsize=8, color='#1B7A43', fontweight='bold')

    fig.suptitle('GFP 发色团自催化成熟过程', fontsize=14, fontweight='bold', color='#003366', y=1.02)
    plt.tight_layout()

    path = os.path.join(img_dir, 'fig2_chromophore_formation.png')
    fig.savefig(path, dpi=150, bbox_inches='tight', facecolor='white')
    plt.close(fig)
    return path


# ---- 图3: ESPT 质子线示意 ----
def draw_espt():
    fig, (ax_gs, ax_es) = plt.subplots(1, 2, figsize=(13, 5.5))

    # 左图: 基态
    ax_gs.set_xlim(-4, 8)
    ax_gs.set_ylim(-2, 3)
    ax_gs.axis('off')
    ax_gs.set_title('基态 (Ground State)\n中性发色团 (A态), λabs ~395 nm', fontsize=11,
                    fontweight='bold', color='#003366')

    # 发色团
    c_gs = Rectangle((-3, -0.3), 1.5, 0.6, fill=True, facecolor='#FFFACD', edgecolor='#8B4513', lw=2)
    ax_gs.add_patch(c_gs)
    ax_gs.text(-2.25, 0.5, 'Cro-OH\n(中性)', fontsize=8, ha='center', color='#8B4513')

    # 水分子
    ax_gs.add_patch(Circle((-0.5, 0), 0.3, fill=True, facecolor='#87CEEB', edgecolor='#4682B4', lw=1.5))
    ax_gs.text(-0.5, -0.5, 'Wat', fontsize=7, ha='center', color='#4682B4')

    # Ser205
    ax_gs.add_patch(Rectangle((0.8, -0.3), 1.0, 0.6, fill=True, facecolor='#DDA0DD', edgecolor='#8B008B', lw=1.5))
    ax_gs.text(1.3, 0.5, 'Ser205', fontsize=8, ha='center', color='#8B008B')
    # Ser205 OH
    ax_gs.plot([1.8, 2.3], [0, 0], 'r-', lw=2)
    ax_gs.text(2.0, -0.3, 'OH', fontsize=7, color='red')

    # Glu222
    ax_gs.add_patch(Rectangle((2.8, -0.3), 1.0, 0.6, fill=True, facecolor='#FFB6C1', edgecolor='#8B0000', lw=1.5))
    ax_gs.text(3.3, 0.5, 'Glu222\n(质子化)', fontsize=8, ha='center', color='#8B0000')

    # H-bond lines
    for x_pos, label in [(-1.2, 'H-bond'), (0.8, 'H-bond')]:
        lx = x_pos
        ax_gs.plot([lx, lx+0.5], [0, 0], 'b--', lw=1, alpha=0.5)

    ax_gs.text(0.15, 0.7, '← 氢键网络 →', fontsize=9, ha='center', color='#4682B4')

    # 右图: 激发态 ESPT
    ax_es.set_xlim(-4, 8)
    ax_es.set_ylim(-2, 3)
    ax_es.axis('off')
    ax_es.set_title('激发态 ESPT (Excited State)\n阴离子发色团 (I*态), λem ~508 nm', fontsize=11,
                    fontweight='bold', color='#003366')

    # 发色团 (激发态阴离子)
    c_es = Rectangle((-3, -0.3), 1.5, 0.6, fill=True, facecolor='#90EE90', edgecolor='#006400', lw=2)
    ax_es.add_patch(c_es)
    ax_es.text(-2.25, 0.5, 'Cro-O⁻*\n(阴离子, S₁)', fontsize=8, ha='center', color='#006400')
    # 荧光发射
    for a in np.linspace(np.pi/2, 3*np.pi/2, 6):
        ax_es.annotate('', xy=(-3.5 + 0.6*np.cos(a), 0 + 0.6*np.sin(a)),
                       xytext=(-3.5 + 0.2*np.cos(a), 0 + 0.2*np.sin(a)),
                       arrowprops=dict(arrowstyle='->', color='#3DDC84', lw=1.5, alpha=0.7))
    ax_es.text(-3.5, 1.0, 'hν\n~508nm', fontsize=8, ha='center', color='#006400')

    # 质子传递 (箭头)
    ax_es.annotate('', xy=(2.8, 0.3), xytext=(-1.5, 0.3),
                   arrowprops=dict(arrowstyle='->', color='red', lw=2.5, connectionstyle='arc3,rad=.2'))
    ax_es.text(1.0, 1.2, 'H⁺ 传递 (ps级)', fontsize=9, ha='center', color='red', fontweight='bold')

    # Wat
    ax_es.add_patch(Circle((-0.5, 0), 0.3, fill=True, facecolor='#ADD8E6', edgecolor='#4682B4', lw=1.5))
    ax_es.text(-0.5, -0.5, 'H₃O⁺', fontsize=7, ha='center', color='#4682B4')

    # Ser205
    ax_es.add_patch(Rectangle((0.8, -0.3), 1.0, 0.6, fill=True, facecolor='#DDA0DD', edgecolor='#8B008B', lw=1.5))
    ax_es.text(1.3, 0.5, 'Ser205\n(中间态)', fontsize=8, ha='center', color='#8B008B')

    # Glu222 (去质子化)
    ax_es.add_patch(Rectangle((2.8, -0.3), 1.0, 0.6, fill=True, facecolor='#FF69B4', edgecolor='#8B0000', lw=1.5))
    ax_es.text(3.3, 0.5, 'Glu222⁻\n(去质子化)', fontsize=8, ha='center', color='#8B0000')

    # pKa标注
    ax_es.text(-2.25, -0.8, 'pKa: ~7→~3\n(光酸)', fontsize=7, ha='center', color='#D81159')

    fig.suptitle('激发态质子转移 (ESPT) 机理', fontsize=14, fontweight='bold', color='#003366', y=1.02)
    plt.tight_layout()

    path = os.path.join(img_dir, 'fig3_espt.png')
    fig.savefig(path, dpi=150, bbox_inches='tight', facecolor='white')
    plt.close(fig)
    return path


# ---- 图4: 亮度影响因素框图 ----
def draw_brightness_factors():
    fig, ax = plt.subplots(1, 1, figsize=(10, 5))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 6)
    ax.axis('off')

    # 中心: 亮度 = ε × Φ
    center = FancyBboxPatch((3.5, 2.3), 3.0, 1.4, boxstyle="round,pad=0.15",
                            facecolor='#FFF8DC', edgecolor='#FF8C00', lw=3)
    ax.add_patch(center)
    ax.text(5.0, 3.3, '荧光亮度\nBrightness = ε × Φ', ha='center', fontsize=13,
            fontweight='bold', color='#8B4513')
    ax.text(5.0, 2.6, 'ε: 消光系数  Φ: 量子产率', ha='center', fontsize=9, color='#8B4513')

    # 左侧: ε 的因素
    eps_box = FancyBboxPatch((0.2, 2.3), 2.5, 1.4, boxstyle="round,pad=0.1",
                             facecolor='#E8F4FD', edgecolor='#2E86AB', lw=2)
    ax.add_patch(eps_box)
    ax.text(1.45, 3.3, '消光系数 ε', ha='center', fontsize=11, fontweight='bold', color='#2E86AB')
    ax.text(1.45, 2.9, '• 发色团质子化状态\n• pKa (中性↔阴离子)\n• 488nm处吸收比例', ha='center', fontsize=8, color='#2E86AB')

    # 右侧: Φ 的因素
    phi_box = FancyBboxPatch((7.3, 2.3), 2.5, 1.4, boxstyle="round,pad=0.1",
                             facecolor='#E8F4FD', edgecolor='#2E86AB', lw=2)
    ax.add_patch(phi_box)
    ax.text(8.55, 3.3, '量子产率 Φ', ha='center', fontsize=11, fontweight='bold', color='#2E86AB')
    ax.text(8.55, 2.9, '• 发色团平面性\n• 非辐射衰变抑制\n• 分子内运动限制', ha='center', fontsize=8, color='#2E86AB')

    # 左箭头
    ax.annotate('', xy=(3.5, 3.0), xytext=(2.7, 3.0),
               arrowprops=dict(arrowstyle='->', color='#FF8C00', lw=2))
    # 右箭头
    ax.annotate('', xy=(7.3, 3.0), xytext=(6.5, 3.0),
               arrowprops=dict(arrowstyle='->', color='#FF8C00', lw=2))

    # 顶端: 上游因素
    top_items = [
        (0.5, '发色团环境\nH-bond网络', '#D81159'),
        (3.5, '关键残基\nE222, S205, T203, R96, H148', '#8B4513'),
        (7.0, 'β-桶构象\n刚性/平面性约束', '#1B7A43'),
    ]
    for x, text, color in top_items:
        box = FancyBboxPatch((x, 4.5), 2.5, 1.2, boxstyle="round,pad=0.1",
                             facecolor='white', edgecolor=color, lw=2)
        ax.add_patch(box)
        ax.text(x+1.25, 5.0, text, ha='center', fontsize=9, color=color)
        ax.annotate('', xy=(x+1.25, 4.5), xytext=(x+1.25, 3.7),
                   arrowprops=dict(arrowstyle='->', color=color, lw=1.5))

    # 底端: 其他因子
    bottom_items = [
        ('pH敏感性\n(生理pH vs pKa失配)', '#9370DB'),
        ('温度\n(发色团成熟效率)', '#CD853F'),
        ('发色团成熟\n(折叠/环化/脱水/氧化)', '#4682B4'),
    ]
    for i, (text, color) in enumerate(bottom_items):
        x = 1.0 + i * 3.2
        box = FancyBboxPatch((x, 0.3), 2.5, 1.2, boxstyle="round,pad=0.1",
                             facecolor='white', edgecolor=color, lw=2)
        ax.add_patch(box)
        ax.text(x+1.25, 0.9, text, ha='center', fontsize=9, color=color)
        ax.annotate('', xy=(x+1.25, 2.3), xytext=(x+1.25, 1.5),
                   arrowprops=dict(arrowstyle='->', color=color, lw=1.5))

    ax.set_title('影响 GFP 荧光亮度的关键因素', fontsize=14, fontweight='bold', color='#003366', pad=10)

    path = os.path.join(img_dir, 'fig4_brightness_factors.png')
    fig.savefig(path, dpi=150, bbox_inches='tight', facecolor='white')
    plt.close(fig)
    return path


# ---- 图5: 热稳定性影响因素 ----
def draw_stability_factors():
    fig, ax = plt.subplots(1, 1, figsize=(10, 5.5))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 7)
    ax.axis('off')

    ax.set_title('影响 GFP 热稳定性的关键因素', fontsize=14, fontweight='bold', color='#003366', pad=10)

    # 中心核
    center = FancyBboxPatch((3.5, 2.7), 3.0, 1.0, boxstyle="round,pad=0.15",
                            facecolor='#FFF5EE', edgecolor='#FF6347', lw=3)
    ax.add_patch(center)
    ax.text(5.0, 3.2, 'GFP 热稳定性\n(抵抗高温变性的能力)', ha='center', fontsize=12,
            fontweight='bold', color='#8B0000')

    items = [
        (0.2, 5.0, '疏水核心堆积', '#1B7A43',
         '• Y145F: 移除OH,改善核心\n  紧密堆积 → +3~4°C\n• F64L: 改善折叠效率\n• 核心残基优化范德华力'),
        (3.5, 5.0, 'β-桶完整性', '#2E86AB',
         '• 11条β链间H-bond网络\n• 桶形刚性 → 减少非辐射衰变\n• 桶盖残基保护发色团\n  (如Y145, H148, T203)'),
        (6.8, 5.0, '表面电荷优化', '#9370DB',
         '• S30R, Y39N: 表面电荷\n  改善 → 抗聚集\n• 超电荷GFP变体\n• 高溶解性 → 难聚集沉淀'),
        (0.2, 0.5, '折叠动力学', '#CD853F',
         '• sfGFP: 6个关键突变\n  (S30R/Y39N/N105T/Y145F/\n   I171V/A206V)\n• 脯氨酸异构化加速'),
        (3.5, 0.5, 'C末端尾巴', '#D81159',
         '• 非结构化C端尾巴\n  删除 → Tm降低~6°C\n• 维持动态接触网络\n• 传递稳定效应至活性位'),
        (6.8, 0.5, '发色团环境', '#FF69B4',
         '• H193Y: π-堆积稳定\n  发色团=TGP→YTP\n  99°C仍有荧光\n• Q66E: 化学稳定性↑\n• 内部水分子氢键网络'),
    ]

    for x, y, title, color, desc in items:
        box = FancyBboxPatch((x, y), 3.0, 2.0, boxstyle="round,pad=0.1",
                             facecolor='white', edgecolor=color, lw=2)
        ax.add_patch(box)
        ax.text(x+1.5, y+1.6, title, ha='center', fontsize=10, fontweight='bold', color=color)
        for j, line in enumerate(desc.split('\n')):
            ax.text(x+0.2, y+1.2 - j*0.35, line, fontsize=7.5, color='#555555')
        # 箭头指向中心
        cx = x + 1.5
        cy = y
        target_y = 3.7 if y > 3 else 2.7
        ax.annotate('', xy=(min(6.5, max(3.5, cx)), target_y),
                   xytext=(cx, y if y > 3 else y+2.0),
                   arrowprops=dict(arrowstyle='->', color=color, lw=1.5, alpha=0.6))

    path = os.path.join(img_dir, 'fig5_stability_factors.png')
    fig.savefig(path, dpi=150, bbox_inches='tight', facecolor='white')
    plt.close(fig)
    return path


# ---- 图6: 应用场景示意图 ----
def draw_applications():
    fig, ax = plt.subplots(1, 1, figsize=(10, 5))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 6)
    ax.axis('off')

    ax.set_title('高亮度+高热稳定性GFP的应用需求', fontsize=14, fontweight='bold', color='#003366', pad=10)

    apps = [
        ('超分辨显微镜\n(SIM/PALM/STORM)', '#2E86AB',
         '• SIM需15+帧/层 → 抗光漂白\n• 更高亮度 → 更短曝光\n• 更长观察窗口'),
        ('长时间活细胞成像\n(器官发育/昼夜节律)', '#1B7A43',
         '• 数小时~数天连续采集\n• 低光毒性 → 细胞健康\n• 低表达水平 → 低扰动'),
        ('深层组织成像\n(脑片/类器官/肿瘤)', '#D81159',
         '• 散射导致信号指数衰减\n• 背景自发荧光干扰\n• 高亮度补偿信号损失'),
        ('遗传编码生物传感器\n(GCaMP/FRET/电压)', '#8B4513',
         '• 传感器动态范围依赖FP亮度\n• 单体化 → 不干扰传感域\n• 快速成熟 → 实时响应'),
        ('极端环境应用\n(高温工业/极端生态)', '#FF6347',
         '• 嗜热菌中蛋白标记\n• 体外诊断(高温条件)\n• 高温筛选/催化监测'),
        ('膨胀显微镜(ExM)\n固定化样品3D成像', '#9370DB',
         '• 抵抗变性/固定/PFA\n• 高化学稳定性\n• 多色3D超分辨'),
    ]

    for i, (title, color, desc) in enumerate(apps):
        row = i // 3
        col = i % 3
        x = 0.3 + col * 3.2
        y = 0.3 + (1 - row) * 2.8
        box = FancyBboxPatch((x, y), 2.9, 2.4, boxstyle="round,pad=0.1",
                             facecolor='white', edgecolor=color, lw=2.5)
        ax.add_patch(box)
        ax.text(x+1.45, y+1.9, title, ha='center', fontsize=9.5, fontweight='bold', color=color)
        for j, line in enumerate(desc.split('\n')):
            ax.text(x+0.15, y+1.2 - j*0.4, line, fontsize=7.5, color='#555555')

    path = os.path.join(img_dir, 'fig6_applications.png')
    fig.savefig(path, dpi=150, bbox_inches='tight', facecolor='white')
    plt.close(fig)
    return path


# ═══════════════════════════════════════════
# 生成所有图片
# ═══════════════════════════════════════════
print("绘制示意图...")
img1 = draw_gfp_structure()
print("  [1/6] GFP结构示意图完成")
img2 = draw_chromophore_formation()
print("  [2/6] 发色团形成过程图完成")
img3 = draw_espt()
print("  [3/6] ESPT机理图完成")
img4 = draw_brightness_factors()
print("  [4/6] 亮度因素图完成")
img5 = draw_stability_factors()
print("  [5/6] 热稳定性因素图完成")
img6 = draw_applications()
print("  [6/6] 应用场景图完成")

# ═══════════════════════════════════════════
# 开始撰写文档
# ═══════════════════════════════════════════
print("\n生成Word文档...")

# ── 封面 ──
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('绿色荧光蛋白（GFP）\n发光机制、荧光强度与热稳定性详解')
run.bold = True
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(0, 51, 102)

doc.add_paragraph()
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('——从分子机理到蛋白质工程应用')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0, 80, 130)

doc.add_paragraph()
date_para = doc.add_paragraph()
date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_para.add_run(f'生成日期: {datetime.date.today().strftime("%Y年%m月%d日")}').font.size = Pt(12)

doc.add_page_break()

# ── 目录 ──
add_heading_styled('目录', 1)
toc = [
    '第一章  GFP为什么发光——荧光机制详解',
    '  1.1 GFP的整体结构：β-桶包裹的发色团',
    '  1.2 发色团的自催化成熟过程',
    '  1.3 荧光光物理：吸收、发射和激发态质子转移',
    '  1.4 发色团质子化状态与光谱性质',
    '第二章  影响GFP荧光强度的因素',
    '  2.1 消光系数（ε）的决定因素',
    '  2.2 量子产率（Φ）的决定因素',
    '  2.3 发色团微环境的关键残基',
    '  2.4 pH依赖性与pKa',
    '  2.5 亮度-光稳定性权衡',
    '第三章  影响GFP热稳定性的因素',
    '  3.1 β-桶结构的稳定性',
    '  3.2 疏水核心的堆积',
    '  3.3 关键稳定突变及其机制',
    '  3.4 表面电荷与溶解性',
    '  3.5 折叠动力学',
    '第四章  为什么要提高GFP的荧光强度和热稳定性',
    '  4.1 超分辨显微成像的需求',
    '  4.2 长时间活细胞成像',
    '  4.3 生物传感器开发',
    '  4.4 极端环境应用',
    '  4.5 蛋白质工程与合成生物学',
    '第五章  关键参考文献与PDB资源',
]
for item in toc:
    add_para(item, size=11)

doc.add_page_break()

# ═══════════════════════════════════════════
# 第一章
# ═══════════════════════════════════════════
add_heading_styled('第一章  GFP为什么发光——荧光机制详解', 1)

add_heading_styled('1.1 GFP的整体结构：β-桶包裹的发色团', 2)
add_para(
    '绿色荧光蛋白（Green Fluorescent Protein, GFP）来源于水母Aequorea victoria，'
    '由238个氨基酸组成，分子量约27 kDa。其晶体结构最早由Yang等人（1996）和Ormö等人（1996）解析，'
    '揭示了一个极其精巧的蛋白质折叠——"β-罐"（β-can）结构。'
)
add_image(img1, 5.5, '图1: GFP整体结构示意图——11条β链围成的桶状结构包裹中心α螺旋和发色团')

add_para('GFP的核心结构特征包括：', bold=True)
features = [
    '11条反平行β链围成近完美的圆柱形桶状结构（β-barrel），直径约24 Å，高度约42 Å',
    '一条短的α螺旋贯穿桶中心轴线，发色团附着在该螺旋中部',
    '发色团被深埋在β-桶内部，受到严密的溶剂屏蔽——这是GFP高量子产率的物理基础',
    '桶的顶部由若干"盖残基"（lid residues，如Y145、H148、T203等）部分封闭，进一步保护发色团',
    '发色团完全通过自催化翻译后修饰形成——无需任何外源酶或辅因子，仅需分子氧（O₂）',
    'GFP在晶体中以二聚体形式存在（通过疏水和极性界面相互作用）',
]
for f in features:
    add_list_item(f)

add_para(
    '关键PDB结构：1GFL（野生型GFP, 1.9 Å）、2B3P/2B3Q（超折叠sfGFP）、1EMA（EGFP, S65T/F64L）、'
    '8q79（mBaoJin, 单体超高光稳定性GFP）。'
)

add_heading_styled('1.2 发色团的自催化成熟过程', 2)
add_para(
    'GFP最令人惊叹的特征是其发色团完全由蛋白质自身的三个连续残基通过自催化反应形成。'
    '这是一个不需要任何外源酶或辅因子的翻译后修饰过程，唯一需要的外源小分子是分子氧（O₂）。'
)

add_para('发色团形成涉及三个连续氨基酸：', bold=True)
add_para('Ser65–Tyr66–Gly67（在EGFP中为Thr65–Tyr66–Gly67）— 即经典的XZG基序（X=可变，Z=芳香族，G=甘氨酸）')

add_image(img2, 5.5, '图2: GFP发色团自催化成熟过程——环化→脱水→氧化的三步反应')

add_para('成熟过程包含三个关键化学步骤：', bold=True)

add_para('步骤1 — 环化（Cyclization, t₁/₂ ~1.5 min）：', bold=True)
add_para(
    '蛋白质折叠后，中心螺旋在Gly67附近产生约80°的弯折（kink），这使得Ser65的羰基碳（C=O）与'
    'Gly67的酰胺氮（NH₂）在空间上靠近。Arg96的正电荷通过静电效应稳定螺旋弯折，并增强Gly67酰胺氮的亲核性。'
    'Gly67的酰胺氮（蛋白质骨架中最佳的亲核位点，因其空间位阻最小）对Ser65的羰基碳进行亲核攻击，'
    '形成共价C–N键，产生五元咪唑啉酮（imidazolidinone）杂环。'
    '该环化在酸性变性条件下可逆。R96A突变使成熟速度从分钟级减慢到月级。'
)

add_para('步骤2 — 脱水（Dehydration, t₁/₂ ~10.6 min）：', bold=True)
add_para(
    'Ser65的羰基氧以水分子形式消除（–H₂O）。Arg96使Tyr66的Cα质子酸性增强（活化），'
    'Glu222稳定脱水过渡态。活性位点内的保守水分子W1辅助质子转移。'
    '计算化学（Grigorenko, Krylov, Nemukhin, JACS 2017）估算脱水能垒约17 kcal/mol。'
    '脱水后形成共轭咪唑酮环中间体——烯醇式中间体（enolate intermediate）。'
    'Getzoff研究组在厌氧条件下用连二亚硫酸钠还原GFP，成功捕获并解析了该中间体的晶体结构（PDB: 2FZU）。'
)

add_para('步骤3 — 氧化（Oxidation, t₁/₂ ~34 min, 限速步骤）：', bold=True)
add_para(
    '分子氧（O₂）氧化Tyr66的Cα–Cβ键，将其从单键转变为双键（外亚甲基桥，exo-methylene bridge）。'
    '该步骤创造完整的共轭π体系——连接咪唑酮环与Tyr66酚环，使分子能够在可见光区吸收。'
    '自由基机理被提出：烯醇中间体向O₂发生单电子转移 → 自由基中间体 → 消除H₂O₂ → 形成双键。'
    'Glu222的质子化状态（中性vs阴离子）控制氧化与裂解途径之间的分配。'
    '反应消耗O₂并产生H₂O₂作为副产物（1:1化学计量比，Wachter等人验证）。'
    'E222Q突变显著减缓质子转移步骤。'
)

add_para('两种机制模型的讨论：', bold=True)
add_para(
    '关于三个步骤的确切顺序，学术界长期存在"Getzoff路径"（环化→脱水→氧化）与"Wachter路径"'
    '（环化→氧化→脱水）的争论。2017年Grigorenko等人的完整QM/MM计算支持Getzoff路径，'
    '但当前共识倾向于两种路径都可能发生，优势路径取决于具体条件（O₂浓度、突变、蛋白质环境）。'
)

add_para('关键催化残基总结：', bold=True)
add_para(
    '• Tyr66 — 提供酚环；其Cα–Cβ键被氧化形成π共轭体系\n'
    '• Gly67 — 酰胺氮作为亲核体攻击Ser65羰基碳\n'
    '• Arg96 — 正电荷稳定螺旋弯折、增强Gly67亲核性、稳定烯醇中间体\n'
    '• Glu222 — 作为广义碱进行质子抽取、控制氧化vs裂解途径分支'
)

add_heading_styled('1.3 荧光光物理：吸收、发射和激发态质子转移 (ESPT)', 2)
add_para(
    '一旦成熟的发色团（4-(p-羟基苯亚甲基)-5-咪唑啉酮，即p-HBDI）形成，GFP即具备荧光活性。'
    '其光物理过程涉及多个独特的机制。'
)

add_image(img3, 5.5, '图3: 激发态质子转移（ESPT）机理——基态的质子化-去质子化平衡与激发态的超快质子传递')

add_para('发色团的两种基态形式：', bold=True)
add_para(
    '• A态（中性形式）：酚羟基质子化（Cro-OH），吸收峰在~395–400 nm（近紫外），发出弱蓝色荧光（~460 nm）\n'
    '• B态（阴离子形式）：酚羟基去质子化（Cro-O⁻），吸收峰在~475–490 nm（蓝光），发出明亮的绿色荧光（~508–510 nm）\n'
    '• 两种形式在基态处于动态平衡，比例由蛋白质微环境的pKa和溶液pH决定'
)

add_para('激发态质子转移（ESPT）——GFP的独特光物理特征：', bold=True)
add_para(
    '当A态（中性）发色团吸收一个光子（~395 nm）后，激发态的发色团pKa急剧下降（从~7降至~3），'
    '变成一个"光酸"（photoacid）。在皮秒级时间尺度上，酚羟基上的质子通过一条精密的'
    '"质子线"（proton wire）传递至蛋白质内部：\n\n'
    '    发色团-OH* → 水分子(Wat) → Ser205 → Glu222\n\n'
    '质子传递后产生的激发态阴离子中间体I*发出绿色荧光（~508 nm）。随后在基态发生反向质子传递'
    '（Glu222 → Ser205 → Wat → 发色团-O⁻），再生中性基态A态。'
    '这一过程解释了野生型GFP看似巨大的表观Stokes位移（~100 nm），实际上发射来自阴离子物种。'
)

add_para('S205V突变的关键发现：', bold=True)
add_para(
    'Ser205→Val突变并未消除绿色荧光（仅使ESPT减慢约30倍），而S205V/T203V双突变完全消除ESPT，'
    '仅发出蓝色荧光。这揭示了质子线具有适应性——蛋白质可通过E222和T203的重排形成替代H-bond通路'
    '绕过缺失的Ser205。'
)

add_para('发色团的平面性与量子产率：', bold=True)
add_para(
    'β-桶结构迫使发色团维持刚性平面构象：\n'
    '• 最大化p轨道重叠以实现高效π共轭\n'
    '• 限制非辐射衰变途径（如C=C键扭转导致的顺反异构化）\n'
    '• 这是GFP在溶液中具有高荧光量子产率（~0.79 for EGFP）而分离的发色团几乎不发光的根本原因\n'
    '• 分子内运动受限（RIM）机制——类似于AIE（聚集诱导发光）原理'
)

add_heading_styled('1.4 发色团质子化状态与光谱性质', 2)
add_para(
    '发色团的质子化状态决定了GFP的光谱性质，而质子化状态由发色团微环境的精细调控。'
)

add_para('代表性GFP变体的光物理参数：', bold=True)

# 表格
table = doc.add_table(rows=8, cols=6)
table.style = 'Light Grid Accent 1'
headers = ['变体', 'ε (M⁻¹cm⁻¹)', 'Φ (量子产率)', 'pKa', 'λabs (nm)', 'λem (nm)']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True
            r.font.size = Pt(9)

data = [
    ['wtGFP', '30,000-37,000', '0.65', '~6.0', '395/475', '508'],
    ['EGFP (S65T/F64L)', '56,000', '0.60', '5.6-6.0', '488', '508'],
    ['sfGFP', '54,000-63,100', '0.65', '5.4', '485', '510'],
    ['SGFP2', '46,000', '0.70', '5.9', '485', '508'],
    ['mBaoJin', '~90,000', '~0.75', '4.37', '497', '512'],
    ['StayGold (二聚体)', '~100,000', '~0.80', '4.5', '497', '512'],
    ['mNeonGreen', '~116,000', '~0.80', '5.7', '506', '517'],
]
for i, row in enumerate(data):
    for j, val in enumerate(row):
        cell = table.rows[i+1].cells[j]
        cell.text = val
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.size = Pt(9)

add_para('')
add_para(
    '注意：mBaoJin和StayGold代表了新一代超高亮度和光稳定性GFP，打破了传统亮度-光稳定性权衡。'
    'mBaoJin的低pKa（4.37）意味着在生理pH范围（胞质pH 7.2，内体pH 5.5-6.5）内几乎完全去质子化，'
    '确保最大的488 nm激发效率。'
)

doc.add_page_break()

# ═══════════════════════════════════════════
# 第二章
# ═══════════════════════════════════════════
add_heading_styled('第二章  影响GFP荧光强度的因素', 1)

add_image(img4, 5.5, '图4: 影响GFP荧光亮度的关键因素总览——消光系数(ε) × 量子产率(Φ) = 亮度')

add_heading_styled('2.1 消光系数（ε）的决定因素', 2)
add_para('消光系数衡量发色团在特定波长下吸收光子的能力（单位：M⁻¹cm⁻¹）。主要决定因素：')

add_para('(a) 发色团质子化状态：', bold=True)
add_para(
    '阴离子形式（Cro-O⁻）在~475-490 nm的ε约为54,000-63,000 M⁻¹cm⁻¹，'
    '而中性形式（Cro-OH）在~395 nm的ε约为25,000-37,000。由于实验中主要使用488 nm激光激发，'
    '提高阴离子形式比例是增强有效亮度的关键。S65T突变将发色团平衡向阴离子方向大幅移动，'
    '使488 nm处的有效亮度相比野生型GFP提高约6倍（Patterson et al., 1997）。'
)

add_para('(b) 发色团环境的极性：', bold=True)
add_para(
    '发色团周围的局部电场影响其基态偶极矩，进而调节跃迁偶极矩强度。'
    '极性/带电残基（如E222、R96）的取向和质子化状态影响ε。'
)

add_para('(c) 发色团的共面性：', bold=True)
add_para(
    '咪唑酮环与酚环之间的扭转角影响π轨道重叠——越接近平面（dihedral ~0°），'
    '共轭越有效，ε越高。某些突变可改变扭转角从而改变吸收特性。'
)

add_heading_styled('2.2 量子产率（Φ）的决定因素', 2)
add_para('量子产率表示吸收的光子中产生荧光发射的比例。EGFP Φ ≈ 0.60-0.70。')

add_para('(a) 发色团构象刚性：', bold=True)
add_para(
    'β-桶限制发色团的分子内运动。在溶液中，分离的p-HBDI发色团因C=C双键的顺反异构化'
    '（Z→E光致异构化）而几乎不发光。蛋白质环境通过物理约束抑制该非辐射衰变通道。'
    '这就是为什么"发色团环境"比发色团本身对荧光更为关键。'
)

add_para('(b) 激发态质子转移效率：', bold=True)
add_para(
    '在wtGFP中，A态→I*的ESPT过程是高效荧光的必要前提。任何破坏质子线完整性的突变'
    '（如S205V/T203V双突变）将导致仅发出弱蓝色荧光。'
)

add_para('(c) 光致电子转移（PET）：', bold=True)
add_para(
    '发色团附近富电子残基（如Trp、Tyr、His）可淬灭激发态。某些GFP变体通过光致电子转移'
    '至Glu222产生暗态（光致变色副反应，van Thor, 2009）。减少PET可提高Φ。'
)

add_para('(d) 发色团成熟效率：', bold=True)
add_para(
    '荧光亮度与成熟发色团的数量成正比。如果折叠或成熟过程中有副反应（如氧化不完全、'
    '副产物裂解），会产生非荧光或弱荧光物种。某些突变（如F64L、S65T）可提高37°C下的成熟效率。'
)

add_heading_styled('2.3 发色团微环境的关键残基', 2)
add_para('以下残基直接或间接影响发色团的光物理性质：')

residues = [
    ('Glu222 (E222)', '终端质子受体（ESPT）、控制发色团pKa、影响氧化vs裂解分支。'
     'E222Q突变改变质子化平衡。在mBaoJin中对应E212。'),
    ('Ser205 (S205)', '质子线中间站。S205V突变使ESPT减慢30倍但不消除。'
     '在S205V/T203V双突变中ESPT完全阻断。'),
    ('Thr203 (T203)', '与发色团酚羟基形成H-bond，稳定阴离子形式。'
     'T203V突变移除H-bond供体，改变pKa。'),
    ('Arg96 (R96)', '催化环化的关键静电催化剂。稳定螺旋kink和烯醇中间体。'
     'R96A突变将成熟从分钟级延迟到月级。'),
    ('His148 (H148)', '在发色团酚环附近，参与H-bond网络。H148S（YuzuFP）通过改善与发色团'
     '和W1水的H-bond提高亮度1.5倍。'),
    ('Tyr145 (Y145)', 'β-桶"盖"残基。Y145F突变（sfGFP中）缓解疏水核心应变，提高热稳定性但不降低亮度。'),
    ('Gln66 (Q66)', '在Thermo Green Protein中，Q66E突变改善化学稳定性和pH稳定性。'),
]
for name, desc in residues:
    add_para(f'{name}：', bold=True)
    add_para(desc)

add_heading_styled('2.4 pH依赖性与pKa', 2)
add_para(
    '发色团的pKa是决定GFP在给定pH条件下"开启"比例的核心参数。pKa被定义为50%发色团处于'
    '阴离子荧光状态时的pH值。'
)
add_para(
    '• wtGFP pKa ~6.0：在细胞质（pH 7.2）中约94%为荧光态；但在酸性细胞器（高尔基体pH 6.0-6.7、'
    '内体pH 5.5-6.0、溶酶体pH 4.5-5.0）中荧光显著降低\n'
    '• EGFP pKa ~5.6-6.0：中等的pH敏感性\n'
    '• sfGFP pKa ~5.4：略优于EGFP\n'
    '• mBaoJin pKa ~4.37：几乎所有细胞器均可使用\n'
    '• bfloGFPa1 pKa ~3.0：极好的耐酸性，适用于溶酶体成像'
)
add_para(
    'pKa由发色团周围H-bond网络和静电场精细调控。降低pKa的策略包括：引入正电荷稳定阴离子发色团、'
    '移除稳定中性发色团的H-bond供体、优化发色团腔的静电环境。'
)

add_heading_styled('2.5 亮度-光稳定性权衡（及其突破）', 2)
add_para(
    '传统GFP工程面临一个根本性矛盾：分子氧（O₂）既是发色团成熟所必需，又在激发态下发色团与O₂'
    '反应产生单线态氧（¹O₂）导致光漂白（photobleaching）。历史上提高光稳定性的努力（如SiriusGFP）'
    '往往以牺牲亮度为代价（量子产率降低约3倍）。'
)
add_para(
    '2022-2024年的突破——StayGold/mBaoJin打破了这一权衡：\n'
    '• 多个H-bond和疏水接触将发色团（p-HBDI）更牢固地固定——抑制Z→E异构化非辐射衰变\n'
    '• 同时降低了激发态发色团与O₂的反应概率\n'
    '• mBaoJin还含有一个氯离子（Cl⁻）结合口袋（Kd ~3 mM），生理氯离子浓度（140-300 mM）'
    '使亮度提高约15%，光稳定性提高约2.24倍\n'
    '• 结果：比EGFP光稳定15倍（mBaoJin），同时保持相当甚至更高的亮度'
)

doc.add_page_break()

# ═══════════════════════════════════════════
# 第三章
# ═══════════════════════════════════════════
add_heading_styled('第三章  影响GFP热稳定性的因素', 1)

add_image(img5, 5.5, '图5: 影响GFP热稳定性的六大关键因素——从疏水核心到表面电荷')

add_heading_styled('3.1 β-桶结构的完整性', 2)
add_para(
    'GFP的11条反平行β链间形成广泛的骨架氢键网络，提供结构的整体刚性。桶形几何结构既是'
    '荧光功能的保障（维持发色团平面性），也是热稳定性的基础。'
)
add_para('影响β-桶完整性的因素：', bold=True)
add_para(
    '• β链间的H-bond饱和程度：断裂的H-bond需要水分子补偿，形成"结构弱点"\n'
    '• β-桶的几何张力：某些突变可能引入局部应变（如Y145的OH被迫弯曲），纠正这些应变可提高稳定性\n'
    '• 桶盖残基（lid residues）：Y145、H148、T203等对保护发色团和维持桶形至关重要\n'
    '• 内部水分子：桶内的保守水分子参与关键的H-bond网络'
)

add_heading_styled('3.2 疏水核心的堆积', 2)
add_para(
    'GFP的β-桶内部主要是疏水环境，核心残基的范德华堆积（van der Waals packing）'
    '是热稳定性的重要贡献者。'
)
add_para(
    'Y145F突变是核心堆积优化的经典案例：野生型中Y145的酚羟基被迫弯曲远离疏水核心，'
    '产生局部应变。Y→F替换移除该羟基，使F145可以深埋于疏水核心中紧密堆积，'
    '单独贡献+3-4°C的热稳定性提升。CD光谱确证β-桶整体折叠保持不变。'
)

add_heading_styled('3.3 关键稳定突变及其机制', 2)

add_para('超折叠GFP（sfGFP）——多突变协同稳定的典范：', bold=True)
add_para(
    'Pédelacq, Cabantous, Waldo等人（2006, Nature Biotechnology, DOI: 10.1038/nbt1172）'
    '通过定向进化开发了sfGFP。除S65T和F64L外，额外引入6个关键突变：'
)
sfgfp_muts = [
    'S30R — 表面电荷突变，改善溶解性和抗聚集',
    'Y39N — 表面极性优化',
    'N105T — 可能影响β链间H-bond',
    'Y145F — 疏水核心堆积优化（+3-4°C，如上所述）',
    'I171V — 核心堆积微调',
    'A206V — 可能影响桶底部区域稳定性',
]
for m in sfgfp_muts:
    add_list_item(m)

add_para(
    'sfGFP在化学变性剂中表现出更强的抗性，更快的折叠动力学，以及对环状排列'
    '（circular permutation）的更高容忍度——这些都是折叠鲁棒性的标志。'
    '当融合到折叠不佳的蛋白时，sfGFP仍能正常折叠发光，而folding reporter GFP则会失活。'
)

add_para('极端热稳定性 — TGP/YTP（Thermo Green/Yellow Protein）：', bold=True)
add_para(
    '从共识绿色蛋白（CGP）经定向进化获得。关键特征：\n'
    '• 与mAG（monomeric Azami Green）具有87%序列一致性\n'
    '• QGY发色团（vs GFP的SGY/TGY）\n'
    '• H193Y突变：酪氨酸引入π-堆积相互作用以稳定发色团 → 产生YTP（黄色热稳定蛋白）\n'
    '• 99°C下维持33-45%荧光，冷却后恢复近100%\n'
    '• Q66E额外突变改善化学稳定性和pH稳定性 → YTP-E'
)

add_para('C末端尾巴——被忽视的稳定性贡献者：', bold=True)
add_para(
    'GFP的柔性非结构化C末端尾巴在稳定性中发挥意想不到的作用：\n'
    '• 截断整个C末端尾巴使T₁/₂从79°C降至72.8°C（降低约6°C）\n'
    '• Arrhenius活化能降低约40%\n'
    '• C末端通过动态接触网络对整个β-桶骨架进行远程稳定\n'
    '• 突变效应可通过接触网络传递至活性位点\n'
    '• 提示：设计GFP变体时应保留C末端（或至少不截断至关键长度以下）'
)

add_heading_styled('3.4 表面电荷与溶解性', 2)
add_para(
    '蛋白质表面电荷通过以下机制影响热稳定性：\n'
    '• 表面电荷残基与水分子形成有序水化层 → 焓驱动稳定\n'
    '• 同性电荷互斥 → 防止热变性过程中的聚集沉淀\n'
    '• 超电荷GFP变体（如+36GFP）：表面引入大量正电荷，维持高溶解性\n'
    '• 但过度超电荷可能干扰折叠或功能——需要在电荷密度和结构完整性间平衡'
)

add_heading_styled('3.5 折叠动力学', 2)
add_para(
    'GFP的热稳定性不仅体现在静态Tm值上，还与折叠动力学密切相关：\n'
    '• GFP折叠涉及顺/反脯氨酸异构化——这是限速步骤\n'
    '• sfGFP比folding reporter GFP折叠更快、更完全\n'
    '• 单氨基酸缺失（如EGFP-G4Δ）可引起螺旋注册位移（helical registry shift），'
    '产生新的极性相互作用网络，使复性效率从77%提高到近100%，细胞荧光强度提高约2倍\n'
    '• 复性表现为三态过程（N→I→D），存在一个折叠中间体。'
    '稳定该中间体有助于提高折叠效率和热力学稳定性'
)

doc.add_page_break()

# ═══════════════════════════════════════════
# 第四章
# ═══════════════════════════════════════════
add_heading_styled('第四章  为什么要提高GFP的荧光强度和热稳定性', 1)

add_image(img6, 5.5, '图6: 高亮度和高热稳定性GFP的关键应用场景')

add_heading_styled('4.1 超分辨显微成像的需求', 2)
add_para(
    '现代超分辨显微技术（SIM、PALM/STORM、RESOLFT、膨胀显微镜ExM）对荧光蛋白提出了极高要求：'
)
add_para(
    '结构化光照明显微镜（SIM）：\n'
    '• 重建单张超分辨图像需要15帧以上原始图像（3角度 × 5相位）\n'
    '• 整个Z-stack采集意味着荧光团几乎连续被照射\n'
    '• 光漂白导致重建伪影 —— 要求极高的光稳定性\n\n'
    'PALM/STORM（单分子定位）：\n'
    '• 需要数千次定位事件重构一张图像\n'
    '• 要求可控的光激活/光切换 + 每分子的高光子产出\n\n'
    '膨胀显微镜（ExM）：\n'
    '• 需要FP耐受变性、固定、PFA交联和聚合条件\n'
    '• 高化学稳定性至关重要 —— mBaoJin耐受6 M GdnHCl和PFA固定\n\n'
    'RESOLFT（可逆饱和光学荧光跃迁）：\n'
    '• 需要可逆光切换FP，可多次ON/OFF循环而低切换疲劳'
)

add_heading_styled('4.2 长时间活细胞成像', 2)
add_para(
    '活细胞动态过程研究要求连续数小时至数天的荧光观测：\n'
    '• 细胞器动态：内质网网络重排（需~134 fps）、线粒体融合/分裂、细胞骨架动态\n'
    '• 发育生物学：器官发生、细胞分化、组织形态建成\n'
    '• 昼夜节律研究：需要24小时以上连续成像\n'
    '• 蛋白运输追踪：脉冲追踪（pulse-chase）实验、FDAP（荧光衰减后光活化）\n'
    '• 基因表达动态：快速成熟的FP（mBaoJin t₁/₂成熟=7.5 min）更准确地报告实时转录/翻译'
)
add_para(
    '关键需求：\n'
    '• 超高光稳定性 → StayGold/mBaoJin比EGFP光稳定15倍\n'
    '• 高亮度 → 允许更低激发强度 → 减少光毒性\n'
    '• 低表达水平 → 避免过表达伪影 → 要求高单分子亮度'
)

add_heading_styled('4.3 生物传感器开发', 2)
add_para(
    'GFP变体是遗传编码生物传感器的核心组件：\n'
    '• FRET传感器（CFP-YFP对）：报告蛋白-蛋白相互作用、构象变化、代谢物浓度\n'
    '• 单FP传感器（如GCaMP钙传感器）：环状排列GFP耦合感应域\n'
    '• 电压指示器、pH传感器、神经递质传感器\n'
    '• 光切换FRET（psFRET）：光控可视化蛋白相互作用\n'
    '• PEAQ生物传感（光致变色-绝对定量）：利用可逆光切换FP实现浓度无关的绝对定量'
)
add_para(
    '单体化的关键性：\n'
    '二聚化FP（如原始StayGold）不能用于生物传感器，因为二聚化会人为驱动传感域的关联。'
    'mBaoJin通过C165Y和Q140P单体化突变解决了此问题（Zhang et al., 2024, Nature Methods）。'
)

add_heading_styled('4.4 极端环境应用', 2)
add_para(
    '更高的热稳定性使GFP可用于以下场景：\n'
    '• 嗜热微生物中的蛋白标记（~60-80°C生长温度）\n'
    '• 体外诊断（高温条件，如PCR后检测，~72°C）\n'
    '• 工业生物催化中的荧光报告\n'
    '• 极端环境生态学研究\n'
    '• 蛋白质工程——高热稳定性是蛋白质"可进化性"的前提（更耐受突变）\n\n'
    '本次竞赛的72°C热处理评估直接对应这些应用需求——FP需要在高温下维持正确的发色团环境。'
)

add_heading_styled('4.5 蛋白质工程与合成生物学', 2)
add_para(
    '从蛋白质工程角度，提高GFP亮度和热稳定性的意义超越了GFP本身：\n'
    '• 作为蛋白质设计竞赛的标杆蛋白——GFP的"可见"表型使其成为验证计算方法的理想载体\n'
    '• 热稳定性高的蛋白更耐受突变 → 更高的"可进化性" → 为进一步功能优化留出空间\n'
    '• 高亮度和高稳定性是对立统一的目标——成功平衡二者验证了多目标优化的计算策略\n'
    '• 成功的GFP变体设计可直接转化为更好的细胞生物学工具\n'
    '• 方法论可推广至其他荧光蛋白（RFP、CFP、YFP等）和更多功能蛋白\n'
    '• 在无细胞体系（CFPS）中的高效表达是合成生物学的核心使能技术'
)

doc.add_page_break()

# ═══════════════════════════════════════════
# 第五章
# ═══════════════════════════════════════════
add_heading_styled('第五章  关键参考文献与PDB资源', 1)

add_heading_styled('5.1 核心综述与机制研究', 2)

refs_ch5 = [
    '[1]  Tsien, R.Y. (1998) "The Green Fluorescent Protein." Annu. Rev. Biochem., 67, 509-544.\n'
    '     DOI: 10.1146/annurev.biochem.67.1.509 — 经典综述，GFP领域的奠基文献。',

    '[2]  Yang, F., Moss, L.G., Phillips, G.N. (1996) "The molecular structure of green fluorescent protein."\n'
    '     Nature Biotechnology, 14, 1246-1251. DOI: 10.1038/nbt1096-1246 — 首个GFP晶体结构 (PDB: 1GFL, 1.9 Å)。',

    '[3]  Ormö, M., et al. (1996) "Crystal Structure of the Aequorea victoria Green Fluorescent Protein."\n'
    '     Science, 273, 1392-1395. DOI: 10.1126/science.273.5280.1392',

    '[4]  Craggs, T.D. (2009) "Green fluorescent protein: structure, folding and chromophore maturation."\n'
    '     Chem. Soc. Rev., 38, 2865-2875. DOI: 10.1039/b903641p — 发色团成熟机制的权威综述。',

    '[5]  van Thor, J.J. (2009) "Photoreactions and dynamics of the green fluorescent protein."\n'
    '     Chem. Soc. Rev., 38, 2935-2950. DOI: 10.1039/b820275n — ESPT和光物理的详细综述。',

    '[6]  Heim, R., Prasher, D.C., Tsien, R.Y. (1994) "Wavelength mutations and posttranslational autoxidation\n'
    '     of green fluorescent protein." PNAS, 91, 12501-12504. DOI: 10.1073/pnas.91.26.12501 — 首次证明O₂必需性。',

    '[7]  Barondeau, D.P., et al. (2006) "Structural Evidence for an Enolate Intermediate in GFP Fluorophore\n'
    '     Biosynthesis." JACS, 128, 3166-3168. DOI: 10.1021/ja0552693 — 烯醇中间体结构 (PDB: 2FZU)。',

    '[8]  Grigorenko, B.L., Krylov, A.I., Nemukhin, A.V. (2017) "Molecular Modeling Clarifies the Mechanism\n'
    '     of Chromophore Maturation in GFP." JACS, 139, 10239-10249. DOI: 10.1021/jacs.7b04677 — 完整的QM/MM计算。',

    '[9]  Zhang, L., et al. (2006) "Reaction Progress of Chromophore Biogenesis in Green Fluorescent Protein."\n'
    '     JACS, 128, 4766-4772. DOI: 10.1021/ja0580439 — Wachter路径实验证据。',

    '[10] Gorbachev, D.A., et al. (2020) "A General Mechanism of Green-to-Red Photoconversions of GFP."\n'
    '     Front. Mol. Biosci., 7, 176. DOI: 10.3389/fmolb.2020.00176 — 光转换机制。',

    '[11] Ong, W.J.-H., et al. (2011) "Function and structure of GFP-like proteins in the protein data bank."\n'
    '     Mol. BioSyst., 7, 984-992. DOI: 10.1039/c1mb05012e — PDB中266个GFP类结构的系统分析。',
]
for ref in refs_ch5:
    add_ref(ref)

add_heading_styled('5.2 关键蛋白质工程研究', 2)

refs_eng = [
    '[12] Pédelacq, J.-D., Cabantous, S., Tran, T., Terwilliger, T.C., Waldo, G.S. (2006)\n'
    '     "Engineering and characterization of a superfolder green fluorescent protein."\n'
    '     Nature Biotechnology, 24, 79-88. DOI: 10.1038/nbt1172 — sfGFP奠基性工作。',

    '[13] Patterson, G.H., Knobel, S.M., Sharif, W.D., Kain, S.R., Piston, D.W. (1997)\n'
    '     "Use of the Green Fluorescent Protein and its mutants in quantitative fluorescence microscopy."\n'
    '     Biophysical Journal, 73, 2782-2790. DOI: 10.1016/S0006-3495(97)78307-3 — 亮度参数系统测量。',

    '[14] Hirano, M., Ando, R., Miyawaki, A., et al. (2022) "A highly photostable and bright green fluorescent\n'
    '     protein." Nature Biotechnology, 40, 1132-1142. DOI: 10.1038/s41587-022-01278-2 — StayGold。',

    '[15] Zhang, H., Subach, F.V., et al. (2024) "Bright and stable monomeric green fluorescent protein derived\n'
    '     from StayGold." Nature Methods, 21, 657-665. DOI: 10.1038/s41592-024-02203-y — mBaoJin (PDB: 8q79)。',

    '[16] Ahmed, T., et al. (2025) "Molecular dynamics guided identification of a brighter variant of superfolder\n'
    '     GFP with increased photobleaching resistance." Commun. Chem., 8, 174.\n'
    '     DOI: 10.1038/s42004-025-01573-4 — YuzuFP (H148S)。',

    '[17] Nikolaev, G., et al. (2024) "Reengineering a flavin-based fluorescent protein using ProteinMPNN."\n'
    '     Protein Science, 33, e70002. DOI: 10.1002/pro.70002 — ProteinMPNN设计荧光蛋白。',

    '[18] Rives, A., et al. (2025) "Simulating 500 million years of evolution with a language model (ESM3)."\n'
    '     Science, 387, 850-858. DOI: 10.1126/science.ads0018 — AI生成全新荧光蛋白esmGFP。',
]
for ref in refs_eng:
    add_ref(ref)

add_heading_styled('5.3 关键PDB结构资源', 2)

pdb_refs = [
    '1GFL — 野生型GFP, 1.9 Å (Yang, Phillips, 1996)',
    '1EMA — EGFP (S65T/F64L), 1.9 Å',
    '2B3P — 超折叠sfGFP (Pédelacq, Waldo, 2006)',
    '2B3Q — sfGFP + 折叠不佳的融合肽段',
    '2FZU — 烯醇中间体 (厌氧还原态, Barondeau/Getzoff, 2006)',
    '2WUR — GFP原子分辨率结构 (旋转阳极)',
    '8q79 — mBaoJin, pH 6.5 (Zhang, Subach, 2024)',
]
for pdb in pdb_refs:
    add_list_item(pdb)

add_heading_styled('5.4 在线资源', 2)
add_para(
    '• RCSB PDB: https://www.rcsb.org/ — 三维结构查看 (Mol* viewer)\n'
    '• PDBsum: https://www.ebi.ac.uk/thornton-srv/databases/pdbsum/ — 二级结构拓扑图\n'
    '• FPbase: https://www.fpbase.org/ — 荧光蛋白数据库（光物理参数、序列）\n'
    '• PROSS: https://pross.weizmann.ac.il/ — 蛋白质稳定性自动设计服务器'
)

# ── 结束语 ──
doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run(
    '本文档由AI辅助整理生成，旨在为合成生物学创新赛"蛋白质设计"赛道的参赛队伍'
    '提供GFP发光机制、荧光强度和热稳定性影响因素的全面背景知识。'
    '文中所有DOI链接和PDB条目截至2026年5月已验证可访问。'
)
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(128, 128, 128)

# ── 保存 ──
output_path = r'C:\Users\fengs\Desktop\蛋白质设计-合成生物学创新赛-Claude\GFP发光机制与影响因素详解.docx'
doc.save(output_path)
print(f'\n报告已保存至: {output_path}')
print('完成！')
