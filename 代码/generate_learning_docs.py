# -*- coding: utf-8 -*-
"""
生成初学者学习与实践文档系列
专题一：总体学习路径与路线图（主文档）
专题二：GFP基础理论与实践
专题三：AI蛋白质设计工具与实战
专题四：蛋白质稳定性与热稳定性工程
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import datetime
import os

base_dir = r'C:\Users\fengs\Desktop\蛋白质设计-合成生物学创新赛-Claude'


def setup_document():
    """Create and configure a new Word document."""
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)
    return doc


def add_heading_styled(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 51, 102)


def add_para(doc, text, bold=False, size=11, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color


def add_bullet(doc, text):
    doc.add_paragraph(text, style='List Bullet')


def add_link(doc, title, url, desc=""):
    p = doc.add_paragraph()
    run_t = p.add_run(f"■ {title}")
    run_t.bold = True
    run_t.font.size = Pt(10)
    if url:
        run_u = p.add_run(f"\n  {url}")
        run_u.font.size = Pt(9)
        run_u.font.color.rgb = RGBColor(0, 80, 160)
    if desc:
        run_d = p.add_run(f"\n  {desc}")
        run_d.font.size = Pt(10)


def add_cover(doc, title_text, subtitle_text):
    for _ in range(4):
        doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = t.add_run(title_text)
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0, 51, 102)

    doc.add_paragraph()
    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = s.add_run(subtitle_text)
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0, 80, 130)

    doc.add_paragraph()
    d = doc.add_paragraph()
    d.alignment = WD_ALIGN_PARAGRAPH.CENTER
    d.add_run(f'生成日期: {datetime.date.today().strftime("%Y年%m月%d日")}').font.size = Pt(12)
    doc.add_page_break()


# ============================================================
# 专题一：总体学习路径与路线图（主文档）
# ============================================================
def build_doc1_overall_learning_path():
    print("生成专题一：总体学习路径与路线图...")
    doc = setup_document()
    add_cover(doc, '蛋白质设计入门学习与实践\n——总体学习路径与路线图',
              '从零基础到合成生物学创新赛蛋白质设计赛道')

    # --- 欢迎页 ---
    add_heading_styled(doc, '欢迎：你将学到什么？', 1)
    add_para(doc, '本系列文档专为零基础或初级水平的学习者设计，目标是帮助你在最短时间内建立起绿色荧光蛋白（GFP）蛋白质设计的完整知识框架，并具备参加合成生物学创新赛"蛋白设计赛道"的实践能力。')
    add_para(doc, '完成全部学习后，你将能够：', bold=True)
    items = [
        '理解GFP的发色团自催化形成机制和荧光光物理原理',
        '掌握影响GFP荧光亮度和热稳定性的关键结构因素',
        '使用蛋白质语言模型（ESM-2/ESM3）提取序列特征并预测突变效应',
        '使用ProteinMPNN进行基于结构的蛋白质序列设计',
        '使用Rosetta/FoldX进行突变稳定性评估（ΔΔG预测）',
        '构建完整的"骨架生成→序列设计→结构验证→功能预测"AI蛋白质设计管线',
        '设计并提交6条符合竞赛要求的高亮度+高热稳定性GFP变体序列',
    ]
    for item in items:
        add_bullet(doc, item)

    doc.add_page_break()

    # --- 学习路线图 ---
    add_heading_styled(doc, '学习路线图总览', 1)
    add_para(doc, '以下路线图按照"基础理论→计算工具→设计实践→竞赛策略"的递进关系组织，建议按顺序学习，每个阶段预计花费时间如下：')
    add_para(doc, '')

    stages = [
        ('第一阶段：GFP基础理论', '2-3天', [
            'GFP的结构：β-桶、中心α螺旋、发色团',
            '发色团自催化成熟三步骤：环化→脱水→氧化',
            '荧光光物理：吸收、激发态质子转移（ESPT）、发射',
            '发色团质子化状态：A态（中性）vs B态（阴离子），pKa的概念',
            '关键残基：Tyr66, Gly67, Arg96, Glu222, Ser205, His148',
            '推荐PDB结构：1GFL, 1EMA, 2B3P, 8q79',
        ]),
        ('第二阶段：AI蛋白质设计工具入门', '3-5天', [
            'ESM-2/ESM3：蛋白质语言模型——序列嵌入、零样本突变预测',
            'ProteinMPNN：逆折叠模型——从结构设计序列',
            'AlphaFold2/ESMFold：结构预测——验证设计序列是否可折叠',
            'RFdiffusion（进阶）：扩散模型——从头生成蛋白质骨架',
            'Colab平台：无需本地GPU即可运行上述所有模型',
        ]),
        ('第三阶段：稳定性与功能预测', '2-3天', [
            'GFP适应度景观：上位效应、Sigmoid截断模型',
            '热稳定性评估：TemBERTure、SPURS、ProteinMPNN-ddG',
            'Rosetta/FoldX ΔΔG计算：预测突变对稳定性的影响',
            'PROSS：自动化蛋白质稳定性设计Web服务器',
        ]),
        ('第四阶段：竞赛设计策略与实践', '3-5天', [
            '理解竞赛评分规则：双能乘积制（亮度×热稳定性保留率）',
            '基准蛋白sfGFP的6个关键突变及其机制',
            '设计策略：表面工程、核心堆积优化、静电网络增强',
            '管线搭建：ESM嵌入→ML模型→候选生成→稳定性筛选',
            '文档撰写：设计思路、Agent逻辑树、关键执行日志',
        ]),
    ]

    for stage_name, time, items in stages:
        add_heading_styled(doc, f'{stage_name}（预计{time}）', 2)
        for item in items:
            add_bullet(doc, item)

    doc.add_page_break()

    # --- 学习资源导航 ---
    add_heading_styled(doc, '学习资源导航（附链接）', 1)
    add_para(doc, '以下所有资源已经过验证，可追溯到原始出处。')

    add_heading_styled(doc, '在线课程与视频', 2)
    add_link(doc, '国家智慧教育平台：AI驱动的蛋白质智能设计', '', '浙江大学于浩然研究员主讲，涵盖MLDE、ProteinMPNN、GFP案例。来源: https://cetc.nwu.edu.cn/info/1151/2836.htm')
    add_link(doc, 'B站：AI驱动蛋白质设计(BindCraft、RFdiffusion、ProteinMPNN、AlphaFold2)', 'https://www.bilibili.com/video/BV1RVopBwEsL/', '全流程实操演示，中文讲解')
    add_link(doc, 'B站：ProteinMPNN定向进化改造蛋白质实操', 'https://www.bilibili.com/video/BV1hN6HYnEPR/', '使用ProteinMPNN进行定向进化的完整操作流程')
    add_link(doc, 'AAAI 2025 Tutorial: AI for Protein Design', 'https://deepgraphlearning.github.io/ProteinTutorial_AAAI2025/', '涵盖ESM、ProteinMPNN、RFdiffusion的学术级教程（英文）')
    add_link(doc, 'Rosetta Commons YouTube: MPNN Lecture', 'https://www.youtube.com/watch?v=6z4XmUAwdNA', 'Rosetta Commons 2024官方讲座，深入讲解ProteinMPNN原理')
    add_link(doc, 'Rosetta Commons YouTube: RFdiffusion Lecture', 'https://www.youtube.com/watch?v=OEnY2yA3jy8', '扩散模型蛋白质骨架生成的原理与实操')
    add_link(doc, 'Rosetta Commons YouTube: AlphaFold Lecture', 'https://www.youtube.com/watch?v=SVrn8_8aKO8', 'AlphaFold用于蛋白质设计的理论与实践')

    add_heading_styled(doc, '交互式Colab Notebook（零安装，浏览器运行）', 2)
    add_link(doc, 'ColabDesign (sokrypton)', 'https://github.com/sokrypton/ColabDesign', 'ProteinMPNN、RFdiffusion、AfDesign的一站式Colab集合，公认最佳入门工具。mpnn/design.ipynb')
    add_link(doc, 'ProteinMPNN官方QuickDemo', 'https://github.com/dauparas/ProteinMPNN', '官方Colab notebook，在colab_notebooks/quickdemo.ipynb')
    add_link(doc, 'DL4Proteins Notebooks (Gray Lab)', 'https://github.com/Graylab/DL4Proteins-notebooks', '完整教学系列，第9章按部就班演示RFdiffusion→ProteinMPNN→AlphaFold全管线')
    add_link(doc, 'ProteinMPNN-ddG Colab', 'https://github.com/PeptoneLtd/proteinmpnn_ddg', '零样本ΔΔG预测，无需训练即可评估突变稳定性')

    add_heading_styled(doc, '核心工具GitHub仓库', 2)
    add_link(doc, 'ESM (EvolutionaryScale/Meta)', 'https://github.com/evolutionaryscale/esm', 'ESM3多模态蛋白质语言模型（基础版免费开放学术界）')
    add_link(doc, 'ESM-2 (Meta FAIR)', 'https://github.com/facebookresearch/esm', 'ESM-2蛋白质语言模型，包括15B参数版本')
    add_link(doc, 'ProteinMPNN (Baker Lab)', 'https://github.com/dauparas/ProteinMPNN', '基于图的逆折叠模型，从结构设计序列')
    add_link(doc, 'RFdiffusion (Baker Lab)', 'https://github.com/RosettaCommons/RFdiffusion', '去噪扩散模型，de novo蛋白质骨架生成')
    add_link(doc, 'ColabFold', 'https://github.com/sokrypton/ColabFold', 'AlphaFold2/3的Colab实现，结构预测验证')
    add_link(doc, 'PROSS Web Server', 'https://pross.weizmann.ac.il', '自动化蛋白质稳定性设计（在线，无需编程）')
    add_link(doc, 'DNAChisel', 'https://github.com/Edinburgh-Genome-Foundry/DNAChisel', '竞赛指定反向翻译工具（了解原理即可，大设施统一执行）')

    add_heading_styled(doc, '关键文献（必读）', 2)
    papers = [
        ('Tsien RY (1998) "The Green Fluorescent Protein"', 'Annu Rev Biochem', 'GFP领域奠基性综述，必读。DOI: 10.1146/annurev.biochem.67.1.509'),
        ('Pédelacq JD et al. (2006) "Superfolder GFP"', 'Nat Biotechnol', 'sfGFP（竞赛基准蛋白）的原始文献。DOI: 10.1038/nbt1172'),
        ('Sarkisyan KS et al. (2016) "Local fitness landscape of GFP"', 'Nature', '理解GFP突变耐受和上位效应的核心文献。DOI: 10.1038/nature17995'),
        ('Hirano M et al. (2022) "StayGold"', 'Nat Biotechnol', '超高光稳定性GFP，mBaoJin的前身。DOI: 10.1038/s41587-022-01278-2'),
        ('Zhang H et al. (2024) "mBaoJin"', 'Nat Methods', '最新单体超亮GFP，赛事设计参考。DOI: 10.1038/s41592-024-02203-y'),
        ('Rives A et al. (2025) "ESM3/esmGFP"', 'Science', 'AI从零生成全新GFP。DOI: 10.1126/science.ads0018'),
        ('Close DW et al. (2015) "TGP"', 'Proteins', '结构导向表面工程创造极端热稳定性。DOI: 10.1002/prot.24699'),
        ('Gonzalez Somermeyer et al. (2022) "Heterogeneity of GFP fitness landscape"', 'eLife', '多种GFP同源蛋白的适应度景观比较。DOI: 10.7554/eLife.75842'),
        ('Church GM et al. (2025) "AI-driven protein design"', 'Nat Rev Bioeng', 'AI蛋白质设计路线图综述。DOI: 10.1038/s44222-025-00349-8'),
        ('Nikolaev G et al. (2024) "ProteinMPNN redesigns FP"', 'Protein Sci', 'ProteinMPNN设计荧光蛋白的直接验证。DOI: 10.1002/pro.70002'),
    ]
    for title, journal, desc in papers:
        add_link(doc, f'{title} — {journal}', '', desc)

    add_heading_styled(doc, '数据库资源', 2)
    add_link(doc, 'RCSB PDB', 'https://www.rcsb.org/', '蛋白质三维结构数据库，使用Mol*查看器交互式观察GFP结构')
    add_link(doc, 'FPbase', 'https://www.fpbase.org/', '荧光蛋白数据库，包含光物理参数和序列信息，竞赛排除列表参照源')
    add_link(doc, 'PDB-101 Molecule of the Month: GFP', 'https://pdb101.rcsb.org/motm/42', 'David Goodsell撰写的GFP科普文章，配有精美插图和交互式JSmol')
    add_link(doc, 'PDBsum', 'https://www.ebi.ac.uk/thornton-srv/databases/pdbsum/', '蛋白质二级结构拓扑图自动生成')

    doc.add_page_break()

    # --- 软件环境配置 ---
    add_heading_styled(doc, '软件环境配置指南', 1)
    add_para(doc, '推荐的环境配置（按优先级排列）：', bold=True)

    add_heading_styled(doc, '方案A：纯云端（推荐初学者）', 2)
    add_para(doc, '无需任何本地GPU，只需Google账号：')
    add_bullet(doc, 'Google Colab：运行ProteinMPNN、ESM-2等模型的Jupyter Notebook')
    add_bullet(doc, 'HuggingFace Spaces：在线运行蛋白质设计模型Demo')
    add_bullet(doc, 'PROSS Web Server：在线蛋白质稳定性自动化设计')
    add_bullet(doc, 'AlphaFold3 Server：在线结构预测')

    add_heading_styled(doc, '方案B：本地轻量环境', 2)
    add_para(doc, '适合有一定编程基础的学习者：')
    add_bullet(doc, 'Python 3.9+ + pip/conda管理包')
    add_bullet(doc, 'pip install fair-esm (ESM-2), transformers (ESM3), biopython')
    add_bullet(doc, 'pip install dnachisel (反向翻译), PyRosetta (如有license)')
    add_bullet(doc, 'pip install foldx (稳定性预测), openmm (MD模拟)')

    add_heading_styled(doc, '方案C：竞赛级环境', 2)
    add_para(doc, '需要GPU（RTX 3090/A100或等效）：')
    add_bullet(doc, 'CUDA 12.x + PyTorch 2.x')
    add_bullet(doc, 'ESM-2 3B/15B嵌入提取（本地GPU加速）')
    add_bullet(doc, 'ProteinMPNN本地推理（批量序列设计）')
    add_bullet(doc, 'ColabFold本地化部署（大规模结构验证）')

    doc.add_page_break()

    # --- 学习时间线 ---
    add_heading_styled(doc, '推荐学习时间线（共约2周）', 1)

    days = [
        ('第1天', 'GFP结构与发色团化学', '阅读PDB-101 GFP文章；用Mol*查看1GFL/2B3P结构；理解β-桶和发色团的三维位置关系'),
        ('第2天', '荧光机制与关键残基', '学习发色团三步成熟过程；理解ESPT质子转移机制；记住Tyr66-Gly67-Arg96-Glu222的功能'),
        ('第3天', '影响亮度与稳定性的因素', '阅读SFGFP原始文献；理解S30R静电网络和Y145F核心堆积优化；总结影响ε和Φ的结构因素'),
        ('第4天', '蛋白质语言模型入门', '在Colab中运行ESM-2 demo notebook；提取GFP序列嵌入；体验零样本突变预测'),
        ('第5天', '逆折叠设计入门', '在ColabDesign中运行ProteinMPNN notebook；用2B3P结构生成新序列；观察温度参数对序列多样性的影响'),
        ('第6天', '结构预测验证', '使用ColabFold预测设计序列的结构；理解pLDDT、PAE、RMSD的含义；筛选高置信度设计'),
        ('第7天', '适应度景观与上位效应', '阅读Nature 2016 GFP适应度景观论文；理解负上位效应和Sigmoid截断函数；计算突变耐受度'),
        ('第8-9天', '稳定性与功能预测', '运行TemBERTure预测Tm；用ProteinMPNN-ddG预测ΔΔG；使用PROSS设计稳定变体'),
        ('第10天', '往届竞赛方案研读', '研究synbiochallenges2025等GitHub项目的管线设计；理解GLEAM、MLDE、交叉熵蒙特卡洛方法'),
        ('第11-12天', '搭建个人设计管线', '整合ESM嵌入→ML模型→候选生成→稳定性筛选→结构验证的完整流程；生成第一批候选序列'),
        ('第13-14天', '优化与文档撰写', '筛选Top 6序列；撰写设计思路PDF；整理GitHub仓库；确保格式符合竞赛提交规范'),
    ]

    for day, topic, detail in days:
        add_heading_styled(doc, f'{day}：{topic}', 3)
        add_para(doc, detail)

    doc.add_page_break()

    # --- 本系列文档导航 ---
    add_heading_styled(doc, '本系列文档导航', 1)
    add_para(doc, '本学习资源系列包含以下专题文档，建议按顺序阅读：')
    add_bullet(doc, '专题一：总体学习路径与路线图（本文档）——全局导航和学习时间规划')
    add_bullet(doc, '专题二：GFP基础理论与实践——从分子结构到发色团自催化成熟，全面理解GFP')
    add_bullet(doc, '专题三：AI蛋白质设计工具与实战——ESM、ProteinMPNN、AlphaFold等工具的实操指南')
    add_bullet(doc, '专题四：蛋白质稳定性与热稳定性工程——ΔΔG预测、表面工程、稳定性设计策略')
    add_para(doc, '')
    add_para(doc, '所有文档中的链接均可在浏览器中打开以追溯信息来源。', color=RGBColor(128, 128, 128))

    output_path = os.path.join(base_dir, '学习文档_专题一_总体学习路径与路线图.docx')
    doc.save(output_path)
    print(f'  保存: {output_path}')
    return output_path


# ============================================================
# 专题二：GFP基础理论与实践
# ============================================================
def build_doc2_gfp_basics():
    print("生成专题二：GFP基础理论与实践...")
    doc = setup_document()
    add_cover(doc, '专题二：GFP基础理论与实践\n——从分子结构到发色团自催化成熟',
              '零基础理解绿色荧光蛋白的核心工作原理')

    # --- 1. GFP的历史与诺贝尔奖 ---
    add_heading_styled(doc, '一、GFP的历史与诺贝尔奖', 1)
    add_para(doc, '绿色荧光蛋白（Green Fluorescent Protein, GFP）的故事是基础科学发现如何改变生物医学研究的经典案例。')
    add_para(doc, '关键里程碑：', bold=True)
    add_bullet(doc, '1962年：下村脩（Osamu Shimomura）从Aequorea victoria水母中分离出GFP和发光蛋白aequorin')
    add_bullet(doc, '1992年：Prasher等人克隆了GFP的cDNA序列（Gene, 111:229-233）')
    add_bullet(doc, '1994年：Chalfie等人在大肠杆菌和秀丽隐杆线虫中表达GFP，证明其可在异源生物中发光（Science, 263:802-805）')
    add_bullet(doc, '1996年：Yang等人和Ormö等人分别解析了GFP的晶体结构（PDB: 1GFL, 1EMA）')
    add_bullet(doc, '2008年：下村脩、Martin Chalfie、钱永健（Roger Y. Tsien）因GFP的发现和发展获诺贝尔化学奖')
    add_bullet(doc, '2022年：Hirano等人报道StayGold——超高光稳定性GFP（Nat Biotechnol）')
    add_bullet(doc, '2024年：mBaoJin——基于StayGold的单体超高亮GFP（Nat Methods）')
    add_bullet(doc, '2025年：ESM3模型生成全新esmGFP，相当于模拟5亿年自然进化（Science）')

    add_link(doc, '诺贝尔奖官方页面：2008年诺贝尔化学奖', 'https://www.nobelprize.org/prizes/chemistry/2008/summary/', '了解GFP相关诺贝尔奖的详细信息')
    add_link(doc, 'PDB-101: Green Fluorescent Protein (双语)', 'https://pdb101.rcsb.org/motm/42', 'David Goodsell撰写的经典GFP科普，配有精美结构图')

    # --- 2. GFP的结构 ---
    add_heading_styled(doc, '二、GFP的三维结构', 1)
    add_para(doc, 'GFP由238个氨基酸组成，分子量约27 kDa，折叠成独特的"β-罐"（β-can）结构。这是自然界最优雅的蛋白质折叠之一。')

    add_heading_styled(doc, '2.1 β-桶（β-barrel）', 2)
    add_para(doc, '11条反平行β链围成近完美的圆柱形桶状结构：')
    add_bullet(doc, '直径约24 Å（2.4纳米），高度约42 Å（4.2纳米）')
    add_bullet(doc, '桶壁由β链间的主链氢键维持，结构极其稳定')
    add_bullet(doc, '桶的内部主要是疏水环境——这对发色团功能至关重要')
    add_bullet(doc, '桶的外部暴露残基与水溶液接触，可进行表面工程')

    add_heading_styled(doc, '2.2 中心α螺旋与发色团', 2)
    add_para(doc, '一条α螺旋沿桶的中心轴线贯穿整个结构。发色团附着在该螺旋的中部位置：')
    add_bullet(doc, '发色团由三个连续残基形成：Ser65-Tyr66-Gly67（在EGFP中为Thr65-Tyr66-Gly67）')
    add_bullet(doc, '发色团被深埋在β-桶内部，受到严密的溶剂屏蔽')
    add_bullet(doc, '溶剂屏蔽是GFP具有高量子产率的物理基础——避免了水分子的碰撞淬灭')

    add_heading_styled(doc, '2.3 关键PDB结构速查', 2)
    add_para(doc, '以下PDB条目是学习GFP结构的核心资源，可在RCSB PDB网站用Mol*查看器交互式观察：')
    pdb_table = [
        ('1GFL', '野生型GFP (wtGFP)', '1.9 Å', 'Yang, Phillips 1996', '最经典的GFP结构'),
        ('1EMA', 'EGFP (S65T/F64L)', '1.9 Å', 'Ormö, Remington 1996', '增强型GFP'),
        ('2B3P', 'Superfolder GFP (sfGFP)', '2.3 Å', 'Pédelacq, Waldo 2006', '竞赛基准蛋白'),
        ('2FZU', 'GFP烯醇中间体', '—', 'Barondeau, Getzoff 2006', '捕获发色团成熟中间体'),
        ('8q79', 'mBaoJin (pH 6.5)', '—', 'Zhang, Subach 2024', '最新单体超亮GFP'),
    ]
    table = doc.add_table(rows=len(pdb_table)+1, cols=5, style='Light Grid Accent 1')
    headers = ['PDB ID', '描述', '分辨率', '来源', '备注']
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        for p in table.rows[0].cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
    for i, row in enumerate(pdb_table):
        for j, val in enumerate(row):
            table.rows[i+1].cells[j].text = val
            for p in table.rows[i+1].cells[j].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    add_link(doc, 'RCSB PDB Mol* Viewer使用教程', 'https://www.rcsb.org/3d-view', '在浏览器中直接查看和旋转三维蛋白质结构')

    doc.add_page_break()

    # --- 3. 发色团自催化成熟 ---
    add_heading_styled(doc, '三、发色团的自催化成熟过程（核心机制）', 1)
    add_para(doc, 'GFP最令人惊叹的特征是其发色团完全由蛋白质自身的三个连续残基通过自催化反应形成——不需要任何外源酶或辅因子，唯一需要的外源小分子是分子氧（O₂）。这一过程被称为"自催化翻译后修饰"。')
    add_para(doc, '')

    add_heading_styled(doc, '3.1 三肽前体', 2)
    add_para(doc, '发色团的三肽前体是XZG基序：')
    add_bullet(doc, 'X = Ser65（wtGFP）或Thr65（EGFP）——提供羰基碳用于亲核攻击')
    add_bullet(doc, 'Z = Tyr66——提供酚环，贡献最终发色团的π共轭体系')
    add_bullet(doc, 'G = Gly67——酰胺氮作为亲核体（蛋白质骨架中最佳亲核位点，空间位阻最小）')

    add_heading_styled(doc, '3.2 三步化学反应', 2)

    add_para(doc, '步骤1：环化（Cyclization, t₁/₂ ≈ 1.5分钟）', bold=True)
    add_para(doc, '蛋白质折叠后，中心螺旋在Gly67附近产生约80°的弯折（kink），使Ser65的羰基碳（C=O）与Gly67的酰胺氮（NH₂）空间靠近。Gly67酰胺氮对Ser65羰基碳进行亲核攻击，形成共价C-N键，产生五元咪唑啉酮（imidazolidinone）杂环。')
    add_bullet(doc, 'Arg96的正电荷稳定螺旋弯折并增强Gly67酰胺氮的亲核性')
    add_bullet(doc, 'R96A突变将成熟速度从分钟级减慢到月级——证明了Arg96的催化作用')
    add_bullet(doc, '该环化反应在酸性变性条件下可逆')

    add_para(doc, '步骤2：脱水（Dehydration, t₁/₂ ≈ 10.6分钟）', bold=True)
    add_para(doc, 'Ser65的羰基氧以水分子形式消除（–H₂O）。Arg96使Tyr66的Cα质子酸性增强（活化），Glu222稳定脱水过渡态。活性位点内的保守水分子W1辅助质子转移。')
    add_bullet(doc, '计算化学（Grigorenko et al., JACS 2017）估算脱水能垒约17 kcal/mol')
    add_bullet(doc, '脱水后形成共轭咪唑酮环中间体（烯醇式中间体）')
    add_bullet(doc, 'Getzoff研究组在厌氧条件下捕获了该中间体（PDB: 2FZU）')

    add_para(doc, '步骤3：氧化（Oxidation, t₁/₂ ≈ 34分钟，限速步骤）', bold=True)
    add_para(doc, '分子氧（O₂）氧化Tyr66的Cα–Cβ键，从单键变为双键，创造完整的共轭π体系——连接咪唑酮环与Tyr66酚环。')
    add_bullet(doc, '自由基机理：烯醇中间体向O₂发生单电子转移→自由基中间体→消除H₂O₂→形成双键')
    add_bullet(doc, '反应消耗O₂并产生H₂O₂作为副产物（1:1化学计量比）')
    add_bullet(doc, 'E222Q突变显著减缓质子转移——Glu222控制氧化与裂解途径的分配')

    add_para(doc, '学术争论：两种路径模型', bold=True)
    add_para(doc, '关于三步的确切顺序，学术界长期存在两种模型：\n  "Getzoff路径"：环化→脱水→氧化（当前多数共识支持）\n  "Wachter路径"：环化→氧化→脱水\n 当前共识倾向于两种路径都可能发生，优势路径取决于O₂浓度、突变和蛋白质环境。')

    add_heading_styled(doc, '3.3 关键催化残基总结', 2)
    add_para(doc, '四个残基是发色团成熟不可或缺的：')
    add_bullet(doc, 'Tyr66：提供酚环；其Cα–Cβ键被氧化形成π共轭体系')
    add_bullet(doc, 'Gly67：酰胺氮作为亲核体攻击Ser65羰基碳（三肽中唯一不可替换的残基）')
    add_bullet(doc, 'Arg96：正电荷稳定螺旋弯折、增强Gly67亲核性、稳定烯醇中间体')
    add_bullet(doc, 'Glu222：作为广义碱进行质子抽取、控制氧化vs裂解途径分叉')

    doc.add_page_break()

    # --- 4. 荧光光物理 ---
    add_heading_styled(doc, '四、荧光光物理原理', 1)

    add_heading_styled(doc, '4.1 吸收与发射', 2)
    add_para(doc, 'GFP发色团（p-HBDI: 4-(p-羟基苯亚甲基)-5-咪唑啉酮）是一个共轭π体系，能够在可见光区吸收和发射光子：')
    add_bullet(doc, '野生型GFP主要吸收峰：~395 nm（近紫外，A态/中性形式）')
    add_bullet(doc, '次要吸收峰：~475 nm（蓝光，B态/阴离子形式）')
    add_bullet(doc, '发射峰：~508 nm（绿色荧光）')
    add_bullet(doc, '表观Stokes位移：~100 nm（从395 nm吸收到508 nm发射）——实际上发射来自阴离子物种')

    add_heading_styled(doc, '4.2 激发态质子转移（ESPT）——GFP的独特光物理', 2)
    add_para(doc, 'ESPT是GFP最具标志性的光物理特征，解释了其巨大表观Stokes位移：')
    add_para(doc, 'ESPT过程（皮秒级时间尺度）：', bold=True)
    add_para(doc, '1. 基态：发色团以中性形式（Cro-OH, A态）存在，吸收~395 nm光子\n'
             '2. 激发态：中性发色团被激发后，pKa从~7急剧下降至~3（成为"光酸"）\n'
             '3. 质子传递：酚羟基上的H⁺沿"质子线"传递：\n'
             '   发色团-OH* → 水分子(Wat) → Ser205 → Glu222\n'
             '4. 发射：去质子化后的阴离子发色团（Cro-O⁻*, I*态）发出~508 nm绿色荧光\n'
             '5. 基态恢复：反向质子传递（Glu222→Ser205→Wat→发色团-O⁻），再生中性基态')

    add_para(doc, '关键验证实验：', bold=True)
    add_bullet(doc, 'S205V单突变：ESPT减慢约30倍，但仍保留绿色荧光')
    add_bullet(doc, 'S205V/T203V双突变：ESPT完全消除，仅发出弱蓝色荧光')
    add_bullet(doc, '结论：质子线具有适应性——蛋白质可通过E222和T203的重排形成替代H-bond通路')

    add_heading_styled(doc, '4.3 发色团的平面性与量子产率', 2)
    add_para(doc, 'β-桶通过物理约束迫使发色团维持刚性平面构象：')
    add_bullet(doc, '最大化p轨道重叠→高效π共轭→450-500 nm强吸收')
    add_bullet(doc, '限制非辐射衰变途径（C=C键扭转导致的顺反异构化）')
    add_bullet(doc, '分离的p-HBDI发色团在溶液中几乎不发光——证明了蛋白质环境的关键作用')
    add_bullet(doc, '分子内运动受限（RIM）机制——类似于AIE（聚集诱导发光）原理')

    add_heading_styled(doc, '4.4 代表性GFP变体的光物理参数', 2)
    fp_table = doc.add_table(rows=8, cols=6, style='Light Grid Accent 1')
    fp_headers = ['变体', 'ε (M⁻¹cm⁻¹)', 'Φ (量子产率)', 'pKa', 'λabs (nm)', 'λem (nm)']
    fp_data = [
        ['wtGFP', '30,000-37,000', '0.65', '~6.0', '395/475', '508'],
        ['EGFP (S65T/F64L)', '56,000', '0.60', '5.6-6.0', '488', '508'],
        ['sfGFP', '54,000-63,100', '0.65', '5.4', '485', '510'],
        ['SGFP2', '46,000', '0.70', '5.9', '485', '508'],
        ['mBaoJin', '~90,000', '~0.75', '4.37', '497', '512'],
        ['StayGold (二聚体)', '~100,000', '~0.80', '4.5', '497', '512'],
        ['mNeonGreen', '~116,000', '~0.80', '5.7', '506', '517'],
    ]
    for i, h in enumerate(fp_headers):
        fp_table.rows[0].cells[i].text = h
        for p in fp_table.rows[0].cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
    for i, row in enumerate(fp_data):
        for j, val in enumerate(row):
            fp_table.rows[i+1].cells[j].text = val
            for p in fp_table.rows[i+1].cells[j].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    add_para(doc, '注意荧光亮度 = ε × Φ，因此mBaoJin (~67,500)和mNeonGreen (~92,800)的理论亮度远超EGFP (~33,600)。')
    add_link(doc, 'FPbase: 荧光蛋白数据库', 'https://www.fpbase.org/', '查询数百种荧光蛋白的光物理参数和序列，竞赛排除列表的参照源')

    doc.add_page_break()

    # --- 5. 影响亮度的因素 ---
    add_heading_styled(doc, '五、影响GFP荧光强度的因素', 1)

    add_heading_styled(doc, '5.1 消光系数（ε）的决定因素', 2)
    add_bullet(doc, '发色团质子化状态：阴离子形式（Cro-O⁻）在488nm的ε约为中性形式的2-3倍')
    add_bullet(doc, 'S65T突变：将发色团平衡向阴离子方向大幅移动，使488nm有效亮度提高约6倍')
    add_bullet(doc, '发色团环境极性：局部电场影响跃迁偶极矩强度')
    add_bullet(doc, '发色团共面性：咪唑酮环与酚环的扭转角越小，π轨道重叠越有效→ε越高')

    add_heading_styled(doc, '5.2 量子产率（Φ）的决定因素', 2)
    add_bullet(doc, '构象刚性：β-桶限制发色团分子内运动→抑制非辐射衰变')
    add_bullet(doc, 'ESPT效率：质子线完整性→高效A→I*转化')
    add_bullet(doc, '光致电子转移（PET）：富电子残基（Trp, Tyr, His）可能淬灭激发态')
    add_bullet(doc, '发色团成熟效率：折叠/成熟过程中的副反应产生弱荧光或非荧光物种')

    add_heading_styled(doc, '5.3 亮度-光稳定性权衡的突破', 2)
    add_para(doc, '传统GFP面临根本矛盾：O₂既是发色团成熟所需，又在激发态下与发色团反应产生¹O₂导致光漂白。')
    add_para(doc, 'StayGold/mBaoJin（2022-2024）打破了这一权衡：')
    add_bullet(doc, '多H-bond和疏水接触将发色团更牢固固定→抑制非辐射衰变')
    add_bullet(doc, '降低激发态发色团与O₂的反应概率')
    add_bullet(doc, 'mBaoJin的氯离子（Cl⁻）结合口袋（Kd ≈ 3 mM）：生理Cl⁻浓度时亮度↑15%，光稳定性↑2.24倍')
    add_bullet(doc, '结果：比EGFP光稳定15倍，同时保持相当或更高的亮度')

    add_link(doc, 'mBaoJin文献 (Nature Methods 2024)', 'https://doi.org/10.1038/s41592-024-02203-y', '最新单体超亮GFP的完整文献')
    add_link(doc, 'StayGold文献 (Nature Biotechnology 2022)', 'https://doi.org/10.1038/s41587-022-01278-2', '超高光稳定性GFP的原始报道')

    # --- 6. 学习检验 ---
    add_heading_styled(doc, '六、自测与实践练习', 1)
    add_para(doc, '完成本章学习后，尝试回答以下问题：')
    add_bullet(doc, 'GFP发色团由哪三个氨基酸形成？写出它们的编号（基于avGFP）')
    add_bullet(doc, '发色团自催化成熟的三步骤是什么？写出每步的化学反应')
    add_bullet(doc, 'Arg96和Glu222在发色团成熟中分别起什么作用？')
    add_bullet(doc, 'ESPT的质子线路径是什么？（四个组分的顺序）')
    add_bullet(doc, 'S205V/T203V双突变为什么消除绿色荧光？')
    add_bullet(doc, 'mBaoJin的pKa=4.37意味着什么？为什么对细胞成像重要？')
    add_para(doc, '')
    add_para(doc, '动手实践：', bold=True)
    add_bullet(doc, '在RCSB PDB网站打开2B3P（sfGFP），找到Tyr66-Gly67的侧链和围绕它们的残基')
    add_bullet(doc, '使用Mol*的测距工具，测量发色团酚羟基到Glu222羧基的距离（应约6-8 Å）')
    add_bullet(doc, '在FPbase上查找sfGFP和mBaoJin的详细光物理参数，比较它们的pKa和量子产率')

    output_path = os.path.join(base_dir, '学习文档_专题二_GFP基础理论与实践.docx')
    doc.save(output_path)
    print(f'  保存: {output_path}')
    return output_path


# ============================================================
# 专题三：AI蛋白质设计工具与实战
# ============================================================
def build_doc3_ai_tools():
    print("生成专题三：AI蛋白质设计工具与实战...")
    doc = setup_document()
    add_cover(doc, '专题三：AI蛋白质设计工具与实战\n——ESM、ProteinMPNN、AlphaFold等工具的实操指南',
              '从零开始使用AI工具进行蛋白质序列设计与评估')

    # --- 1. 工具生态总览 ---
    add_heading_styled(doc, '一、AI蛋白质设计工具生态系统', 1)
    add_para(doc, '2024-2025年，AI蛋白质设计领域呈现出"四大类工具协同工作"的格局：')

    add_para(doc, '工具分类与功能对应关系：', bold=True)
    eco_table = [
        ('蛋白质语言模型（PLM）', 'ESM-2/ESM3, ProtBERT, SaProt', '从序列提取特征表示、零样本突变效应预测、序列生成'),
        ('逆折叠模型', 'ProteinMPNN, ESM-IF', '给定骨架结构，设计氨基酸序列'),
        ('结构预测模型', 'AlphaFold2/3, ESMFold, Chai-1', '验证设计序列是否折叠为目标结构'),
        ('骨架生成模型（扩散）', 'RFdiffusion, FrameBuilder', '从头生成蛋白质骨架结构'),
    ]
    table = doc.add_table(rows=len(eco_table)+1, cols=4, style='Light Grid Accent 1')
    for i, h in enumerate(['工具类别', '代表工具', '核心功能']):
        table.rows[0].cells[i].text = h
        for p in table.rows[0].cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    for i, row in enumerate(eco_table):
        for j, val in enumerate(row[:3]):
            table.rows[i+1].cells[j].text = val
            for p in table.rows[i+1].cells[j].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)

    add_para(doc, '')
    add_para(doc, '典型设计管线（被所有2025年workshop采用）：', bold=True)
    add_para(doc, '  RFdiffusion (骨架生成) → ProteinMPNN (序列设计) → AlphaFold2/3 (结构验证) → 迭代筛选')
    add_para(doc, '对于GFP竞赛，由于已有高分辨率结构（2B3P），通常不需要从头生成骨架，可直接使用ProteinMPNN在现有骨架上设计序列变体。')

    doc.add_page_break()

    # --- 2. ESM ---
    add_heading_styled(doc, '二、ESM-2/ESM3：蛋白质语言模型', 1)

    add_heading_styled(doc, '2.1 ESM是什么？', 2)
    add_para(doc, 'ESM（Evolutionary Scale Modeling）是Meta AI开发的蛋白质语言模型系列，灵感来自NLP领域的BERT/GPT。核心思想：蛋白质氨基酸序列中的"语法"和"语义"编码了结构和功能信息。')
    add_bullet(doc, 'ESM-2（2022）：15B参数版本，在数亿条蛋白质序列上训练。可生成序列嵌入（embedding），用于下游预测任务')
    add_bullet(doc, 'ESM3（2025）：980亿参数的多模态（序列+结构+功能）生成式模型。2025年Science封面文章报道了用ESM3生成全新GFP（esmGFP）')
    add_bullet(doc, 'ESM-1v：仅使用UniRef50序列训练的版本，适合零样本突变效应预测')

    add_heading_styled(doc, '2.2 如何在Colab中使用ESM-2提取嵌入？', 2)
    add_para(doc, '以下是核心Python代码框架（完整代码参见Colab notebook）：', bold=True)
    code_text = """# 安装ESM
!pip install fair-esm

import torch
import esm

# 加载ESM-2模型（35M参数，适合Colab免费GPU）
model, alphabet = esm.pretrained.load_model_and_alphabet("esm2_t12_35M_UR50D")
model.eval()
batch_converter = alphabet.get_batch_converter()

# 准备序列数据
sequences = [
    ("sfGFP", "MSKGEELFTGVVPILVELDGDVNGHKFSVRGEGEGDATNGKLTLKFICTTGKLPVP..."),
    ("mutant1", "MSKGEELFTGVVPILVELDGDVNGHKFSVRGEGEGDATNGKLTLKFICTTGKLPVP..."),
]

# 提取嵌入
batch_labels, batch_strs, batch_tokens = batch_converter(sequences)
with torch.no_grad():
    results = model(batch_tokens, repr_layers=[model.num_layers])
    token_repr = results["representations"][model.num_layers]
    # 平均池化得到序列级嵌入
    seq_embedding = token_repr[:, 1:-1, :].mean(dim=1)

print(f"嵌入形状: {seq_embedding.shape}")  # [2, 480]"""
    p = doc.add_paragraph()
    run = p.add_run(code_text)
    run.font.size = Pt(9)
    run.font.name = 'Courier New'

    add_para(doc, '')
    add_link(doc, 'ESM-2 GitHub (Meta)', 'https://github.com/facebookresearch/esm', '官方仓库，含预训练模型下载和Colab示例')
    add_link(doc, 'ESM3 GitHub (EvolutionaryScale)', 'https://github.com/evolutionaryscale/esm', 'ESM3开源版，包括esmGFP生成代码')
    add_link(doc, 'ESM-2 Notebook on Colab (Gray Lab)', 'https://github.com/Graylab/DL4Proteins-notebooks', 'DL4Proteins系列的教学Notebook，适合入门')

    add_heading_styled(doc, '2.3 ESM在GFP竞赛中的应用', 2)
    add_bullet(doc, '提取训练序列的ESM嵌入作为ML模型的输入特征（如随机森林、XGBoost、注意力网络）')
    add_bullet(doc, '使用ESM-1v进行零样本突变效应预测（log-likelihood ratio）')
    add_bullet(doc, '使用ESM3的条件生成能力：给定结构条件和亮度/稳定性prompt生成新序列')
    add_bullet(doc, 'ESM-2嵌入可用于构建序列相似性矩阵和t-SNE/UMAP可视化')

    doc.add_page_break()

    # --- 3. ProteinMPNN ---
    add_heading_styled(doc, '三、ProteinMPNN：逆折叠蛋白质序列设计', 1)

    add_heading_styled(doc, '3.1 ProteinMPNN是什么？', 2)
    add_para(doc, 'ProteinMPNN（Dauparas et al., Science 2022）是基于消息传递神经网络的"逆折叠"模型。给定一个蛋白质骨架结构，它能够设计出折叠为该结构的氨基酸序列。')
    add_bullet(doc, '输入：蛋白质骨架坐标（N, Cα, C原子+可选Cβ方向）')
    add_bullet(doc, '输出：每个位置的氨基酸概率分布→采样获得序列')
    add_bullet(doc, '支持单链和多链蛋白设计')
    add_bullet(doc, '支持固定特定位置（如发色团、催化残基）不动')
    add_bullet(doc, '温度（temperature）参数控制序列多样性：T=0.1保守，T=1.0多样')
    add_link(doc, 'ProteinMPNN原始论文 (Science 2022)', 'https://doi.org/10.1126/science.add2187', '了解模型架构和创新点')
    add_link(doc, 'ProteinMPNN官方GitHub', 'https://github.com/dauparas/ProteinMPNN', '包含Colab notebook，零安装即可运行')

    add_heading_styled(doc, '3.2 在Colab中运行ProteinMPNN（零安装）', 2)
    add_para(doc, '最快上手方式：', bold=True)
    add_bullet(doc, '步骤1：打开 https://github.com/dauparas/ProteinMPNN')
    add_bullet(doc, '步骤2：找到colab_notebooks/quickdemo.ipynb，点击"Open in Colab"')
    add_bullet(doc, '步骤3：上传PDB文件（如2B3P——sfGFP的结构）')
    add_bullet(doc, '步骤4：运行所有单元格→获得设计序列')

    add_para(doc, 'ProteinMPNN设计荧光蛋白的关键参数设置：', bold=True)
    code_text2 = """# 固定发色团附近的关键残基
fixed_positions = {
    "A": [65, 66, 67, 96, 222],  # Tyr66, Gly67, Gly68, Arg96, Glu222
}
# 设置采样温度
temperature = 0.3  # 较低温度→较保守设计→保持更多原始折叠信息
# 生成多个候选
num_sequences = 100  # 生成100条候选序列"""
    p = doc.add_paragraph()
    run = p.add_run(code_text2)
    run.font.size = Pt(9)
    run.font.name = 'Courier New'

    add_para(doc, '')
    add_para(doc, 'ProteinMPNN设计GFP变体的关键约束（确保荧光功能保留）：', bold=True)
    add_bullet(doc, '必须固定发色团三肽（Tyr66-Gly67-Gly68 / sfGFP编号）')
    add_bullet(doc, '建议固定Arg96（发色团成熟催化）和Glu222（质子转移）')
    add_bullet(doc, '考虑固定His148和Thr203（桶盖残基，保护发色团）')
    add_bullet(doc, '允许其余β-桶表面和核心残基自由设计')
    add_link(doc, 'Nikolaev et al. (2024) ProteinMPNN设计荧光蛋白验证', 'https://doi.org/10.1002/pro.70002', 'ProteinMPNN成功重设计荧光蛋白的实验验证')

    doc.add_page_break()

    # --- 4. AlphaFold ---
    add_heading_styled(doc, '四、AlphaFold2/ESMFold：结构预测验证', 1)

    add_heading_styled(doc, '4.1 为什么需要结构验证？', 2)
    add_para(doc, '设计序列必须能正确折叠为预期的β-桶结构。如果设计序列折叠失败，荧光功能将不存在。结构预测模型提供了高通量的计算机筛选手段。')

    add_heading_styled(doc, '4.2 ColabFold快速上手', 2)
    add_bullet(doc, 'GitHub: https://github.com/sokrypton/ColabFold')
    add_bullet(doc, '点击AlphaFold2.ipynb → "Open in Colab"')
    add_bullet(doc, '输入设计的氨基酸序列')
    add_bullet(doc, '等待预测完成（通常5-15分钟，取决于序列长度和模型设置）')
    add_bullet(doc, '下载PDB文件和pLDDT评分')

    add_heading_styled(doc, '4.3 理解质量指标', 2)
    add_bullet(doc, 'pLDDT（predicted Local Distance Difference Test）：0-100分，>90为高置信度，<50为低置信度。发色团周围区域通常需要>80分')
    add_bullet(doc, 'PAE（Predicted Aligned Error）：残基间距离预测误差矩阵。低PAE表示结构预测准确')
    add_bullet(doc, 'pTM（predicted TM-score）：整体折叠质量分数，>0.8表示预测质量高')
    add_bullet(doc, 'RMSD（Root Mean Square Deviation）：与参考结构的偏差。对sfGFP模板应有<2Å的Cα RMSD')
    add_para(doc, '筛选标准建议：pLDDT > 80，pTM > 0.7，与2B3P的Cα RMSD < 3Å。')

    add_heading_styled(doc, '4.4 AlphaFold3（2024年发布）', 2)
    add_bullet(doc, '支持蛋白质+小分子+DNA/RNA+离子的联合预测')
    add_bullet(doc, '可通过 https://alphafoldserver.com 在线使用')
    add_bullet(doc, '对GFP设计，可用于验证发色团周围结构完整性')

    doc.add_page_break()

    # --- 5. RFdiffusion ---
    add_heading_styled(doc, '五、RFdiffusion（进阶）：从头蛋白质骨架生成', 1)
    add_para(doc, 'RFdiffusion（Watson et al., Nature 2023）是基于去噪扩散模型的蛋白质骨架生成工具。')
    add_para(doc, '对GFP竞赛的适用性：', bold=True)
    add_para(doc, '由于sfGFP已有高分辨率结构（2B3P），通常不需要从头生成骨架。但如果希望在β-桶的几何参数（桶直径、链数、桶高度）上进行创新，RFdiffusion提供了这种可能。')
    add_link(doc, 'RFdiffusion GitHub', 'https://github.com/RosettaCommons/RFdiffusion', 'de novo蛋白质骨架生成，需要PyRosetta license')
    add_link(doc, 'Baker Lab 2024: Parametrically guided β-barrel design', 'https://doi.org/10.1101/2024.07.22.604663', '参数化β-桶从头设计，72/96可溶，14个95°C稳定')

    add_para(doc, '')
    add_para(doc, 'ColabDesign中集成了RFdiffusion的简化版（无需license）：', bold=True)
    add_link(doc, 'ColabDesign RFdiffusion Notebook', 'https://github.com/sokrypton/ColabDesign', 'rfp/design.ipynb，在Colab中直接运行骨架生成')

    doc.add_page_break()

    # --- 6. 实践项目 ---
    add_heading_styled(doc, '六、实践项目：搭建你的首个GFP设计管线', 1)
    add_para(doc, '以下是一个最小可行的GFP设计管线，分为5个步骤：')

    add_heading_styled(doc, '步骤1：准备模板结构', 2)
    add_bullet(doc, '从RCSB PDB下载2B3P（sfGFP结构）')
    add_bullet(doc, '在PyMOL中删除水分子和配体，保留蛋白链')
    add_bullet(doc, '如有必要，进行能量最小化（Rosetta relax或OpenMM）')

    add_heading_styled(doc, '步骤2：用ProteinMPNN生成候选序列', 2)
    add_bullet(doc, '加载2B3P到ProteinMPNN Colab')
    add_bullet(doc, '固定残基65-67（发色团）、96（Arg96）、222（Glu222）')
    add_bullet(doc, '设置温度T=0.3-0.5，生成50-100条候选序列')
    add_bullet(doc, '记录每条序列的ProteinMPNN得分（负对数概率，越低越好）')

    add_heading_styled(doc, '步骤3：用ColabFold验证结构', 2)
    add_bullet(doc, '选择ProteinMPNN得分前20的序列')
    add_bullet(doc, '在ColabFold中预测它们的结构')
    add_bullet(doc, '计算每条序列与2B3P的Cα RMSD')
    add_bullet(doc, '筛选pLDDT > 80 且 RMSD < 2Å的序列')

    add_heading_styled(doc, '步骤4：用ESM-2评估功能保留', 2)
    add_bullet(doc, '提取通过结构筛选的序列和原始sfGFP的ESM-2嵌入')
    add_bullet(doc, '计算余弦相似度——高相似度意味着设计保留了类似的功能特征')
    add_bullet(doc, '选择ESM-2嵌入相似度最高的5条序列')

    add_heading_styled(doc, '步骤5：稳定性评估', 2)
    add_bullet(doc, '使用TemBERTure或ProteinMPNN-ddG预测突变对稳定性的影响')
    add_bullet(doc, '排除ΔΔG > 2 kcal/mol的变体（可能显著失稳）')
    add_bullet(doc, '最终选出6条序列作为候选提交')

    output_path = os.path.join(base_dir, '学习文档_专题三_AI蛋白质设计工具与实战.docx')
    doc.save(output_path)
    print(f'  保存: {output_path}')
    return output_path


# ============================================================
# 专题四：蛋白质稳定性与热稳定性工程
# ============================================================
def build_doc4_stability():
    print("生成专题四：蛋白质稳定性与热稳定性工程...")
    doc = setup_document()
    add_cover(doc, '专题四：蛋白质稳定性与热稳定性工程\n——ΔΔG预测、表面工程与稳定性设计策略',
              '掌握提升GFP热稳定性的计算方法和实验策略')

    # --- 1. 稳定性基础 ---
    add_heading_styled(doc, '一、蛋白质稳定性的基础概念', 1)
    add_para(doc, '蛋白质的稳定性可以从多个维度理解：')

    add_heading_styled(doc, '1.1 热力学稳定性', 2)
    add_bullet(doc, 'ΔG_folding = G_folded - G_unfolded：折叠态与去折叠态的自由能差')
    add_bullet(doc, 'ΔΔG = ΔG_mutant - ΔG_WT：突变对折叠自由能的改变（负值=稳定，正值=不稳定）')
    add_bullet(doc, 'Tm（熔解温度）：50%蛋白去折叠的温度')
    add_bullet(doc, 'GFP的典型Tm值：wtGFP ~78°C, sfGFP ~83°C, TGP >95°C')

    add_heading_styled(doc, '1.2 影响GFP稳定性的结构因素', 2)
    add_bullet(doc, '疏水核心堆积：紧密的范德华接触→显著的焓贡献。Y145F突变是核心优化的经典案例（移除极性OH→改善堆积→+3-4°C）')
    add_bullet(doc, '静电网络：跨β-链的离子对。S30R在sfGFP中创建了E32-R30-E17-R122-E115五元网络')
    add_bullet(doc, '表面电荷：防止热聚集。同性电荷互斥→减少变性后的不可逆聚集')
    add_bullet(doc, 'β-桶刚性：完整的H-bond网络提供结构完整性')
    add_bullet(doc, 'C末端尾巴：非结构化尾部通过动态接触网络远程稳定β-桶（截断→Tm降低~6°C）')
    add_bullet(doc, '发色团环境：H193Y（TGP中）→π-堆积稳定发色团→99°C仍维持荧光')

    doc.add_page_break()

    # --- 2. 计算方法 ---
    add_heading_styled(doc, '二、ΔΔG预测的计算方法', 1)

    add_heading_styled(doc, '2.1 方法对比', 2)
    methods_table = [
        ('Rosetta cartesian_ddg', '力场+骨架柔性', '~0.69（单体）', '中等（30-60 min/变体）', '结构导向精确设计'),
        ('FoldX', '经验力场', '~0.5-0.7', '快（1-5 min/变体）', '快速扫描大量突变'),
        ('ProteinMPNN-ddG', '零样本ML', '~0.77（Megascale）', '快（GPU加速）', '高通量稳定性筛选'),
        ('SPURS', 'ProteinMPNN+ESM融合', '~0.83（Megascale）', '快（单次前向传播）', '对所有20×L单突变同时预测'),
        ('TemBERTure', 'ProtBERT+适配器', '分类+回归', '快', '嗜热/非嗜热分类+Tm预测'),
        ('ThermoMPNN', 'ProteinMPNN微调', '—', '快', '在Megascale上微调的特化版'),
    ]
    table = doc.add_table(rows=len(methods_table)+1, cols=5, style='Light Grid Accent 1')
    for i, h in enumerate(['方法', '原理', '精度/相关性', '速度', '适用场景']):
        table.rows[0].cells[i].text = h
        for p in table.rows[0].cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
    for i, row in enumerate(methods_table):
        for j, val in enumerate(row):
            table.rows[i+1].cells[j].text = val
            for p in table.rows[i+1].cells[j].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)

    add_para(doc, '')
    add_link(doc, 'Rosetta cartesian_ddg 文档', 'https://docs.rosettacommons.org/docs/latest/cartesian-ddG', '官方教程，含完整命令行示例')
    add_link(doc, 'ProteinMPNN-ddG GitHub', 'https://github.com/PeptoneLtd/proteinmpnn_ddg', '零样本ΔΔG预测，含Colab notebook')
    add_link(doc, 'SPURS GitHub', 'https://github.com/luo-group/SPURS', '最新ProteinMPNN+ESM融合方法')
    add_link(doc, 'TemBERTure GitHub', 'https://github.com/ibmm-unibe-ch/TemBERTure', '基于ProtBERT的温度稳定性预测')
    add_link(doc, 'PROSS Web Server', 'https://pross.weizmann.ac.il', '自动化稳定性设计（在线，推荐初学者首选）')

    doc.add_page_break()

    # --- 3. 表面工程策略 ---
    add_heading_styled(doc, '三、表面工程策略（TGP案例深度解析）', 1)
    add_para(doc, 'TGP（Thermal Green Protein, Close et al., 2015）提供了结构导向表面工程提升热稳定性的典范案例：')

    add_heading_styled(doc, '3.1 两项核心原则', 2)
    add_bullet(doc, '原则一：识别并破坏晶体晶格中的分子间接触界面——这些界面在溶液中可能导致聚集')
    add_bullet(doc, '原则二：增加表面净负电荷——通过将表面碱性残基替换为谷氨酸（E），电荷排斥减少聚集')

    add_heading_styled(doc, '3.2 TGP的具体突变', 2)
    add_bullet(doc, 'K45E, K73E, K117E, R149E, N158E——5个表面正电荷/中性残基→谷氨酸')
    add_bullet(doc, 'C末端219-225（MLPSQAK）→GGGSGGG柔性连接序列——去除C末端介导的分子间接触')
    add_bullet(doc, '结果：90°C下半衰期从175分钟（eCGP123）提升到380分钟（TGP）→约2.2倍提升')

    add_heading_styled(doc, '3.3 对GFP竞赛设计的启示', 2)
    add_para(doc, '在sfGFP上应用类似策略：', bold=True)
    add_bullet(doc, '检查sfGFP（2B3P）的晶体堆积——识别分子间接触热点')
    add_bullet(doc, '在不破坏S30R静电网络的前提下，优化β-桶外表面电荷')
    add_bullet(doc, '考虑在sfGFP的C末端区域引入有利修饰（但不删除——C末端有助于稳定性！）')

    doc.add_page_break()

    # --- 4. 关键突变 ---
    add_heading_styled(doc, '四、已知的GFP稳定化突变汇总', 1)

    add_para(doc, '以下是已知的GFP热稳定性提升突变（按来源蛋白分类）：', bold=True)
    muts_table = [
        ('S30R', 'sfGFP', '在S1/S2/S5/S6间形成五元静电网络', '+显著（折叠和稳定性）', '最关键的sfGFP突变，不可逆改'),
        ('Y39N', 'sfGFP', '表面极性优化', '+适度', '改善溶解性'),
        ('N105T', 'sfGFP', '可能影响β链间H-bond', '+适度', '改善折叠动力学'),
        ('Y145F', 'sfGFP', '移除OH改善核心疏水堆积', '+3-4°C', '可在其他骨架中独立迁移'),
        ('I171V', 'sfGFP', '核心堆积微调', '+轻微', '与Y145F协同'),
        ('A206V', 'sfGFP', '减少二聚化', '+轻微', '改善单体性'),
        ('H193Y', 'TGP', 'π-堆积稳定发色团', '+显著（Tm和化学稳定性）', '可迁移至sfGFP'),
        ('Q66E', 'TGP', '发色团环境优化', '+化学稳定性和pH稳定性', '改善耐酸性'),
        ('C165Y', 'mBaoJin', '单体化+局部结构', '+光稳定性和单体性', '减少二聚化倾向'),
        ('H148S', 'YuzuFP', '改善与发色团和W1水的H-bond', '亮度↑1.5倍+光稳定性↑3倍', 'MD引导理性设计'),
    ]
    table = doc.add_table(rows=len(muts_table)+1, cols=5, style='Light Grid Accent 1')
    for i, h in enumerate(['突变', '来源', '机制', '效果', '可迁移性']):
        table.rows[0].cells[i].text = h
        for p in table.rows[0].cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(8)
    for i, row in enumerate(muts_table):
        for j, val in enumerate(row):
            table.rows[i+1].cells[j].text = val
            for p in table.rows[i+1].cells[j].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(8)

    doc.add_page_break()

    # --- 5. 适应度景观 ---
    add_heading_styled(doc, '五、GFP适应度景观与上位效应——设计中的"陷阱"', 1)
    add_para(doc, '这是参加GFP竞赛必须理解的核心概念。Sarkisyan等人（Nature 2016）的研究提供了系统性证据：')

    add_heading_styled(doc, '5.1 关键数字', 2)
    add_bullet(doc, '75%的单突变导致荧光降低——GFP对突变高度敏感')
    add_bullet(doc, '仅约23%的单突变是中性的')
    add_bullet(doc, '50%的含4个突变的基因型完全丧失荧光')
    add_bullet(doc, '30%的多突变基因型呈现显著负上位效应（epistasis）')

    add_heading_styled(doc, '5.2 上位效应是什么？', 2)
    add_para(doc, '上位效应指：多个突变组合在一起的效应不等于单个突变效应的简单加和。')
    add_para(doc, 'GFP中主要是负上位效应：两个各自中性的突变组合在一起可能导致荧光丧失。机制是：', bold=True)
    add_para(doc, '单个突变略微降低蛋白质稳定性（仍在折叠阈值之上）→多个微失稳突变累积→总稳定性低于折叠阈值→蛋白质去折叠→发色团无法成熟→荧光丧失。')
    add_para(doc, '适应度景观可用Sigmoid截断函数描述：简单线性非上位模型仅解释70%方差；引入Sigmoid函数后解释率85%。')

    add_heading_styled(doc, '5.3 对设计的启示', 2)
    add_bullet(doc, '不能简单组合"好的"单点突变——需要全局评估每个组合')
    add_bullet(doc, '利用超稳定骨架（如sfGFP, TGP）作为"突变缓冲器"吸收多突变的影响')
    add_bullet(doc, '设计时使用计算工具（FoldX, Rosetta ddG, SPURS）评估全局ΔΔG')
    add_bullet(doc, '保守策略：每次引入1-2个突变，验证后再叠加（迭代式设计）')
    add_bullet(doc, '在序列空间中搜索时，使用能显式建模上位效应的方法（交叉熵蒙特卡洛、贝叶斯优化）')
    add_link(doc, 'Sarkisyan et al. (2016) Nature', 'https://doi.org/10.1038/nature17995', 'GFP适应度景观经典研究')
    add_link(doc, 'eLife Digest: Lighting up protein design', 'https://elifesciences.org/articles/79310', 'Gonzalez Somermeyer 2022研究的科普解读')

    doc.add_page_break()

    # --- 6. 竞赛应用 ---
    add_heading_styled(doc, '六、竞赛热稳定性设计策略总结', 1)

    add_para(doc, '基于以上所有知识，以下是为合成生物学创新赛蛋白设计赛道推荐的热稳定性优化策略：')

    add_para(doc, '策略A：保守优化路径（低风险）', bold=True)
    add_bullet(doc, '以sfGFP为模板，保留其6个关键突变（S30R等）')
    add_bullet(doc, '在β-桶表面引入TGP风格的电荷优化（K/R→E替换）')
    add_bullet(doc, '引入H193Y（π-堆积）和Q66E（化学稳定性）增强发色团环境')
    add_bullet(doc, '使用FoldX/SPURS验证每个突变的ΔΔG，确保单独和累积稳定性')
    add_bullet(doc, '最终突变数控制在8-12个以内（减少上位效应风险）')

    add_para(doc, '策略B：数据驱动路径（中风险）', bold=True)
    add_bullet(doc, '利用竞赛提供的GFP_data.xlsx训练数据训练ML亮度预测模型（如GLEAM式架构）')
    add_bullet(doc, '使用TemBERTure/SPURS为每条训练序列标注预测Tm/稳定性')
    add_bullet(doc, '多目标优化：联合搜索高亮度+高稳定性的序列空间')
    add_bullet(doc, '使用交叉熵蒙特卡洛或贝叶斯优化处理上位效应')
    add_bullet(doc, '最终通过AlphaFold2结构验证+Rosetta ddG确认')

    add_para(doc, '策略C：生成式设计路径（高风险·高回报）', bold=True)
    add_bullet(doc, '使用ESM3或GeoEvoBuilder在GFP结构条件下生成大量候选序列')
    add_bullet(doc, '使用ProteinMPNN+固定发色团约束生成序列变体')
    add_bullet(doc, '通过ColabFold+稳定性评分工具进行多轮迭代筛选')
    add_bullet(doc, '可能发现超越现有GFP家族的全新功能序列')
    add_bullet(doc, '风险：生成序列的折叠和功能不确定性更高，需要更多验证')

    add_para(doc, '无论选择哪条路径，最终提交的6条序列应具有互补性——覆盖不同的突变策略和序列多样性——以最大化至少一条序列获得高分的概率。', bold=True)

    output_path = os.path.join(base_dir, '学习文档_专题四_蛋白质稳定性与热稳定性工程.docx')
    doc.save(output_path)
    print(f'  保存: {output_path}')
    return output_path


# ============================================================
# 主函数
# ============================================================
if __name__ == '__main__':
    print("=" * 60)
    print("生成初学者学习与实践文档系列")
    print("=" * 60)
    build_doc1_overall_learning_path()
    build_doc2_gfp_basics()
    build_doc3_ai_tools()
    build_doc4_stability()
    print("\n全部文档生成完成！")
