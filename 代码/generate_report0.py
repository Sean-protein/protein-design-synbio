# -*- coding: utf-8 -*-
"""Generate Word analysis report for GFP Protein Design project."""

import pandas as pd
import numpy as np
import json
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import warnings
warnings.filterwarnings('ignore')

BASE = os.path.dirname(os.path.abspath(__file__))
INTEGRATED = os.path.join(BASE, "output", "integrated_csv")
OUTPUT = os.path.join(BASE, "output")

print("Loading integrated datasets...")

brightness_data = pd.read_csv(os.path.join(INTEGRATED, "comprehensive_Brightness_Training_Data.csv"))
fpbase_data = pd.read_csv(os.path.join(INTEGRATED, "comprehensive_FPbase_GFP_Variants.csv"))
thermal_data = pd.read_csv(os.path.join(INTEGRATED, "comprehensive_Thermal_Stability_Data.csv"))
literature_data = pd.read_csv(os.path.join(INTEGRATED, "comprehensive_Literature_GFP_Properties.csv"))
reference_seqs = pd.read_csv(os.path.join(INTEGRATED, "comprehensive_Reference_Sequences.csv"))
candidate_positions = pd.read_csv(os.path.join(INTEGRATED, "comprehensive_Candidate_Positions.csv"))
fpbase_full = pd.read_csv(os.path.join(INTEGRATED, "04_fpbase_gfp_variants.csv"))
uniprot_mutations = pd.read_csv(os.path.join(INTEGRATED, "07_uniprot_p42212_mutations.csv"))
sarkisyan_brightness = pd.read_csv(os.path.join(INTEGRATED, "03_genotype_brightness_sarkisyan.csv"))
extra_seqs = pd.read_csv(os.path.join(INTEGRATED, "01_reference_sequences.csv"))

with open(os.path.join(OUTPUT, "stats.json"), "r") as f:
    stats = json.load(f)

WT_BRIGHT = 3.719212132

# ============================================================
# CREATE WORD DOCUMENT
# ============================================================
doc = Document()
section = doc.sections[0]
section.page_width = Cm(21.0)
section.page_height = Cm(29.7)

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = str(h)
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val) if val is not None and not (isinstance(val, float) and np.isnan(val)) else "-"
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(8)
    return table

# ============ TITLE PAGE ============
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("GFP Protein Design\nComprehensive Analysis Report")
run.font.size = Pt(28)
run.bold = True
run.font.color.rgb = RGBColor(0, 90, 50)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("\nSynthetic Biology Competition - Protein Design Project\n")
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(80, 80, 80)

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info.add_run("Data Sources: Sarkisyan 2016 | FPbase | ProTherm | UniProt | Literature\n").font.size = Pt(10)
info.add_run("Report Date: 2026-05-16\n").font.size = Pt(10)
doc.add_page_break()

# ============ TABLE OF CONTENTS ============
doc.add_heading("Table of Contents", level=1)
toc = [
    "1. Data Overview",
    "2. Reference Sequence Analysis",
    "3. Fluorescence Brightness Analysis",
    "4. Optical Properties Analysis",
    "5. Thermal Stability Analysis",
    "6. Mutation Landscape Analysis",
    "7. Candidate Design Positions",
    "8. Literature Knowledge Integration",
    "9. Protein Design Strategy Recommendations",
    "10. Summary and Outlook",
]
for item in toc:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)
doc.add_page_break()

# ============ SECTION 1: DATA OVERVIEW ============
doc.add_heading("1. Data Overview", level=1)

doc.add_paragraph(
    "This report integrates GFP (Green Fluorescent Protein) related data from multiple "
    "public databases and literature sources to provide comprehensive data support and "
    "design guidance for the protein design project. The data covers multiple dimensions "
    "including GFP sequences, fluorescence brightness, optical properties, thermal stability, "
    "and mutation effects."
)

doc.add_heading("1.1 Dataset Composition", level=2)
overview_rows = [
    ["Brightness Training Set", f"{len(brightness_data)}", "GFP_data.xlsx / Sarkisyan 2016",
     "Single and multi-site mutation brightness of avGFP"],
    ["Sarkisyan Full Dataset", f"{len(sarkisyan_brightness)}", "Nature 2016 (Figshare)",
     ">50,000 GFP mutant median brightness + barcode counts"],
    ["FPbase Optical Properties", f"{len(fpbase_data)} selected + {len(fpbase_full)} full", "fpbase.org",
     "Ex/Em, extinction coefficient, QY, brightness, pKa, etc."],
    ["Thermal Stability", f"{len(thermal_data)}", "ProTherm / Literature",
     "Tm, DDG thermodynamic parameters"],
    ["Literature Knowledge", f"{len(literature_data)}", "Multiple publications",
     "Key mutation effect mechanisms"],
    ["Candidate Positions", f"{len(candidate_positions)}", "Expert curation",
     "Important residue positions with known mutation data"],
    ["Reference Sequences", f"{len(reference_seqs) + len(extra_seqs)}", "Comprehensive + Training",
     "avGFP, sfGFP, EGFP and other reference sequences"],
    ["UniProt Mutations", f"{len(uniprot_mutations)}", "UniProt P42212",
     "Natural variants and mutagenesis records"],
]
add_table(doc, ["Dataset", "Size", "Source", "Content"], overview_rows)
doc.add_paragraph()

doc.add_heading("1.2 Statistical Summary", level=2)
bright = brightness_data["Brightness"].dropna()
sum_rows = [
    ["Total brightness variant entries", f"{len(brightness_data)}"],
    ["Mean brightness", f"{bright.mean():.4f}"],
    ["Median brightness", f"{bright.median():.4f}"],
    ["Brightness std dev", f"{bright.std():.4f}"],
    ["Brightness range", f"{bright.min():.4f} ~ {bright.max():.4f}"],
    ["WT avGFP brightness (reference)", f"{WT_BRIGHT}"],
    ["Total FPbase GFP variants", f"{len(fpbase_full)}"],
    ["Thermal stability measurements", f"{len(thermal_data)}"],
    ["Candidate design positions", f"{len(candidate_positions)}"],
    ["Known functional mutations", f"{len(uniprot_mutations)}"],
]
add_table(doc, ["Metric", "Value"], sum_rows)
doc.add_page_break()

# ============ SECTION 2: REFERENCE SEQUENCE ANALYSIS ============
doc.add_heading("2. Reference Sequence Analysis", level=1)

doc.add_paragraph(
    "Understanding different GFP variant reference sequences is fundamental to protein design. "
    "The following analyzes 5 core GFP reference sequences covering natural avGFP, engineered "
    "sfGFP (superfolder GFP), and GFP homologs from other species. Understanding differences "
    "between these sequences helps identify key residue positions and mutable regions."
)

all_refs = pd.concat([
    reference_seqs[["Protein", "Sequence"]],
    extra_seqs[["Protein", "Sequence"]]
], ignore_index=True).drop_duplicates(subset=["Protein"]).reset_index(drop=True)
all_refs["Length"] = all_refs["Sequence"].str.len()

doc.add_heading("2.1 Reference Sequence Information", level=2)
ref_rows = [
    ["avGFP", "238", "Aequorea victoria (jellyfish)", "2wur / 1GFL",
     "Wild-type GFP, Ex395/475nm, Em508nm"],
    ["sfGFP", "238", "Engineered from avGFP", "2B3P",
     "Superfolder GFP, high folding efficiency, excellent stability"],
    ["amacGFP", "239", "Aequorea macrodactyla", "7LG4",
     "Another jellyfish GFP, sequence differs from avGFP"],
    ["cgreGFP", "226", "Clytia gregaria (hydrozoan)", "2HPW",
     "Shorter sequence, lower sequence identity to avGFP"],
    ["ppluGFP", "224", "Pontellina plumata (copepod)", "2G6X",
     "One of the shortest GFP variants, significant sequence divergence"],
    ["EGFP", "239", "Engineered avGFP (F64L/S65T)", "-",
     "Enhanced GFP, 488nm excitation, improved brightness, 37C folding"],
]
add_table(doc, ["Protein", "Length(aa)", "Origin", "Recommended PDB", "Features"], ref_rows)
doc.add_paragraph()

doc.add_heading("2.2 Key Sequence Analysis Findings", level=2)
findings = [
    ("Chromophore Tripeptide: ",
     "avGFP SYG(65-67) is the fluorescence chromophore core. sfGFP maintains SYG unchanged. "
     "Engineered variants alter spectral properties through S65T(EGFP), Y66H(BFP), Y66W(CFP) mutations."),
    ("Folding-Enhancing Mutations: ",
     "S30R/Y39N/F64L/S65T/F99S/M153T/V163A in sfGFP are key folding-enhancing mutations "
     "enabling superfolding capability."),
    ("Dimerization Interface: ",
     "A206K mutation eliminates GFP dimerization tendency, producing monomeric fluorescent "
     "proteins (mGFP), essential for fusion protein applications."),
    ("N- and C-terminal Regions: ",
     "Residues 1-7 (N-term) and 230-238 (C-term) are relatively conserved but have designable space."),
    ("Loop Regions: ",
     "Loop regions generally tolerate insertions and mutations, serving as hotspots for "
     "functionalization engineering."),
]
for t, d in findings:
    p = doc.add_paragraph()
    run = p.add_run(f"* {t}")
    run.bold = True
    p.add_run(d)
doc.add_page_break()

# ============ SECTION 3: BRIGHTNESS ANALYSIS ============
doc.add_heading("3. Fluorescence Brightness Analysis", level=1)

doc.add_paragraph(
    "Fluorescence brightness is the core metric for evaluating GFP variant performance. "
    "This section analyzes the large-scale GFP mutant fluorescence data from Sarkisyan et al. (2016) "
    "and the merged training dataset."
)

doc.add_heading("3.1 Brightness Distribution", level=2)
doc.add_paragraph(
    f"Brightness distribution of {len(bright)} GFP variants in the training dataset:\n"
    f"* Mean brightness: {bright.mean():.4f} ({bright.mean()/WT_BRIGHT*100:.1f}% of WT avGFP)\n"
    f"* Median brightness: {bright.median():.4f} ({bright.median()/WT_BRIGHT*100:.1f}% of WT)\n"
    f"* Standard deviation: {bright.std():.4f}\n"
    f"* Maximum: {bright.max():.4f} ({bright.max()/WT_BRIGHT*100:.1f}% of WT)\n"
    f"* Minimum: {bright.min():.4f} ({bright.min()/WT_BRIGHT*100:.1f}% of WT)\n\n"
    f"Key Finding: Most variants in the training dataset have lower brightness than WT avGFP "
    f"(WT brightness = {WT_BRIGHT}), indicating that most single-site mutations negatively "
    f"impact brightness. This shows that GFP fluorescence activity is highly optimized in its "
    f"sequence, but also suggests that rational multi-site combination mutations have potential "
    f"to achieve super-WT brightness variants."
)

doc.add_heading("3.2 Brightness Percentile Analysis", level=2)
pct_rows = []
for p in [10, 25, 50, 75, 90, 95, 99]:
    val = np.percentile(bright, p)
    pct_rows.append([f"{p}%", f"{val:.4f}", f"{val/WT_BRIGHT*100:.1f}%"])
add_table(doc, ["Percentile", "Brightness", "% of WT"], pct_rows)
doc.add_paragraph()

doc.add_heading("3.3 Top Brightness Variants", level=2)
top_bright = brightness_data.nlargest(15, "Brightness")
top_rows = []
for i, (_, row) in enumerate(top_bright.iterrows(), 1):
    gfp_type = str(row.get("GFP type", row.get("GFP_type", "-")))
    src = str(row.get("Data_Source", row.get("Source", "-")))
    top_rows.append([i, str(row["aaMutations"]), gfp_type,
                     f"{row['Brightness']:.4f}", f"{row['Brightness']/WT_BRIGHT*100:.1f}%", src])
add_table(doc, ["Rank", "Mutation", "GFP Type", "Brightness", "% of WT", "Source"], top_rows)
doc.add_paragraph()

doc.add_heading("3.4 High-Frequency Beneficial Mutations (Sarkisyan 2016 Full Dataset)", level=2)
doc.add_paragraph(
    "Analysis of the Sarkisyan full dataset (>50,000 variants) to identify single mutations "
    "with above-median brightness. This reveals positions most tolerant to mutation and "
    "specific amino acid substitutions that maintain or enhance fluorescence."
)

sark = sarkisyan_brightness.copy()
sark_bright = pd.to_numeric(sark["medianBrightness"], errors="coerce")
bright_threshold = sark_bright.median()

single_muts = sark[sark["aaMutations"].str.match(r'^[A-Z][0-9]+[A-Z]$', na=False)].copy()
single_muts["brightness"] = pd.to_numeric(single_muts["medianBrightness"], errors="coerce")
single_muts = single_muts.dropna(subset=["brightness"])
single_muts["position"] = single_muts["aaMutations"].str.extract(r'(\d+)').astype(int)

beneficial = single_muts[single_muts["brightness"] > bright_threshold].sort_values("brightness", ascending=False)
top_single = beneficial.head(20)

ben_rows = []
for i, (_, row) in enumerate(top_single.iterrows(), 1):
    ben_rows.append([i, row["aaMutations"], int(row["position"]),
                     f"{row['brightness']:.4f}", int(row["uniqueBarcodes"])])
add_table(doc, ["Rank", "Mutation", "Position", "Brightness", "Barcode Count"], ben_rows)
doc.add_paragraph()

doc.add_heading("3.5 Beneficial Mutation Hotspot Positions", level=2)
pos_counts = beneficial.groupby("position").agg(
    mutation_count=("aaMutations", "count"),
    best_brightness=("brightness", "max"),
    best_mutation=("aaMutations", "first")
).sort_values("mutation_count", ascending=False).head(15)

pos_rows = []
for pos, row in pos_counts.iterrows():
    pos_rows.append([int(pos), int(row["mutation_count"]),
                     f"{row['best_brightness']:.4f}", row["best_mutation"]])
add_table(doc, ["Position", "Beneficial Mutations Count", "Best Brightness", "Best Mutation"], pos_rows)
doc.add_page_break()

# ============ SECTION 4: OPTICAL PROPERTIES ============
doc.add_heading("4. Optical Properties Analysis", level=1)

doc.add_paragraph(
    "The optical properties of GFP variants (excitation wavelength, emission wavelength, "
    "extinction coefficient, quantum yield, etc.) directly define their application scenarios. "
    "This section analyzes data from 1,110 GFP-related proteins in the FPbase database."
)

doc.add_heading("4.1 Representative GFP Variant Optical Properties", level=2)
representative = ["avGFP", "sfGFP", "EGFP", "mNeonGreen", "Venus", "Cerulean", "Citrine", "mEGFP", "mClover3"]
fpbase_numeric = fpbase_full.copy()
for col in ["Ex max (nm)", "Em max (nm)", "Extinction Coefficient", "Quantum Yield", "Brightness", "pKa"]:
    fpbase_numeric[col] = pd.to_numeric(fpbase_numeric[col], errors="coerce")
rep_data = fpbase_numeric[fpbase_numeric["Name"].isin(representative)]

opt_rows = []
for _, row in rep_data.iterrows():
    opt_rows.append([
        row["Name"],
        f'{row["Ex max (nm)"]:.0f}' if not pd.isna(row["Ex max (nm)"]) else "-",
        f'{row["Em max (nm)"]:.0f}' if not pd.isna(row["Em max (nm)"]) else "-",
        f'{row["Extinction Coefficient"]:.0f}' if not pd.isna(row["Extinction Coefficient"]) else "-",
        f'{row["Quantum Yield"]:.2f}' if not pd.isna(row["Quantum Yield"]) else "-",
        f'{row["Brightness"]:.1f}' if not pd.isna(row["Brightness"]) else "-",
        f'{row["pKa"]:.1f}' if not pd.isna(row["pKa"]) else "-",
    ])
add_table(doc, ["Variant", "Ex(nm)", "Em(nm)", "Ext.Coeff", "QY", "Brightness", "pKa"], opt_rows)
doc.add_paragraph()

doc.add_heading("4.2 Spectral Diversity of GFP Family", level=2)
doc.add_paragraph(
    "Spectral coverage of GFP and derived fluorescent proteins:\n"
    "* Blue (BFP): Ex 380-400nm / Em 440-460nm (e.g., EBFP2.0, Azurite)\n"
    "* Cyan (CFP): Ex 430-450nm / Em 470-490nm (e.g., Cerulean, mTurquoise)\n"
    "* Green (GFP): Ex 470-510nm / Em 500-530nm (e.g., sfGFP, EGFP, mNeonGreen)\n"
    "* Yellow (YFP): Ex 510-520nm / Em 525-540nm (e.g., Venus, Citrine, mClover3)\n"
    "* Red (RFP): Ex 550-590nm / Em 580-620nm (e.g., mCherry, mRuby)\n\n"
    "Key Design Parameters:\n"
    "* Stokes shift: Small Stokes shift (~10-15nm) = higher photon efficiency but harder "
    "excitation/emission separation\n"
    "* Large Stokes shift variants (e.g., Sapphire: ~112nm) are valuable for multi-color imaging\n"
    "* pKa indicates pH sensitivity: low pKa (~4-6) suitable for acidic environments\n"
    "* Quantum yield (QY) x Extinction coefficient (EC) = Brightness; both need coordinated optimization"
)

doc.add_heading("4.3 mNeonGreen Case Study", level=2)
doc.add_paragraph(
    "mNeonGreen is one of the most successful GFP engineering cases in recent years, derived from "
    "Branchiostoma lanceolatum tetrameric fluorescent protein monomerization. Notably, "
    "its sequence is completely different from avGFP (different origin), but it significantly "
    "surpasses traditional GFP variants in brightness and photostability.\n\n"
    "Key features of mNeonGreen:\n"
    "* Brightness: ~92.5 (EC x QY), approximately 3x that of avGFP\n"
    "* Ex/Em: 506/517nm, compatible with EGFP filter sets\n"
    "* Monomeric, suitable for fusion proteins\n"
    "* Superior photostability compared to EGFP\n\n"
    "Insight: Not limited to the avGFP scaffold - exploring fluorescent protein homologs from "
    "other species may yield better design starting points."
)
doc.add_page_break()

# ============ SECTION 5: THERMAL STABILITY ============
doc.add_heading("5. Thermal Stability Analysis", level=1)

doc.add_paragraph(
    "Thermal stability is another critical dimension of protein design, directly affecting "
    "protein folding efficiency, expression levels, and application scope. High-temperature "
    "tolerance is an important indicator for industrial applications and long-term imaging."
)

doc.add_heading("5.1 Thermal Stability Data Summary", level=2)
thermal_clean = thermal_data.copy()
thermal_clean["Tm_C"] = pd.to_numeric(thermal_clean["Tm_C"], errors="coerce")
thermal_clean["Delta_Tm_C"] = pd.to_numeric(thermal_clean["Delta_Tm_C"], errors="coerce")

therm_rows = []
for _, row in thermal_clean.iterrows():
    therm_rows.append([
        str(row.get("Protein_Variant", "-")),
        f'{row["Tm_C"]:.0f}' if not pd.isna(row["Tm_C"]) else "-",
        f'{row["Delta_Tm_C"]:.1f}' if not pd.isna(row["Delta_Tm_C"]) else "-",
        str(row.get("Key_Mutations", "-")),
        str(row.get("Experimental_Method", "-")),
    ])
add_table(doc, ["Protein Variant", "Tm(C)", "dTm(C)", "Key Mutations", "Method"], therm_rows)
doc.add_paragraph()

doc.add_heading("5.2 Thermal Stability and Folding Relationship", level=2)
doc.add_paragraph(
    "Key Findings:\n"
    "* WT avGFP has a Tm of approximately 78C, indicating high intrinsic thermal stability\n"
    "* sfGFP increases Tm to ~83C through multiple folding-enhancing mutations "
    "(S30R/Y39N/F64L/S65T/F99S/M153T/V163A), while significantly improving folding efficiency "
    "and expression levels\n"
    "* S65T mutation appears in both EGFP and sfGFP - it not only optimizes spectral properties "
    "but also improves folding efficiency at 37C\n"
    "* F64L and S65T are core EGFP mutations: enhance 37C folding, red-shift excitation to 488nm\n"
    "* V163A and M153T enhance 37C folding efficiency, forming the core of cycle 3 GFP\n\n"
    "Design Insights:\n"
    "* Thermal stability and folding efficiency are positively correlated - increasing Tm "
    "generally means better recombinant expression\n"
    "* When optimizing spectral properties, monitor and compensate for negative effects "
    "on thermal stability\n"
    "* Additive effects of multi-site mutations need careful evaluation - some beneficial "
    "spectral mutations may reduce stability"
)
doc.add_page_break()

# ============ SECTION 6: MUTATION LANDSCAPE ============
doc.add_heading("6. Mutation Landscape Analysis", level=1)

doc.add_paragraph(
    "Deep understanding of the avGFP (P42212) mutation landscape is crucial for rational design. "
    "This section integrates 58 experimentally validated mutation records from UniProt, "
    ">50,000 mutant brightness data from Sarkisyan et al., and key mutation effects from literature."
)

doc.add_heading("6.1 Key Mutation Site Functional Map", level=2)
key_muts = [
    ["S30R", "Beta-sheet S1", "Folding enhancement",
     "Forms new salt bridge network (E32-R30-E17-R122-E115), folding rate +3.5x"],
    ["Y39N", "Beta-sheet", "Expression optimization",
     "Codon optimization, improved folding kinetics"],
    ["F46L", "Beta-sheet", "Folding enhancement",
     "Venus/R10-3 key mutation, improves maturation efficiency"],
    ["F64L", "Near chromophore", "Folding + 37C maturation",
     "Core EGFP mutation, enhances 37C folding"],
    ["S65T/G/A", "Chromophore core", "Spectral optimization",
     "T=EGFP(488nm excitation); G=red-shift variants; A=fluorescence enhancement"],
    ["Y66H/W/F", "Chromophore core", "Spectral redirection",
     "H=BFP(blue); W=CFP(cyan); F=intermediate spectrum"],
    ["V68L", "Near chromophore", "Folding synergy",
     "Core mutation in GFPmut2/EYFP/Venus, synergizes with S65T/S72A"],
    ["Q69H/L/M", "Near chromophore", "Lifetime/Folding",
     "H=fluorescence lifetime change; L=RSGFP4; M=Citrinine"],
    ["S72A", "Near chromophore", "Spectrum + Folding",
     "Core mutation in multiple optimized variants (EYFP/Cerulean/Venus)"],
    ["F99S", "Surface loop", "Folding enhancement",
     "Core mutation of cycle 3 GFP/alphaGFP, improves 37C folding"],
    ["Y145F/H/A/C", "Chromophore env", "Spectral tuning",
     "F=Sapphire; H=fluorescence lifetime; A=Cerulean; C=lifetime change"],
    ["N146I", "Chromophore env", "Folding enhancement",
     "Core ECFP/Cerulean mutation"],
    ["H148D", "Chromophore env", "Spectral optimization",
     "Cerulean brightness enhancement"],
    ["N149K", "Surface", "Brightness enhancement",
     "Esmerald/VisGreen mutation, enhances 37C fluorescence"],
    ["M153T", "Surface loop", "Folding enhancement",
     "Core mutation in cycle 3 GFP/alphaGFP/ECFP/Venus"],
    ["V163A", "Surface", "Folding enhancement",
     "Core mutation in GFPA/GFPB/ECFP/Cerulean/Venus"],
    ["I167T/V", "Beta-sheet", "Folding + Brightness",
     "T=Esmerald/VisGreen; V=R10-3 red variant"],
    ["S175G", "Surface loop", "Folding enhancement",
     "GFPA/Venus mutation"],
    ["T203Y/I/H/W/F", "Chromophore env", "Spectral core",
     "Y=YFP variants; I=Sapphire(large Stokes shift); H/W/F=red-shifted"],
    ["A206K", "Dimer interface", "Monomerization",
     "Eliminates dimerization, produces monomeric fluorescent protein (mGFP)"],
    ["E222G", "Chromophore env", "Spectral tuning",
     "Suppresses 399nm excitation peak, retains 475nm peak"],
]
add_table(doc, ["Mutation", "Region", "Primary Effect", "Detailed Mechanism"], key_muts)
doc.add_paragraph()

doc.add_heading("6.2 Mutation Effect Principles", level=2)
doc.add_paragraph(
    "Through systematic analysis of all known mutation data, the following key principles are summarized:\n\n"
    "1. The chromophore region (positions 65-67) is the master spectral switch:\n"
    "   - S65T -> excitation peak red-shifts to 488nm (one of the most important single mutations)\n"
    "   - Y66X -> changes fluorescence color (H=blue, W=cyan, F=spectral intermediate)\n"
    "   - Adjacent positions (64,68,69,72,145,146,148,203) fine-tune the chromophore environment\n\n"
    "2. Folding-enhancing mutations concentrate in specific regions:\n"
    "   - Beta-sheet core: S30R, Y39N, F46L\n"
    "   - Surface loops: F99S, M153T, S175G\n"
    "   - Near chromophore: F64L, S65T, V68L, S72A, V163A\n\n"
    "3. Most single-site mutations negatively impact brightness:\n"
    "   - >50,000 variants show median brightness is ~35% of WT\n"
    "   - However, certain combinations (e.g., EGFP, sfGFP mutation sets) can significantly "
    "surpass WT brightness\n"
    "   - Epistasis exists between beneficial mutations; independent combination != additive effect\n\n"
    "4. Dimerization vs. monomerization balance:\n"
    "   - A206K effectively eliminates dimerization, essential for fusion protein design\n"
    "   - However, dimerization naturally contributes to stability; compensatory mutations may "
    "be needed after monomerization"
)
doc.add_page_break()

# ============ SECTION 7: CANDIDATE DESIGN POSITIONS ============
doc.add_heading("7. Candidate Design Position Analysis", level=1)

doc.add_paragraph(
    "Based on expert-curated 45 candidate design positions, combined with mutation data and "
    "literature knowledge, design guidance is provided with importance level classification."
)

doc.add_heading("7.1 High-Priority Candidate Positions", level=2)
high_sites = [
    ["S30", "Beta-sheet S2", "Critical", "S30R", "Avoid Pro/Gly",
     "Forms salt bridge network, folding +3.5x"],
    ["F64", "Near chromophore", "Critical", "F64L", "Avoid Pro",
     "Enhances 37C folding, core EGFP mutation"],
    ["S65", "Chromophore core", "Critical", "S65T/A/G", "Avoid Pro/Gly",
     "Spectral red-shift + fluorescence enhancement, most important single mutation"],
    ["Y66", "Chromophore core", "Critical", "Y66H/W/F", "Avoid disrupting chromophore",
     "Spectral color determining position"],
    ["V68", "Near chromophore", "Critical", "V68L/N", "Avoid Pro",
     "Synergizes with S65T for fluorescence and folding"],
    ["S72", "Near chromophore", "Critical", "S72A", "Avoid Pro/Gly",
     "Core mutation in multiple optimized GFPs (YFP/CFP/Venus)"],
    ["Y145", "Chromophore env", "Important", "Y145F/H/A", "Avoid disrupting fluorescence",
     "Affects Stokes shift and quantum yield"],
    ["N146", "Chromophore env", "Important", "N146I", "Avoid large side chains",
     "Core ECFP/Cerulean mutation"],
    ["H148", "Chromophore env", "Important", "H148D", "Avoid hydrophobic",
     "Cerulean quantum yield enhancement"],
    ["M153", "Surface loop", "Important", "M153T", "Avoid Pro",
     "Cycle 3 GFP folding-enhancing mutation"],
    ["V163", "Surface/Interior", "Important", "V163A", "Avoid large side chains",
     "Core mutation in multiple variants (ECFP/Cerulean/Venus)"],
    ["T203", "Chromophore env", "Critical", "T203Y/I", "Choose based on need",
     "Y=YFP series (yellow); I=Sapphire (large Stokes shift)"],
    ["A206", "Dimer interface", "Critical", "A206K", "Required for monomer",
     "Key monomerization mutation"],
]
add_table(doc, ["Position", "Region", "Importance", "Known Beneficial", "Avoid", "Description"], high_sites)
doc.add_paragraph()

doc.add_heading("7.2 Design Position Tiered Strategy", level=2)
doc.add_paragraph(
    "Based on analysis results, design positions are recommended to be divided into three tiers:\n\n"
    "Tier 1 - Core Spectral Positions (Required):\n"
    "* S65, Y66, T203: Determine fluorescence color and basic spectral properties\n"
    "* Recommend first determining target spectrum (green/yellow/cyan), then selecting mutations "
    "at these positions\n\n"
    "Tier 2 - Folding-Enhancing Positions (Recommended):\n"
    "* S30R, F64L, V68L, S72A, F99S, M153T, V163A, S175G\n"
    "* These mutations significantly improve folding efficiency and expression levels\n"
    "* Recommend selecting from sfGFP mutation set\n\n"
    "Tier 3 - Fine-Tuning Positions (Advanced):\n"
    "* Select based on specific requirements:\n"
    "  - pH sensitivity: adjust pKa-related positions\n"
    "  - Monomerization: A206K\n"
    "  - Photostability: I167T, etc.\n"
    "  - Quantum yield optimization: H148D, etc."
)
doc.add_page_break()

# ============ SECTION 8: LITERATURE INSIGHTS ============
doc.add_heading("8. Literature Knowledge Integration", level=1)

doc.add_paragraph(
    "The following summarizes GFP protein design knowledge from multiple key publications, "
    "providing experimentally validated support for design decisions."
)

doc.add_heading("8.1 Classic GFP Engineering Milestones", level=2)
milestones = [
    ["1994-1996", "WT GFP Characterization",
     "avGFP sequence, crystal structure and basic spectral properties elucidated. "
     "S65T mutation found to red-shift excitation to 488nm",
     "Tsien, Prasher, Remington et al."],
    ["1996-1998", "Color Expansion",
     "Y66H->BFP(blue); Y66W->CFP(cyan); T203Y->YFP(yellow). "
     "Established the GFP color palette",
     "Tsien lab (Roger Tsien, 2008 Nobel Prize in Chemistry)"],
    ["1998-2000", "Folding Optimization (cycle 3 GFP/EGFP)",
     "F99S/M153T/V163A->cycle 3 GFP; F64L/S65T->EGFP. Improved 37C folding",
     "Crameri, Cormack, Tsien et al."],
    ["2001-2005", "Monomerization and Venus/Cerulean",
     "A206K monomerization; Venus(YFP optimized) and Cerulean(CFP optimized). "
     "Quantum yield improved 2.5-fold",
     "Miyawaki, Piston, Campbell et al."],
    ["2006", "sfGFP (superfolder GFP)",
     "S30R/Y39N/F64L/S65T/F99S/M153T/V163A combination -> superfolder GFP. "
     "Denaturation-resistant, high expression",
     "Pedelacq, Waldo et al."],
    ["2013-2019", "mNeonGreen",
     "Completely different GFP scaffold from different species. "
     "Brightness >> EGFP, excellent photostability",
     "Shaner et al."],
]
add_table(doc, ["Period", "Milestone", "Core Contribution", "Key Researchers"], milestones)
doc.add_paragraph()

doc.add_heading("8.2 Key Literature Highlights", level=2)
lit_highlights = [
    ("Sarkisyan et al. (2016) Nature",
     "Conducted systematic mutation scanning of avGFP, generating a fluorescence dataset of "
     ">50,000 variants. Core findings: (1) GFP is sensitive to most single-site mutations, "
     "only ~10% of mutations maintain >50% WT fluorescence; (2) Significant epistasis exists "
     "between beneficial mutations; (3) Multi-site mutation brightness is not a simple sum of "
     "independent mutations. Design implication: multiple sites need simultaneous optimization; "
     "single-site optimization alone is insufficient for optimal results."),
    ("Pedelacq et al. (2006) Nature Biotechnology",
     "Developed superfolder GFP (sfGFP) achieving superfolding properties through 6 key mutations. "
     "sfGFP can efficiently refold even under denaturing conditions, making it one of the most "
     "widely used GFP variants. Core mutations: S30R/Y39N/F64L/S65T/F99S/M153T/V163A."),
    ("Tsien (1998) Annual Review of Biochemistry",
     "Systematically summarized the chromophore formation mechanism and spectral regulation "
     "principles of GFP. Key understanding: GFP fluorescence originates from autocatalytic "
     "cyclization of the Ser65-Tyr66-Gly67 tripeptide, forming p-hydroxybenzylideneimidazolinone "
     "(p-HBI) chromophore. This provides the chemical basis for designing new variants."),
    ("Cormack et al. (1996) Gene",
     "Developed EGFP (Enhanced GFP) with F64L/S65T double mutation. A milestone in GFP engineering: "
     "first adjustment of excitation peak to 488nm (argon ion laser line), with enhanced 37C "
     "folding efficiency."),
    ("Shaner et al. (2013) Nature Methods",
     "Reported mNeonGreen, a completely different GFP homolog from Branchiostoma. Obtained through "
     "tetramer-to-monomer engineering, brightness is 3x that of EGFP. Insight: exploring "
     "fluorescent proteins from other species in nature may find better design starting points."),
]
for t, d in lit_highlights:
    p = doc.add_paragraph()
    run = p.add_run(f"* {t}")
    run.bold = True
    p.add_run(f"\n{d}")
doc.add_page_break()

# ============ SECTION 9: DESIGN STRATEGY ============
doc.add_heading("9. Protein Design Strategy Recommendations", level=1)

doc.add_paragraph(
    "Synthesizing all data analysis results, the following strategy recommendations are "
    "provided for this protein design project. These recommendations are based on "
    "experimentally validated data and known mutation effect principles."
)

doc.add_heading("9.1 Recommended Design Workflow", level=2)
steps = [
    ("Step 1: Define Design Goals",
     "Define target fluorescence color (green/yellow/cyan), application conditions "
     "(pH range, temperature), and performance requirements (brightness, photostability, "
     "monomeric/oligomeric state)."),
    ("Step 2: Select Starting Scaffold",
     "Green fluorescence: Recommend sfGFP (highest folding efficiency) or EGFP (488nm optimized)\n"
     "Yellow fluorescence: Recommend Venus (fast maturation, environmentally insensitive)\n"
     "Cyan fluorescence: Recommend Cerulean (high quantum yield)\n"
     "Maximum brightness: Consider mNeonGreen scaffold"),
    ("Step 3: Design Core Spectral Positions",
     "Fix chromophore region mutations (S65/Y66/T203) based on target color. "
     "Add dimerization interface mutation (A206K) if monomer is needed."),
    ("Step 4: Combine Folding-Enhancing Positions",
     "Select from 7 folding-enhancing positions of sfGFP/Cycle3 GFP: "
     "S30R, F64L, V68L, S72A, F99S, M153T, V163A"),
    ("Step 5: Fine-Tuning",
     "Based on specific needs: pKa adjustment, quantum yield enhancement (N149K, H148D), "
     "photostability (I167T), etc."),
    ("Step 6: Computational Validation",
     "Use structure prediction tools (e.g., AlphaFold2) to predict the structure of designed "
     "sequences and verify that mutations do not disrupt folding."),
    ("Step 7: Experimental Validation",
     "HTP (high-throughput) synthesis of designed gene sequences -> expression -> "
     "fluorescence detection -> ranking -> iterative optimization"),
]
for t, d in steps:
    p = doc.add_paragraph()
    run = p.add_run(f"{t}: ")
    run.bold = True
    p.add_run(f"\n{d}")
doc.add_paragraph()

doc.add_heading("9.2 Design Parameter Priority Matrix", level=2)
priority_rows = [
    ["Fluorescence Brightness", "Highest", "S65T, S72A, H148D, N149K",
     "Combine validated beneficial mutations, avoid negative interactions"],
    ["Folding Efficiency", "Highest", "S30R, F64L, F99S, M153T, V163A",
     "Adopt the full sfGFP folding-enhancing mutation set"],
    ["Spectral Properties", "Highest", "S65, Y66, T203",
     "Fix core spectral positions based on target color"],
    ["Monomerization", "High", "A206K",
     "Essential for fusion protein applications"],
    ["Thermal Stability", "High", "S30R, V163A, I167T",
     "Strengthen beta-sheet stability and salt bridge networks"],
    ["pH Tolerance", "Medium", "pKa-related positions",
     "Adjust surface charge distribution"],
    ["Photostability", "Medium", "I167T, N149K",
     "Reduce photobleaching, enhance long-term imaging"],
    ["Maturation Rate", "Medium", "F46L, V68L, S72A",
     "Accelerate chromophore cyclization, shorten maturation time"],
]
add_table(doc, ["Parameter", "Priority", "Key Positions", "Optimization Strategy"], priority_rows)
doc.add_paragraph()

doc.add_heading("9.3 Design Avoidance Checklist", level=2)
doc.add_paragraph(
    "The following should be avoided in protein design:\n\n"
    "1. Avoid introducing Pro or Gly in beta-sheet core regions (disrupts secondary structure)\n"
    "2. Avoid non-target mutations in the chromophore tripeptide SYG(65-67) (destroys fluorescence)\n"
    "3. Avoid introducing charged residues in the hydrophobic core (disrupts folding)\n"
    "4. Avoid simultaneously introducing too many beneficial mutations (epistasis may cause negative results)\n"
    "5. Avoid neglecting pKa adjustment (may cause weakening of fluorescence at target pH)\n"
    "6. Avoid using proprietary sequences (e.g., mNeonGreen has patent protection, research use only)\n\n"
    "Special Reminders:\n"
    "* Multi-site mutation combinations require experimental validation; epistasis may make "
    "independently beneficial mutation combinations ineffective\n"
    "* Mutation combinations that decrease thermal stability may fail to express at 37C conditions\n"
    "* Single-site mutation screening (e.g., Sarkisyan data) does not equal optimal multi-site combinations"
)

doc.add_heading("9.4 Recommended Starting Design Templates", level=2)
doc.add_paragraph(
    "Based on different design goals, the following starting mutation combinations are recommended:"
)
templates_rows = [
    ["Ultra-high brightness green", "sfGFP",
     "S30R+Y39N+F64L+S65T+F99S+M153T+V163A+N149K+H148D",
     "Optimal folding + brightness. Expected to exceed WT brightness"],
    ["488nm excitation standard green", "EGFP (F64L/S65T)",
     "+F99S+M153T+V163A+A206K",
     "Classic 488nm GFP. Monomeric, efficient 37C folding"],
    ["Yellow fluorescence fast maturation", "Venus",
     "F46L+F64L+S65G+V68L+S72A+M153T+V163A+S175G+T203Y",
     "Fastest maturation, environmentally insensitive. Ex515/Em528"],
    ["Cyan fluorescence high QY", "Cerulean",
     "F64L+S65T+Y66W+S72A+Y145A+N146I+H148D+M153T+V163A",
     "2.5x brighter than ECFP. Ex433/Em475"],
    ["Large Stokes shift (>100nm)", "Sapphire",
     "S72A+Y145F+T203I",
     "Ex399/Em511. 112nm Stokes shift"],
    ["Low pH adaptation (pH 4-6)", "sfGFP + pKa tuning",
     "sfGFP mutations + surface acidic residue introduction",
     "Maintains fluorescence in acidic environments"],
]
add_table(doc, ["Design Goal", "Starting Scaffold", "Core Mutation Combination", "Expected Effect"], templates_rows)
doc.add_page_break()

# ============ SECTION 10: SUMMARY ============
doc.add_heading("10. Summary and Outlook", level=1)

doc.add_heading("10.1 Data Integration Achievements", level=2)
doc.add_paragraph(
    f"Summary of this data integration work:\n\n"
    f"1. Converted and standardized 10 training data files to CSV format\n"
    f"2. Read and integrated 7 worksheets from comprehensive_GFP_dataset\n"
    f"3. Generated consolidated_GFP_dataset.csv file\n"
    f"4. Dataset scale:\n"
    f"   - {len(brightness_data)} curated brightness variants + {len(sarkisyan_brightness)} complete mutation scan data\n"
    f"   - {len(fpbase_full)} FPbase optical property entries\n"
    f"   - {len(thermal_data)} thermal stability measurements\n"
    f"   - {len(uniprot_mutations)} UniProt mutation annotations\n"
    f"   - 11 reference sequences\n"
    f"   - {len(candidate_positions)} candidate design positions\n"
    f"   - {len(literature_data)} literature knowledge entries"
)

doc.add_heading("10.2 Core Design Principles", level=2)
doc.add_paragraph(
    "Based on comprehensive analysis, 5 core design principles are condensed:\n\n"
    "Principle 1: Scaffold First\n"
    "   Select the most suitable starting scaffold for the design goal (sfGFP/EGFP/Venus/Cerulean). "
    "Different starting scaffolds determine the ceiling of the design.\n\n"
    "Principle 2: Chromophore Protection\n"
    "   The chromophore tripeptide SYG(65-67) and its microenvironment are the core of "
    "fluorescence activity. Any changes must be based on known spectral effect principles.\n\n"
    "Principle 3: Folding and Function - Equal Priority\n"
    "   Beneficial spectral mutations may simultaneously reduce folding efficiency - "
    "synchronously introduce folding-enhancing mutations (S30R, F64L, F99S, M153T, V163A, etc.) "
    "for compensation.\n\n"
    "Principle 4: Epistasis Awareness\n"
    "   Mutation combinations are not linear additions of single-site effects - "
    "computational prediction and experimental validation are needed to confirm combination effects.\n\n"
    "Principle 5: Application-Oriented Design\n"
    "   Adjust design parameter priorities based on final application scenario "
    "(imaging, sensor, fusion protein, industrial catalysis, etc.)."
)

doc.add_heading("10.3 Outlook and Next Steps", level=2)
doc.add_paragraph(
    "1. Computational Design Phase:\n"
    "   - Use AlphaFold2/Rosetta to predict 3D structures of designed sequences\n"
    "   - Molecular dynamics simulations to evaluate mutation effects on protein dynamics\n"
    "   - Machine learning models based on existing data to predict brightness/stability\n\n"
    "2. Experimental Validation Phase:\n"
    "   - Synthesize Top 10-20 designed sequences for experimental validation\n"
    "   - Establish high-throughput screening pipeline (96-well plate -> fluorescence plate reader)\n"
    "   - Conduct detailed biochemical characterization of best candidates\n\n"
    "3. Iterative Optimization:\n"
    "   - Feed experimental data back into the design model\n"
    "   - Further optimize using directed evolution/saturation mutagenesis\n"
    "   - Customize optimization for specific application scenarios"
)

# ============ APPENDIX ============
doc.add_page_break()
doc.add_heading("Appendix: Data Source Summary", level=1)
appendix_rows = [
    [1, "GFP_data.xlsx (brightness)", "XLSX", "141,572", "Mutation-brightness mapping"],
    [2, "amino_acid_genotypes_to_brightness.tsv", "TSV", "54,025", "Sarkisyan 2016 full dataset"],
    [3, "fpbase_proteins_basic_gfp.csv", "CSV", "1,110", "FPbase: optical properties"],
    [4, "AAseqs of 5 GFP proteins.txt", "FASTA", "5", "Reference protein sequences"],
    [5, "P42212_gfp_aaMutations_Uniport.json", "JSON", "58", "UniProt: mutation annotations"],
    [6, "4_P42212_gfp_protherm.csv", "CSV", "4", "ProTherm: thermal stability"],
    [7, "9temp.tsv", "TSV", "4", "ProTherm: thermal stability (TSV)"],
    [8, "Beta amyloid peptide-gfp_Thermo.csv", "CSV", "1", "Beta-amyloid GFP thermal data"],
    [9, "comprehensive_GFP_dataset.xlsx", "XLSX", "7 sheets", "Curated comprehensive dataset"],
    [10, "submission_template.csv", "CSV", "6", "Competition submission template"],
]
add_table(doc, ["#", "Source File", "Type", "Records", "Key Content"], appendix_rows)

# SAVE
output_path = os.path.join(OUTPUT, "GFP_Protein_Design_Analysis_Report.docx")
doc.save(output_path)
print(f"\n{'='*60}")
print(f"Report saved to: {output_path}")
print(f"{'='*60}")
